from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.utils import timezone
from django.contrib import messages
from django.http import FileResponse, Http404, JsonResponse, HttpResponse, HttpResponseNotFound
from django.core.paginator import Paginator
from django.db.models import Q, Count
from PROTECHAPP.models import CustomUser, Student, Section, Grade, Guardian, Attendance, UserRole, UserStatus, ExcusedAbsence, AdvisoryAssignment
from django.views.decorators.http import require_http_methods, require_GET
from django.contrib.auth.hashers import make_password
import json
import os
import io
from datetime import datetime
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from django.views.decorators.csrf import csrf_exempt

# Define profile pics directory
PROFILE_PICS_DIR = os.path.join(settings.BASE_DIR, 'private_profile_pics')
os.makedirs(PROFILE_PICS_DIR, exist_ok=True)

# Define excuse letters directory
PRIVATE_EXCUSE_LETTERS_DIR = os.path.join(settings.BASE_DIR, 'private_excuse_letters')
os.makedirs(PRIVATE_EXCUSE_LETTERS_DIR, exist_ok=True)

def is_registrar(user):
    """Check if the logged-in user is a registrar"""
    return user.is_authenticated and user.role == UserRole.REGISTRAR

def is_admin_or_registrar(user):
    """Check if the logged-in user is an admin or registrar"""
    return user.is_authenticated and user.role in [UserRole.ADMIN, UserRole.REGISTRAR]

def get_pagination_range(paginator, current_page, num_pages=5):
    """Helper function to get page range for pagination UI"""
    total_pages = paginator.num_pages
    start_page = max(1, current_page - num_pages // 2)
    end_page = min(total_pages, start_page + num_pages - 1)
    
    if end_page - start_page < num_pages - 1:
        start_page = max(1, end_page - num_pages + 1)
    
    return list(range(start_page, end_page + 1))

@login_required
@user_passes_test(is_registrar)
def registrar_dashboard(request):
    """View for registrar dashboard"""
    current_date = timezone.now()
    
    # Example data for dashboard
    student_count = Student.objects.count()
    face_enrolled_count = Student.objects.filter(face_path__isnull=False).count()
    pending_enrollments = 42  # This would be calculated based on actual data
    
    face_enrolled_percentage = (face_enrolled_count / student_count * 100) if student_count > 0 else 0
    
    context = {
        'current_date': current_date,
        'student_count': student_count,
        'face_enrolled_count': face_enrolled_count,
        'face_enrolled_percentage': round(face_enrolled_percentage),
        'pending_enrollments': pending_enrollments
    }
    
    return render(request, 'registrar/dashboard.html', context)



@login_required
@user_passes_test(is_registrar)
def registrar_announcements(request):
    """View for announcements"""
    return render(request, 'registrar/announcements.html')

@login_required
@user_passes_test(is_registrar)
def registrar_messages(request):
    """View for messages"""
    from django.conf import settings
    context = {
        'firebase_config': settings.FIREBASE_WEB_CONFIG
    }
    return render(request, 'registrar/messages.html', context)

@login_required
@user_passes_test(is_registrar)
def registrar_settings(request):
    """View for settings"""
    return render(request, 'registrar/settings.html')


# ==========================
#  GRADE MANAGEMENT VIEWS
# ==========================

@login_required
@user_passes_test(is_registrar)
def registrar_grades(request):
    """View for grade management"""
    # Get query parameters
    search_query = request.GET.get('search', '')
    page_number = request.GET.get('page', 1)
    
    # Base queryset - order by created_at DESC (newest first)
    grades = Grade.objects.all().order_by('-created_at')
    
    # Apply search if provided
    if search_query:
        grades = grades.filter(name__icontains=search_query)
    
    # Annotate with section and student counts
    grades = grades.annotate(
        sections_count=Count('sections'),
        students_count=Count('sections__students', distinct=True)
    )
    
    # Get stats for dashboard
    total_grades = Grade.objects.count()
    grades_with_sections = Grade.objects.filter(sections__isnull=False).distinct().count()
    total_sections = Section.objects.count()
    total_students = Student.objects.count()
    
    # Pagination
    paginator = Paginator(grades, 10)
    page_obj = paginator.get_page(page_number)
    
    # Calculate page range for pagination UI
    page_range = get_pagination_range(paginator, page_obj.number, 5)
    
    context = {
        'grades': page_obj,
        'search_query': search_query,
        'total_grades': total_grades,
        'grades_with_sections': grades_with_sections,
        'total_sections': total_sections,
        'total_students': total_students,
        'page_range': page_range,
    }
    return render(request, 'registrar/grades.html', context)

@require_http_methods(["GET"])
@login_required
@user_passes_test(is_registrar)
def registrar_search_grades(request):
    """AJAX view for searching grades with pagination"""
    try:
        # Get query parameters
        search_query = request.GET.get('search', '')
        page_number = request.GET.get('page', 1)
        items_per_page = int(request.GET.get('items_per_page', 10))
        
        # Base queryset - order by created_at DESC (newest first)
        grades = Grade.objects.all().order_by('-created_at')
        
        # Apply search if provided
        if search_query:
            grades = grades.filter(name__icontains=search_query)
        
        # Annotate with section and student counts
        grades = grades.annotate(
            sections_count=Count('sections'),
            students_count=Count('sections__students', distinct=True)
        )
        
        # Pagination
        paginator = Paginator(grades, items_per_page)
        page_obj = paginator.get_page(page_number)
        
        # Prepare data for JSON response
        grades_data = []
        for grade in page_obj:
            grades_data.append({
                'id': grade.id,
                'name': grade.name,
                'description': grade.description or '',
                'sections_count': grade.sections_count,
                'students_count': grade.students_count,
                'created_at': grade.created_at.strftime('%b %d, %Y'),
            })
        
        return JsonResponse({
            'success': True,
            'grades': grades_data,
            'pagination': {
                'current_page': page_obj.number,
                'total_pages': paginator.num_pages,
                'total_items': paginator.count,
                'has_previous': page_obj.has_previous(),
                'has_next': page_obj.has_next(),
                'previous_page': page_obj.previous_page_number() if page_obj.has_previous() else None,
                'next_page': page_obj.next_page_number() if page_obj.has_next() else None,
            }
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@require_http_methods(["POST"])
@login_required
@user_passes_test(is_registrar)
def registrar_create_grade(request):
    """AJAX view for creating a new grade"""
    try:
        data = json.loads(request.body)
        
        # Validate required fields
        if not data.get('name'):
            return JsonResponse({
                'success': False,
                'error': 'Grade name is required'
            }, status=400)
        
        # Check if grade with this name already exists
        if Grade.objects.filter(name=data['name']).exists():
            return JsonResponse({
                'success': False,
                'error': 'A grade with this name already exists'
            }, status=400)
        
        # Create the grade
        grade = Grade.objects.create(
            name=data['name'],
            description=data.get('description', '')
        )
        
        messages.success(request, f'Grade "{grade.name}" created successfully!')
        
        return JsonResponse({
            'success': True,
            'message': f'Grade "{grade.name}" created successfully!',
            'grade': {
                'id': grade.id,
                'name': grade.name,
                'description': grade.description or '',
                'sections_count': 0,
                'students_count': 0,
                'created_at': grade.created_at.strftime('%b %d, %Y'),
            }
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@require_http_methods(["POST"])
@login_required
@user_passes_test(is_registrar)  
def registrar_update_grade(request, grade_id):
    """AJAX view for updating a grade"""
    try:
        grade = get_object_or_404(Grade, id=grade_id)
        data = json.loads(request.body)
        
        # Validate required fields
        if not data.get('name'):
            return JsonResponse({
                'success': False,
                'error': 'Grade name is required'
            }, status=400)
        
        # Check if another grade with this name already exists
        if Grade.objects.filter(name=data['name']).exclude(id=grade_id).exists():
            return JsonResponse({
                'success': False,
                'error': 'A grade with this name already exists'
            }, status=400)
        
        # Update the grade
        grade.name = data['name']
        grade.description = data.get('description', '')
        grade.save()
        
        messages.success(request, f'Grade "{grade.name}" updated successfully!')
        
        return JsonResponse({
            'success': True,
            'message': f'Grade "{grade.name}" updated successfully!'
        })
        
    except Grade.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Grade not found'
        }, status=404)
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@require_http_methods(["DELETE"])
@login_required
@user_passes_test(is_registrar)
def registrar_delete_grade(request, grade_id):
    """AJAX view for deleting a grade"""
    try:
        grade = get_object_or_404(Grade, id=grade_id)
        
        # Check if grade has sections
        if grade.sections.exists():
            return JsonResponse({
                'success': False,
                'error': 'Cannot delete grade with existing sections'
            }, status=400)
        
        grade_name = grade.name
        grade.delete()
        
        messages.success(request, f'Grade "{grade_name}" deleted successfully!')
        
        return JsonResponse({
            'success': True,
            'message': f'Grade "{grade_name}" deleted successfully!'
        })
        
    except Grade.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Grade not found'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@require_http_methods(["GET"])
@login_required
@user_passes_test(is_registrar)
def registrar_get_grade_sections(request, grade_id):
    """AJAX view for getting sections of a specific grade"""
    try:
        grade = get_object_or_404(Grade, id=grade_id)
        sections = grade.sections.all().annotate(
            students_count=Count('students')
        )
        
        sections_data = []
        for section in sections:
            sections_data.append({
                'id': section.id,
                'name': section.name,
                'capacity': section.capacity,
                'students_count': section.students_count,
                'room_number': section.room_number or '',
            })
        
        return JsonResponse({
            'success': True,
            'grade_name': grade.name,
            'sections': sections_data
        })
        
    except Grade.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Grade not found'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


# ==========================
#  SECTION MANAGEMENT VIEWS - EXACT CARBON COPY OF ADMIN
# ==========================

@login_required
@user_passes_test(is_registrar)
def registrar_sections(request):
    """View for section management"""
    # Get query parameters
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    advisor_filter = request.GET.get('advisor', '')
    page_number = request.GET.get('page', 1)
    
    # Base queryset - order by created_at DESC (newest first)
    sections = Section.objects.select_related('grade').order_by('-created_at')
    
    # Apply search if provided
    if search_query:
        sections = sections.filter(
            Q(name__icontains=search_query) | 
            Q(grade__name__icontains=search_query)
        )
    
    # Apply grade filter if provided
    if grade_filter:
        try:
            grade_id = int(grade_filter)
            sections = sections.filter(grade_id=grade_id)
        except (ValueError, TypeError):
            pass
    
    # Apply advisor filter if provided
    if advisor_filter:
        if advisor_filter == 'with_advisor':
            sections = sections.filter(advisors__isnull=False).distinct()
        elif advisor_filter == 'without_advisor':
            sections = sections.filter(advisors__isnull=True)
    
    # Annotate with student counts and teacher info
    sections = sections.annotate(
        students_count=Count('students')
    ).prefetch_related('advisors')
    
    # Get stats for dashboard
    total_sections = Section.objects.count()
    sections_with_students = Section.objects.filter(students__isnull=False).distinct().count()
    sections_with_advisors = Section.objects.filter(advisors__isnull=False).distinct().count()
    total_students = Student.objects.count()
    
    # Get all grades for filter dropdown
    all_grades = Grade.objects.all().order_by('name')
    
    # Pagination
    paginator = Paginator(sections, 10)
    page_obj = paginator.get_page(page_number)
    
    # Calculate page range for pagination UI
    page_range = get_pagination_range(paginator, page_obj.number, 5)
    
    context = {
        'sections': page_obj,
        'grades': all_grades,
        'search_query': search_query,
        'grade_filter': grade_filter,
        'advisor_filter': advisor_filter,
        'total_sections': total_sections,
        'sections_with_students': sections_with_students,
        'sections_with_advisors': sections_with_advisors,
        'total_students': total_students,
        'page_range': page_range,
    }
    return render(request, 'registrar/sections.html', context)

@login_required
@user_passes_test(is_registrar)
@require_GET
def registrar_search_sections(request):
    """AJAX search/filter for sections"""
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    advisor_filter = request.GET.get('advisor', '')
    page_number = request.GET.get('page', 1)
    items_per_page = request.GET.get('items_per_page', 10)
    
    # Base queryset - order by created_at DESC (newest first)
    sections = Section.objects.select_related('grade').order_by('-created_at')
    
    # Apply search if provided
    if search_query:
        sections = sections.filter(
            Q(name__icontains=search_query) | 
            Q(grade__name__icontains=search_query)
        )
    
    # Apply grade filter if provided
    if grade_filter:
        try:
            grade_id = int(grade_filter)
            sections = sections.filter(grade_id=grade_id)
        except (ValueError, TypeError):
            pass
    
    # Apply advisor filter if provided
    if advisor_filter:
        if advisor_filter == 'with_advisor':
            sections = sections.filter(advisors__isnull=False).distinct()
        elif advisor_filter == 'without_advisor':
            sections = sections.filter(advisors__isnull=True)
    
    # Annotate with counts and prefetch related data
    sections = sections.annotate(
        students_count=Count('students')
    ).prefetch_related('advisors')
    
    # Get total count for pagination
    total_count = sections.count()
    
    # Parse page number and items_per_page
    try:
        page_number = int(page_number)
        items_per_page = int(items_per_page)
    except (ValueError, TypeError):
        page_number = 1
        items_per_page = 10
    
    # Create paginator
    paginator = Paginator(sections, items_per_page)
    page_obj = paginator.get_page(page_number)

    # Prepare section data for JSON response
    section_data = []
    for section in page_obj:
        # Get advisor info
        advisor = None
        advisor_user = CustomUser.objects.filter(section=section, role=UserRole.TEACHER).first()
        if advisor_user:
            advisor = {
                'id': advisor_user.id,
                'name': f"{advisor_user.first_name} {advisor_user.last_name}",
                'username': advisor_user.username
            }
        
        section_data.append({
            'id': section.id,
            'name': section.name,
            'grade_id': section.grade.id,
            'grade_name': section.grade.name,
            'students_count': section.students_count,
            'advisor': advisor,
            'created_at': section.created_at.strftime('%Y-%m-%d'),
            'created_at_display': section.created_at.strftime('%b %d, %Y'),
        })
    
    # Prepare pagination data
    pagination = {
        'current_page': page_obj.number,
        'num_pages': paginator.num_pages,
        'has_next': page_obj.has_next(),
        'has_previous': page_obj.has_previous(),
        'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
        'start_index': page_obj.start_index(),
        'end_index': page_obj.end_index(),
        'total_count': total_count,
    }
    
    return JsonResponse({
        'status': 'success',
        'sections': section_data,
        'pagination': pagination,
        'total_count': total_count,
    })

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["POST"])
def registrar_create_section(request):
    """Create a new section"""
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        grade_id = data.get('grade_id')
        
        if not name:
            return JsonResponse({'status': 'error', 'message': 'Section name is required.'}, status=400)
        
        if not grade_id:
            return JsonResponse({'status': 'error', 'message': 'Grade is required.'}, status=400)
        
        # Check if grade exists
        try:
            grade = Grade.objects.get(id=grade_id)
        except Grade.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Selected grade does not exist.'}, status=400)
        
        # Check if section with same name exists in the same grade
        if Section.objects.filter(name__iexact=name, grade=grade).exists():
            return JsonResponse({'status': 'error', 'message': f'Section "{name}" already exists in {grade.name}.'}, status=400)
        
        section = Section.objects.create(name=name, grade=grade)
        
        return JsonResponse({
            'status': 'success',
            'message': f'Section "{name}" created successfully in {grade.name}.',
            'section': {
                'id': section.id,
                'name': section.name,
                'grade_id': section.grade.id,
                'grade_name': section.grade.name,
                'students_count': 0,
                'advisor': None,
                'created_at': section.created_at.strftime('%Y-%m-%d'),
                'created_at_display': section.created_at.strftime('%b %d, %Y'),
            }
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["PATCH"])
def registrar_update_section(request, section_id):
    """Update an existing section"""
    try:
        section = get_object_or_404(Section, id=section_id)
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        grade_id = data.get('grade_id')
        
        if not name:
            return JsonResponse({'status': 'error', 'message': 'Section name is required.'}, status=400)
        
        if not grade_id:
            return JsonResponse({'status': 'error', 'message': 'Grade is required.'}, status=400)
        
        # Check if grade exists
        try:
            grade = Grade.objects.get(id=grade_id)
        except Grade.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Selected grade does not exist.'}, status=400)
        
        # Check if section with same name exists in the same grade (excluding current section)
        if Section.objects.filter(name__iexact=name, grade=grade).exclude(id=section_id).exists():
            return JsonResponse({'status': 'error', 'message': f'Section "{name}" already exists in {grade.name}.'}, status=400)
        
        old_name = section.name
        old_grade = section.grade.name
        section.name = name
        section.grade = grade
        section.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Section updated from "{old_name}" ({old_grade}) to "{name}" ({grade.name}) successfully.'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["DELETE"])
def registrar_delete_section(request, section_id):
    """Delete a section"""
    try:
        section = get_object_or_404(Section, id=section_id)
        
        # Check if section has students
        students_count = section.students.count()
        
        if students_count > 0:
            return JsonResponse({
                'status': 'error', 
                'message': f'Cannot delete section "{section.name}" because it has {students_count} student(s). Please move or remove all students first.'
            }, status=400)
        
        section_name = section.name
        grade_name = section.grade.name
        section.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Section "{section_name}" from {grade_name} deleted successfully.'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
@require_GET
def registrar_get_section_students(request, section_id):
    """Get students for a specific section"""
    try:
        section = get_object_or_404(Section, id=section_id)
        students = Student.objects.filter(section=section).order_by('last_name', 'first_name')
        
        students_data = []
        for student in students:
            students_data.append({
                'id': student.id,
                'lrn': student.lrn,
                'first_name': student.first_name,
                'last_name': student.last_name,
                'full_name': f"{student.first_name} {student.last_name}",
                'status': student.status,
                'created_at': student.created_at.strftime('%b %d, %Y'),
            })
        
        return JsonResponse({
            'status': 'success',
            'section': {
                'id': section.id,
                'name': section.name,
                'grade_name': section.grade.name
            },
            'students': students_data,
            'total_students': len(students_data)
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
# ADDITIONAL REGISTRAR VIEWS - Part 2
# This contains the remaining critical registrar views needed

# ==========================
#  STUDENT MANAGEMENT VIEWS
# ==========================

@login_required
@user_passes_test(is_registrar)
def registrar_students(request):
    """View for student management"""
    # Get query parameters
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    page_number = request.GET.get('page', 1)

    # Base queryset
    students = Student.objects.select_related('grade', 'section').all().order_by('-created_at')

    # Apply search
    if search_query:
        students = students.filter(
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(lrn__icontains=search_query) |
            Q(email__icontains=search_query)
        )

    # Apply grade filter
    if grade_filter:
        students = students.filter(grade_id=grade_filter)

    # Apply section filter
    if section_filter:
        students = students.filter(section_id=section_filter)

    # Apply status filter
    if status_filter:
        students = students.filter(status=status_filter)

    # Dashboard stats
    total_students = Student.objects.count()
    active_students = Student.objects.filter(status='ACTIVE').count()
    inactive_students = Student.objects.filter(status='INACTIVE').count()
    face_enrolled_count = Student.objects.exclude(face_path__isnull=True).exclude(face_path__exact='').count()

    # For filter dropdowns
    grades = Grade.objects.all().order_by('name')
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    # Pagination
    paginator = Paginator(students, 10)
    page_obj = paginator.get_page(page_number)
    page_range = get_pagination_range(paginator, page_obj.number, 5)

    context = {
        'students': page_obj,
        'grades': grades,
        'sections': sections,
        'search_query': search_query,
        'grade_filter': grade_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'total_students': total_students,
        'active_students': active_students,
        'inactive_students': inactive_students,
        'face_enrolled_count': face_enrolled_count,
        'page_range': page_range,
    }
    return render(request, 'registrar/students.html', context)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["GET"])
def registrar_search_students(request):
    """AJAX search/filter for students"""
    try:
        search_query = request.GET.get('search', '')
        grade_filter = request.GET.get('grade', '')
        section_filter = request.GET.get('section', '')
        status_filter = request.GET.get('status', '')
        page_number = request.GET.get('page', 1)
        items_per_page = request.GET.get('items_per_page', 10)

        students = Student.objects.select_related('grade', 'section').all().order_by('-created_at')

        if search_query:
            students = students.filter(
                Q(first_name__icontains=search_query) |
                Q(last_name__icontains=search_query) |
                Q(lrn__icontains=search_query)
            )
        if grade_filter:
            students = students.filter(grade_id=grade_filter)
        if section_filter:
            students = students.filter(section_id=section_filter)
        if status_filter:
            students = students.filter(status=status_filter)

        total_count = students.count()
        try:
            page_number = int(page_number)
            items_per_page = int(items_per_page)
        except (ValueError, TypeError):
            page_number = 1
            items_per_page = 10

        paginator = Paginator(students, items_per_page)
        page_obj = paginator.get_page(page_number)

        student_data = []
        for student in page_obj:
            # Complete the missing profile_pic handling code
            profile_pic = None
            if student.profile_pic:
                profile_pic = student.profile_pic

            student_data.append({
                'id': student.id,
                'lrn': student.lrn,
                'first_name': student.first_name,
                'last_name': student.last_name,
                'full_name': f"{student.first_name} {student.last_name}",
                'grade': student.grade.name if student.grade else '',
                'grade_id': student.grade.id if student.grade else None,
                'section': student.section.name if student.section else '',
                'section_id': student.section.id if student.section else None,
                'status': student.status,
                'profile_pic': profile_pic,
                'created_at': student.created_at.strftime('%Y-%m-%d'),
                'created_at_display': student.created_at.strftime('%b %d, %Y'),
            })

        pagination = {
            'current_page': page_obj.number,
            'num_pages': paginator.num_pages,
            'has_next': page_obj.has_next(),
            'has_previous': page_obj.has_previous(),
            'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
            'start_index': page_obj.start_index(),
            'end_index': page_obj.end_index(),
            'total_count': total_count,
        }

        return JsonResponse({
            'status': 'success',
            'students': student_data,
            'pagination': pagination,
            'total_count': total_count,
        })
    except Exception as e:
        import traceback
        return JsonResponse({
            'status': 'error',
            'message': str(e),
            'traceback': traceback.format_exc()
        }, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["POST"])
def registrar_create_student(request):
    """Create a new student"""
    try:
        # Handle form data with file upload
        lrn = request.POST.get('lrn', '')
        first_name = request.POST.get('first_name', '')
        middle_name = request.POST.get('middle_name', '')
        last_name = request.POST.get('last_name', '')
        grade_id = request.POST.get('grade_id', '')
        section_id = request.POST.get('section_id', '')
        status = request.POST.get('status', 'ACTIVE')
        
        # Validate required fields
        if not lrn or not first_name or not last_name or not grade_id or not section_id:
            return JsonResponse({'status': 'error', 'message': 'Missing required fields'}, status=400)
        
        # Check if LRN is unique
        if Student.objects.filter(lrn=lrn).exists():
            return JsonResponse({'status': 'error', 'message': 'LRN already exists'}, status=400)
        
        # Get grade and section objects
        try:
            grade = Grade.objects.get(id=grade_id)
            section = Section.objects.get(id=section_id)
        except (Grade.DoesNotExist, Section.DoesNotExist):
            return JsonResponse({'status': 'error', 'message': 'Invalid grade or section'}, status=400)
        
        # Create the student
        student = Student(
            lrn=lrn,
            first_name=first_name,
            middle_name=middle_name,
            last_name=last_name,
            grade=grade,
            section=section,
            status=status
        )
        
        # Handle profile picture if provided
        if 'profile_pic' in request.FILES:
            profile_pic = request.FILES['profile_pic']
            
            # Validate file size and type
            if profile_pic.size > 2 * 1024 * 1024:  # 2MB limit
                return JsonResponse({'status': 'error', 'message': 'File size must be less than 2MB'}, status=400)
            
            if not profile_pic.content_type.startswith('image/'):
                return JsonResponse({'status': 'error', 'message': 'Only image files are allowed'}, status=400)
            
            # Save the profile picture
            filename = f"student_{int(timezone.now().timestamp())}_{profile_pic.name}"
            fs = FileSystemStorage(location=settings.MEDIA_ROOT + '/private_profile_pics')
            filename = fs.save(filename, profile_pic)
            student.profile_pic = filename
        
        student.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Student {first_name} {last_name} created successfully',
            'student_id': student.id
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
def registrar_get_student(request, student_id):
    """
    Endpoint to retrieve a specific student's data for editing
    """
    try:
        student = get_object_or_404(Student, id=student_id)
        
        # Build the student data for response
        student_data = {
            'id': student.id,
            'lrn': student.lrn,
            'first_name': student.first_name,
            'middle_name': student.middle_name or '',
            'last_name': student.last_name,
            'grade_id': str(student.grade.id) if student.grade else '',
            'grade': student.grade.name if student.grade else '',
            'section_id': str(student.section.id) if student.section else '',
            'section': student.section.name if student.section else '',
            'status': student.status,
            'profile_pic': student.profile_pic if student.profile_pic else '',
        }
        
        return JsonResponse({
            'status': 'success',
            'student': student_data
        })
    except Student.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Student not found'}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["POST", "PATCH"])
def registrar_update_student(request, student_id):
    """Update a student - handles both PATCH and POST requests"""
    try:
        student = get_object_or_404(Student, id=student_id)
        
        # Handle PATCH requests (usually JSON data)
        if request.method == 'PATCH':
            if request.content_type == 'application/json':
                # Handle JSON data
                data = json.loads(request.body)
                
                # Update student fields from JSON data
                if 'lrn' in data:
                    student.lrn = data.get('lrn')
                if 'first_name' in data:
                    student.first_name = data.get('first_name')
                if 'middle_name' in data:
                    student.middle_name = data.get('middle_name', '')
                if 'last_name' in data:
                    student.last_name = data.get('last_name')
                if 'grade_id' in data and data['grade_id']:
                    student.grade = get_object_or_404(Grade, id=data['grade_id'])
                if 'section_id' in data and data['section_id']:
                    student.section = get_object_or_404(Section, id=data['section_id'])
                if 'status' in data:
                    student.status = data.get('status')
                
                student.save()
                
                return JsonResponse({
                    'status': 'success',
                    'message': f'Student {student.first_name} {student.last_name} updated successfully',
                    'student_id': student.id
                })
            else:
                # Unsupported content-type
                return JsonResponse({
                    'status': 'error',
                    'message': 'Unsupported content type. Use multipart/form-data for file uploads or application/json for JSON data.'
                }, status=400)
        
        # Handle POST requests (usually form data with files)
        # Get values from POST data
        lrn = request.POST.get('lrn', student.lrn)
        first_name = request.POST.get('first_name', student.first_name)
        middle_name = request.POST.get('middle_name', '')
        last_name = request.POST.get('last_name', student.last_name)
        grade_id = request.POST.get('grade_id', student.grade.id if student.grade else None)
        section_id = request.POST.get('section_id', student.section.id if student.section else None)
        status = request.POST.get('status', student.status)
        
        # Validate required fields
        if not lrn or not first_name or not last_name or not grade_id or not section_id:
            return JsonResponse({
                'status': 'error',
                'message': 'Missing required fields'
            }, status=400)
        
        # Check if LRN already exists for another student
        if Student.objects.filter(lrn=lrn).exclude(id=student_id).exists():
            return JsonResponse({
                'status': 'error',
                'message': f'Student with LRN {lrn} already exists'
            }, status=400)
        
        # Get grade and section objects
        try:
            grade = Grade.objects.get(id=grade_id)
            section = Section.objects.get(id=section_id)
        except (Grade.DoesNotExist, Section.DoesNotExist):
            return JsonResponse({
                'status': 'error',
                'message': 'Invalid grade or section'
            }, status=400)
        
        # Update student fields
        student.lrn = lrn
        student.first_name = first_name
        student.middle_name = middle_name
        student.last_name = last_name
        student.grade = grade
        student.section = section
        student.status = status
        
        # Handle profile picture if provided
        if 'profile_pic' in request.FILES:
            profile_pic = request.FILES['profile_pic']
            
            # Validate file size and type
            if profile_pic.size > 2 * 1024 * 1024:  # 2MB limit
                return JsonResponse({'status': 'error', 'message': 'File size must be less than 2MB'}, status=400)
            
            if not profile_pic.content_type.startswith('image/'):
                return JsonResponse({'status': 'error', 'message': 'Only image files are allowed'}, status=400)
            
            # Delete old profile picture if exists
            if student.profile_pic:
                old_profile_pic_path = os.path.join(PROFILE_PICS_DIR, student.profile_pic)
                if os.path.exists(old_profile_pic_path):
                    os.remove(old_profile_pic_path)
        
            # Save the new profile picture
            filename = f"student_{int(timezone.now().timestamp())}_{profile_pic.name}"
            fs = FileSystemStorage(location=PROFILE_PICS_DIR)
            filename = fs.save(filename, profile_pic)
            student.profile_pic = filename
        
        student.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Student {first_name} {last_name} updated successfully',
            'student_id': student.id
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@require_http_methods(["DELETE"])
@login_required
@user_passes_test(is_registrar)
def registrar_delete_student(request, student_id):
    """AJAX view for deleting a student"""
    try:
        student = get_object_or_404(Student, id=student_id)
        student_name = f"{student.first_name} {student.last_name}"
        student.delete()
        
        messages.success(request, f'Student "{student_name}" deleted successfully!')
        
        return JsonResponse({
            'success': True,
            'message': f'Student "{student_name}" deleted successfully!'
        })
        
    except Student.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Student not found'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@require_http_methods(["POST"])
@login_required
@user_passes_test(is_registrar)
def registrar_reset_student_password(request, student_id):
    """AJAX view for resetting student password"""
    try:
        student = get_object_or_404(Student, id=student_id)
        
        # Reset password to default
        student.password = make_password('student123')
        student.save()
        
        messages.success(request, f'Password reset for student "{student.first_name} {student.last_name}"!')
        
        return JsonResponse({
            'success': True,
            'message': f'Password reset for student "{student.first_name} {student.last_name}"!'
        })
        
    except Student.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Student not found'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


# ==========================
#  STUDENT EXCEL IMPORT FUNCTIONALITY - COMPLETE SYSTEM
# ==========================

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["GET"])
def registrar_download_student_template(request):
    """Download Excel template for student import with guardian fields"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.utils import get_column_letter
        from django.http import HttpResponse
        import io
        
        # Create workbook
        wb = openpyxl.Workbook()
        
        # Create Students sheet
        ws_students = wb.active
        ws_students.title = "Students"
        
        # Student headers
        student_headers = [
            'ID', 'LRN', 'First Name', 'Middle Name', 'Last Name', 
            'Grade', 'Section', 'Status'
        ]
        
        # Guardian headers  
        guardian_headers = [
            'Guardian First Name', 'Guardian Middle Name', 'Guardian Last Name',
            'Guardian Email', 'Guardian Phone', 'Guardian Relationship'
        ]
        
        all_headers = student_headers + guardian_headers
        
        # Add headers with styling
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        for col, header in enumerate(all_headers, 1):
            cell = ws_students.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            ws_students.column_dimensions[get_column_letter(col)].width = 15
        
        # Add instruction row
        instruction_text = "Leave ID blank for new students. Rows with existing IDs will be SKIPPED (not imported)."
        ws_students.merge_cells('A2:N2')
        instruction_cell = ws_students.cell(row=2, column=1, value=instruction_text)
        instruction_cell.font = Font(italic=True, color="666666")
        instruction_cell.alignment = Alignment(horizontal='left', wrap_text=True)
        ws_students.row_dimensions[2].height = 30
        
        # Add sample data row
        sample_data = [
            '', '123456789012', 'John', 'Middle', 'Doe', 
            'Grade 1', 'A', 'ACTIVE',
            'Jane', '', 'Doe', 'jane.doe@email.com', '09123456789', 'MOTHER'
        ]
        
        for col, data in enumerate(sample_data, 1):
            ws_students.cell(row=3, column=col, value=data)
        
        # Create Reference Data sheet
        ws_ref = wb.create_sheet(title="Reference Data")
        ws_ref.cell(row=1, column=1, value="AVAILABLE GRADES").font = Font(bold=True)
        ws_ref.cell(row=1, column=3, value="SECTIONS BY GRADE").font = Font(bold=True)
        ws_ref.cell(row=1, column=5, value="STATUS OPTIONS").font = Font(bold=True)
        ws_ref.cell(row=1, column=7, value="RELATIONSHIP OPTIONS").font = Font(bold=True)
        
        # Get grades and sections from database
        grades = Grade.objects.all().order_by('name')
        sections = Section.objects.select_related('grade').order_by('grade__name', 'name')
        
        # Add grades
        for i, grade in enumerate(grades, 2):
            ws_ref.cell(row=i, column=1, value=grade.name)
        
        # Add sections organized by grade
        current_row = 2
        for grade in grades:
            # Add grade header
            grade_cell = ws_ref.cell(row=current_row, column=3, value=f"--- {grade.name} ---")
            grade_cell.font = Font(bold=True, color="0000FF")
            current_row += 1
            
            # Add sections for this grade
            grade_sections = sections.filter(grade=grade)
            for section in grade_sections:
                ws_ref.cell(row=current_row, column=3, value=f"  {section.name}")
                current_row += 1
            
            # Add blank line between grades
            current_row += 1
        
        # Add status options
        status_options = ['ACTIVE', 'INACTIVE']
        for i, status in enumerate(status_options, 2):
            ws_ref.cell(row=i, column=5, value=status)
        
        # Add relationship options
        relationship_options = ['FATHER', 'MOTHER', 'GUARDIAN', 'GRANDMOTHER', 'GRANDFATHER', 'AUNT', 'UNCLE', 'SIBLING', 'OTHER']
        for i, rel in enumerate(relationship_options, 2):
            ws_ref.cell(row=i, column=7, value=rel)
        
        # Create Instructions sheet
        ws_instructions = wb.create_sheet(title="Instructions")
        instructions = [
            "STUDENT IMPORT TEMPLATE - INSTRUCTIONS",
            "",
            "IMPORTANT REQUIREMENTS:",
            "1. LRN must be exactly 12 digits and unique",
            "2. First Name and Last Name are required",
            "3. Grade must match exactly from 'Reference Data' sheet",
            "4. Section must match exactly from 'Reference Data' sheet", 
            "5. Status must be either 'ACTIVE' or 'INACTIVE'",
            "",
            "GRADE AND SECTION MATCHING:",
            "• Each section belongs to a specific grade",
            "• Check 'Reference Data' sheet to see sections organized by grade",
            "• You must select a section that belongs to the grade you specified",
            "• Example: If Grade = 'Grade 1', Section must be from Grade 1 sections",
            "",
            "GUARDIAN INFORMATION (OPTIONAL):",
            "• Guardian fields can be left empty",
            "• If providing guardian info, First Name, Last Name, Phone are required",
            "• Guardian Relationship must match options in 'Reference Data' sheet",
            "• Phone should be valid format (09XXXXXXXXX)",
            "",
            "VALIDATION RULES:",
            "• Duplicate LRNs will be rejected",
            "• Invalid Grade/Section combinations will be rejected",
            "• Sections that don't belong to the specified grade will be rejected",
            "• Email format will be validated if provided",
            "• All required fields must be filled",
            "",
            "NOTES:",
            "• Check 'Reference Data' sheet for valid options",
            "• Sections are organized by grade for easy reference",
            "• Sample data is provided in row 2 of Students sheet",
            "• Delete sample data before importing your actual data"
        ]
        
        for i, instruction in enumerate(instructions, 1):
            cell = ws_instructions.cell(row=i, column=1, value=instruction)
            if i == 1:
                cell.font = Font(bold=True, size=14)
            elif instruction.startswith(("IMPORTANT", "GUARDIAN", "VALIDATION", "NOTES")):
                cell.font = Font(bold=True)
        
        # Save to memory
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        # Create response
        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename="student_import_template.xlsx"'
        
        return response
        
    except ImportError:
        return JsonResponse({
            'status': 'error',
            'message': 'Excel functionality not available. Please install openpyxl.'
        }, status=500)
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Error generating template: {str(e)}'
        }, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["POST"])
def registrar_import_students(request):
    """Import students from Excel file with complete validation"""
    try:
        import openpyxl
        import re
        from django.db import transaction
        
        if 'excel_file' not in request.FILES:
            return JsonResponse({
                'status': 'error',
                'message': 'No file uploaded'
            }, status=400)
        
        excel_file = request.FILES['excel_file']
        
        # Validate file type
        if not excel_file.name.endswith(('.xlsx', '.xls')):
            return JsonResponse({
                'status': 'error',
                'message': 'Invalid file type. Please upload Excel (.xlsx or .xls) files only.'
            }, status=400)
        
        # Validate file size (max 5MB)
        if excel_file.size > 5 * 1024 * 1024:
            return JsonResponse({
                'status': 'error',
                'message': 'File too large. Maximum file size is 5MB.'
            }, status=400)
        
        try:
            # Read Excel file
            wb = openpyxl.load_workbook(excel_file)
            if 'Students' not in wb.sheetnames:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Invalid template. "Students" sheet not found.'
                }, status=400)
            
            ws = wb['Students']
            
            # Get header row
            headers = []
            for col in range(1, 15):  # Expecting 14 columns (added ID)
                cell_value = ws.cell(row=1, column=col).value
                if cell_value:
                    headers.append(str(cell_value).strip())
            
            # Validate headers
            expected_headers = [
                'ID', 'LRN', 'First Name', 'Middle Name', 'Last Name', 
                'Grade', 'Section', 'Status',
                'Guardian First Name', 'Guardian Middle Name', 'Guardian Last Name',
                'Guardian Email', 'Guardian Phone', 'Guardian Relationship'
            ]
            
            if headers != expected_headers:
                return JsonResponse({
                    'status': 'error',
                    'message': f'Invalid template headers. Expected: {", ".join(expected_headers)}'
                }, status=400)
            
            # Process rows
            students_data = []
            errors = []
            
            # Get valid grades and sections for validation
            valid_grades = {grade.name: grade for grade in Grade.objects.all()}
            valid_sections = {section.name: section for section in Section.objects.select_related('grade').all()}
            valid_relationships = ['FATHER', 'MOTHER', 'GUARDIAN', 'GRANDMOTHER', 'GRANDFATHER', 'AUNT', 'UNCLE', 'SIBLING', 'OTHER']
            existing_lrns = set(Student.objects.values_list('lrn', flat=True))
            
            for row_num in range(2, ws.max_row + 1):
                row_data = []
                for col in range(1, 15):
                    cell_value = ws.cell(row=row_num, column=col).value
                    row_data.append(str(cell_value).strip() if cell_value else '')
                
                # Skip empty rows
                if not any(row_data[:8]):  # Check if student data is empty
                    continue
                
                # Extract data
                student_id, lrn, first_name, middle_name, last_name, grade_name, section_name, status = row_data[:8]
                guardian_first, guardian_middle, guardian_last, guardian_email, guardian_phone, guardian_relationship = row_data[8:]
                
                row_errors = []
                
                # Check if ID is provided - skip if exists
                if student_id and student_id != 'nan' and student_id != '':
                    try:
                        if Student.objects.filter(id=int(float(student_id))).exists():
                            # Skip rows with existing ID
                            continue
                    except ValueError:
                        pass
                
                # Validate required student fields (only for new students)
                if not lrn:
                    row_errors.append("LRN is required")
                elif not re.match(r'^\d{12}$', lrn):
                    row_errors.append("LRN must be exactly 12 digits")
                elif lrn in existing_lrns:
                    row_errors.append("LRN already exists")
                
                if not first_name:
                    row_errors.append("First Name is required")
                if not last_name:
                    row_errors.append("Last Name is required")
                if not grade_name:
                    row_errors.append("Grade is required")
                elif grade_name not in valid_grades:
                    row_errors.append(f"Invalid Grade '{grade_name}'. Check Reference Data sheet")
                
                if not section_name:
                    row_errors.append("Section is required")
                elif section_name not in valid_sections:
                    row_errors.append(f"Invalid Section '{section_name}'. Check Reference Data sheet")
                elif grade_name and grade_name in valid_grades and section_name in valid_sections:
                    # Validate that section belongs to the specified grade
                    section_obj = valid_sections[section_name]
                    grade_obj = valid_grades[grade_name]
                    if section_obj.grade.id != grade_obj.id:
                        row_errors.append(f"Section '{section_name}' does not belong to grade '{grade_name}'. Please check Reference Data sheet.")
                
                if status not in ['ACTIVE', 'INACTIVE']:
                    row_errors.append("Status must be 'ACTIVE' or 'INACTIVE'")
                
                # Validate guardian data if provided
                guardian_data = None
                has_guardian_data = any([guardian_first, guardian_last, guardian_phone, guardian_relationship])
                
                if has_guardian_data:
                    if not guardian_first:
                        row_errors.append("Guardian First Name is required if guardian data provided")
                    if not guardian_last:
                        row_errors.append("Guardian Last Name is required if guardian data provided")
                    if not guardian_phone:
                        row_errors.append("Guardian Phone is required if guardian data provided")
                    elif not re.match(r'^09\d{9}$', guardian_phone):
                        row_errors.append("Guardian Phone must be valid format (09XXXXXXXXX)")
                    if not guardian_relationship:
                        row_errors.append("Guardian Relationship is required if guardian data provided")
                    elif guardian_relationship not in valid_relationships:
                        row_errors.append(f"Invalid Guardian Relationship '{guardian_relationship}'. Check Reference Data sheet")
                    
                    # Validate email format if provided
                    if guardian_email and not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', guardian_email):
                        row_errors.append("Invalid Guardian Email format")
                    
                    if not row_errors:  # Only create guardian data if no errors
                        guardian_data = {
                            'first_name': guardian_first,
                            'middle_name': guardian_middle,
                            'last_name': guardian_last,
                            'email': guardian_email if guardian_email else '',
                            'phone': guardian_phone,
                            'relationship': guardian_relationship
                        }
                
                if row_errors:
                    errors.append(f"Row {row_num}: {'; '.join(row_errors)}")
                else:
                    students_data.append({
                        'lrn': lrn,
                        'first_name': first_name,
                        'middle_name': middle_name,
                        'last_name': last_name,
                        'grade': valid_grades[grade_name],
                        'section': valid_sections[section_name],
                        'status': status,
                        'guardian': guardian_data
                    })
                    existing_lrns.add(lrn)  # Add to prevent duplicates in same file
            
            # Return errors if any
            if errors:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Validation failed',
                    'errors': errors
                }, status=400)
            
            if not students_data:
                return JsonResponse({
                    'status': 'error',
                    'message': 'No valid student data found in file'
                }, status=400)
            
            # Import data in transaction
            with transaction.atomic():
                created_students = 0
                created_guardians = 0
                
                for data in students_data:
                    # Create new student
                    student = Student.objects.create(
                        lrn=data['lrn'],
                        first_name=data['first_name'],
                        middle_name=data['middle_name'],
                        last_name=data['last_name'],
                        grade=data['grade'],
                        section=data['section'],
                        status=data['status']
                    )
                    created_students += 1
                    
                    # Create guardian if data provided
                    if data['guardian']:
                        Guardian.objects.create(
                            first_name=data['guardian']['first_name'],
                            middle_name=data['guardian']['middle_name'],
                            last_name=data['guardian']['last_name'],
                            email=data['guardian']['email'],
                            phone=data['guardian']['phone'],
                            relationship=data['guardian']['relationship'],
                            student=student
                        )
                        created_guardians += 1
                
                return JsonResponse({
                    'status': 'success',
                    'message': f'Successfully imported {created_students} students and {created_guardians} guardians',
                    'data': {
                        'students_created': created_students,
                        'guardians_created': created_guardians
                    }
                })
        
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': f'Error processing Excel file: {str(e)}'
            }, status=500)
            
    except ImportError:
        return JsonResponse({
            'status': 'error',
            'message': 'Excel functionality not available. Please install openpyxl.'
        }, status=500)
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Import failed: {str(e)}'
        }, status=500)


# ==========================
#  GUARDIAN MANAGEMENT VIEWS - EXACT CARBON COPY OF ADMIN
# ==========================

@login_required
@user_passes_test(is_registrar)
def registrar_guardians(request):
    """Guardian management view with dashboard cards"""
    # Calculate dashboard statistics
    total_guardians = Guardian.objects.count()
    
    # Count guardians with children (since each guardian has exactly one student)
    guardians_with_children = Guardian.objects.exclude(student=None).count()
    
    # Count mother guardians
    mother_count = Guardian.objects.filter(relationship='MOTHER').count()
    
    # Count father guardians
    father_count = Guardian.objects.filter(relationship='FATHER').count()
    
    # Get relationship choices for filter dropdown
    relationship_choices = [
        ('FATHER', 'Father'),
        ('MOTHER', 'Mother'),
        ('GUARDIAN', 'Guardian'),
        ('GRANDMOTHER', 'Grandmother'),
        ('GRANDFATHER', 'Grandfather'),
        ('AUNT', 'Aunt'),
        ('UNCLE', 'Uncle'),
        ('SIBLING', 'Sibling'),
        ('OTHER', 'Other'),
    ]
    
    # Get all grades for filter dropdown
    grades = Grade.objects.all().order_by('name')
    
    # Get all sections for filter dropdown
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')
    
    # Get students for children dropdown
    students = Student.objects.select_related('grade', 'section').filter(status='ACTIVE')
    
    context = {
        'total_guardians': total_guardians,
        'guardians_with_children': guardians_with_children,
        'mother_count': mother_count,
        'father_count': father_count,
        'relationship_choices': relationship_choices,
        'grades': grades,
        'sections': sections,
        'students': students,
    }
    
    return render(request, 'registrar/guardians.html', context)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["GET"])
def registrar_search_guardians(request):
    """API endpoint to search and filter guardians for AJAX requests"""
    try:
        # Get query parameters
        search_query = request.GET.get('search', '')
        relationship_filter = request.GET.get('relationship', '')
        grade_filter = request.GET.get('grade', '')
        section_filter = request.GET.get('section', '')
        page_number = request.GET.get('page', 1)
        items_per_page = request.GET.get('items_per_page', 10)
        
        # Base queryset
        guardians = Guardian.objects.all().order_by('-created_at')
        
        # Apply search if provided
        if search_query:
            guardians = guardians.filter(
                Q(first_name__icontains=search_query) |
                Q(middle_name__icontains=search_query) |
                Q(last_name__icontains=search_query) |
                Q(email__icontains=search_query) |
                Q(phone__icontains=search_query)  # Use phone instead of phone_number
            )
        
        # Apply relationship filter if provided
        if relationship_filter:
            guardians = guardians.filter(relationship=relationship_filter)
        
        # Apply grade filter if provided
        if grade_filter:
            # Find students in this grade, then filter guardians related to them
            student_ids = Student.objects.filter(grade_id=grade_filter).values_list('id', flat=True)
            guardians = guardians.filter(student_id__in=student_ids)
        
        # Apply section filter if provided
        if section_filter:
            # Find students in this section, then filter guardians related to them
            student_ids = Student.objects.filter(section_id=section_filter).values_list('id', flat=True)
            guardians = guardians.filter(student_id__in=student_ids)
        
        # Get total count for pagination
        total_count = guardians.count()
        
        # Parse page number and items_per_page
        try:
            page_number = int(page_number)
            items_per_page = int(items_per_page)
        except (ValueError, TypeError):
            page_number = 1
            items_per_page = 10
        
        # Create paginator
        paginator = Paginator(guardians, items_per_page)
        page_obj = paginator.get_page(page_number)
        
        # Prepare guardian data for JSON response
        guardian_data = []
        for guardian in page_obj:
            try:
                # Create guardian data dict with safe fallbacks
                data = {
                    'id': guardian.id,
                    'full_name': f"{guardian.first_name} {guardian.last_name}",
                    'first_name': guardian.first_name,
                    'middle_name': guardian.middle_name or '',
                    'last_name': guardian.last_name,
                    'email': guardian.email or 'No email',
                    'phone_number': guardian.phone or 'No phone',  # Use phone instead of phone_number
                    'relationship': guardian.relationship,
                    'created_at_display': guardian.created_at.strftime('%B %d, %Y') if guardian.created_at else 'N/A'
                }
                
                # Safely handle relationship display and children data
                try:
                    data['relationship_display'] = guardian.get_relationship_display()
                except (AttributeError, TypeError):
                    data['relationship_display'] = guardian.relationship
                
                # Child info with safe fallbacks
                try:
                    student = Student.objects.filter(id=guardian.student_id).first() if hasattr(guardian, 'student_id') else None
                    if student:
                        data['children_count'] = 1
                        data['children_preview'] = f"{student.first_name} {student.last_name}"
                    else:
                        data['children_count'] = 0
                        data['children_preview'] = 'No children'

                except Exception as e:
                    data['children_count'] = 0
                    data['children_preview'] = 'Error loading children'
                    print(f"Error loading children for guardian {guardian.id}: {str(e)}")
                
                guardian_data.append(data)
            except Exception as e:
                print(f"Error processing guardian {guardian.id if hasattr(guardian, 'id') else 'unknown'}: {str(e)}")
                continue
        
        # Prepare pagination data
        pagination = {
            'current_page': page_obj.number,
            'num_pages': paginator.num_pages,
            'has_next': page_obj.has_next(),
            'has_previous': page_obj.has_previous(),
            'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
            'start_index': page_obj.start_index(),
            'end_index': page_obj.end_index(),
            'total_count': total_count,
        }
        
        return JsonResponse({
            'status': 'success',
            'guardians': guardian_data,
            'pagination': pagination,
            'total_count': total_count,
        })
    
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        print(f"Error in registrar_search_guardians: {str(e)}")
        print(error_traceback)
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}',
            'traceback': error_traceback
        }, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["POST"])
def registrar_create_guardian(request):
    """Create a new guardian"""
    try:
        data = json.loads(request.body)
        
        # Validate required fields
        required_fields = ['first_name', 'last_name', 'phone_number', 'relationship', 'student_id']
        for field in required_fields:
            if not field in data or not data[field]:
                return JsonResponse({
                    'status': 'error',
                    'message': f'Field {field} is required'
                }, status=400)
        
        # Check if student exists and doesn't already have a guardian
        try:
            student = Student.objects.get(id=data['student_id'])
        except Student.DoesNotExist:
            return JsonResponse({
                'status': 'error',
                'message': 'Student not found'
            }, status=404)
        
        # Validate email if provided
        email = data.get('email', '').strip()
        if email:
            # Add email validation if needed
            pass
        
        # Create guardian - Note: using phone_number from form but storing in phone field
        guardian = Guardian.objects.create(
            first_name=data['first_name'].strip(),
            middle_name=data.get('middle_name', '').strip(),
            last_name=data['last_name'].strip(),
            email=email,
            phone=data['phone_number'].strip(),  # Changed to map phone_number to phone
            relationship=data['relationship'],
            student=student
        )
        
        return JsonResponse({
            'status': 'success',
            'message': 'Guardian created successfully',
            'guardian': {
                'id': guardian.id,
                'full_name': f"{guardian.first_name} {guardian.last_name}",
                'email': guardian.email or '',
                'phone_number': guardian.phone,  # Return as phone_number for consistency with form
                'relationship': guardian.relationship,
                'student_name': f"{student.first_name} {student.last_name}",
            }
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'status': 'error',
            'message': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}'
        }, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["PATCH"])
def registrar_update_guardian(request, guardian_id):
    """Update an existing guardian"""
    try:
        data = json.loads(request.body)
        
        guardian = get_object_or_404(Guardian, id=guardian_id)
        
        # Validate required fields
        required_fields = ['first_name', 'last_name', 'phone_number', 'relationship', 'student_id']
        for field in required_fields:
            if not field in data or not data[field]:
                return JsonResponse({
                    'status': 'error',
                    'message': f'Field {field} is required'
                }, status=400)
        
        # Check if student exists and validate assignment
        try:
            student = Student.objects.get(id=data['student_id'])
        except Student.DoesNotExist:
            return JsonResponse({
                'status': 'error',
                'message': 'Student not found'
            }, status=404)
        
        # Validate email if provided
        email = data.get('email', '').strip()
        # Email validation could go here if needed
        
        # Update guardian fields
        guardian.first_name = data['first_name'].strip()
        guardian.middle_name = data.get('middle_name', '').strip()
        guardian.last_name = data['last_name'].strip()
        guardian.email = email
        guardian.phone = data['phone_number'].strip()  # Map phone_number from form to phone field in model
        guardian.relationship = data['relationship']
        guardian.student = student
        guardian.save()
        
        return JsonResponse({
            'status': 'success',
            'message': 'Guardian updated successfully',
            'guardian': {
                'id': guardian.id,
                'full_name': f"{guardian.first_name} {guardian.last_name}",
                'email': guardian.email or '',
                'phone_number': guardian.phone,  # Return as phone_number for consistency with form
                'relationship': guardian.relationship,
                'student_name': f"{student.first_name} {student.last_name}"
            }
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'status': 'error',
            'message': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        print(f"Error in registrar_update_guardian: {str(e)}")
        print(error_traceback)
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}',
            'traceback': error_traceback
        }, status=500)  # Changed to 500 for server errors

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["DELETE"])
def registrar_delete_guardian(request, guardian_id):
    """Delete a guardian"""
    try:
        guardian = get_object_or_404(Guardian, id=guardian_id)
        guardian_name = f"{guardian.first_name} {guardian.last_name}"  # Use direct concatenation instead of get_full_name
        guardian.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Guardian {guardian_name} deleted successfully'
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}'
        }, status=500)

@login_required
@user_passes_test(is_registrar)
@require_GET
def registrar_get_guardian_children(request, guardian_id):
    """Get children for a specific guardian"""
    try:
        guardian = get_object_or_404(Guardian, id=guardian_id)
        
        children = []
        if guardian.student:
            children.append({
                'id': guardian.student.id,
                'full_name': f"{guardian.student.first_name} {guardian.student.last_name}",  # Direct concatenation
                'lrn': guardian.student.lrn,
                'grade': guardian.student.grade.name if guardian.student.grade else 'N/A',
                'section': guardian.student.section.name if guardian.student.section else 'N/A',
                'status': guardian.student.status
            })
        
        return JsonResponse({
            'status': 'success',
            'guardian': {
                'name': f"{guardian.first_name} {guardian.last_name}",  # Direct concatenation
                'relationship': guardian.get_relationship_display() if hasattr(guardian, 'get_relationship_display') else guardian.relationship
            },
            'children': children
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}'
        }, status=500)

@login_required
@user_passes_test(is_registrar)
@require_GET
def registrar_get_sections_by_grade(request):
    """Get sections for a specific grade"""
    try:
        grade_id = request.GET.get('grade_id')
        if not grade_id:
            return JsonResponse({
                'status': 'error',
                'message': 'Grade ID is required'
            }, status=400)
        
        sections = Section.objects.filter(grade_id=grade_id).order_by('name')
        sections_data = [
            {
                'id': section.id,
                'name': section.name,
                'grade_id': section.grade.id,
                'grade_name': section.grade.name
            }
            for section in sections
        ]
        
        return JsonResponse({
            'status': 'success',
            'sections': sections_data
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}'
        }, status=500)

@login_required
@user_passes_test(is_registrar)
@require_GET
def registrar_get_students_by_section(request):
    """Get students for a specific section"""
    try:
        section_id = request.GET.get('section_id')
        if not section_id:
            return JsonResponse({
                'status': 'error',
                'message': 'Section ID is required'
            }, status=400)
        
        # Get students who don't already have a guardian assigned (for create) or are assigned to the current guardian (for edit)
        students = Student.objects.filter(
            section_id=section_id,
            status='ACTIVE'
        ).select_related('grade', 'section').order_by('first_name', 'last_name')
        
        students_data = [
            {
                'id': student.id,
                'first_name': student.first_name,
                'last_name': student.last_name,
                'lrn': student.lrn,
                'grade': student.grade.name if student.grade else 'N/A',
                'section': student.section.name if student.section else 'N/A',
                'grade_id': student.grade.id if student.grade else None,
                'section_id': student.section.id if student.section else None,
                'has_guardian': hasattr(student, 'guardian') and student.guardian is not None
            }
            for student in students
        ]
        
        return JsonResponse({
            'status': 'success',
            'students': students_data
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}'
        }, status=500)

@login_required
@user_passes_test(is_registrar)
@require_GET
def registrar_get_guardian_details(request, guardian_id):
    """
    Endpoint to retrieve a specific guardian's data for editing
    """
    try:
        guardian = get_object_or_404(Guardian, id=guardian_id)
        
        # Prepare children data
        children = []
        if guardian.student:
            child = guardian.student
            children.append({
                'id': child.id,
                'grade_id': child.grade.id if child.grade else '',
                'grade': child.grade.name if child.grade else '',
                'section_id': child.section.id if child.section else '',
                'section': child.section.name if child.section else '',
                'first_name': child.first_name,
                'last_name': child.last_name,
                'full_name': f"{child.first_name} {child.last_name}"  # Direct concatenation
            })
        
        guardian_data = {
            'id': guardian.id,
            'first_name': guardian.first_name,
            'middle_name': guardian.middle_name or '',
            'last_name': guardian.last_name,
            'email': guardian.email or '',
            'phone_number': guardian.phone,  # Return phone field as phone_number for form
            'relationship': guardian.relationship,
            'children': children,
            'created_at': guardian.created_at.isoformat() if guardian.created_at else None
        }
        
        return JsonResponse({
            'status': 'success',
            'guardian': guardian_data
        })
        
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        print(f"Error in registrar_get_guardian_details: {str(e)}")
        print(error_traceback)
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}',
            'traceback': error_traceback
        }, status=500)


# ==========================
#  ATTENDANCE MANAGEMENT VIEWS - EXACT CARBON COPY OF ADMIN
# ==========================

@login_required
@user_passes_test(is_registrar)
def registrar_attendance_records(request):
    """Attendance management view with dashboard cards and filters"""
    from django.db.models import Q
    from datetime import date

    # Dashboard stats
    total_records = Attendance.objects.count()
    present_count = Attendance.objects.filter(status='ON TIME').count() + Attendance.objects.filter(status='LATE').count()
    absent_count = Attendance.objects.filter(status='ABSENT').count()
    excused_count = Attendance.objects.filter(status='EXCUSED').count()

    grades = Grade.objects.all().order_by('name')
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    # Filters (for initial page load)
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')

    context = {
        'grades': grades,
        'sections': sections,
        'grade_filter': grade_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'date_filter': date_filter,
        'total_records': total_records,
        'present_count': present_count,
        'absent_count': absent_count,
        'excused_count': excused_count,
    }
    return render(request, 'registrar/attendance.html', context)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["GET"])
def registrar_search_attendance_records(request):
    """AJAX search/filter for attendance records"""
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')
    page_number = request.GET.get('page', 1)
    items_per_page = request.GET.get('items_per_page', 10)

    records = Attendance.objects.select_related('student', 'student__grade', 'student__section').all().order_by('-date', '-student__last_name')

    if search_query:
        records = records.filter(
            Q(student__first_name__icontains=search_query) |
            Q(student__last_name__icontains=search_query) |
            Q(student__lrn__icontains=search_query) |
            Q(student__section__name__icontains=search_query)
        )
    if grade_filter:
        records = records.filter(student__grade_id=grade_filter)
    if section_filter:
        records = records.filter(student__section_id=section_filter)
    if status_filter:
        records = records.filter(status=status_filter)
    if date_filter:
        records = records.filter(date=date_filter)

    total_count = records.count()
    try:
        page_number = int(page_number)
        items_per_page = int(items_per_page)
    except (ValueError, TypeError):
        page_number = 1
        items_per_page = 10

    paginator = Paginator(records, items_per_page)
    page_obj = paginator.get_page(page_number)

    attendance_data = []
    for record in page_obj:
        attendance_data.append({
            'id': record.id,
            'lrn': record.student.lrn,
            'full_name': f"{record.student.first_name} {record.student.last_name}",
            'grade': record.student.grade.name if record.student.grade else '',
            'section': record.student.section.name if record.student.section else '',
            'date': record.date.strftime('%Y-%m-%d'),
            'time_in': record.time_in.strftime('%H:%M:%S') if record.time_in else '',
            'time_out': record.time_out.strftime('%H:%M:%S') if record.time_out else '',
            'status': record.status,
            'profile_pic': record.student.profile_pic or None,
        })

    pagination = {
        'current_page': page_obj.number,
        'num_pages': paginator.num_pages,
        'has_next': page_obj.has_next(),
        'has_previous': page_obj.has_previous(),
        'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
        'start_index': page_obj.start_index(),
        'end_index': page_obj.end_index(),
        'total_count': total_count,
    }

    return JsonResponse({
        'status': 'success',
        'records': attendance_data,
        'pagination': pagination,
        'total_count': total_count,
    })

@login_required
@user_passes_test(is_registrar)
@require_GET
def registrar_get_attendance_record(request, attendance_id):
    """API endpoint to get details of a specific attendance record"""
    try:
        record = get_object_or_404(Attendance, id=attendance_id)
        
        data = {
            'id': record.id,
            'student_id': record.student.id,
            'student_name': f"{record.student.first_name} {record.student.last_name}",
            'lrn': record.student.lrn,
            'date': record.date.strftime('%Y-%m-%d'),
            'time_in': record.time_in.strftime('%H:%M:%S') if record.time_in else None,
            'time_out': record.time_out.strftime('%H:%M:%S') if record.time_out else None,
            'status': record.status,
            'notes': record.notes or '',
            'section': record.student.section.name if record.student.section else '',
            'section_id': record.student.section.id if record.student.section else None,
            'grade': record.student.grade.name if record.student.grade else '',
            'grade_id': record.student.grade.id if record.student.grade else None,
        }
        
        return JsonResponse({
            'status': 'success',
            'record': data
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["POST"])
def registrar_create_attendance_record(request):
    """API endpoint to create a new attendance record"""
    try:
        data = json.loads(request.body)
        
        # Validate required fields
        required_fields = ['student_id', 'date', 'status']
        for field in required_fields:
            if field not in data or not data[field]:
                return JsonResponse({
                    'status': 'error',
                    'message': f'{field} is required'
                }, status=400)
        
        # Check if student exists
        try:
            student = Student.objects.get(id=data['student_id'])
        except Student.DoesNotExist:
            return JsonResponse({
                'status': 'error',
                'message': 'Student not found'
            }, status=404)
        
        # Check if an attendance record already exists for this student on this date
        if Attendance.objects.filter(student=student, date=data['date']).exists():
            return JsonResponse({
                'status': 'error',
                'message': f'Attendance record already exists for {student.first_name} {student.last_name} on {data["date"]}'
            }, status=400)
        
        # Parse time fields if provided
        time_in = None
        if data.get('time_in'):
            try:
                from datetime import datetime
                time_in = datetime.strptime(data['time_in'], '%H:%M:%S').time()
            except ValueError:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Invalid time_in format. Use HH:MM:SS'
                }, status=400)
                
        time_out = None
        if data.get('time_out'):
            try:
                from datetime import datetime
                time_out = datetime.strptime(data['time_out'], '%H:%M:%S').time()
            except ValueError:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Invalid time_out format. Use HH:MM:SS'
                }, status=400)
        
        # Create attendance record
        record = Attendance(
            student=student,
            date=data['date'],
            time_in=time_in,
            time_out=time_out,
            status=data['status']
        )
        record.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Attendance record for {student.first_name} {student.last_name} created successfully',
            'record_id': record.id
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'status': 'error',
            'message': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        import traceback
        print("Error in registrar_create_attendance_record:", str(e))
        print(traceback.format_exc())
        return JsonResponse({
            'status': 'error',
            'message': str(e),
            'traceback': traceback.format_exc()
        }, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["PUT", "PATCH"])
def registrar_update_attendance_record(request, attendance_id):
    """API endpoint to update an attendance record"""
    try:
        record = get_object_or_404(Attendance, id=attendance_id)
        data = json.loads(request.body)
        
        # Update status if provided
        if 'status' in data:
            record.status = data['status']
        
        # Update date if provided
        if 'date' in data:
            # Check if another record exists for this student on the new date (excluding current record)
            if Attendance.objects.filter(student=record.student, date=data['date']).exclude(id=attendance_id).exists():
                return JsonResponse({
                    'status': 'error',
                    'message': f'Another attendance record exists for {record.student.first_name} {record.student.last_name} on {data["date"]}'
                }, status=400)
            record.date = data['date']
        
        # Update time_in if provided
        if 'time_in' in data:
            if data['time_in']:
                try:
                    from datetime import datetime
                    record.time_in = datetime.strptime(data['time_in'], '%H:%M:%S').time()
                except ValueError:
                    return JsonResponse({
                        'status': 'error',
                        'message': 'Invalid time_in format. Use HH:MM:SS'
                    }, status=400)
            else:
                record.time_in = None
        
        # Update time_out if provided
        if 'time_out' in data:
            if data['time_out']:
                try:
                    from datetime import datetime
                    record.time_out = datetime.strptime(data['time_out'], '%H:%M:%S').time()
                except ValueError:
                    return JsonResponse({
                        'status': 'error',
                        'message': 'Invalid time_out format. Use HH:MM:SS'
                    }, status=400)
            else:
                record.time_out = None
        
        # Update notes if provided
        if 'notes' in data:
            record.notes = data.get('notes', '')
        
        # Save the updated record
        record.save()
        
        return JsonResponse({
            'status': 'success',
            'message': 'Attendance record updated successfully'
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'status': 'error',
            'message': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["DELETE"])
def registrar_delete_attendance_record(request, attendance_id):
    """API endpoint to delete an attendance record"""
    try:
        record = get_object_or_404(Attendance, id=attendance_id)
        student_name = f"{record.student.first_name} {record.student.last_name}"
        record_date = record.date.strftime('%Y-%m-%d')
        
        record.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Attendance record for {student_name} on {record_date} deleted successfully'
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

@login_required
@user_passes_test(is_registrar)
@require_GET
def registrar_get_students_for_attendance(request):
    """API endpoint to get students for attendance form"""
    try:
        # Get filters
        section_id = request.GET.get('section', None)
        grade_id = request.GET.get('grade', None)
        
        # Base queryset - only active students
        students = Student.objects.filter(status='ACTIVE').select_related('grade', 'section')
        
        # Apply filters if provided
        if section_id:
            students = students.filter(section_id=section_id)
        elif grade_id:
            students = students.filter(grade_id=grade_id)
        
        # Order by name
        students = students.order_by('last_name', 'first_name')
        
        # Prepare data for response
        students_data = []
        for student in students:
            students_data.append({
                'id': student.id,
                'lrn': student.lrn,
                'name': f"{student.first_name} {student.last_name}",
                'grade': student.grade.name if student.grade else '',
                'section': student.section.name if student.section else '',
                'grade_id': student.grade.id if student.grade else None,
                'section_id': student.section.id if student.section else None
            })
        
        return JsonResponse({
            'status': 'success',
            'students': students_data,
            'total': len(students_data)
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}'
        }, status=500)

# ==========================
#  EXCUSED MANAGEMENT VIEWS - EXACT CARBON COPY OF ADMIN
# ==========================

@login_required
@user_passes_test(is_registrar)
def registrar_excused(request):
    """View for excused absences management"""
    from datetime import date
    total_excused = ExcusedAbsence.objects.count()
    today = date.today()
    active_excused = ExcusedAbsence.objects.filter(effective_date__lte=today, end_date__gte=today).count()
    expired_excused = ExcusedAbsence.objects.filter(end_date__lt=today).count()
    upcoming_excused = ExcusedAbsence.objects.filter(effective_date__gt=today).count()

    grades = Grade.objects.all().order_by('name')
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')

    context = {
        'grades': grades,
        'sections': sections,
        'grade_filter': grade_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'date_filter': date_filter,
        'total_excused': total_excused,
        'active_excused': active_excused,
        'expired_excused': expired_excused,
        'upcoming_excused': upcoming_excused,
    }
    return render(request, 'registrar/excused.html', context)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["GET"])
def registrar_search_excused_absences(request):
    """AJAX search/filter for excused absences"""
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')
    page_number = request.GET.get('page', 1)
    items_per_page = request.GET.get('items_per_page', 10)

    excused = ExcusedAbsence.objects.select_related('student', 'student__grade', 'student__section').all().order_by('-date_absent', '-student__last_name')

    if search_query:
        excused = excused.filter(
            Q(student__first_name__icontains=search_query) |
            Q(student__last_name__icontains=search_query) |
            Q(student__lrn__icontains=search_query) |
            Q(student__section__name__icontains=search_query)
        )
    if grade_filter:
        excused = excused.filter(student__grade_id=grade_filter)
    if section_filter:
        excused = excused.filter(student__section_id=section_filter)
    if status_filter:
        from datetime import date
        today = date.today()
        if status_filter == 'ACTIVE':
            excused = excused.filter(effective_date__lte=today, end_date__gte=today)
        elif status_filter == 'EXPIRED':
            excused = excused.filter(end_date__lt=today)
        elif status_filter == 'UPCOMING':
            excused = excused.filter(effective_date__gt=today)
    if date_filter:
        excused = excused.filter(date_absent=date_filter)

    total_count = excused.count()
    try:
        page_number = int(page_number)
        items_per_page = int(items_per_page)
    except (ValueError, TypeError):
        page_number = 1
        items_per_page = 10

    paginator = Paginator(excused, items_per_page)
    page_obj = paginator.get_page(page_number)

    excused_data = []
    for record in page_obj:
        # Determine status
        from datetime import date
        today = date.today()
        if record.effective_date > today:
            status = 'UPCOMING'
        elif record.end_date < today:
            status = 'EXPIRED'
        else:
            status = 'ACTIVE'
        excused_data.append({
            'id': record.id,
            'lrn': record.student.lrn,
            'full_name': f"{record.student.first_name} {record.student.last_name}",
            'grade': record.student.grade.name if record.student.grade else '',
            'section': record.student.section.name if record.student.section else '',
            'date_absent': record.date_absent.strftime('%Y-%m-%d'),
            'effective_date': record.effective_date.strftime('%Y-%m-%d'),
            'end_date': record.end_date.strftime('%Y-%m-%d'),
            'excuse_letter': record.excuse_letter,
            'status': status,
            'profile_pic': record.student.profile_pic or None,
        })

    pagination = {
        'current_page': page_obj.number,
        'num_pages': paginator.num_pages,
        'has_next': page_obj.has_next(),
        'has_previous': page_obj.has_previous(),
        'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
        'start_index': page_obj.start_index(),
        'end_index': page_obj.end_index(),
        'total_count': total_count,
    }

    return JsonResponse({
        'status': 'success',
        'records': excused_data,
        'pagination': pagination,
        'total_count': total_count,
    })

# Add placeholder functions for create/update/delete excused absences
# These would follow similar patterns to the above implementations

# File upload/delete functionality is handled directly in create/update/delete operations above

# Basic implementations for guardian update/delete
@require_http_methods(["POST"])
@login_required
@user_passes_test(is_registrar)
def registrar_update_guardian(request, guardian_id):
    """Simplified guardian update"""
    try:
        guardian = get_object_or_404(Guardian, id=guardian_id)
        data = json.loads(request.body)
        
        guardian.first_name = data.get('first_name', guardian.first_name)
        guardian.last_name = data.get('last_name', guardian.last_name)
        guardian.email = data.get('email', guardian.email)
        guardian.phone = data.get('phone', guardian.phone)
        guardian.relationship = data.get('relationship', guardian.relationship)
        guardian.save()
        
        messages.success(request, f'Guardian "{guardian.first_name} {guardian.last_name}" updated successfully!')
        
        return JsonResponse({
            'success': True,          
            'message': f'Guardian "{guardian.first_name} {guardian.last_name}" updated successfully!'
        })
        
    except Guardian.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Guardian not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@require_http_methods(["DELETE"])
@login_required
@user_passes_test(is_registrar)
def registrar_delete_guardian(request, guardian_id):
    """Simplified guardian deletion"""
    try:
        guardian = get_object_or_404(Guardian, id=guardian_id)
        guardian_name = f"{guardian.first_name} {guardian.last_name}"
        guardian.delete()
        
        messages.success(request, f'Guardian "{guardian_name}" deleted successfully!')
        
        return JsonResponse({
            'success': True,
            'message': f'Guardian "{guardian_name}" deleted successfully!'
        })
        
    except Guardian.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Guardian not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@require_http_methods(["GET"])
@login_required
@user_passes_test(is_registrar)
def registrar_get_guardian_details(request, guardian_id):
    """Get guardian details"""
    try:
        guardian = get_object_or_404(Guardian, id=guardian_id)
        
        guardian_data = {
            'id': guardian.id,
            'first_name': guardian.first_name,
            'middle_name': guardian.middle_name or '',
            'last_name': guardian.last_name,
            'email': guardian.email or '',
            'phone': guardian.phone or '',
            'relationship': guardian.relationship or '',
            'student_id': guardian.student.id if guardian.student else None,
            'student_name': f"{guardian.student.first_name} {guardian.student.last_name}" if guardian.student else '',
            'grade_id': guardian.student.grade.id if guardian.student and guardian.student.grade else None,
            'section_id': guardian.student.section.id if guardian.student and guardian.student.section else None,
        }
        
        return JsonResponse({
            'success': True,
            'guardian': guardian_data
        })
        
    except Guardian.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Guardian not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@require_http_methods(["GET"])
@login_required
@user_passes_test(is_registrar)
def registrar_get_guardian_children(request, guardian_id):
    """Get children for a guardian"""
    try:
        guardian = get_object_or_404(Guardian, id=guardian_id)
        
        # For simplified implementation, assume one child per guardian
        children_data = []
        if guardian.student:
            children_data.append({
                'id': guardian.student.id,
                'name': f"{guardian.student.first_name} {guardian.student.last_name}",
                'student_id': guardian.student.lrn,
                'grade': guardian.student.grade.name if guardian.student.grade else '',
                'section': guardian.student.section.name if guardian.student.section else '',
            })
        
        return JsonResponse({
            'success': True,
            'children': children_data
        })
        
    except Guardian.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Guardian not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

# Basic attendance CRUD operations
@require_http_methods(["POST"])
@login_required
@user_passes_test(is_registrar)
def registrar_create_attendance_record(request):
    """Simplified attendance record creation"""
    try:
        data = json.loads(request.body)
        
        student = get_object_or_404(Student, id=data['student_id'])
        
        attendance = Attendance.objects.create(
            student=student,
            date=data.get('date', timezone.now().date()),
            time_in=data.get('time_in'),
            time_out=data.get('time_out'),
            status=data.get('status', 'present')
        )
        
        messages.success(request, 'Attendance record created successfully!')
        
        return JsonResponse({
            'success': True,
            'message': 'Attendance record created successfully!'
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

# Add other attendance CRUD operations following similar patterns...

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["POST"])
def registrar_create_excused_absence(request):
    """API endpoint to create a new excused absence"""
    try:
        # Accept multipart/form-data for file upload
        student_id = request.POST.get('student_id')
        date_absent = request.POST.get('date_absent')
        effective_date = request.POST.get('effective_date')
        end_date = request.POST.get('end_date')
        excuse_letter_file = request.FILES.get('excuse_letter')

        # Validate required fields
        if not student_id or not date_absent or not effective_date or not end_date or not excuse_letter_file:
            return JsonResponse({'status': 'error', 'message': 'All fields are required'}, status=400)

        # Validate student
        try:
            student = Student.objects.get(id=student_id)
        except Student.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Student not found'}, status=404)

        # Validate file type and size
        if excuse_letter_file.size > 5 * 1024 * 1024:
            return JsonResponse({'status': 'error', 'message': 'File too large. Max 5MB.'}, status=400)
        if not excuse_letter_file.content_type.startswith('image/'):
            return JsonResponse({'status': 'error', 'message': 'Only image files allowed.'}, status=400)

        # Save file in private_excuse_letters
        import uuid
        filename = f"{uuid.uuid4()}_{excuse_letter_file.name}"
        fs = FileSystemStorage(location=PRIVATE_EXCUSE_LETTERS_DIR)
        filename = fs.save(filename, excuse_letter_file)

        # Create record
        record = ExcusedAbsence.objects.create(
            student=student,
            date_absent=date_absent,
            excuse_letter=filename,
            effective_date=effective_date,
            end_date=end_date
        )

        return JsonResponse({
            'status': 'success',
            'message': f'Excused absence for {student.first_name} {student.last_name} created successfully',
            'record_id': record.id
        })
    except Exception as e:
        import traceback
        print("Error in registrar_create_excused_absence:", str(e))
        print(traceback.format_exc())
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

# Add other excused absence CRUD operations following similar patterns...

# Helper function for getting students for attendance
@require_http_methods(["GET"])
@login_required
@user_passes_test(is_registrar)  
def registrar_get_students_for_attendance(request):
    """Get students for attendance record creation"""
    grade_id = request.GET.get('grade_id')
    section_id = request.GET.get('section_id')
    
    students = Student.objects.filter(status='ACTIVE')
    
    if grade_id:
        students = students.filter(grade_id=grade_id)
    if section_id:
        students = students.filter(section_id=section_id)
    
    students_data = []
    for student in students.order_by('first_name', 'last_name'):
        students_data.append({
            'id': student.id,
            'name': f"{student.first_name} {student.last_name}",
            'student_id': student.lrn,
            'grade': student.grade.name if student.grade else '',
            'section': student.section.name if student.section else '',
        })
    
    return JsonResponse({
        'success': True,
        'students': students_data
    })

# Placeholder implementations for remaining attendance/excused CRUD operations
@require_http_methods(["GET"])
@login_required
@user_passes_test(is_registrar)
def registrar_get_attendance_record(request, attendance_id):
    """Get attendance record details"""
    try:
        attendance = get_object_or_404(Attendance, id=attendance_id)
        
        attendance_data = {
            'id': attendance.id,
            'student_id': attendance.student.id,
            'student_name': f"{attendance.student.first_name} {attendance.student.last_name}",
            'date': attendance.date.strftime('%Y-%m-%d'),
            'time_in': attendance.time_in.strftime('%H:%M') if attendance.time_in else '',
            'time_out': attendance.time_out.strftime('%H:%M') if attendance.time_out else '',
            'status': attendance.status,
        }
        
        return JsonResponse({
            'success': True,
            'attendance': attendance_data
        })
        
    except Attendance.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Attendance record not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["GET"])
def registrar_get_excused_absence(request, excused_id):
    """API endpoint to get details of a specific excused absence"""
    try:
        record = get_object_or_404(ExcusedAbsence, id=excused_id)
        
        # Extract just the filename from the excuse_letter path
        excuse_letter_filename = record.excuse_letter
        if excuse_letter_filename and excuse_letter_filename.startswith('private_excuse_letters/'):
            excuse_letter_filename = excuse_letter_filename.replace('private_excuse_letters/', '')
        
        data = {
            'id': record.id,
            'student_id': record.student.id,
            'student_name': f"{record.student.first_name} {record.student.last_name}",
            'lrn': record.student.lrn,
            'date_absent': record.date_absent.strftime('%Y-%m-%d'),
            'effective_date': record.effective_date.strftime('%Y-%m-%d'),
            'end_date': record.end_date.strftime('%Y-%m-%d'),
            'excuse_letter': excuse_letter_filename,  # Return just the filename
            'section': record.student.section.name if record.student.section else '',
            'section_id': record.student.section.id if record.student.section else None,
            'grade': record.student.grade.name if record.student.grade else '',
            'grade_id': record.student.grade.id if record.student.grade else None,
        }
        return JsonResponse({'status': 'success', 'record': data})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

# Add update/delete functions for attendance and excused absences following similar patterns
@require_http_methods(["POST"])
@login_required
@user_passes_test(is_registrar)
def registrar_update_attendance_record(request, attendance_id):
    """Update attendance record"""
    try:
        attendance = get_object_or_404(Attendance, id=attendance_id)
        data = json.loads(request.body)
        
        attendance.time_in = data.get('time_in', attendance.time_in)
        attendance.time_out = data.get('time_out', attendance.time_out)
        attendance.status = data.get('status', attendance.status)
        attendance.save()
        
        messages.success(request, 'Attendance record updated successfully!')
        
        return JsonResponse({
            'success': True,
            'message': 'Attendance record updated successfully!'
        })
        
    except Attendance.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Attendance record not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@require_http_methods(["DELETE"])
@login_required
@user_passes_test(is_registrar)
def registrar_delete_attendance_record(request, attendance_id):
    """Delete attendance record"""
    try:
        attendance = get_object_or_404(Attendance, id=attendance_id)
        attendance.delete()
        
        messages.success(request, 'Attendance record deleted successfully!')
        
        return JsonResponse({
            'success': True,
            'message': 'Attendance record deleted successfully!'
        })
        
    except Attendance.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Attendance record not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["POST", "PATCH"])
def registrar_update_excused_absence(request, excused_id):
    """API endpoint to update an excused absence"""
    try:
        record = get_object_or_404(ExcusedAbsence, id=excused_id)
        # Accept multipart/form-data for file upload
        student_id = request.POST.get('student_id', record.student.id)
        date_absent = request.POST.get('date_absent', record.date_absent)
        effective_date = request.POST.get('effective_date', record.effective_date)
        end_date = request.POST.get('end_date', record.end_date)
        excuse_letter_file = request.FILES.get('excuse_letter', None)

        # Validate student
        try:
            student = Student.objects.get(id=student_id)
        except Student.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Student not found'}, status=404)
        
        record.student = student
        record.date_absent = date_absent
        record.effective_date = effective_date
        record.end_date = end_date

        # Handle new file upload
        if excuse_letter_file:
            if excuse_letter_file.size > 5 * 1024 * 1024:
                return JsonResponse({'status': 'error', 'message': 'File too large. Max 5MB.'}, status=400)
            if not excuse_letter_file.content_type.startswith('image/'):
                return JsonResponse({'status': 'error', 'message': 'Only image files allowed.'}, status=400)
            # Delete old file
            if record.excuse_letter:
                old_path = os.path.join(PRIVATE_EXCUSE_LETTERS_DIR, record.excuse_letter)
                if os.path.exists(old_path):
                    os.remove(old_path)
            # Save new file
            import uuid
            filename = f"{uuid.uuid4()}_{excuse_letter_file.name}"
            fs = FileSystemStorage(location=PRIVATE_EXCUSE_LETTERS_DIR)
            filename = fs.save(filename, excuse_letter_file)
            record.excuse_letter = filename
        record.save()
        return JsonResponse({'status': 'success', 'message': 'Excused absence updated successfully'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["DELETE"])
def registrar_delete_excused_absence(request, excused_id):
    """API endpoint to delete an excused absence"""
    try:
        record = get_object_or_404(ExcusedAbsence, id=excused_id)
        # Delete file
        if record.excuse_letter:
            file_path = os.path.join(PRIVATE_EXCUSE_LETTERS_DIR, record.excuse_letter)
            if os.path.exists(file_path):
                os.remove(file_path)
        record.delete()
        return JsonResponse({'status': 'success', 'message': 'Excused absence deleted successfully'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


# ==========================
#  GRADE MANAGEMENT VIEWS - EXACT CARBON COPY OF ADMIN
# ==========================

@login_required
@user_passes_test(is_registrar)
def registrar_grades(request):
    """View for grade management"""
    # Get query parameters
    search_query = request.GET.get('search', '')
    page_number = request.GET.get('page', 1)
    
    # Base queryset - order by created_at DESC (newest first)
    grades = Grade.objects.all().order_by('-created_at')
    
    # Apply search if provided
    if search_query:
        grades = grades.filter(name__icontains=search_query)
    
    # Annotate with section and student counts
    grades = grades.annotate(
        sections_count=Count('sections'),
        students_count=Count('sections__students', distinct=True)
    )
    
    # Get stats for dashboard
    total_grades = Grade.objects.count()
    grades_with_sections = Grade.objects.filter(sections__isnull=False).distinct().count()
    total_sections = Section.objects.count()
    total_students = Student.objects.count()
    
    # Pagination
    paginator = Paginator(grades, 10)
    page_obj = paginator.get_page(page_number)
    
    # Calculate page range for pagination UI
    page_range = get_pagination_range(paginator, page_obj.number, 5)
    
    context = {
        'grades': page_obj,
        'search_query': search_query,
        'total_grades': total_grades,
        'grades_with_sections': grades_with_sections,
        'total_sections': total_sections,
        'total_students': total_students,
        'page_range': page_range,
    }
    return render(request, 'registrar/grades.html', context)

@login_required
@user_passes_test(is_registrar)
@require_GET
def registrar_search_grades(request):
    """AJAX search/filter for grades"""
    search_query = request.GET.get('search', '')
    page_number = request.GET.get('page', 1)
    items_per_page = request.GET.get('items_per_page', 10)
    
    # Base queryset - order by created_at DESC (newest first)
    grades = Grade.objects.all().order_by('-created_at')
    
    # Apply search if provided
    if search_query:
        grades = grades.filter(name__icontains=search_query)
    
    # Annotate with counts (fix: use distinct for sections_count)
    grades = grades.annotate(
        sections_count=Count('sections', distinct=True),
        students_count=Count('sections__students', distinct=True)
    )
    
    # Get total count for pagination
    total_count = grades.count()
    
    # Parse page number and items_per_page
    try:
        page_number = int(page_number)
        items_per_page = int(items_per_page)
    except (ValueError, TypeError):
        page_number = 1
        items_per_page = 10
    
    # Create paginator
    paginator = Paginator(grades, items_per_page)
    page_obj = paginator.get_page(page_number)
    
    # Prepare grade data for JSON response
    grade_data = []
    for grade in page_obj:
        grade_data.append({
            'id': grade.id,
            'name': grade.name,
            'sections_count': grade.sections_count,
            'students_count': grade.students_count,
            'created_at': grade.created_at.strftime('%Y-%m-%d'),
            'created_at_display': grade.created_at.strftime('%b %d, %Y'),
        })
    
    # Prepare pagination data
    pagination = {
        'current_page': page_obj.number,
        'num_pages': paginator.num_pages,
        'has_next': page_obj.has_next(),
        'has_previous': page_obj.has_previous(),
        'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
        'start_index': page_obj.start_index(),
        'end_index': page_obj.end_index(),
        'total_count': total_count,
    }
    
    return JsonResponse({
        'status': 'success',
        'grades': grade_data,
        'pagination': pagination,
        'total_count': total_count,
    })

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["POST"])
def registrar_create_grade(request):
    """Create a new grade"""
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        
        if not name:
            return JsonResponse({'status': 'error', 'message': 'Grade name is required.'}, status=400)
        
        if Grade.objects.filter(name__iexact=name).exists():
            return JsonResponse({'status': 'error', 'message': 'Grade with this name already exists.'}, status=400)
        
        grade = Grade.objects.create(name=name)
        
        return JsonResponse({
            'status': 'success',
            'message': f'Grade "{name}" created successfully.',
            'grade': {
                'id': grade.id,
                'name': grade.name,
                'sections_count': 0,
                'students_count': 0,
                'created_at': grade.created_at.strftime('%Y-%m-%d'),
                'created_at_display': grade.created_at.strftime('%b %d, %Y'),
            }
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["PATCH"])
def registrar_update_grade(request, grade_id):
    """Update an existing grade"""
    try:
        grade = get_object_or_404(Grade, id=grade_id)
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        
        if not name:
            return JsonResponse({'status': 'error', 'message': 'Grade name is required.'}, status=400)
        
        if Grade.objects.filter(name__iexact=name).exclude(id=grade_id).exists():
            return JsonResponse({'status': 'error', 'message': 'Another grade with this name already exists.'}, status=400)
        
        old_name = grade.name
        grade.name = name
        grade.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Grade "{old_name}" updated to "{name}" successfully.'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["DELETE"])
def registrar_delete_grade(request, grade_id):
    """Delete a grade"""
    try:
        grade = get_object_or_404(Grade, id=grade_id)
        
        # Check if grade has sections
        sections_count = grade.sections.count()
        students_count = Student.objects.filter(grade=grade).count()
        
        if sections_count > 0 or students_count > 0:
            return JsonResponse({
                'status': 'error', 
                'message': f'Cannot delete grade "{grade.name}" because it has {sections_count} section(s) and {students_count} student(s). Please remove all sections and students first.'
            }, status=400)
        
        grade_name = grade.name
        grade.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Grade "{grade_name}" deleted successfully.'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
@require_GET
def registrar_get_grade_sections(request, grade_id):
    """Get sections for a specific grade"""
    try:
        grade = get_object_or_404(Grade, id=grade_id)
        sections = Section.objects.filter(grade=grade).annotate(
            students_count=Count('students')
        ).order_by('name')
        
        sections_data = []
        for section in sections:
            # Get the advisor for this section
            advisor = CustomUser.objects.filter(section=section, role=UserRole.TEACHER).first()
            advisor_name = f"{advisor.first_name} {advisor.last_name}" if advisor else "No advisor"
            
            sections_data.append({
                'id': section.id,
                'name': section.name,
                'students_count': section.students_count,
                'advisor': advisor_name,
                'created_at': section.created_at.strftime('%b %d, %Y'),
            })
        
        return JsonResponse({
            'status': 'success',
            'grade': {
                'id': grade.id,
                'name': grade.name
            },
            'sections': sections_data,
            'total_sections': len(sections_data)
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


# ==========================
#  FACE ENROLLMENT MANAGEMENT VIEWS
# ==========================

@login_required
@user_passes_test(is_admin_or_registrar)
def registrar_face_enroll(request):
    """Face enrollment page - same as students page but for face enrollment"""
    # Get query parameters
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    page_number = request.GET.get('page', 1)

    # Base queryset
    students = Student.objects.select_related('grade', 'section').all().order_by('-created_at')

    # Apply search
    if search_query:
        students = students.filter(
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(lrn__icontains=search_query) |
            Q(email__icontains=search_query)
        )

    # Apply grade filter
    if grade_filter:
        students = students.filter(grade_id=grade_filter)

    # Apply section filter
    if section_filter:
        students = students.filter(section_id=section_filter)

    # Apply status filter
    if status_filter:
        students = students.filter(status=status_filter)

    # Dashboard stats
    total_students = Student.objects.count()
    active_students = Student.objects.filter(status='ACTIVE').count()
    inactive_students = Student.objects.filter(status='INACTIVE').count()
    face_enrolled_count = Student.objects.exclude(face_path__isnull=True).exclude(face_path__exact='').count()
    not_enrolled_count = Student.objects.filter(status='ACTIVE').filter(Q(face_path__isnull=True) | Q(face_path__exact='')).count()

    # For filter dropdowns
    grades = Grade.objects.all().order_by('name')
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    # Pagination
    from django.core.paginator import Paginator
    paginator = Paginator(students, 10)
    page_obj = paginator.get_page(page_number)
    page_range = get_pagination_range(paginator, page_obj.number, 5)

    context = {
        'students': page_obj,
        'grades': grades,
        'sections': sections,
        'search_query': search_query,
        'grade_filter': grade_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'total_students': total_students,
        'active_students': active_students,
        'inactive_students': inactive_students,
        'face_enrolled_count': face_enrolled_count,
        'not_enrolled_count': not_enrolled_count,
        'page_range': page_range,
    }
    return render(request, 'registrar/face_enroll.html', context)


def get_pagination_range(paginator, current_page, display_range=5):
    """
    Calculate pagination range to display in the template
    This creates a better UX by showing ellipsis for large page sets
    """
    total_pages = paginator.num_pages
    
    # If total pages is less than display range, show all pages
    if total_pages <= display_range:
        return range(1, total_pages + 1)
    
    # Calculate the range to display
    start_page = max(current_page - display_range // 2, 1)
    end_page = min(start_page + display_range - 1, total_pages)
    
    # Adjust if we're near the end
    if end_page - start_page + 1 < display_range:
        start_page = max(end_page - display_range + 1, 1)
    
    return range(start_page, end_page + 1)

@login_required
@user_passes_test(is_admin_or_registrar)
@require_http_methods(["GET"])
def registrar_search_students_for_face_enroll(request):
    """AJAX search/filter for students in face enroll page"""
    try:
        search_query = request.GET.get('search', '')
        grade_filter = request.GET.get('grade', '')
        section_filter = request.GET.get('section', '')
        status_filter = request.GET.get('status', '')
        page_number = request.GET.get('page', 1)
        items_per_page = request.GET.get('items_per_page', 10)

        # Base filter: Only show ACTIVE students
        students = Student.objects.filter(status='ACTIVE').select_related('grade', 'section').order_by('-created_at')

        if search_query:
            students = students.filter(
                Q(first_name__icontains=search_query) |
                Q(last_name__icontains=search_query) |
                Q(lrn__icontains=search_query)
            )
        if grade_filter:
            students = students.filter(grade_id=grade_filter)
        if section_filter:
            students = students.filter(section_id=section_filter)
        
        # Handle enrollment status filter
        if status_filter == 'enrolled':
            students = students.exclude(face_path__isnull=True).exclude(face_path__exact='')
        elif status_filter == 'not_enrolled':
            students = students.filter(Q(face_path__isnull=True) | Q(face_path__exact=''))

        total_count = students.count()
        try:
            page_number = int(page_number)
            items_per_page = int(items_per_page)
        except (ValueError, TypeError):
            page_number = 1
            items_per_page = 10

        from django.core.paginator import Paginator
        paginator = Paginator(students, items_per_page)
        page_obj = paginator.get_page(page_number)

        student_data = []
        for student in page_obj:
            profile_pic = None
            if student.profile_pic:
                profile_pic = student.profile_pic

            student_data.append({
                'id': student.id,
                'lrn': student.lrn,
                'first_name': student.first_name,
                'last_name': student.last_name,
                'full_name': f"{student.first_name} {student.last_name}",
                'grade': student.grade.name if student.grade else '',
                'grade_id': student.grade.id if student.grade else None,
                'section': student.section.name if student.section else '',
                'section_id': student.section.id if student.section else None,
                'status': student.status,
                'is_face_enrolled': bool(student.face_path and student.face_path.strip()),
                'profile_pic': profile_pic,
                'created_at': student.created_at.strftime('%Y-%m-%d'),
                'created_at_display': student.created_at.strftime('%b %d, %Y'),
            })

        pagination = {
            'current_page': page_obj.number,
            'num_pages': paginator.num_pages,
            'has_next': page_obj.has_next(),
            'has_previous': page_obj.has_previous(),
            'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
            'start_index': page_obj.start_index(),
            'end_index': page_obj.end_index(),
            'total_count': total_count,
        }

        return JsonResponse({
            'status': 'success',
            'students': student_data,
            'pagination': pagination,
            'total_count': total_count,
        })
    except Exception as e:
        import traceback
        return JsonResponse({
            'status': 'error',
            'message': str(e),
            'traceback': traceback.format_exc()
        }, status=500)


@login_required
@user_passes_test(is_admin_or_registrar)
@require_http_methods(["POST"])
def registrar_enroll_face(request):
    """Handle face enrollment submission"""
    try:
        import json
        import numpy as np
        from django.core.files.base import ContentFile
        from django.conf import settings
        import os
        import base64
        
        student_id = request.POST.get('student_id')
        embeddings_json = request.POST.get('embeddings')
        
        if not student_id or not embeddings_json:
            return JsonResponse({
                'status': 'error',
                'message': 'Missing required data'
            }, status=400)
        
        try:
            student = Student.objects.get(id=student_id)
        except Student.DoesNotExist:
            return JsonResponse({
                'status': 'error',
                'message': 'Student not found'
            }, status=404)
        
        # Check if this is an update (student already has face_path)
        is_update = bool(student.face_path)
        
        # Parse embeddings
        embeddings = json.loads(embeddings_json)
        
        # Average the three embeddings (front, left, right)
        embeddings_array = np.array(embeddings)
        averaged_embedding = np.mean(embeddings_array, axis=0)
        
        # Save embedding to file
        face_embeddings_dir = os.path.join(settings.BASE_DIR, 'face_embeddings')
        os.makedirs(face_embeddings_dir, exist_ok=True)
        
        from datetime import datetime
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        embedding_filename = f'student_{student.lrn}_{timestamp}.npy'
        embedding_path = os.path.join(face_embeddings_dir, embedding_filename)
        
        # Save the numpy array
        np.save(embedding_path, averaged_embedding)
        
        # Create student folder for images using LRN
        student_images_dir = os.path.join(settings.MEDIA_ROOT, 'face_enrollment_images', f'student_{student.lrn}')
        os.makedirs(student_images_dir, exist_ok=True)
        
        # Save images (front, left, right)
        for position in ['front', 'left', 'right']:
            image_file = request.FILES.get(f'image_{position}')
            if image_file:
                image_filename = f'{position}_{timestamp}.jpg'
                image_path = os.path.join(student_images_dir, image_filename)
                with open(image_path, 'wb') as f:
                    for chunk in image_file.chunks():
                        f.write(chunk)
        
        # Update student record with face_path
        student.face_path = embedding_filename
        student.save()
        
        # Force reload embeddings cache immediately so new face is available for recognition
        from PROTECHAPP.face_recognition_engine import face_engine
        face_engine.load_all_embeddings()
        
        action_message = 'updated' if is_update else 'enrolled'
        return JsonResponse({
            'status': 'success',
            'message': f'Face {action_message} successfully for {student.first_name} {student.last_name}'
        })
        
    except Exception as e:
        import traceback
        return JsonResponse({
            'status': 'error',
            'message': str(e),
            'traceback': traceback.format_exc()
        }, status=500)
    
    # Apply enrollment status filter
    if enrollment_status == 'enrolled':
        students = students.filter(face_path__isnull=False).exclude(face_path='')
    elif enrollment_status == 'not_enrolled':
        students = students.filter(Q(face_path__isnull=True) | Q(face_path=''))

    total_count = students.count()
    
    try:
        page_number = int(page_number)
        items_per_page = int(items_per_page)
    except (ValueError, TypeError):
        page_number = 1
        items_per_page = 12

    paginator = Paginator(students, items_per_page)
    page_obj = paginator.get_page(page_number)

    students_data = []
    for student in page_obj:
        has_face = bool(student.face_path and student.face_path.strip())
        students_data.append({
            'id': student.id,
            'lrn': student.lrn,
            'full_name': f"{student.first_name} {student.middle_name or ''} {student.last_name}".strip(),
            'first_name': student.first_name,
            'last_name': student.last_name,
            'grade': student.grade.name if student.grade else '',
            'grade_id': student.grade.id if student.grade else None,
            'section': student.section.name if student.section else '',
            'section_id': student.section.id if student.section else None,
            'profile_pic': student.profile_pic or None,
            'has_face_enrolled': has_face,
        })

    pagination = {
        'current_page': page_obj.number,
        'num_pages': paginator.num_pages,
        'has_next': page_obj.has_next(),
        'has_previous': page_obj.has_previous(),
        'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
        'start_index': page_obj.start_index(),
        'end_index': page_obj.end_index(),
        'total_count': total_count,
    }

    return JsonResponse({
        'status': 'success',
        'students': students_data,
        'pagination': pagination,
        'total_count': total_count,
    })

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["GET"])
def registrar_search_students_for_face_enrollment(request):
    """Search students for face enrollment by name or LRN"""
    try:
        query = request.GET.get('q', '').strip()
        
        if not query or len(query) < 2:
            return JsonResponse({
                'status': 'success',
                'students': []
            })
        
        # Search by first name, last name, or LRN
        students = Student.objects.filter(
            Q(first_name__icontains=query) |
            Q(last_name__icontains=query) |
            Q(lrn__icontains=query),
            status='ACTIVE'
        ).select_related('grade', 'section').order_by('first_name', 'last_name')[:20]
        
        students_data = []
        for student in students:
            has_face = bool(student.face_path and student.face_path.strip())
            students_data.append({
                'id': student.id,
                'lrn': student.lrn,
                'first_name': student.first_name,
                'middle_name': student.middle_name or '',
                'last_name': student.last_name,
                'full_name': f"{student.first_name} {student.middle_name or ''} {student.last_name}".replace('  ', ' '),
                'grade': student.grade.name,
                'section': student.section.name,
                'grade_section': f"{student.grade.name} - {student.section.name}",
                'has_face_enrolled': has_face,
                'profile_pic': student.profile_pic if student.profile_pic else None
            })
        
        return JsonResponse({
            'status': 'success',
            'students': students_data
        })
        
    except Exception as e:
        import traceback
        print("Error in registrar_search_students_for_face_enrollment:", str(e))
        print(traceback.format_exc())
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["POST"])
def registrar_save_face_embedding(request):
    """Save face embedding for a student"""
    try:
        import json
        import base64
        import numpy as np
        from datetime import datetime
        
        data = json.loads(request.body)
        student_id = data.get('student_id')
        face_embeddings = data.get('face_embeddings')  # Array of embeddings from multiple angles
        
        if not student_id or not face_embeddings:
            return JsonResponse({'status': 'error', 'message': 'Student ID and face embeddings are required'}, status=400)
        
        # Get student
        try:
            student = Student.objects.get(id=student_id, status='ACTIVE')
        except Student.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Student not found'}, status=404)
        
        # Create face embeddings directory if it doesn't exist
        face_embeddings_dir = os.path.join(settings.BASE_DIR, 'face_embeddings')
        os.makedirs(face_embeddings_dir, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"student_{student.id}_{timestamp}.npy"
        file_path = os.path.join(face_embeddings_dir, filename)
        
        # Delete old face embedding file if exists
        if student.face_path and student.face_path.strip():
            old_file_path = os.path.join(settings.BASE_DIR, student.face_path)
            if os.path.exists(old_file_path):
                try:
                    os.remove(old_file_path)
                except Exception as e:
                    print(f"Error deleting old face embedding: {str(e)}")
        
        # Convert embeddings to numpy array and save
        embeddings_array = np.array(face_embeddings)
        np.save(file_path, embeddings_array)
        
        # Update student record with relative path
        relative_path = os.path.join('face_embeddings', filename)
        student.face_path = relative_path
        student.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Face enrolled successfully for {student.first_name} {student.last_name}',
            'student_id': student.id,
            'face_path': relative_path
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON data'}, status=400)
    except Exception as e:
        import traceback
        print("Error in registrar_save_face_embedding:", str(e))
        print(traceback.format_exc())
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["GET"])
def registrar_get_student_face_status(request, student_id):
    """Check if a student has face enrolled"""
    try:
        student = get_object_or_404(Student, id=student_id, status='ACTIVE')
        
        has_face = bool(student.face_path and student.face_path.strip())
        face_file_exists = False
        
        if has_face:
            face_file_path = os.path.join(settings.BASE_DIR, student.face_path)
            face_file_exists = os.path.exists(face_file_path)
        
        return JsonResponse({
            'status': 'success',
            'student_id': student.id,
            'student_name': f"{student.first_name} {student.last_name}",
            'has_face_enrolled': has_face,
            'face_file_exists': face_file_exists,
            'face_path': student.face_path if has_face else None,
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
@require_http_methods(["DELETE"])
def registrar_delete_face_embedding(request, student_id):
    """Delete face embedding for a student"""
    try:
        student = get_object_or_404(Student, id=student_id, status='ACTIVE')
        
        if not student.face_path or not student.face_path.strip():
            return JsonResponse({'status': 'error', 'message': 'No face enrollment found for this student'}, status=404)
        
        # Delete the face embedding file
        face_file_path = os.path.join(settings.BASE_DIR, student.face_path)
        if os.path.exists(face_file_path):
            try:
                os.remove(face_file_path)
            except Exception as e:
                print(f"Error deleting face file: {str(e)}")
        
        # Clear face_path from student record
        student.face_path = None
        student.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Face enrollment deleted for {student.first_name} {student.last_name}'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

# ==========================
# EXPORT FUNCTIONS
# ==========================

@login_required
@user_passes_test(is_registrar)
def export_registrar_students(request):
    """Export students to Excel, PDF, or Word"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError as e:
        return JsonResponse({'success': False, 'error': f'Required library not installed: {str(e)}'}, status=500)
    
    export_format = request.GET.get('format', 'excel')
    students = Student.objects.all().select_related('grade', 'section').order_by('id')
    
    headers = ['ID', 'LRN', 'First Name', 'Middle Name', 'Last Name', 'Email', 'Grade', 'Section', 'Status', 'Face Enrolled']
    data_rows = []
    
    for student in students:
        data_rows.append([
            student.id,
            student.lrn or '',
            student.first_name,
            student.middle_name or '',
            student.last_name,
            student.email or '',
            student.grade.name if student.grade else '',
            student.section.name if student.section else '',
            student.status,
            'Yes' if student.face_path else 'No'
        ])
    
    if export_format == 'pdf':
        return export_registrar_data_to_pdf(headers, data_rows, "Students")
    elif export_format == 'word':
        return export_registrar_data_to_word(headers, data_rows, "Students")
    else:
        return export_registrar_data_to_excel(headers, data_rows, "Students")

@login_required
@user_passes_test(is_registrar)
def export_registrar_guardians(request):
    """Export guardians to Excel, PDF, or Word"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError as e:
        return JsonResponse({'success': False, 'error': f'Required library not installed: {str(e)}'}, status=500)
    
    export_format = request.GET.get('format', 'excel')
    guardians = Guardian.objects.all().select_related('student').order_by('id')
    
    headers = ['ID', 'First Name', 'Middle Name', 'Last Name', 'Relationship', 'Contact Number', 'Email', 'Student Name']
    data_rows = []
    
    for guardian in guardians:
        data_rows.append([
            guardian.id,
            guardian.first_name,
            guardian.middle_name or '',
            guardian.last_name,
            guardian.relationship or '',
            guardian.contact_number or '',
            guardian.email or '',
            f"{guardian.student.first_name} {guardian.student.last_name}" if guardian.student else ''
        ])
    
    if export_format == 'pdf':
        return export_registrar_data_to_pdf(headers, data_rows, "Guardians")
    elif export_format == 'word':
        return export_registrar_data_to_word(headers, data_rows, "Guardians")
    else:
        return export_registrar_data_to_excel(headers, data_rows, "Guardians")

@login_required
@user_passes_test(is_registrar)
def export_registrar_grades(request):
    """Export grades to Excel, PDF, or Word"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError as e:
        return JsonResponse({'success': False, 'error': f'Required library not installed: {str(e)}'}, status=500)
    
    export_format = request.GET.get('format', 'excel')
    grades = Grade.objects.all().order_by('id')
    
    headers = ['ID', 'Grade Name']
    data_rows = [[grade.id, grade.name] for grade in grades]
    
    if export_format == 'pdf':
        return export_registrar_data_to_pdf(headers, data_rows, "Grades")
    elif export_format == 'word':
        return export_registrar_data_to_word(headers, data_rows, "Grades")
    else:
        return export_registrar_data_to_excel(headers, data_rows, "Grades")

@login_required
@user_passes_test(is_registrar)
def export_registrar_sections(request):
    """Export sections to Excel, PDF, or Word"""
    try:
        export_format = request.GET.get('format', 'excel')
        search_query = request.GET.get('search', '').strip()
        grade_filter = request.GET.get('grade', '').strip()
        advisor_filter = request.GET.get('advisor', '').strip()
        
        # Use select_related for grade and prefetch advisory assignments
        sections = Section.objects.select_related('grade').prefetch_related(
            'advisory_assignments__teacher'
        ).all().order_by('id')
        
        if search_query:
            sections = sections.filter(
                Q(name__icontains=search_query) |
                Q(grade__name__icontains=search_query)
            )
        
        if grade_filter:
            sections = sections.filter(grade__id=grade_filter)
        
        if advisor_filter:
            if advisor_filter == 'with_advisor':
                sections = sections.filter(advisory_assignments__isnull=False).distinct()
            elif advisor_filter == 'without_advisor':
                sections = sections.filter(advisory_assignments__isnull=True)
        
        headers = ['Section Name', 'Grade', 'Students', 'Advisor']
        
        data_rows = []
        for section in sections:
            # Get student count
            student_count = section.students.count()
            
            # Get advisor name
            advisor = section.advisory_assignments.first()
            advisor_name = ''
            if advisor and advisor.teacher:
                advisor_name = f"{advisor.teacher.first_name} {advisor.teacher.last_name}"

            data_rows.append([
                section.name,
                section.grade.name if section.grade else "",
                str(student_count),
                advisor_name
            ])
        
        if export_format == 'pdf':
            return export_registrar_data_to_pdf(headers, data_rows, "Sections")
        elif export_format == 'word':
            return export_registrar_data_to_word(headers, data_rows, "Sections")
        else:
            return export_registrar_data_to_excel(headers, data_rows, "Sections")
    except Exception as e:
        import traceback
        print(f"Export sections error: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
def export_registrar_attendance(request):
    """Export attendance records to Excel, PDF, or Word"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError as e:
        return JsonResponse({'success': False, 'error': f'Required library not installed: {str(e)}'}, status=500)
    
    try:
        export_format = request.GET.get('format', 'excel')
        date_filter = request.GET.get('date', str(timezone.now().date()))
        
        attendance_records = Attendance.objects.filter(date=date_filter).select_related('student').order_by('id')
        
        headers = ['ID', 'Student Name', 'LRN', 'Date', 'Time In', 'Time Out', 'Status']
        data_rows = []
        
        for record in attendance_records:
            data_rows.append([
                str(record.id),
                f"{record.student.first_name} {record.student.last_name}",
                record.student.lrn or '',
                str(record.date),
                str(record.time_in) if record.time_in else '',
                str(record.time_out) if record.time_out else '',
                record.status or ''
            ])
        
        title = f"Attendance {date_filter}"
        
        if export_format == 'pdf':
            return export_registrar_data_to_pdf(headers, data_rows, title)
        elif export_format == 'word':
            return export_registrar_data_to_word(headers, data_rows, title)
        else:
            return export_registrar_data_to_excel(headers, data_rows, title)
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"Export error: {error_trace}")
        return JsonResponse({'success': False, 'error': f'Export failed: {str(e)}'}, status=500)

@login_required
@user_passes_test(is_registrar)
def export_registrar_excused(request):
    """Export excused absences to Excel, PDF, or Word"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError as e:
        return JsonResponse({'success': False, 'error': f'Required library not installed: {str(e)}'}, status=500)
    
    export_format = request.GET.get('format', 'excel')
    
    excused_absences = ExcusedAbsence.objects.all().select_related('student').order_by('id')

    headers = ['ID', 'Student Name', 'LRN', 'Date Absent', 'Effective Date', 'End Date', 'Has Letter']
    data_rows = []

    for excuse in excused_absences:
        student_name = f"{excuse.student.first_name} {excuse.student.last_name}"
        has_letter = 'Yes' if excuse.excuse_letter else 'No'
        data_rows.append([
            str(excuse.id),
            student_name,
            excuse.student.lrn or '',
            excuse.date_absent.strftime('%Y-%m-%d') if excuse.date_absent else '',
            excuse.effective_date.strftime('%Y-%m-%d') if excuse.effective_date else '',
            excuse.end_date.strftime('%Y-%m-%d') if excuse.end_date else '',
            has_letter
        ])
    
    if export_format == 'pdf':
        return export_registrar_data_to_pdf(headers, data_rows, "Excused Absences")
    elif export_format == 'word':
        return export_registrar_data_to_word(headers, data_rows, "Excused Absences")
    else:
        return export_registrar_data_to_excel(headers, data_rows, "Excused Absences")

def export_registrar_data_to_excel(headers, data_rows, title):
    """Export data to Excel format"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    
    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]
    
    # Write headers
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = thin_border
    
    # Write data rows
    for row_num, row_data in enumerate(data_rows, 2):
        for col_num, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='left', vertical='center')
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column = [cell for cell in column]
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[get_column_letter(column[0].column)].width = adjusted_width
    
    # Save to HTTP response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{title}.xlsx"'
    wb.save(response)
    return response

def export_registrar_data_to_pdf(headers, data_rows, title):
    """Export data to PDF format"""
    try:
        from reportlab.lib.pagesizes import letter, landscape
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        
        # Ensure all data is string
        headers = [str(h) for h in headers]
        data_rows = [[str(cell) if cell is not None else "" for cell in row] for row in data_rows]
        
        # If no data, add a message row
        if not data_rows:
            data_rows = [["No records found"] + [""] * (len(headers) - 1)]
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), topMargin=0.5*inch, bottomMargin=0.5*inch)
        elements = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = styles['Heading1']
        title_style.textColor = colors.HexColor('#1F4E78')
        elements.append(Paragraph(str(title), title_style))
        elements.append(Spacer(1, 0.3*inch))
        
        # Table
        table_data = [headers] + data_rows
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E78')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F0F0')]),
        ]))
        elements.append(table)
        
        doc.build(elements)
        buffer.seek(0)
        
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{title}.pdf"'
        return response
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"PDF export error: {error_trace}")
        # Re-raise the exception to be caught by the parent function
        raise

def export_registrar_data_to_word(headers, data_rows, title):
    """Export data to Word format"""
    from docx import Document
    from docx.shared import Inches, RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.color.rgb = RGBColor(31, 78, 120)
    
    # Table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(31, 78, 120)
    
    # Data rows
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{title}.docx"'
    return response

# ==========================
# IMPORT FUNCTIONS
# ==========================

@login_required
@user_passes_test(is_registrar)
def import_registrar_grades(request):
    """Import grades from CSV or Excel file"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Invalid request method'}, status=400)
    
    try:
        import pandas as pd
    except ImportError:
        return JsonResponse({'success': False, 'error': 'pandas library not installed'}, status=500)
    
    if 'import_file' not in request.FILES:
        return JsonResponse({'success': False, 'error': 'No file uploaded'}, status=400)
    
    file = request.FILES['import_file']
    
    try:
        # Read file
        if file.name.endswith('.csv'):
            df = pd.read_csv(file)
        elif file.name.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(file)
        else:
            return JsonResponse({'success': False, 'error': 'Invalid file format. Use CSV or Excel'}, status=400)
        
        # Normalize column names
        df.columns = df.columns.str.strip().str.lower().str.replace(' ', '_')
        
        # Map column variations
        column_map = {
            'id': ['id', 'grade_id'],
            'name': ['name', 'grade_name', 'grade']
        }
        
        # Rename columns
        for standard, variations in column_map.items():
            for var in variations:
                if var in df.columns:
                    df.rename(columns={var: standard}, inplace=True)
                    break
        
        # Validate required columns
        if 'name' not in df.columns:
            return JsonResponse({'success': False, 'error': 'Missing required column: Name'}, status=400)
        
        created_count = 0
        updated_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                grade_name = str(row['name']).strip()
                
                if not grade_name or grade_name.lower() == 'nan':
                    continue
                
                # Check if ID is provided in Excel
                grade_id = row.get('id') if 'id' in row and pd.notna(row['id']) else None
                
                if grade_id:
                    # Skip rows with existing ID - don't import or update
                    try:
                        if Grade.objects.filter(id=int(grade_id)).exists():
                            skipped_count += 1
                            continue
                    except ValueError:
                        pass
                
                # Only import new grades (blank ID)
                Grade.objects.create(name=grade_name)
                created_count += 1
                        
            except Exception as e:
                errors.append(f"Row {index + 2}: {str(e)}")
        
        message = f"Successfully imported grades. Created: {created_count}, Updated: {updated_count}"
        if errors:
            message += f". Errors: {len(errors)}"
        
        return JsonResponse({
            'success': True,
            'message': message,
            'created': created_count,
            'updated': updated_count,
            'errors': errors[:10]  # Limit errors shown
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
def download_registrar_grades_template(request):
    """Download Excel template for grade import"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Grades"
    
    # Headers
    headers = ['ID', 'Grade Name']
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 20
    
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Add instruction row
    ws.cell(row=2, column=1, value="Instructions:")
    ws.cell(row=2, column=2, value="Leave ID blank for new grades. Rows with existing IDs will be SKIPPED (not imported).")
    ws.merge_cells('B2:B2')
    ws.cell(row=2, column=1).font = Font(italic=True, color="666666", size=10)
    ws.cell(row=2, column=2).font = Font(italic=True, color="666666", size=10)
    ws.cell(row=2, column=2).alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
    ws.row_dimensions[2].height = 30
    
    # Sample data
    sample_data = [
        ['', 'Grade 7'],
        ['', 'Grade 8'],
        ['', 'Grade 9']
    ]
    
    for row_num, row_data in enumerate(sample_data, 3):
        for col_num, value in enumerate(row_data, 1):
            ws.cell(row=row_num, column=col_num, value=value)
    
    # Save to response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="grades_template.xlsx"'
    wb.save(response)
    return response

@login_required
@user_passes_test(is_registrar)
def import_registrar_sections(request):
    """Import sections from CSV or Excel file"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Invalid request method'}, status=400)
    
    try:
        import pandas as pd
    except ImportError:
        return JsonResponse({'success': False, 'error': 'pandas library not installed'}, status=500)
    
    if 'import_file' not in request.FILES:
        return JsonResponse({'success': False, 'error': 'No file uploaded'}, status=400)
    
    file = request.FILES['import_file']
    
    try:
        # Read file
        if file.name.endswith('.csv'):
            df = pd.read_csv(file)
        elif file.name.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(file)
        else:
            return JsonResponse({'success': False, 'error': 'Invalid file format'}, status=400)
        
        # Normalize columns
        df.columns = df.columns.str.strip().str.lower().str.replace(' ', '_')
        
        # Map column variations
        column_map = {
            'id': ['id', 'section_id'],
            'name': ['name', 'section_name', 'section'],
            'grade': ['grade', 'grade_name'],
            'room_number': ['room_number', 'room', 'room_no'],
            'capacity': ['capacity', 'max_capacity']
        }
        
        for standard, variations in column_map.items():
            for var in variations:
                if var in df.columns:
                    df.rename(columns={var: standard}, inplace=True)
                    break
        
        # Validate required columns
        if 'name' not in df.columns or 'grade' not in df.columns:
            return JsonResponse({'success': False, 'error': 'Missing required columns: Name, Grade'}, status=400)
        
        created_count = 0
        updated_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                section_name = str(row['name']).strip()
                grade_name = str(row['grade']).strip()
                
                if not section_name or not grade_name or section_name.lower() == 'nan' or grade_name.lower() == 'nan':
                    continue
                
                # Get grade
                try:
                    grade = Grade.objects.get(name=grade_name)
                except Grade.DoesNotExist:
                    errors.append(f"Row {index + 2}: Grade '{grade_name}' not found")
                    continue
                
                section_data = {
                    'name': section_name,
                    'grade': grade
                }
                
                if 'room_number' in df.columns and pd.notna(row.get('room_number')):
                    section_data['room_number'] = str(row['room_number']).strip()
                
                if 'capacity' in df.columns and pd.notna(row.get('capacity')):
                    try:
                        section_data['capacity'] = int(row['capacity'])
                    except:
                        pass
                
                # Check if ID is provided in Excel
                section_id = row.get('id') if 'id' in row and pd.notna(row['id']) else None
                
                if section_id:
                    # Skip rows with existing ID - don't import or update
                    try:
                        if Section.objects.filter(id=int(section_id)).exists():
                            skipped_count += 1
                            continue
                    except ValueError:
                        pass
                
                # Only import new sections (blank ID)
                Section.objects.create(**section_data)
                created_count += 1
                        
            except Exception as e:
                errors.append(f"Row {index + 2}: {str(e)}")
        
        message = f"Successfully imported {created_count} sections. Skipped {skipped_count} existing records."
        if errors:
            message += f" Errors: {len(errors)}"
        
        return JsonResponse({
            'success': True,
            'message': message,
            'created': created_count,
            'skipped': skipped_count,
            'errors': errors[:10]
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
def download_registrar_sections_template(request):
    """Download Excel template for section import"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Sections"
    
    # Headers
    headers = ['ID', 'Section Name', 'Grade', 'Room Number', 'Capacity']
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Add instruction row
    ws.cell(row=2, column=1, value="Instructions:")
    ws.cell(row=2, column=2, value="Leave ID blank for new sections. Rows with existing IDs will be SKIPPED (not imported).")
    ws.cell(row=2, column=1).font = Font(italic=True, color="666666", size=10)
    ws.cell(row=2, column=2).font = Font(italic=True, color="666666", size=10)
    ws.cell(row=2, column=2).alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
    ws.merge_cells('B2:E2')
    ws.row_dimensions[2].height = 30
    
    # Sample data
    sample_data = [
        ['', 'A', 'Grade 7', 'Room 101', '40'],
        ['', 'B', 'Grade 7', 'Room 102', '40'],
        ['', 'A', 'Grade 8', 'Room 201', '40']
    ]
    
    for row_num, row_data in enumerate(sample_data, 3):
        for col_num, value in enumerate(row_data, 1):
            ws.cell(row=row_num, column=col_num, value=value)
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 15
    ws.column_dimensions['E'].width = 12
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="sections_template.xlsx"'
    wb.save(response)
    return response

@login_required
@user_passes_test(is_registrar)
def import_registrar_guardians(request):
    """Import guardians from CSV or Excel file"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Invalid request method'}, status=400)
    
    try:
        import pandas as pd
    except ImportError:
        return JsonResponse({'success': False, 'error': 'pandas library not installed'}, status=500)
    
    if 'import_file' not in request.FILES:
        return JsonResponse({'success': False, 'error': 'No file uploaded'}, status=400)
    
    file = request.FILES['import_file']
    
    try:
        # Read file
        if file.name.endswith('.csv'):
            df = pd.read_csv(file)
        elif file.name.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(file)
        else:
            return JsonResponse({'success': False, 'error': 'Invalid file format'}, status=400)
        
        # Normalize columns
        df.columns = df.columns.str.strip().str.lower().str.replace(' ', '_')
        
        # Map column variations
        column_map = {
            'id': ['id', 'guardian_id'],
            'first_name': ['first_name', 'firstname', 'fname'],
            'middle_name': ['middle_name', 'middlename', 'mname'],
            'last_name': ['last_name', 'lastname', 'lname'],
            'relationship': ['relationship', 'relation'],
            'contact_number': ['contact_number', 'phone', 'mobile', 'contact'],
            'email': ['email', 'email_address'],
            'student_name': ['student_name', 'student']
        }
        
        for standard, variations in column_map.items():
            for var in variations:
                if var in df.columns:
                    df.rename(columns={var: standard}, inplace=True)
                    break
        
        # Validate required columns
        required = ['first_name', 'last_name', 'student_name']
        missing = [col for col in required if col not in df.columns]
        if missing:
            return JsonResponse({'success': False, 'error': f'Missing required columns: {", ".join(missing)}'}, status=400)
        
        created_count = 0
        updated_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                first_name = str(row['first_name']).strip()
                last_name = str(row['last_name']).strip()
                student_name = str(row['student_name']).strip()
                
                if not first_name or not last_name or not student_name:
                    continue
                
                # Find student
                student_parts = student_name.split()
                if len(student_parts) < 2:
                    errors.append(f"Row {index + 2}: Invalid student name format")
                    continue
                
                student_fname = student_parts[0]
                student_lname = student_parts[-1]
                
                try:
                    student = Student.objects.get(
                        first_name__icontains=student_fname,
                        last_name__icontains=student_lname
                    )
                except Student.DoesNotExist:
                    errors.append(f"Row {index + 2}: Student '{student_name}' not found")
                    continue
                except Student.MultipleObjectsReturned:
                    errors.append(f"Row {index + 2}: Multiple students found for '{student_name}'")
                    continue
                
                guardian_data = {
                    'first_name': first_name,
                    'last_name': last_name,
                    'student': student
                }
                
                if 'middle_name' in df.columns and pd.notna(row.get('middle_name')):
                    guardian_data['middle_name'] = str(row['middle_name']).strip()
                
                if 'relationship' in df.columns and pd.notna(row.get('relationship')):
                    guardian_data['relationship'] = str(row['relationship']).strip()
                
                if 'contact_number' in df.columns and pd.notna(row.get('contact_number')):
                    guardian_data['contact_number'] = str(row['contact_number']).strip()
                
                if 'email' in df.columns and pd.notna(row.get('email')):
                    guardian_data['email'] = str(row['email']).strip()
                
                # Check if ID is provided in Excel
                guardian_id = row.get('id') if 'id' in row and pd.notna(row['id']) else None
                
                if guardian_id:
                    # Skip rows with existing ID - don't import or update
                    try:
                        if Guardian.objects.filter(id=int(guardian_id)).exists():
                            skipped_count += 1
                            continue
                    except ValueError:
                        pass
                
                # Only import new guardians (blank ID)
                Guardian.objects.create(**guardian_data)
                created_count += 1
                        
            except Exception as e:
                errors.append(f"Row {index + 2}: {str(e)}")
        
        message = f"Successfully imported {created_count} guardians. Skipped {skipped_count} existing records."
        if errors:
            message += f" Errors: {len(errors)}"
        
        return JsonResponse({
            'success': True,
            'message': message,
            'created': created_count,
            'skipped': skipped_count,
            'errors': errors[:10]
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@login_required
@user_passes_test(is_registrar)
def download_registrar_guardians_template(request):
    """Download Excel template for guardian import"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Guardians"
    
    # Headers
    headers = ['ID', 'First Name', 'Middle Name', 'Last Name', 'Relationship', 'Contact Number', 'Email', 'Student Name']
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Add instruction row
    ws.cell(row=2, column=1, value="Instructions:")
    ws.cell(row=2, column=2, value="Leave ID blank for new guardians. Rows with existing IDs will be SKIPPED (not imported).")
    ws.cell(row=2, column=1).font = Font(italic=True, color="666666", size=10)
    ws.cell(row=2, column=2).font = Font(italic=True, color="666666", size=10)
    ws.cell(row=2, column=2).alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
    ws.merge_cells('B2:H2')
    ws.row_dimensions[2].height = 30
    
    # Sample data
    sample_data = [
        ['', 'John', '', 'Doe', 'FATHER', '09123456789', 'john.doe@email.com', 'Jane Doe'],
        ['', 'Maria', '', 'Smith', 'MOTHER', '09987654321', 'maria.smith@email.com', 'John Smith']
    ]
    
    for row_num, row_data in enumerate(sample_data, 3):
        for col_num, value in enumerate(row_data, 1):
            ws.cell(row=row_num, column=col_num, value=value)
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 8
    for i in range(2, 9):
        ws.column_dimensions[get_column_letter(i)].width = 18
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="guardians_template.xlsx"'
    wb.save(response)
    return response

