from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.views.decorators.http import require_http_methods
from django.core.paginator import Paginator
from PROTECHAPP.views.admin_views import get_pagination_range
from django.utils import timezone
from django.contrib import messages
from django.http import HttpResponse, JsonResponse
from django.db.models import Q
from PROTECHAPP.models import CustomUser, Student, Attendance, Guardian, UserRole, Section
import io
from datetime import datetime


@login_required
@require_http_methods(["GET"])
def teacher_search_attendance_records(request):
    """AJAX search/filter for advisory teacher attendance records (paginated)"""
    # Runtime authorization check to avoid referencing is_advisory_teacher at import time
    try:
        if not is_advisory_teacher(request.user):
            return JsonResponse({'status': 'error', 'message': 'Unauthorized', 'records': [], 'pagination': {}, 'total_count': 0}, status=403)
    except Exception:
        return JsonResponse({'status': 'error', 'message': 'Unauthorized', 'records': [], 'pagination': {}, 'total_count': 0}, status=403)
    try:
        section = request.user.section
        if not section:
            return JsonResponse({'status': 'error', 'records': [], 'pagination': {}, 'total_count': 0}, status=400)

        search_query = request.GET.get('search', '')
        status_filter = request.GET.get('status', '')
        date_start = request.GET.get('date_start', '')
        date_end = request.GET.get('date_end', '')
        date_single = request.GET.get('date', '')
        page_number = request.GET.get('page', 1)
        items_per_page = request.GET.get('items_per_page', 10)

        records = Attendance.objects.select_related('student', 'student__grade', 'student__section').filter(student__section=section).order_by('-date', '-student__last_name')

        if search_query:
            records = records.filter(
                Q(student__first_name__icontains=search_query) |
                Q(student__last_name__icontains=search_query) |
                Q(student__lrn__icontains=search_query)
            )
        if status_filter:
            records = records.filter(status=status_filter)

        # date range or single date
        if date_start and date_end:
            try:
                start = datetime.strptime(date_start, '%Y-%m-%d').date()
                end = datetime.strptime(date_end, '%Y-%m-%d').date()
                records = records.filter(date__range=(start, end))
            except Exception:
                pass
        elif date_single:
            try:
                d = datetime.strptime(date_single, '%Y-%m-%d').date()
                records = records.filter(date=d)
            except Exception:
                pass

        total_count = records.count()
        try:
            page_number = int(page_number)
            items_per_page = int(items_per_page)
        except (ValueError, TypeError):
            page_number = 1
            items_per_page = 10

        paginator = Paginator(records, items_per_page)
        page_obj = paginator.get_page(page_number)

        attendance_data = []
        for record in page_obj:
            attendance_data.append({
                'id': record.id,
                'lrn': record.student.lrn,
                'full_name': f"{record.student.first_name} {record.student.last_name}",
                'grade': record.student.grade.name if record.student.grade else '',
                'section': record.student.section.name if record.student.section else '',
                'date': record.date.strftime('%Y-%m-%d'),
                'time_in': record.time_in.strftime('%H:%M:%S') if record.time_in else '',
                'time_out': record.time_out.strftime('%H:%M:%S') if record.time_out else '',
                'status': record.status,
                'profile_pic': record.student.profile_pic or None,
            })

        pagination = {
            'current_page': page_obj.number,
            'num_pages': paginator.num_pages,
            'has_next': page_obj.has_next(),
            'has_previous': page_obj.has_previous(),
            'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
            'start_index': page_obj.start_index(),
            'end_index': page_obj.end_index(),
            'total_count': total_count,
        }

        return JsonResponse({
            'status': 'success',
            'records': attendance_data,
            'pagination': pagination,
            'total_count': total_count,
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e),
            'records': [],
            'pagination': {
                'current_page': 1,
                'num_pages': 1,
                'has_next': False,
                'has_previous': False,
                'page_range': [1],
                'start_index': 0,
                'end_index': 0,
                'total_count': 0,
            },
            'total_count': 0
        }, status=500)
def is_advisory_teacher(user):
    """Check if the logged-in user is an advisory teacher"""
    return user.is_authenticated and user.role == UserRole.TEACHER and user.section is not None

@login_required
@user_passes_test(is_advisory_teacher)
def teacher_advisory_dashboard(request):
    """View for advisory teacher dashboard"""
    # Get current class attendance statistics
    current_date = timezone.now()
    section = request.user.section
    
    # Get students in the teacher's advisory class
    students_in_section = Student.objects.filter(section=section).count() if section else 0
    
    # Example data for dashboard
    attendance_percentage = 96  # This would come from a database query
    
    context = {
        'current_date': current_date,
        'students_count': students_in_section,
        'attendance_percentage': attendance_percentage
    }
    
    return render(request, 'teacher/advisory/dashboard.html', context)

@login_required
@user_passes_test(is_advisory_teacher)
def teacher_advisory_students(request):
    """View for advisory students listing"""
    section = request.user.section
    students_list = Student.objects.filter(section=section).select_related('grade', 'section').order_by('last_name', 'first_name') if section else []
    
    # Calculate statistics
    total_students = students_list.count() if section else 0
    active_students = students_list.filter(status='ACTIVE').count() if section else 0
    face_enrolled = students_list.exclude(face_path__isnull=True).exclude(face_path='').count() if section else 0
    
    # Send ALL students to template - pagination handled by JavaScript
    context = {
        'students': students_list,
        'section': section,
        'total_students': total_students,
        'active_students': active_students,
        'face_enrolled': face_enrolled,
    }
    return render(request, 'teacher/advisory/students.html', context)

@login_required
@user_passes_test(is_advisory_teacher)
def teacher_advisory_attendance(request):
    """View for advisory class attendance records"""
    section = request.user.section
    today = timezone.now().date()
    # Support either a single `date` param or a `date_start`/`date_end` range (used by the UI)
    date_param = request.GET.get('date')
    date_start = request.GET.get('date_start')
    date_end = request.GET.get('date_end')

    # Build queryset for attendance records in the requested range
    if section:
        qs = Attendance.objects.filter(student__section=section).select_related('student')
        if date_start and date_end:
            # Expect format YYYY-MM-DD from the UI
            try:
                # Simple validation/parsing
                start = datetime.strptime(date_start, '%Y-%m-%d').date()
                end = datetime.strptime(date_end, '%Y-%m-%d').date()
                attendance_records = qs.filter(date__range=(start, end)).order_by('-date', '-time_in')
                # expose selected range to template
                date_filter = f"{start} to {end}"
            except Exception:
                # fallback to today if parsing fails
                attendance_records = qs.filter(date=today).order_by('-time_in')
                date_filter = str(today)
        elif date_param:
            try:
                d = datetime.strptime(date_param, '%Y-%m-%d').date()
                attendance_records = qs.filter(date=d).order_by('-time_in')
                date_filter = date_param
            except Exception:
                attendance_records = qs.filter(date=today).order_by('-time_in')
                date_filter = str(today)
        else:
            attendance_records = qs.filter(date=today).order_by('-time_in')
            date_filter = str(today)
    else:
        attendance_records = []
    
    # Get statistics
    total_students = Student.objects.filter(section=section).count() if section else 0
    present_count = attendance_records.filter(status__in=['On Time', 'Late']).count()
    absent_count = total_students - present_count
    late_count = attendance_records.filter(status='Late').count()
    
    context = {
        'attendance_records': attendance_records,
        'section': section,
        'date_filter': date_filter,
        'today': today,
        'total_students': total_students,
        'present_count': present_count,
        'absent_count': absent_count,
        'late_count': late_count,
    }
    return render(request, 'teacher/advisory/attendance.html', context)

@login_required
@user_passes_test(is_advisory_teacher)
def teacher_advisory_excused(request):
    """View for excused absences for advisory class"""
    from PROTECHAPP.models import ExcusedAbsence
    from django.db.models import Q
    
    section = request.user.section
    today = timezone.now().date()
    
    # Get excused absences for students in teacher's section
    excused_absences = ExcusedAbsence.objects.filter(
        student__section=section
    ).select_related('student').order_by('-effective_date') if section else []
    
    # Calculate status for each excuse
    for excuse in excused_absences:
        if excuse.effective_date > today:
            excuse.current_status = 'Upcoming'
        elif excuse.end_date and excuse.end_date < today:
            excuse.current_status = 'Expired'
        else:
            excuse.current_status = 'Active'
    
    # Get statistics
    total_excused = excused_absences.count()
    active_count = sum(1 for e in excused_absences if e.current_status == 'Active')
    upcoming_count = sum(1 for e in excused_absences if e.current_status == 'Upcoming')
    expired_count = sum(1 for e in excused_absences if e.current_status == 'Expired')
    
    context = {
        'excused_absences': excused_absences,
        'section': section,
        'total_excused': total_excused,
        'active_count': active_count,
        'upcoming_count': upcoming_count,
        'expired_count': expired_count,
    }
    return render(request, 'teacher/advisory/excused.html', context)

@login_required
@user_passes_test(is_advisory_teacher)
def teacher_advisory_messages(request):
    """View for advisory messages - now using PostgreSQL with polling"""
    return render(request, 'teacher/advisory/messages.html')

@login_required
@user_passes_test(is_advisory_teacher)
def teacher_advisory_settings(request):
    """View for advisory settings"""
    return render(request, 'teacher/advisory/settings.html')

# Export/Import Functions
@login_required
@user_passes_test(is_advisory_teacher)
def export_advisory_students(request):
    """Export advisory students to Excel, PDF, or Word"""
    try:
        import pandas as pd
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from reportlab.lib.pagesizes import letter, landscape
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from docx import Document
        from docx.shared import Inches, RGBColor, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        return JsonResponse({'success': False, 'error': f'Required library not installed: {str(e)}'}, status=500)
    
    section = request.user.section
    if not section:
        return JsonResponse({'success': False, 'error': 'No advisory section assigned'}, status=400)
    
    export_format = request.GET.get('format', 'excel')
    students = Student.objects.filter(section=section).select_related('grade', 'section').order_by('id')
    
    headers = ['ID', 'LRN', 'First Name', 'Middle Name', 'Last Name', 'Email', 'Grade', 'Section', 'Status', 'Face Enrolled']
    data_rows = []
    
    for student in students:
        data_rows.append([
            student.id,
            student.lrn or '',
            student.first_name,
            student.middle_name or '',
            student.last_name,
            student.email or '',
            student.grade.name if student.grade else '',
            student.section.name if student.section else '',
            student.status,
            'Yes' if student.face_path else 'No'
        ])
    
    if export_format == 'pdf':
        return export_students_to_pdf(headers, data_rows, f"{section.grade.name} - {section.name} Students")
    elif export_format == 'word':
        return export_students_to_word(headers, data_rows, f"{section.grade.name} - {section.name} Students")
    else:
        return export_students_to_excel_format(headers, data_rows, f"{section.grade.name} - {section.name} Students")

def export_students_to_excel_format(headers, data_rows, title):
    """Export students to Excel format"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Students"
    
    # Write headers
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = thin_border
    
    # Write data rows
    for row_num, row_data in enumerate(data_rows, 2):
        for col_num, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='left', vertical='center')
    
    # Auto-adjust column widths
    column_widths = [8, 15, 15, 15, 15, 25, 12, 12, 12, 15]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = width
    
    # Save to HTTP response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{title}.xlsx"'
    wb.save(response)
    return response

def export_students_to_pdf(headers, data_rows, title):
    """Export students to PDF format"""
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), topMargin=0.5*inch, bottomMargin=0.5*inch)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = styles['Heading1']
    title_style.textColor = colors.HexColor('#1F4E78')
    elements.append(Paragraph(title, title_style))
    elements.append(Spacer(1, 0.3*inch))
    
    # Table
    table_data = [headers] + data_rows
    table = Table(table_data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E78')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F0F0')]),
    ]))
    elements.append(table)
    
    doc.build(elements)
    buffer.seek(0)
    
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{title}.pdf"'
    return response

def export_students_to_word(headers, data_rows, title):
    """Export students to Word format"""
    from docx import Document
    from docx.shared import Inches, RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.color.rgb = RGBColor(31, 78, 120)
    
    # Table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(31, 78, 120)
    
    # Data rows
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{title}.docx"'
    return response

@login_required
@user_passes_test(is_advisory_teacher)
def export_advisory_attendance(request):
    """Export advisory attendance records to Excel, PDF, or Word"""
    try:
        import pandas as pd
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from reportlab.lib.pagesizes import letter, landscape
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from docx import Document
        from docx.shared import Inches, RGBColor, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        return JsonResponse({'success': False, 'error': f'Required library not installed: {str(e)}'}, status=500)
    
    section = request.user.section
    if not section:
        return JsonResponse({'success': False, 'error': 'No advisory section assigned'}, status=400)
    
    export_format = request.GET.get('format', 'excel')
    # Accept same filters as the table: search, status, date or date_start/date_end
    search_query = request.GET.get('search', '').strip()
    status_filter = request.GET.get('status', '').strip()
    date_filter = request.GET.get('date', '')
    date_start = request.GET.get('date_start', '')
    date_end = request.GET.get('date_end', '')

    # Build base queryset for the teacher's section
    qs = Attendance.objects.select_related('student').filter(student__section=section).order_by('id')

    if search_query:
        qs = qs.filter(
            Q(student__first_name__icontains=search_query) |
            Q(student__last_name__icontains=search_query) |
            Q(student__lrn__icontains=search_query)
        )
    if status_filter:
        qs = qs.filter(status=status_filter)

    # Date handling: prefer range if both provided
    if date_start and date_end:
        try:
            start = datetime.strptime(date_start, '%Y-%m-%d').date()
            end = datetime.strptime(date_end, '%Y-%m-%d').date()
            qs = qs.filter(date__range=(start, end))
            date_label = f"{start} to {end}"
        except Exception:
            date_label = ''
    elif date_filter:
        try:
            d = datetime.strptime(date_filter, '%Y-%m-%d').date()
            qs = qs.filter(date=d)
            date_label = str(d)
        except Exception:
            date_label = ''
    else:
        # No date filter provided: export all matching records
        date_label = 'all_dates'

    attendance_records = qs
    
    headers = ['ID', 'Student Name', 'LRN', 'Date', 'Time In', 'Time Out', 'Status']
    data_rows = []
    
    for record in attendance_records:
        data_rows.append([
            record.id,
            f"{record.student.first_name} {record.student.last_name}",
            record.student.lrn or '',
            str(record.date),
            str(record.time_in) if record.time_in else '',
            str(record.time_out) if record.time_out else '',
            record.status or ''
        ])
    
    title = f"{section.grade.name} - {section.name} Attendance {date_label}"
    
    if export_format == 'pdf':
        return export_students_to_pdf(headers, data_rows, title)
    elif export_format == 'word':
        return export_students_to_word(headers, data_rows, title)
    else:
        return export_students_to_excel_format(headers, data_rows, title)

@login_required
@user_passes_test(is_advisory_teacher)
def export_advisory_excused(request):
    """Export advisory excused absences to Excel, PDF, or Word"""
    try:
        import pandas as pd
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from reportlab.lib.pagesizes import letter, landscape
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from docx import Document
        from docx.shared import Inches, RGBColor, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        return JsonResponse({'success': False, 'error': f'Required library not installed: {str(e)}'}, status=500)
    
    from PROTECHAPP.models import ExcusedAbsence
    section = request.user.section
    if not section:
        return JsonResponse({'success': False, 'error': 'No advisory section assigned'}, status=400)
    
    export_format = request.GET.get('format', 'excel')
    
    excused_absences = ExcusedAbsence.objects.filter(
        student__section=section
    ).select_related('student').order_by('id')
    
    # Use the same exported columns as admin's excused export (model does not have 'reason' or 'uploaded_by')
    headers = ['ID', 'Student Name', 'LRN', 'Date Absent', 'Effective Date', 'End Date', 'Has Letter']
    data_rows = []

    for excuse in excused_absences:
        student_name = f"{excuse.student.first_name} {excuse.student.last_name}"
        has_letter = 'Yes' if excuse.excuse_letter else 'No'
        data_rows.append([
            str(excuse.id),
            student_name,
            excuse.student.lrn or '',
            excuse.date_absent.strftime('%Y-%m-%d') if excuse.date_absent else '',
            excuse.effective_date.strftime('%Y-%m-%d') if excuse.effective_date else '',
            excuse.end_date.strftime('%Y-%m-%d') if excuse.end_date else '',
            has_letter
        ])
    
    title = f"{section.grade.name} - {section.name} Excused Absences"
    
    if export_format == 'pdf':
        return export_students_to_pdf(headers, data_rows, title)
    elif export_format == 'word':
        return export_students_to_word(headers, data_rows, title)
    else:
        return export_students_to_excel_format(headers, data_rows, title)

@login_required
@user_passes_test(is_advisory_teacher)
def download_advisory_student_template(request):
    """Generate Excel template for importing students"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        from PROTECHAPP.models import Grade
        
        section = request.user.section
        if not section:
            return JsonResponse({'status': 'error', 'message': 'No advisory section assigned'}, status=400)
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Students"
        
        # Headers
        headers = ['ID', 'LRN', 'First Name', 'Middle Name', 'Last Name', 'Status', 
                   'Guardian First Name', 'Guardian Middle Name', 'Guardian Last Name',
                   'Guardian Email', 'Guardian Phone', 'Guardian Relationship']
        
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.value = header
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Add instruction row
        instruction_values = [
            'Leave blank for new',
            '12 digits',
            'Required',
            'Optional',
            'Required',
            'ACTIVE or INACTIVE',
            'Required if guardian data',
            'Optional',
            'Required if guardian data',
            'Optional',
            '09XXXXXXXXX',
            'FATHER, MOTHER, etc.'
        ]
        
        for col_num, instruction in enumerate(instruction_values, 1):
            cell = ws.cell(row=2, column=col_num)
            cell.value = instruction
            cell.font = Font(italic=True, color="808080")
            cell.fill = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 10
        ws.column_dimensions['B'].width = 15
        for col in ['C', 'D', 'E', 'G', 'H', 'I']:
            ws.column_dimensions[col].width = 18
        ws.column_dimensions['F'].width = 12
        ws.column_dimensions['J'].width = 25
        ws.column_dimensions['K'].width = 15
        ws.column_dimensions['L'].width = 20
        
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        response = HttpResponse(
            output.read(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="student_import_template_{section.name}.xlsx"'
        return response
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Error generating template: {str(e)}'}, status=500)

@login_required
@user_passes_test(is_advisory_teacher)
def import_advisory_students(request):
    """Import students for advisory teacher's section"""
    try:
        import openpyxl
        import re
        from django.db import transaction
        from django.views.decorators.http import require_http_methods
        
        if request.method != 'POST':
            return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=400)
        
        section = request.user.section
        if not section:
            return JsonResponse({'status': 'error', 'message': 'No advisory section assigned'}, status=400)
        
        if 'import_file' not in request.FILES:
            return JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)
        
        excel_file = request.FILES['import_file']
        
        # Validate file type
        if not excel_file.name.endswith(('.xlsx', '.xls', '.csv')):
            return JsonResponse({
                'status': 'error',
                'message': 'Invalid file type. Please upload Excel (.xlsx, .xls) or CSV files only.'
            }, status=400)
        
        # Validate file size (max 10MB)
        if excel_file.size > 10 * 1024 * 1024:
            return JsonResponse({
                'status': 'error',
                'message': 'File too large. Maximum file size is 10MB.'
            }, status=400)
        
        try:
            # Read Excel file
            wb = openpyxl.load_workbook(excel_file)
            ws = wb.active
            
            # Get header row (read all columns in the first row)
            headers = []
            for col in range(1, ws.max_column + 1):
                cell_value = ws.cell(row=1, column=col).value
                if cell_value:
                    headers.append(str(cell_value).strip())

            # Normalize helper to compare headers permissively (case/spacing/punctuation tolerant)
            def _normalize(h):
                return re.sub(r'[^a-z0-9]', '', h.strip().lower()) if h else ''

            # Canonical expected headers (human-friendly)
            expected_headers = ['ID', 'LRN', 'First Name', 'Middle Name', 'Last Name', 'Status',
                                'Guardian First Name', 'Guardian Middle Name', 'Guardian Last Name',
                                'Guardian Email', 'Guardian Phone', 'Guardian Relationship']

            normalized_expected = {_normalize(h): h for h in expected_headers}
            normalized_headers = set(_normalize(h) for h in headers)

            # Common synonyms mapping (normalized form -> expected normalized key)
            synonyms = {
                'studentid': 'id', 'id': 'id',
                'lrn': 'lrn', 'learnerreference': 'lrn', 'learnerref': 'lrn',
                'firstname': 'firstname', 'givenname': 'firstname', 'first': 'firstname',
                'middlename': 'middlename', 'middle': 'middlename',
                'lastname': 'lastname', 'surname': 'lastname', 'last': 'lastname',
                'status': 'status',
                'guardianfirstname': 'guardianfirstname', 'guardianfirst': 'guardianfirstname',
                'guardianmiddlename': 'guardianmiddlename', 'guardianmiddle': 'guardianmiddlename',
                'guardianlastname': 'guardianlastname', 'guardianlast': 'guardianlastname',
                'guardianemail': 'guardianemail', 'guardianemailaddress': 'guardianemail', 'guardianemailaddr': 'guardianemail',
                'guardianphone': 'guardianphone', 'guardiancontact': 'guardianphone', 'guardianphonenumber': 'guardianphone',
                'guardianrelationship': 'guardianrelationship', 'relationship': 'guardianrelationship'
            }

            # Determine which expected headers are present, accepting synonyms
            found_expected = set()
            for exp_nk in normalized_expected.keys():
                if exp_nk in normalized_headers:
                    found_expected.add(exp_nk)

            for h in normalized_headers:
                if h in synonyms:
                    mapped = synonyms[h]
                    # if mapped corresponds to an expected normalized key, mark it found
                    if mapped in normalized_expected:
                        found_expected.add(mapped)

            # Only require a subset of headers for import to proceed.
            # Guardian columns are OPTIONAL.
            required_headers = ['LRN', 'First Name', 'Last Name', 'Status']
            normalized_required = set(_normalize(h) for h in required_headers)

            required_missing = [v for k, v in normalized_expected.items() if k in normalized_required and k not in found_expected]

            if required_missing:
                return JsonResponse({
                    'status': 'error',
                    'message': f'Invalid template headers. Missing required: {", ".join(required_missing)}.\nExpected headers (examples): {", ".join(expected_headers)}'
                }, status=400)
            
            # Process rows using header->column mapping (order independent)
            students_data = []
            errors = []

            valid_relationships = ['FATHER', 'MOTHER', 'GUARDIAN', 'GRANDMOTHER', 'GRANDFATHER', 'AUNT', 'UNCLE', 'SIBLING', 'OTHER']
            existing_lrns = set(Student.objects.values_list('lrn', flat=True))

            # Build a mapping from normalized expected keys to column indices
            header_to_col = {}
            for col in range(1, ws.max_column + 1):
                raw = ws.cell(row=1, column=col).value
                if not raw:
                    continue
                nk = _normalize(str(raw))
                header_to_col[nk] = col

            # Common synonyms mapping (normalized form -> expected normalized key)
            synonyms = {
                'studentid': 'id', 'id': 'id',
                'lrn': 'lrn', 'learnerreference': 'lrn', 'learnerref': 'lrn',
                'firstname': 'firstname', 'givenname': 'firstname', 'first': 'firstname',
                'middlename': 'middlename', 'middle': 'middlename',
                'lastname': 'lastname', 'surname': 'lastname', 'last': 'lastname',
                'status': 'status',
                'guardianfirstname': 'guardianfirstname', 'guardianfirst': 'guardianfirstname',
                'guardianmiddlename': 'guardianmiddlename', 'guardianmiddle': 'guardianmiddlename',
                'guardianlastname': 'guardianlastname', 'guardianlast': 'guardianlastname',
                'guardianemail': 'guardianemail', 'guardianemailaddress': 'guardianemail', 'guardianemailaddr': 'guardianemail',
                'guardianphone': 'guardianphone', 'guardiancontact': 'guardianphone', 'guardianphonenumber': 'guardianphone', 'guardianphoneNo': 'guardianphone',
                'guardianrelationship': 'guardianrelationship', 'relationship': 'guardianrelationship'
            }

            # Map expected canonical normalized keys to actual column indices
            canonical_keys = {
                'id': 'ID', 'lrn': 'LRN', 'firstname': 'First Name', 'middlename': 'Middle Name', 'lastname': 'Last Name', 'status': 'Status',
                'guardianfirstname': 'Guardian First Name', 'guardianmiddlename': 'Guardian Middle Name', 'guardianlastname': 'Guardian Last Name',
                'guardianemail': 'Guardian Email', 'guardianphone': 'Guardian Phone', 'guardianrelationship': 'Guardian Relationship'
            }

            col_for = {}
            # try direct normalized_expected matches first
            for nk in normalized_expected.keys():
                if nk in header_to_col:
                    col_for[nk] = header_to_col[nk]

            # fill remaining from synonyms
            for raw_nk, col in header_to_col.items():
                if raw_nk in col_for:
                    continue
                if raw_nk in synonyms:
                    mapped = synonyms[raw_nk]
                    # map synonyms target to normalized_expected key name
                    # attempt to find matching normalized_expected key that endswith mapped (best-effort)
                    for exp_nk in normalized_expected.keys():
                        if mapped == exp_nk or mapped == exp_nk.replace('guardian', '').lstrip():
                            if exp_nk not in col_for:
                                col_for[exp_nk] = col
                                break

            # Helper to get cell value by canonical key
            def get_val(row_num, canonical_key):
                # canonical_key here is like 'id', 'lrn', 'firstname', etc.
                col = None
                # find matching normalized_expected key for canonical_key
                for nk, human in normalized_expected.items():
                    if canonical_key == nk or canonical_key == nk.replace('guardian', '').lstrip():
                        col = col_for.get(nk)
                        break
                if not col:
                    # try synonyms mapping directly
                    mapped = synonyms.get(canonical_key)
                    if mapped and mapped in col_for:
                        col = col_for[mapped]
                if not col:
                    return ''
                cell_value = ws.cell(row=row_num, column=col).value
                return str(cell_value).strip() if cell_value is not None else ''

            for row_num in range(2, ws.max_row + 1):
                # Extract data using header mapping
                student_id = get_val(row_num, 'id')
                lrn = get_val(row_num, 'lrn')
                first_name = get_val(row_num, 'firstname')
                middle_name = get_val(row_num, 'middlename')
                last_name = get_val(row_num, 'lastname')
                status = get_val(row_num, 'status')

                guardian_first = get_val(row_num, 'guardianfirstname')
                guardian_middle = get_val(row_num, 'guardianmiddlename')
                guardian_last = get_val(row_num, 'guardianlastname')
                guardian_email = get_val(row_num, 'guardianemail')
                guardian_phone = get_val(row_num, 'guardianphone')
                guardian_relationship = get_val(row_num, 'guardianrelationship')

                # Skip empty rows (no identifying student data)
                if not (lrn or first_name or last_name):
                    continue

                row_errors = []
                
                # Check if ID is provided - skip if exists
                if student_id and student_id != 'nan' and student_id != '':
                    try:
                        if Student.objects.filter(id=int(float(student_id))).exists():
                            continue  # Skip rows with existing ID
                    except ValueError:
                        pass
                
                # Validate required student fields
                if not lrn:
                    row_errors.append("LRN is required")
                elif not re.match(r'^\d{12}$', lrn):
                    row_errors.append("LRN must be exactly 12 digits")
                elif lrn in existing_lrns:
                    row_errors.append("LRN already exists")
                
                if not first_name:
                    row_errors.append("First Name is required")
                if not last_name:
                    row_errors.append("Last Name is required")
                
                if status not in ['ACTIVE', 'INACTIVE']:
                    row_errors.append("Status must be 'ACTIVE' or 'INACTIVE'")
                
                # Validate guardian data if provided
                guardian_data = None
                has_guardian_data = any([guardian_first, guardian_last, guardian_phone, guardian_relationship])
                
                if has_guardian_data:
                    if not guardian_first:
                        row_errors.append("Guardian First Name is required if guardian data provided")
                    if not guardian_last:
                        row_errors.append("Guardian Last Name is required if guardian data provided")
                    if not guardian_phone:
                        row_errors.append("Guardian Phone is required if guardian data provided")
                    elif not re.match(r'^09\d{9}$', guardian_phone):
                        row_errors.append("Guardian Phone must be valid format (09XXXXXXXXX)")
                    if not guardian_relationship:
                        row_errors.append("Guardian Relationship is required if guardian data provided")
                    elif guardian_relationship not in valid_relationships:
                        row_errors.append(f"Invalid Guardian Relationship '{guardian_relationship}'")
                    
                    if guardian_email and not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', guardian_email):
                        row_errors.append("Invalid Guardian Email format")
                    
                    if not row_errors:
                        guardian_data = {
                            'first_name': guardian_first,
                            'middle_name': guardian_middle,
                            'last_name': guardian_last,
                            'email': guardian_email if guardian_email else '',
                            'phone': guardian_phone,
                            'relationship': guardian_relationship
                        }
                
                if row_errors:
                    errors.append(f"Row {row_num}: {'; '.join(row_errors)}")
                else:
                    students_data.append({
                        'lrn': lrn,
                        'first_name': first_name,
                        'middle_name': middle_name,
                        'last_name': last_name,
                        'grade': section.grade,
                        'section': section,
                        'status': status,
                        'guardian': guardian_data
                    })
                    existing_lrns.add(lrn)
            
            # Return errors if any
            if errors:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Validation failed',
                    'errors': errors[:10]  # Limit to first 10 errors
                }, status=400)
            
            if not students_data:
                return JsonResponse({
                    'status': 'error',
                    'message': 'No valid student data found in file'
                }, status=400)
            
            # Import data in transaction
            with transaction.atomic():
                created_students = 0
                created_guardians = 0
                
                for student_info in students_data:
                    guardian_info = student_info.pop('guardian', None)
                    
                    # Create student
                    student = Student.objects.create(**student_info)
                    created_students += 1
                    
                    # Create guardian if data provided
                    if guardian_info:
                        Guardian.objects.create(
                            student=student,
                            **guardian_info
                        )
                        created_guardians += 1
            
            return JsonResponse({
                'status': 'success',
                'message': f'Successfully imported {created_students} student(s) and {created_guardians} guardian(s)',
                'students_created': created_students,
                'guardians_created': created_guardians
            })
            
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': f'Error processing file: {str(e)}'
            }, status=500)
            
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Error importing students: {str(e)}'
        }, status=500)
