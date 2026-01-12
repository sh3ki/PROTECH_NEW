from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.views.decorators.http import require_http_methods, require_GET
from django.core.files.storage import FileSystemStorage
from django.utils import timezone
from django.http import JsonResponse, Http404, FileResponse, HttpResponse
from django.core.paginator import Paginator
from django.db.models import Q
from django.conf import settings
import os
import json
import io
import traceback
from datetime import datetime, date

from PROTECHAPP.models import (
    Student,
    Attendance,
    Guardian,
    Grade,
    Section,
    ExcusedAbsence,
    UserRole,
)

# Directories used for file storage
PROFILE_PICS_DIR = getattr(settings, 'PROFILE_PICS_DIR', os.path.join(settings.MEDIA_ROOT, 'private_profile_pics'))
PRIVATE_EXCUSE_LETTERS_DIR = getattr(settings, 'PRIVATE_EXCUSE_LETTERS_DIR', os.path.join(settings.MEDIA_ROOT, 'private_excuse_letters'))
os.makedirs(PROFILE_PICS_DIR, exist_ok=True)
os.makedirs(PRIVATE_EXCUSE_LETTERS_DIR, exist_ok=True)


def is_teacher(user):
    """Check if the logged-in user is a teacher."""
    return user.is_authenticated and user.role == UserRole.TEACHER


def get_pagination_range(paginator, current_page, window=5):
    """Helper to generate a page range with a sliding window."""
    start = max(current_page - window, 1)
    end = min(current_page + window, paginator.num_pages)
    return range(start, end + 1)


# ======================
# DASHBOARD / MISC PAGES
# ======================

@login_required
@user_passes_test(is_teacher)
def teacher_non_advisory_dashboard(request):
    return render(request, 'teacher/non_advisory/dashboard.html')


@login_required
@user_passes_test(is_teacher)
def teacher_non_advisory_messages(request):
    return render(request, 'teacher/non_advisory/messages.html')


@login_required
@user_passes_test(is_teacher)
def teacher_non_advisory_settings(request):
    return render(request, 'teacher/non_advisory/settings.html')


@login_required
@user_passes_test(is_teacher)
def teacher_non_advisory_calendar(request):
    """View for non-advisory teacher calendar"""
    return render(request, 'teacher/non_advisory/calendar.html')


@login_required
@user_passes_test(is_teacher)
def get_teacher_non_advisory_calendar_events(request):
    """API endpoint to get all calendar events for non-advisory teachers"""
    from PROTECHAPP.models import CalendarEvent
    
    events = CalendarEvent.objects.all()
    events_data = []
    
    for event in events:
        events_data.append({
            'id': event.id,
            'title': event.title,
            'start': event.start_date.isoformat(),
            'end': event.end_date.isoformat(),
            'description': event.description,
            'location': event.location,
            'event_type': event.event_type,
            'target_role': event.target_role,
            'target_role_display': event.get_target_role_display() if event.target_role else None,
            'created_by': f"{event.created_by.first_name} {event.created_by.last_name}",
        })
    
    return JsonResponse(events_data, safe=False)


@login_required
@user_passes_test(is_teacher)
def add_teacher_non_advisory_calendar_event(request):
    """API endpoint to add a new event for non-advisory teachers"""
    if request.method == 'POST':
        from PROTECHAPP.models import CalendarEvent
        from datetime import datetime
        
        try:
            title = request.POST.get('title')
            description = request.POST.get('description', '')
            location = request.POST.get('location', '')
            start_date = request.POST.get('start_date')
            end_date = request.POST.get('end_date')
            
            # Convert datetime strings to datetime objects
            start_dt = datetime.fromisoformat(start_date)
            end_dt = datetime.fromisoformat(end_date)
            
            event = CalendarEvent.objects.create(
                title=title,
                description=description,
                location=location,
                event_type='EVENT',
                start_date=start_dt,
                end_date=end_dt,
                created_by=request.user
            )
            
            return JsonResponse({'success': True, 'event_id': event.id})
        except Exception as e:
            return JsonResponse({'success': False, 'message': str(e)})
    
    return JsonResponse({'success': False, 'message': 'Invalid request method'})


@login_required
@user_passes_test(is_teacher)
def add_teacher_non_advisory_calendar_announcement(request):
    """API endpoint to add a new announcement for non-advisory teachers"""
    if request.method == 'POST':
        from PROTECHAPP.models import CalendarEvent
        from datetime import datetime
        
        try:
            title = request.POST.get('title')
            description = request.POST.get('description', '')
            start_date = request.POST.get('start_date')
            end_date = request.POST.get('end_date')
            target_role = request.POST.get('target_role')
            
            # Convert datetime strings to datetime objects
            start_dt = datetime.fromisoformat(start_date)
            end_dt = datetime.fromisoformat(end_date)
            
            event = CalendarEvent.objects.create(
                title=title,
                description=description,
                event_type='ANNOUNCEMENT',
                start_date=start_dt,
                end_date=end_dt,
                target_role=target_role,
                created_by=request.user
            )
            
            return JsonResponse({'success': True, 'event_id': event.id})
        except Exception as e:
            return JsonResponse({'success': False, 'message': str(e)})
    
    return JsonResponse({'success': False, 'message': 'Invalid request method'})


@login_required
@user_passes_test(is_teacher)
def delete_teacher_non_advisory_calendar_event(request, event_id):
    """API endpoint to delete an event for non-advisory teachers"""
    if request.method == 'DELETE':
        from PROTECHAPP.models import CalendarEvent
        
        try:
            event = CalendarEvent.objects.get(id=event_id)
            event.delete()
            return JsonResponse({'success': True})
        except CalendarEvent.DoesNotExist:
            return JsonResponse({'success': False, 'message': 'Event not found'})
        except Exception as e:
            return JsonResponse({'success': False, 'message': str(e)})
    
    return JsonResponse({'success': False, 'message': 'Invalid request method'})


# ======================
# STUDENTS
# ======================

@login_required
@user_passes_test(is_teacher)
def teacher_non_advisory_students(request):
    """Student management view (admin clone)."""
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    page_number = request.GET.get('page', 1)

    students = Student.objects.select_related('grade', 'section').all().order_by('-created_at')

    if search_query:
        students = students.filter(
            Q(first_name__icontains=search_query)
            | Q(last_name__icontains=search_query)
            | Q(lrn__icontains=search_query)
            | Q(email__icontains=search_query)
        )

    if grade_filter:
        students = students.filter(grade_id=grade_filter)

    if section_filter:
        students = students.filter(section_id=section_filter)

    if status_filter:
        students = students.filter(status=status_filter)

    total_students = Student.objects.count()
    active_students = Student.objects.filter(status='ACTIVE').count()
    inactive_students = Student.objects.filter(status='INACTIVE').count()
    face_enrolled_count = Student.objects.exclude(face_path__isnull=True).exclude(face_path__exact='').count()

    grades = Grade.objects.all().order_by('name')
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    paginator = Paginator(students, 10)
    page_obj = paginator.get_page(page_number)
    page_range = get_pagination_range(paginator, page_obj.number, 5)

    context = {
        'students': page_obj,
        'grades': grades,
        'sections': sections,
        'search_query': search_query,
        'grade_filter': grade_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'total_students': total_students,
        'active_students': active_students,
        'inactive_students': inactive_students,
        'face_enrolled_count': face_enrolled_count,
        'page_range': page_range,
    }
    return render(request, 'teacher/non_advisory/students.html', context)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["GET"])
def teacher_non_advisory_search_students(request):
    try:
        search_query = request.GET.get('search', '')
        grade_filter = request.GET.get('grade', '')
        section_filter = request.GET.get('section', '')
        status_filter = request.GET.get('status', '')
        page_number = request.GET.get('page', 1)
        items_per_page = request.GET.get('items_per_page', 10)

        students = Student.objects.select_related('grade', 'section').all().order_by('-created_at')

        if search_query:
            students = students.filter(
                Q(first_name__icontains=search_query)
                | Q(last_name__icontains=search_query)
                | Q(lrn__icontains=search_query)
            )
        if grade_filter:
            students = students.filter(grade_id=grade_filter)
        if section_filter:
            students = students.filter(section_id=section_filter)
        if status_filter:
            students = students.filter(status=status_filter)

        total_count = students.count()
        try:
            page_number = int(page_number)
            items_per_page = int(items_per_page)
        except (ValueError, TypeError):
            page_number = 1
            items_per_page = 10

        paginator = Paginator(students, items_per_page)
        page_obj = paginator.get_page(page_number)

        student_data = []
        for student in page_obj:
            student_data.append({
                'id': student.id,
                'lrn': student.lrn,
                'first_name': student.first_name,
                'last_name': student.last_name,
                'full_name': f"{student.first_name} {student.last_name}",
                'grade': student.grade.name if student.grade else '',
                'grade_id': student.grade.id if student.grade else None,
                'section': student.section.name if student.section else '',
                'section_id': student.section.id if student.section else None,
                'status': student.status,
                'profile_pic': student.profile_pic if student.profile_pic else None,
                'created_at': student.created_at.strftime('%Y-%m-%d'),
                'created_at_display': student.created_at.strftime('%b %d, %Y'),
            })

        pagination = {
            'current_page': page_obj.number,
            'num_pages': paginator.num_pages,
            'has_next': page_obj.has_next(),
            'has_previous': page_obj.has_previous(),
            'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
            'start_index': page_obj.start_index(),
            'end_index': page_obj.end_index(),
            'total_count': total_count,
        }

        return JsonResponse({
            'status': 'success',
            'students': student_data,
            'pagination': pagination,
            'total_count': total_count,
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e),
            'traceback': traceback.format_exc()
        }, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["POST"])
def teacher_non_advisory_create_student(request):
    try:
        lrn = request.POST.get('lrn', '')
        first_name = request.POST.get('first_name', '')
        middle_name = request.POST.get('middle_name', '')
        last_name = request.POST.get('last_name', '')
        grade_id = request.POST.get('grade_id', '')
        section_id = request.POST.get('section_id', '')
        status = request.POST.get('status', 'ACTIVE')

        if not lrn or not first_name or not last_name or not grade_id or not section_id:
            return JsonResponse({'status': 'error', 'message': 'Missing required fields'}, status=400)

        if Student.objects.filter(lrn=lrn).exists():
            return JsonResponse({'status': 'error', 'message': 'LRN already exists'}, status=400)

        try:
            grade = Grade.objects.get(id=grade_id)
            section = Section.objects.get(id=section_id)
        except (Grade.DoesNotExist, Section.DoesNotExist):
            return JsonResponse({'status': 'error', 'message': 'Invalid grade or section'}, status=400)

        student = Student(
            lrn=lrn,
            first_name=first_name,
            middle_name=middle_name,
            last_name=last_name,
            grade=grade,
            section=section,
            status=status,
        )

        if 'profile_pic' in request.FILES:
            profile_pic = request.FILES['profile_pic']
            if profile_pic.size > 2 * 1024 * 1024:
                return JsonResponse({'status': 'error', 'message': 'File size must be less than 2MB'}, status=400)
            if not profile_pic.content_type.startswith('image/'):
                return JsonResponse({'status': 'error', 'message': 'Only image files are allowed'}, status=400)
            filename = f"student_{int(timezone.now().timestamp())}_{profile_pic.name}"
            fs = FileSystemStorage(location=PROFILE_PICS_DIR)
            filename = fs.save(filename, profile_pic)
            student.profile_pic = filename

        student.save()

        has_guardian = request.POST.get('has_guardian') == 'true'
        if has_guardian:
            guardians_data_json = request.POST.get('guardians_data')
            if guardians_data_json:
                try:
                    guardians_data = json.loads(guardians_data_json)
                    for guardian_data in guardians_data:
                        g_first = guardian_data.get('first_name', '')
                        g_last = guardian_data.get('last_name', '')
                        g_middle = guardian_data.get('middle_name', '')
                        g_email = guardian_data.get('email', '')
                        g_phone = guardian_data.get('phone_number', '')
                        g_rel = guardian_data.get('relationship', '')
                        if g_first and g_last and g_email and g_phone and g_rel:
                            Guardian.objects.create(
                                first_name=g_first,
                                middle_name=g_middle,
                                last_name=g_last,
                                email=g_email,
                                phone=g_phone,
                                relationship=g_rel,
                                student=student,
                            )
                except json.JSONDecodeError:
                    pass
            else:
                g_first = request.POST.get('guardian_first_name', '')
                g_last = request.POST.get('guardian_last_name', '')
                g_middle = request.POST.get('guardian_middle_name', '')
                g_email = request.POST.get('guardian_email', '')
                g_phone = request.POST.get('guardian_phone_number', '')
                g_rel = request.POST.get('guardian_relationship', '')
                if g_first and g_last and g_email and g_phone and g_rel:
                    Guardian.objects.create(
                        first_name=g_first,
                        middle_name=g_middle,
                        last_name=g_last,
                        email=g_email,
                        phone=g_phone,
                        relationship=g_rel,
                        student=student,
                    )

        has_face_enrollment = request.POST.get('has_face_enrollment') == 'true'
        if has_face_enrollment:
            student.face_path = f"pending_{student.id}"
            student.save()

        return JsonResponse({
            'status': 'success',
            'message': f'Student {first_name} {last_name} created successfully',
            'student_id': student.id,
        })

    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_teacher)
def teacher_non_advisory_get_student(request, student_id):
    try:
        student = get_object_or_404(Student, id=student_id)
        student_data = {
            'id': student.id,
            'lrn': student.lrn,
            'first_name': student.first_name,
            'middle_name': student.middle_name or '',
            'last_name': student.last_name,
            'grade_id': str(student.grade.id) if student.grade else '',
            'grade': student.grade.name if student.grade else '',
            'section_id': str(student.section.id) if student.section else '',
            'section': student.section.name if student.section else '',
            'status': student.status,
            'full_name': f"{student.first_name} {student.last_name}",
            'created_at': student.created_at.strftime('%Y-%m-%d'),
            'profile_pic': student.profile_pic or None,
        }
        return JsonResponse({'status': 'success', 'student': student_data})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["POST", "PATCH"])
def teacher_non_advisory_update_student(request, student_id):
    try:
        student = get_object_or_404(Student, id=student_id)

        if request.method in ['PATCH', 'POST']:
            if 'multipart/form-data' in request.content_type:
                lrn = request.POST.get('lrn', student.lrn)
                first_name = request.POST.get('first_name', student.first_name)
                middle_name = request.POST.get('middle_name', '')
                last_name = request.POST.get('last_name', student.last_name)
                grade_id = request.POST.get('grade_id', student.grade.id if student.grade else None)
                section_id = request.POST.get('section_id', student.section.id if student.section else None)
                status = request.POST.get('status', student.status)

                if not lrn or not first_name or not last_name or not grade_id or not section_id:
                    return JsonResponse({'status': 'error', 'message': 'Missing required fields'}, status=400)

                if Student.objects.filter(lrn=lrn).exclude(id=student_id).exists():
                    return JsonResponse({'status': 'error', 'message': f'Student with LRN {lrn} already exists'}, status=400)

                try:
                    grade = Grade.objects.get(id=grade_id)
                    section = Section.objects.get(id=section_id)
                except (Grade.DoesNotExist, Section.DoesNotExist):
                    return JsonResponse({'status': 'error', 'message': 'Invalid grade or section'}, status=400)

                student.lrn = lrn
                student.first_name = first_name
                student.middle_name = middle_name
                student.last_name = last_name
                student.grade = grade
                student.section = section
                student.status = status

                if 'profile_pic' in request.FILES:
                    profile_pic = request.FILES['profile_pic']
                    if profile_pic.size > 2 * 1024 * 1024:
                        return JsonResponse({'status': 'error', 'message': 'File size must be less than 2MB'}, status=400)
                    if not profile_pic.content_type.startswith('image/'):
                        return JsonResponse({'status': 'error', 'message': 'Only image files are allowed'}, status=400)
                    if student.profile_pic:
                        old_profile_pic_path = os.path.join(PROFILE_PICS_DIR, student.profile_pic)
                        if os.path.exists(old_profile_pic_path):
                            try:
                                os.remove(old_profile_pic_path)
                            except OSError:
                                pass
                    filename = f"student_{int(timezone.now().timestamp())}_{profile_pic.name}"
                    fs = FileSystemStorage(location=PROFILE_PICS_DIR)
                    filename = fs.save(filename, profile_pic)
                    student.profile_pic = filename

                student.save()

                return JsonResponse({'status': 'success', 'message': f'Student {first_name} {last_name} updated successfully', 'student_id': student.id})

        return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=405)

    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["POST", "DELETE"])
def teacher_non_advisory_delete_student(request, student_id):
    try:
        student = get_object_or_404(Student, id=student_id)
        student.delete()
        return JsonResponse({'status': 'success', 'message': 'Student deleted successfully'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["POST"])
def teacher_non_advisory_toggle_student_status(request):
    try:
        student_id = request.POST.get('student_id')
        if not student_id:
            return JsonResponse({'status': 'error', 'message': 'Student ID is required'}, status=400)
        student = get_object_or_404(Student, id=student_id)
        student.status = 'INACTIVE' if student.status == 'ACTIVE' else 'ACTIVE'
        student.save()
        return JsonResponse({'status': 'success', 'new_status': student.status})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


# ======================
# GUARDIANS
# ======================

@login_required
@user_passes_test(is_teacher)
def teacher_non_advisory_guardians(request):
    total_guardians = Guardian.objects.count()
    guardians_with_children = Guardian.objects.exclude(student=None).count()
    mother_count = Guardian.objects.filter(relationship='MOTHER').count()
    father_count = Guardian.objects.filter(relationship='FATHER').count()

    relationship_choices = [
        ('FATHER', 'Father'),
        ('MOTHER', 'Mother'),
        ('GUARDIAN', 'Guardian'),
        ('GRANDMOTHER', 'Grandmother'),
        ('GRANDFATHER', 'Grandfather'),
        ('AUNT', 'Aunt'),
        ('UNCLE', 'Uncle'),
        ('SIBLING', 'Sibling'),
        ('OTHER', 'Other'),
    ]

    grades = Grade.objects.all().order_by('name')
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')
    students = Student.objects.select_related('grade', 'section').filter(status='ACTIVE')

    context = {
        'total_guardians': total_guardians,
        'guardians_with_children': guardians_with_children,
        'mother_count': mother_count,
        'father_count': father_count,
        'relationship_choices': relationship_choices,
        'grades': grades,
        'sections': sections,
        'students': students,
    }
    return render(request, 'teacher/non_advisory/guardians.html', context)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["GET"])
def teacher_non_advisory_search_guardians(request):
    try:
        search_query = request.GET.get('search', '')
        relationship_filter = request.GET.get('relationship', '')
        grade_filter = request.GET.get('grade', '')
        section_filter = request.GET.get('section', '')
        page_number = request.GET.get('page', 1)
        items_per_page = request.GET.get('items_per_page', 10)

        guardians = Guardian.objects.all().order_by('-created_at')

        if search_query:
            guardians = guardians.filter(
                Q(first_name__icontains=search_query)
                | Q(middle_name__icontains=search_query)
                | Q(last_name__icontains=search_query)
                | Q(email__icontains=search_query)
                | Q(phone__icontains=search_query)
            )
        if relationship_filter:
            guardians = guardians.filter(relationship=relationship_filter)
        if grade_filter:
            student_ids = Student.objects.filter(grade_id=grade_filter).values_list('id', flat=True)
            guardians = guardians.filter(student_id__in=student_ids)
        if section_filter:
            student_ids = Student.objects.filter(section_id=section_filter).values_list('id', flat=True)
            guardians = guardians.filter(student_id__in=student_ids)

        total_count = guardians.count()
        try:
            page_number = int(page_number)
            items_per_page = int(items_per_page)
        except (ValueError, TypeError):
            page_number = 1
            items_per_page = 10

        paginator = Paginator(guardians, items_per_page)
        page_obj = paginator.get_page(page_number)

        guardian_data = []
        for guardian in page_obj:
            data = {
                'id': guardian.id,
                'full_name': f"{guardian.first_name} {guardian.last_name}",
                'first_name': guardian.first_name,
                'middle_name': guardian.middle_name or '',
                'last_name': guardian.last_name,
                'email': guardian.email or 'No email',
                'phone_number': guardian.phone or 'No phone',
                'relationship': guardian.relationship,
                'created_at_display': guardian.created_at.strftime('%B %d, %Y') if guardian.created_at else 'N/A',
            }
            try:
                data['relationship_display'] = guardian.get_relationship_display()
            except Exception:
                data['relationship_display'] = guardian.relationship
            try:
                student = Student.objects.filter(id=guardian.student_id).first() if hasattr(guardian, 'student_id') else None
                if student:
                    data['children_count'] = 1
                    data['children_preview'] = f"{student.first_name} {student.last_name}"
                else:
                    data['children_count'] = 0
                    data['children_preview'] = 'No children'
            except Exception:
                data['children_count'] = 0
                data['children_preview'] = 'Error loading children'
            guardian_data.append(data)

        pagination = {
            'current_page': page_obj.number,
            'num_pages': paginator.num_pages,
            'has_next': page_obj.has_next(),
            'has_previous': page_obj.has_previous(),
            'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
            'start_index': page_obj.start_index(),
            'end_index': page_obj.end_index(),
            'total_count': total_count,
        }

        return JsonResponse({
            'status': 'success',
            'guardians': guardian_data,
            'pagination': pagination,
            'total_count': total_count,
        })

    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}',
            'traceback': traceback.format_exc(),
        }, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["POST"])
def teacher_non_advisory_create_guardian(request):
    try:
        data = json.loads(request.body)
        required_fields = ['first_name', 'last_name', 'phone_number', 'relationship', 'student_id']
        for field in required_fields:
            if field not in data or not data[field]:
                return JsonResponse({'status': 'error', 'message': f'Field {field} is required'}, status=400)

        try:
            student = Student.objects.get(id=data['student_id'])
        except Student.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Student not found'}, status=404)

        guardian = Guardian.objects.create(
            first_name=data['first_name'],
            middle_name=data.get('middle_name', ''),
            last_name=data['last_name'],
            email=data.get('email', ''),
            phone=data['phone_number'],
            relationship=data['relationship'],
            student=student,
        )

        return JsonResponse({'status': 'success', 'message': 'Guardian created successfully', 'guardian_id': guardian.id})

    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["GET"])
def teacher_non_advisory_get_guardian(request, guardian_id):
    try:
        guardian = get_object_or_404(Guardian, id=guardian_id)
        data = {
            'id': guardian.id,
            'first_name': guardian.first_name,
            'middle_name': guardian.middle_name or '',
            'last_name': guardian.last_name,
            'email': guardian.email or '',
            'phone_number': guardian.phone or '',
            'relationship': guardian.relationship,
            'student_id': guardian.student.id if guardian.student else None,
        }
        return JsonResponse({'status': 'success', 'guardian': data})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["POST", "PATCH"])
def teacher_non_advisory_update_guardian(request, guardian_id):
    try:
        guardian = get_object_or_404(Guardian, id=guardian_id)
        data = json.loads(request.body)

        guardian.first_name = data.get('first_name', guardian.first_name)
        guardian.middle_name = data.get('middle_name', guardian.middle_name)
        guardian.last_name = data.get('last_name', guardian.last_name)
        guardian.email = data.get('email', guardian.email)
        guardian.phone = data.get('phone_number', guardian.phone)
        guardian.relationship = data.get('relationship', guardian.relationship)

        if 'student_id' in data:
            try:
                student = Student.objects.get(id=data['student_id'])
                guardian.student = student
            except Student.DoesNotExist:
                return JsonResponse({'status': 'error', 'message': 'Student not found'}, status=404)

        guardian.save()
        return JsonResponse({'status': 'success', 'message': 'Guardian updated successfully'})

    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["POST", "DELETE"])
def teacher_non_advisory_delete_guardian(request, guardian_id):
    try:
        guardian = get_object_or_404(Guardian, id=guardian_id)
        guardian.delete()
        return JsonResponse({'status': 'success', 'message': 'Guardian deleted successfully'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


# ======================
# ATTENDANCE RECORDS
# ======================

@login_required
@user_passes_test(is_teacher)
def teacher_non_advisory_attendance(request):
    total_records = Attendance.objects.count()
    present_count = Attendance.objects.filter(status='ON TIME').count() + Attendance.objects.filter(status='LATE').count()
    absent_count = Attendance.objects.filter(status='ABSENT').count()
    excused_count = Attendance.objects.filter(status='EXCUSED').count()

    grades = Grade.objects.all().order_by('name')
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')

    context = {
        'grades': grades,
        'sections': sections,
        'grade_filter': grade_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'date_filter': date_filter,
        'total_records': total_records,
        'present_count': present_count,
        'absent_count': absent_count,
        'excused_count': excused_count,
    }
    return render(request, 'teacher/non_advisory/attendance.html', context)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["GET"])
def teacher_non_advisory_search_attendance_records(request):
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')
    page_number = request.GET.get('page', 1)
    items_per_page = request.GET.get('items_per_page', 10)

    records = Attendance.objects.select_related('student', 'student__grade', 'student__section').all().order_by('-date', '-student__last_name')

    if search_query:
        records = records.filter(
            Q(student__first_name__icontains=search_query)
            | Q(student__last_name__icontains=search_query)
            | Q(student__lrn__icontains=search_query)
            | Q(student__section__name__icontains=search_query)
        )
    if grade_filter:
        records = records.filter(student__grade_id=grade_filter)
    if section_filter:
        records = records.filter(student__section_id=section_filter)
    if status_filter:
        records = records.filter(status=status_filter)
    if date_filter:
        records = records.filter(date=date_filter)

    total_count = records.count()
    try:
        page_number = int(page_number)
        items_per_page = int(items_per_page)
    except (ValueError, TypeError):
        page_number = 1
        items_per_page = 10

    paginator = Paginator(records, items_per_page)
    page_obj = paginator.get_page(page_number)

    attendance_data = []
    for record in page_obj:
        attendance_data.append({
            'id': record.id,
            'lrn': record.student.lrn,
            'full_name': f"{record.student.first_name} {record.student.last_name}",
            'grade': record.student.grade.name if record.student.grade else '',
            'section': record.student.section.name if record.student.section else '',
            'date': record.date.strftime('%Y-%m-%d'),
            'time_in': record.time_in.strftime('%H:%M:%S') if record.time_in else '',
            'time_out': record.time_out.strftime('%H:%M:%S') if record.time_out else '',
            'status': record.status,
            'profile_pic': record.student.profile_pic or None,
        })

    pagination = {
        'current_page': page_obj.number,
        'num_pages': paginator.num_pages,
        'has_next': page_obj.has_next(),
        'has_previous': page_obj.has_previous(),
        'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
        'start_index': page_obj.start_index(),
        'end_index': page_obj.end_index(),
        'total_count': total_count,
    }

    return JsonResponse({
        'status': 'success',
        'records': attendance_data,
        'pagination': pagination,
        'total_count': total_count,
    })


@login_required
@user_passes_test(is_teacher)
@require_GET
def teacher_non_advisory_get_attendance_record(request, attendance_id):
    try:
        record = get_object_or_404(Attendance, id=attendance_id)
        data = {
            'id': record.id,
            'student_id': record.student.id,
            'student_name': f"{record.student.first_name} {record.student.last_name}",
            'lrn': record.student.lrn,
            'date': record.date.strftime('%Y-%m-%d'),
            'time_in': record.time_in.strftime('%H:%M:%S') if record.time_in else None,
            'time_out': record.time_out.strftime('%H:%M:%S') if record.time_out else None,
            'status': record.status,
            'notes': record.notes or '',
            'section': record.student.section.name if record.student.section else '',
            'section_id': record.student.section.id if record.student.section else None,
            'grade': record.student.grade.name if record.student.grade else '',
            'grade_id': record.student.grade.id if record.student.grade else None,
        }
        return JsonResponse({'status': 'success', 'record': data})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["POST"])
def teacher_non_advisory_create_attendance_record(request):
    try:
        data = json.loads(request.body)
        required_fields = ['student_id', 'date', 'status']
        for field in required_fields:
            if field not in data or not data[field]:
                return JsonResponse({'status': 'error', 'message': f'{field} is required'}, status=400)

        try:
            student = Student.objects.get(id=data['student_id'])
        except Student.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Student not found'}, status=404)

        if Attendance.objects.filter(student=student, date=data['date']).exists():
            return JsonResponse({'status': 'error', 'message': f'Attendance record already exists for {student.first_name} {student.last_name} on {data["date"]}'}, status=400)

        time_in = None
        if data.get('time_in'):
            time_in = datetime.strptime(data['time_in'], '%H:%M:%S').time()
        time_out = None
        if data.get('time_out'):
            time_out = datetime.strptime(data['time_out'], '%H:%M:%S').time()

        record = Attendance.objects.create(
            student=student,
            date=data['date'],
            time_in=time_in,
            time_out=time_out,
            status=data['status'],
            notes=data.get('notes', '')
        )

        return JsonResponse({'status': 'success', 'message': 'Attendance record created successfully', 'record_id': record.id})
    except ValueError:
        return JsonResponse({'status': 'error', 'message': 'Invalid date or time format'}, status=400)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["POST", "PATCH"])
def teacher_non_advisory_update_attendance_record(request, attendance_id):
    try:
        record = get_object_or_404(Attendance, id=attendance_id)
        data = json.loads(request.body)

        if 'student_id' in data:
            try:
                student = Student.objects.get(id=data['student_id'])
                record.student = student
            except Student.DoesNotExist:
                return JsonResponse({'status': 'error', 'message': 'Student not found'}, status=404)

        if 'date' in data:
            record.date = data['date']

        if 'time_in' in data:
            if data['time_in']:
                record.time_in = datetime.strptime(data['time_in'], '%H:%M:%S').time()
            else:
                record.time_in = None

        if 'time_out' in data:
            if data['time_out']:
                record.time_out = datetime.strptime(data['time_out'], '%H:%M:%S').time()
            else:
                record.time_out = None

        if 'status' in data:
            record.status = data['status']

        if 'notes' in data:
            record.notes = data['notes']

        record.save()
        return JsonResponse({'status': 'success', 'message': 'Attendance record updated successfully'})
    except ValueError:
        return JsonResponse({'status': 'error', 'message': 'Invalid date or time format'}, status=400)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["POST", "DELETE"])
def teacher_non_advisory_delete_attendance_record(request, attendance_id):
    try:
        record = get_object_or_404(Attendance, id=attendance_id)
        record.delete()
        return JsonResponse({'status': 'success', 'message': 'Attendance record deleted successfully'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


# ======================
# EXCUSED ABSENCES
# ======================

@login_required
@user_passes_test(is_teacher)
def teacher_non_advisory_excused(request):
    total_excused = ExcusedAbsence.objects.count()
    today = date.today()
    active_excused = ExcusedAbsence.objects.filter(effective_date__lte=today, end_date__gte=today).count()
    expired_excused = ExcusedAbsence.objects.filter(end_date__lt=today).count()
    upcoming_excused = ExcusedAbsence.objects.filter(effective_date__gt=today).count()

    grades = Grade.objects.all().order_by('name')
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')

    context = {
        'grades': grades,
        'sections': sections,
        'grade_filter': grade_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'date_filter': date_filter,
        'total_excused': total_excused,
        'active_excused': active_excused,
        'expired_excused': expired_excused,
        'upcoming_excused': upcoming_excused,
    }
    return render(request, 'teacher/non_advisory/excused.html', context)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["GET"])
def teacher_non_advisory_search_excused_absences(request):
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')
    page_number = request.GET.get('page', 1)
    items_per_page = request.GET.get('items_per_page', 10)

    excused = ExcusedAbsence.objects.select_related('student', 'student__grade', 'student__section').all().order_by('-date_absent', '-student__last_name')

    if search_query:
        excused = excused.filter(
            Q(student__first_name__icontains=search_query)
            | Q(student__last_name__icontains=search_query)
            | Q(student__lrn__icontains=search_query)
            | Q(student__section__name__icontains=search_query)
        )
    if grade_filter:
        excused = excused.filter(student__grade_id=grade_filter)
    if section_filter:
        excused = excused.filter(student__section_id=section_filter)
    if status_filter:
        today = date.today()
        if status_filter == 'ACTIVE':
            excused = excused.filter(effective_date__lte=today, end_date__gte=today)
        elif status_filter == 'EXPIRED':
            excused = excused.filter(end_date__lt=today)
        elif status_filter == 'UPCOMING':
            excused = excused.filter(effective_date__gt=today)
    if date_filter:
        excused = excused.filter(date_absent=date_filter)

    total_count = excused.count()
    try:
        page_number = int(page_number)
        items_per_page = int(items_per_page)
    except (ValueError, TypeError):
        page_number = 1
        items_per_page = 10

    paginator = Paginator(excused, items_per_page)
    page_obj = paginator.get_page(page_number)

    excused_data = []
    for record in page_obj:
        today = date.today()
        if record.effective_date > today:
            status = 'UPCOMING'
        elif record.end_date < today:
            status = 'EXPIRED'
        else:
            status = 'ACTIVE'
        excused_data.append({
            'id': record.id,
            'lrn': record.student.lrn,
            'full_name': f"{record.student.first_name} {record.student.last_name}",
            'grade': record.student.grade.name if record.student.grade else '',
            'section': record.student.section.name if record.student.section else '',
            'date_absent': record.date_absent.strftime('%Y-%m-%d'),
            'effective_date': record.effective_date.strftime('%Y-%m-%d'),
            'end_date': record.end_date.strftime('%Y-%m-%d'),
            'excuse_letter': record.excuse_letter,
            'status': status,
            'profile_pic': record.student.profile_pic or None,
        })

    pagination = {
        'current_page': page_obj.number,
        'num_pages': paginator.num_pages,
        'has_next': page_obj.has_next(),
        'has_previous': page_obj.has_previous(),
        'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
        'start_index': page_obj.start_index(),
        'end_index': page_obj.end_index(),
        'total_count': total_count,
    }

    return JsonResponse({
        'status': 'success',
        'records': excused_data,
        'pagination': pagination,
        'total_count': total_count,
    })


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["POST"])
def teacher_non_advisory_create_excused_absence(request):
    try:
        student_id = request.POST.get('student_id')
        date_absent = request.POST.get('date_absent')
        effective_date = request.POST.get('effective_date')
        end_date = request.POST.get('end_date')
        excuse_letter_file = request.FILES.get('excuse_letter')

        if not student_id or not date_absent or not effective_date or not end_date or not excuse_letter_file:
            return JsonResponse({'status': 'error', 'message': 'All fields are required'}, status=400)

        try:
            student = Student.objects.get(id=student_id)
        except Student.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Student not found'}, status=404)

        if excuse_letter_file.size > 5 * 1024 * 1024:
            return JsonResponse({'status': 'error', 'message': 'File too large. Max 5MB.'}, status=400)
        if not excuse_letter_file.content_type.startswith('image/'):
            return JsonResponse({'status': 'error', 'message': 'Only image files allowed.'}, status=400)

        filename = f"{timezone.now().strftime('%Y%m%d%H%M%S')}_{excuse_letter_file.name}"
        fs = FileSystemStorage(location=PRIVATE_EXCUSE_LETTERS_DIR)
        filename = fs.save(filename, excuse_letter_file)

        record = ExcusedAbsence.objects.create(
            student=student,
            date_absent=date_absent,
            excuse_letter=filename,
            effective_date=effective_date,
            end_date=end_date,
        )

        return JsonResponse({
            'status': 'success',
            'message': f'Excused absence for {student.first_name} {student.last_name} created successfully',
            'record_id': record.id,
        })
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["GET"])
def teacher_non_advisory_get_excused_absence(request, excused_id):
    try:
        record = get_object_or_404(ExcusedAbsence, id=excused_id)
        data = {
            'id': record.id,
            'student_id': record.student.id,
            'student_name': f"{record.student.first_name} {record.student.last_name}",
            'lrn': record.student.lrn,
            'date_absent': record.date_absent.strftime('%Y-%m-%d'),
            'effective_date': record.effective_date.strftime('%Y-%m-%d'),
            'end_date': record.end_date.strftime('%Y-%m-%d'),
            'excuse_letter': record.excuse_letter,
            'section': record.student.section.name if record.student.section else '',
            'section_id': record.student.section.id if record.student.section else None,
            'grade': record.student.grade.name if record.student.grade else '',
            'grade_id': record.student.grade.id if record.student.grade else None,
        }
        return JsonResponse({'status': 'success', 'record': data})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["POST", "PATCH"])
def teacher_non_advisory_update_excused_absence(request, excused_id):
    try:
        record = get_object_or_404(ExcusedAbsence, id=excused_id)
        student_id = request.POST.get('student_id', record.student.id)
        date_absent = request.POST.get('date_absent', record.date_absent)
        effective_date = request.POST.get('effective_date', record.effective_date)
        end_date = request.POST.get('end_date', record.end_date)
        excuse_letter_file = request.FILES.get('excuse_letter', None)

        try:
            student = Student.objects.get(id=student_id)
        except Student.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Student not found'}, status=404)

        record.student = student
        record.date_absent = date_absent
        record.effective_date = effective_date
        record.end_date = end_date

        if excuse_letter_file:
            if excuse_letter_file.size > 5 * 1024 * 1024:
                return JsonResponse({'status': 'error', 'message': 'File too large. Max 5MB.'}, status=400)
            if not excuse_letter_file.content_type.startswith('image/'):
                return JsonResponse({'status': 'error', 'message': 'Only image files allowed.'}, status=400)
            if record.excuse_letter:
                old_path = os.path.join(PRIVATE_EXCUSE_LETTERS_DIR, record.excuse_letter)
                if os.path.exists(old_path):
                    try:
                        os.remove(old_path)
                    except OSError:
                        pass
            filename = f"{timezone.now().strftime('%Y%m%d%H%M%S')}_{excuse_letter_file.name}"
            fs = FileSystemStorage(location=PRIVATE_EXCUSE_LETTERS_DIR)
            filename = fs.save(filename, excuse_letter_file)
            record.excuse_letter = filename

        record.save()
        return JsonResponse({'status': 'success', 'message': 'Excused absence updated successfully'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["POST", "DELETE"])
def teacher_non_advisory_delete_excused_absence(request, excused_id):
    try:
        record = get_object_or_404(ExcusedAbsence, id=excused_id)
        if record.excuse_letter:
            old_path = os.path.join(PRIVATE_EXCUSE_LETTERS_DIR, record.excuse_letter)
            if os.path.exists(old_path):
                try:
                    os.remove(old_path)
                except OSError:
                    pass
        record.delete()
        return JsonResponse({'status': 'success', 'message': 'Excused absence deleted successfully'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["GET"])
def teacher_non_advisory_serve_excuse_letter(request, filename):
    file_path = os.path.join(PRIVATE_EXCUSE_LETTERS_DIR, filename)
    if not os.path.exists(file_path):
        raise Http404('File not found')
    response = FileResponse(open(file_path, 'rb'), content_type='application/octet-stream')
    response['Content-Disposition'] = "inline; filename=" + '"' + filename + '"'
    return response


# ======================
# SUPPORTING HELPERS
# ======================

@login_required
@user_passes_test(is_teacher)
@require_http_methods(["GET"])
def teacher_non_advisory_get_sections_by_grade(request):
    grade_id = request.GET.get('grade_id')
    if not grade_id:
        return JsonResponse({'status': 'error', 'message': 'Grade ID is required'}, status=400)
    sections = Section.objects.filter(grade_id=grade_id).order_by('name')
    data = [{'id': section.id, 'name': section.name} for section in sections]
    return JsonResponse({'status': 'success', 'sections': data})


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["GET"])
def teacher_non_advisory_get_students_by_section(request):
    section_id = request.GET.get('section_id')
    if not section_id:
        return JsonResponse({'status': 'error', 'message': 'Section ID is required'}, status=400)
    students = Student.objects.filter(section_id=section_id, status='ACTIVE').select_related('grade', 'section').order_by('last_name', 'first_name')
    data = [
        {
            'id': student.id,
            'first_name': student.first_name,
            'last_name': student.last_name,
            'lrn': student.lrn,
            'grade': student.grade.name if student.grade else '',
            'grade_id': student.grade.id if student.grade else None,
            'section': student.section.name if student.section else '',
            'section_id': student.section.id if student.section else None,
            'status': student.status,
        }
        for student in students
    ]
    return JsonResponse({'status': 'success', 'students': data})


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["GET"])
def teacher_non_advisory_get_guardian_details(request, guardian_id):
    guardian = get_object_or_404(Guardian, id=guardian_id)
    child = guardian.student
    children_data = []
    if child:
        children_data.append({
            'id': child.id,
            'first_name': child.first_name,
            'last_name': child.last_name,
            'full_name': f"{child.first_name} {child.last_name}",
            'lrn': child.lrn,
            'grade': child.grade.name if child.grade else '',
            'grade_id': child.grade.id if child.grade else None,
            'section': child.section.name if child.section else '',
            'section_id': child.section.id if child.section else None,
            'status': child.status,
        })

    guardian_data = {
        'id': guardian.id,
        'first_name': guardian.first_name,
        'middle_name': guardian.middle_name or '',
        'last_name': guardian.last_name,
        'email': guardian.email or '',
        'phone_number': guardian.phone or '',
        'relationship': guardian.relationship,
        'children': children_data,
    }
    return JsonResponse({'status': 'success', 'guardian': guardian_data})


@login_required
@user_passes_test(is_teacher)
@require_http_methods(["GET"])
def teacher_non_advisory_get_guardian_children(request, guardian_id):
    guardian = get_object_or_404(Guardian, id=guardian_id)
    child = guardian.student
    children = []
    if child:
        children.append({
            'id': child.id,
            'full_name': f"{child.first_name} {child.last_name}",
            'lrn': child.lrn,
            'grade': child.grade.name if child.grade else '',
            'section': child.section.name if child.section else '',
            'status': child.status,
        })
    guardian_info = {
        'id': guardian.id,
        'name': f"{guardian.first_name} {guardian.last_name}",
        'relationship': guardian.get_relationship_display() if hasattr(guardian, 'get_relationship_display') else guardian.relationship,
    }
    return JsonResponse({'status': 'success', 'guardian': guardian_info, 'children': children})


# ======================
# EXPORT HELPERS
# ======================

@login_required
@user_passes_test(is_teacher)
def export_non_advisory_students(request):
    try:
        import pandas as pd  # noqa: F401
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from reportlab.lib.pagesizes import letter, landscape
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from docx import Document
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        return JsonResponse({'success': False, 'error': f'Required library not installed: {str(e)}'}, status=500)

    export_format = request.GET.get('format', 'excel')
    students = Student.objects.all().select_related('grade', 'section').order_by('id')

    headers = ['ID', 'LRN', 'First Name', 'Middle Name', 'Last Name', 'Email', 'Grade', 'Section', 'Status', 'Face Enrolled']
    data_rows = []

    for student in students:
        data_rows.append([
            student.id,
            student.lrn or '',
            student.first_name,
            student.middle_name or '',
            student.last_name,
            student.email or '',
            student.grade.name if student.grade else '',
            student.section.name if student.section else '',
            student.status,
            'Yes' if student.face_path else 'No',
        ])

    if export_format == 'pdf':
        return export_teacher_data_to_pdf(headers, data_rows, "All Students")
    if export_format == 'word':
        return export_teacher_data_to_word(headers, data_rows, "All Students")
    return export_teacher_data_to_excel(headers, data_rows, "All Students")


@login_required
@user_passes_test(is_teacher)
def export_non_advisory_guardians(request):
    try:
        import pandas as pd  # noqa: F401
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from reportlab.lib.pagesizes import letter, landscape
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from docx import Document
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        return JsonResponse({'success': False, 'error': f'Required library not installed: {str(e)}'}, status=500)

    export_format = request.GET.get('format', 'excel')
    guardians = Guardian.objects.all().select_related('student').order_by('-created_at')

    headers = ['ID', 'First Name', 'Middle Name', 'Last Name', 'Relationship', 'Contact Number', 'Email', 'Student Name']
    data_rows = []

    for guardian in guardians:
        data_rows.append([
            guardian.id,
            guardian.first_name,
            guardian.middle_name or '',
            guardian.last_name,
            guardian.relationship or '',
            guardian.phone or '',
            guardian.email or '',
            f"{guardian.student.first_name} {guardian.student.last_name}" if guardian.student else '',
        ])

    if export_format == 'pdf':
        return export_teacher_data_to_pdf(headers, data_rows, "All Guardians")
    if export_format == 'word':
        return export_teacher_data_to_word(headers, data_rows, "All Guardians")
    return export_teacher_data_to_excel(headers, data_rows, "All Guardians")


def export_teacher_data_to_excel(headers, data_rows, title):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]

    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = thin_border

    for row_num, row_data in enumerate(data_rows, 2):
        for col_num, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='left', vertical='center')

    for column in ws.columns:
        max_length = 0
        column = [cell for cell in column]
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except Exception:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[get_column_letter(column[0].column)].width = adjusted_width

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{title}.xlsx"'
    wb.save(response)
    return response


def export_teacher_data_to_pdf(headers, data_rows, title):
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), topMargin=0.5 * inch, bottomMargin=0.5 * inch)
    elements = []
    styles = getSampleStyleSheet()

    title_style = styles['Heading1']
    title_style.textColor = colors.HexColor('#1F4E78')
    elements.append(Paragraph(title, title_style))
    elements.append(Spacer(1, 0.3 * inch))

    table_data = [headers] + data_rows
    table = Table(table_data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E78')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F0F0')]),
    ]))
    elements.append(table)

    doc.build(elements)
    buffer.seek(0)

    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{title}.pdf"'
    return response


def export_teacher_data_to_word(headers, data_rows, title):
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.color.rgb = RGBColor(31, 78, 120)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(31, 78, 120)

    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{title}.docx"'
    return response
