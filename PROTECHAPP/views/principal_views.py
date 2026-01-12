from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.utils import timezone
from django.contrib import messages
from django.db.models import Q, Count
from django.http import HttpResponse, JsonResponse
from PROTECHAPP.models import CustomUser, Student, Section, Grade, Guardian, Attendance, ExcusedAbsence, UserRole, AdvisoryAssignment
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import io

def is_principal(user):
    """Check if the logged-in user is a principal"""
    return user.is_authenticated and user.role == UserRole.PRINCIPAL

@login_required
@user_passes_test(is_principal)
def principal_grades_redirect(request):
    """Redirect old combined grades&sections URL to new grades page"""
    return redirect('principal_grades')

@login_required
@user_passes_test(is_principal)
def principal_dashboard(request):
    """View for principal dashboard"""
    current_date = timezone.now()
    
    # Example data for dashboard
    attendance_rate = 94.8
    student_count = Student.objects.count()
    teacher_count = CustomUser.objects.filter(role=UserRole.TEACHER).count()
    
    context = {
        'current_date': current_date,
        'attendance_rate': attendance_rate,
        'student_count': student_count,
        'teacher_count': teacher_count
    }
    
    return render(request, 'principal/dashboard.html', context)

@login_required
@user_passes_test(is_principal)
def principal_teachers(request):
    """View for teacher management with filtering and section assignment"""
    # Get query parameters
    search_query = request.GET.get('search', '')
    advisory_filter = request.GET.get('advisory', '')
    status_filter = request.GET.get('status', '')
    section_filter = request.GET.get('section', '')
    page_number = request.GET.get('page', 1)

    # Base queryset - only teachers, order by created_at DESC (newest first)
    teachers = CustomUser.objects.filter(role=UserRole.TEACHER).select_related('section', 'section__grade').order_by('-created_at')

    # Apply search if provided
    if search_query:
        teachers = teachers.filter(
            Q(username__icontains=search_query) | 
            Q(email__icontains=search_query) |
            Q(first_name__icontains=search_query) | 
            Q(last_name__icontains=search_query)
        )

    # Apply status filter using is_active
    if status_filter:
        if status_filter == 'active':
            teachers = teachers.filter(is_active=True)
        elif status_filter == 'inactive':
            teachers = teachers.filter(is_active=False)

    # Get all advisory assignments for later use
    advisory_assignments = AdvisoryAssignment.objects.select_related('section', 'section__grade')

    # Create a dictionary of teacher_id -> section for quick lookup
    teacher_sections = {}
    for assignment in advisory_assignments:
        teacher_sections[assignment.teacher_id] = {
            'section': assignment.section,
            'section_id': assignment.section_id,
            'section_name': assignment.section.name,
            'grade_id': assignment.section.grade.id,
            'grade_name': assignment.section.grade.name
        }

    # Get all available sections for assignment dropdown
    all_sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    # Filter by advisory status - check if teacher has advisory assignment
    if advisory_filter == 'advisory':
        teachers = teachers.filter(id__in=teacher_sections.keys())
    elif advisory_filter == 'non-advisory':
        teachers = teachers.exclude(id__in=teacher_sections.keys())

    # Filter by section (if needed)
    if section_filter:
        try:
            section_id = int(section_filter)
            teacher_ids = [teacher_id for teacher_id, data in teacher_sections.items() 
                           if data['section_id'] == section_id]
            teachers = teachers.filter(Q(id__in=teacher_ids) | Q(section_id=section_id))
        except Exception:
            pass

    # Get counts for dashboard
    total_teachers = CustomUser.objects.filter(role=UserRole.TEACHER).count()
    advisory_teacher_ids = AdvisoryAssignment.objects.values_list('teacher_id', flat=True).distinct()
    advisory_teachers_count = len(advisory_teacher_ids)
    non_advisory_teachers_count = total_teachers - advisory_teachers_count
    active_teachers_count = CustomUser.objects.filter(role=UserRole.TEACHER, is_active=True).count()

    # Pagination
    from django.core.paginator import Paginator
    paginator = Paginator(teachers, 10)
    page_obj = paginator.get_page(page_number)

    # Calculate page range for pagination UI
    from .admin_views import get_pagination_range
    page_range = get_pagination_range(paginator, page_obj.number, 5)

    # Annotate each teacher in the page with section/grade info and advisory status
    for teacher in page_obj:
        section_info = teacher_sections.get(teacher.id)
        # is_advisory is now a property that auto-calculates from teacher.section
        if section_info:
            teacher.section_name = section_info['section_name']
            teacher.grade_name = section_info['grade_name']
            teacher.section_id = section_info['section_id']
            teacher.grade_id = section_info['grade_id']
        else:
            teacher.section_name = None
            teacher.grade_name = None
            teacher.section_id = None
            teacher.grade_id = None
        # Add status for frontend
        teacher.status = 'Active' if teacher.is_active else 'Inactive'

    context = {
        'teachers': page_obj,
        'teacher_sections': teacher_sections,
        'sections': all_sections,
        'grades': Grade.objects.all().order_by('name'), 
        'status_choices': [('active', 'Active'), ('inactive', 'Inactive')],
        'search_query': search_query,
        'advisory_filter': advisory_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'total_teachers': total_teachers,
        'advisory_teachers_count': advisory_teachers_count,
        'non_advisory_teachers_count': non_advisory_teachers_count,
        'active_teachers_count': active_teachers_count,
        'page_range': page_range
    }
    return render(request, 'principal/teachers.html', context)

@login_required
@user_passes_test(is_principal)
def principal_grades(request):
    """View for grade management"""
    from django.db.models import Count
    
    # Get query parameters
    search_query = request.GET.get('search', '')
    page_number = request.GET.get('page', 1)
    
    # Base queryset - order by created_at DESC (newest first)
    grades = Grade.objects.all().order_by('-created_at')
    
    # Apply search if provided
    if search_query:
        grades = grades.filter(name__icontains=search_query)
    
    # Annotate with section and student counts
    grades = grades.annotate(
        sections_count=Count('sections'),
        students_count=Count('sections__students', distinct=True)
    )
    
    # Get stats for dashboard
    total_grades = Grade.objects.count()
    grades_with_sections = Grade.objects.filter(sections__isnull=False).distinct().count()
    total_sections = Section.objects.count()
    total_students = Student.objects.count()
    
    # Pagination
    from django.core.paginator import Paginator
    paginator = Paginator(grades, 10)
    page_obj = paginator.get_page(page_number)
    
    from .admin_views import get_pagination_range
    page_range = get_pagination_range(paginator, page_obj.number, 5)
    
    context = {
        'grades': page_obj,
        'search_query': search_query,
        'total_grades': total_grades,
        'grades_with_sections': grades_with_sections,
        'total_sections': total_sections,
        'total_students': total_students,
        'page_range': page_range,
    }
    return render(request, 'principal/grades.html', context)

@login_required
@user_passes_test(is_principal)
def principal_sections(request):
    """View for section management"""
    from django.db.models import Count, Q, Exists, OuterRef
    
    # Get query parameters
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    advisor_filter = request.GET.get('advisor', '')
    page_number = request.GET.get('page', 1)
    
    # Base queryset - order by grade name, then section name
    sections = Section.objects.select_related('grade').all().order_by('grade__name', 'name')
    
    # Apply search if provided
    if search_query:
        sections = sections.filter(
            Q(name__icontains=search_query) |
            Q(grade__name__icontains=search_query)
        )
    
    # Apply grade filter
    if grade_filter:
        sections = sections.filter(grade_id=grade_filter)
    
    # Apply advisor filter
    if advisor_filter == 'with_advisor':
        sections = sections.filter(advisory_assignments__isnull=False)
    elif advisor_filter == 'without_advisor':
        sections = sections.filter(advisory_assignments__isnull=True)
    
    # Annotate with student count
    sections = sections.annotate(students_count=Count('students'))
    
    # Get advisor information for each section
    for section in sections:
        advisor = AdvisoryAssignment.objects.filter(section=section).select_related('teacher').first()
        section.advisor = advisor.teacher if advisor else None
    
    # Get stats for dashboard
    total_sections = Section.objects.count()
    sections_with_students = Section.objects.filter(students__isnull=False).distinct().count()
    sections_with_advisors = Section.objects.filter(advisory_assignments__isnull=False).distinct().count()
    total_students = Student.objects.count()
    
    # For filter dropdowns
    grades = Grade.objects.all().order_by('name')
    
    # Pagination
    from django.core.paginator import Paginator
    paginator = Paginator(sections, 10)
    page_obj = paginator.get_page(page_number)
    
    from .admin_views import get_pagination_range
    page_range = get_pagination_range(paginator, page_obj.number, 5)
    
    context = {
        'sections': page_obj,
        'grades': grades,
        'search_query': search_query,
        'grade_filter': grade_filter,
        'advisor_filter': advisor_filter,
        'total_sections': total_sections,
        'sections_with_students': sections_with_students,
        'sections_with_advisors': sections_with_advisors,
        'total_students': total_students,
        'page_range': page_range,
    }
    return render(request, 'principal/sections.html', context)

@login_required
@user_passes_test(is_principal)
def get_sections_by_grade(request):
    """API endpoint to get sections for a specific grade"""
    from django.http import JsonResponse
    from django.db.models import Count
    
    grade_id = request.GET.get('grade')
    
    if not grade_id:
        return JsonResponse({'error': 'Grade ID is required'}, status=400)
    
    try:
        grade = Grade.objects.get(id=grade_id)
        sections = Section.objects.filter(grade=grade).select_related('grade').annotate(
            students_count=Count('students')
        ).order_by('name')
        
        sections_data = []
        for section in sections:
            # Get advisor
            advisor = AdvisoryAssignment.objects.filter(section=section).select_related('teacher').first()
            advisor_name = f"{advisor.teacher.first_name} {advisor.teacher.last_name}" if advisor and advisor.teacher else None
            
            sections_data.append({
                'id': section.id,
                'name': section.name,
                'students_count': section.students_count,
                'advisor': advisor_name
            })
        
        return JsonResponse({
            'grade_name': grade.name,
            'sections': sections_data
        })
    except Grade.DoesNotExist:
        return JsonResponse({'error': 'Grade not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@login_required
@user_passes_test(is_principal)
def principal_students(request):
    """View for student management"""
    # Get query parameters
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    page_number = request.GET.get('page', 1)

    # Base queryset
    students = Student.objects.select_related('grade', 'section').all().order_by('-created_at')

    # Apply search
    if search_query:
        students = students.filter(
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(lrn__icontains=search_query) |
            Q(email__icontains=search_query)
        )

    # Apply grade filter
    if grade_filter:
        students = students.filter(grade_id=grade_filter)

    # Apply section filter
    if section_filter:
        students = students.filter(section_id=section_filter)

    # Apply status filter
    if status_filter:
        students = students.filter(status=status_filter)

    # Dashboard stats
    total_students = Student.objects.count()
    active_students = Student.objects.filter(status='ACTIVE').count()
    inactive_students = Student.objects.filter(status='INACTIVE').count()
    face_enrolled_count = Student.objects.exclude(face_path__isnull=True).exclude(face_path__exact='').count()

    # For filter dropdowns
    grades = Grade.objects.all().order_by('name')
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    # Pagination
    from django.core.paginator import Paginator
    paginator = Paginator(students, 10)
    page_obj = paginator.get_page(page_number)
    
    from .admin_views import get_pagination_range
    page_range = get_pagination_range(paginator, page_obj.number, 5)

    context = {
        'students': page_obj,
        'grades': grades,
        'sections': sections,
        'search_query': search_query,
        'grade_filter': grade_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'total_students': total_students,
        'active_students': active_students,
        'inactive_students': inactive_students,
        'face_enrolled_count': face_enrolled_count,
        'page_range': page_range,
    }
    return render(request, 'principal/students.html', context)

@login_required
@user_passes_test(is_principal)
def principal_guardians(request):
    """Guardian management view with dashboard cards"""
    # Get query parameters
    search_query = request.GET.get('search', '')
    relationship_filter = request.GET.get('relationship', '')
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    page_number = request.GET.get('page', 1)

    # Base queryset - order by created_at DESC (newest first)
    guardians = Guardian.objects.select_related('student', 'student__grade', 'student__section').order_by('-created_at')

    # Apply search
    if search_query:
        guardians = guardians.filter(
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(email__icontains=search_query) |
            Q(phone__icontains=search_query)
        )

    # Apply relationship filter
    if relationship_filter:
        guardians = guardians.filter(relationship=relationship_filter)

    # Apply grade filter (filter by student's grade)
    if grade_filter:
        guardians = guardians.filter(student__grade_id=grade_filter)

    # Apply section filter (filter by student's section)
    if section_filter:
        guardians = guardians.filter(student__section_id=section_filter)

    # Calculate dashboard statistics
    total_guardians = Guardian.objects.count()
    
    # Count guardians with children (since each guardian has exactly one student)
    guardians_with_children = Guardian.objects.exclude(student=None).count()
    
    # Count mother guardians
    mother_count = Guardian.objects.filter(relationship='MOTHER').count()
    
    # Count father guardians
    father_count = Guardian.objects.filter(relationship='FATHER').count()
    
    # Get relationship choices for filter dropdown
    relationship_choices = [
        ('FATHER', 'Father'),
        ('MOTHER', 'Mother'),
        ('GUARDIAN', 'Guardian'),
        ('GRANDMOTHER', 'Grandmother'),
        ('GRANDFATHER', 'Grandfather'),
        ('AUNT', 'Aunt'),
        ('UNCLE', 'Uncle'),
        ('SIBLING', 'Sibling'),
        ('OTHER', 'Other'),
    ]
    
    # Get all grades for filter dropdown
    grades = Grade.objects.all().order_by('name')
    
    # Get all sections for filter dropdown
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')
    
    # Get students for children dropdown
    students = Student.objects.select_related('grade', 'section').filter(status='ACTIVE')

    # Pagination
    from django.core.paginator import Paginator
    paginator = Paginator(guardians, 10)
    page_obj = paginator.get_page(page_number)
    
    from .admin_views import get_pagination_range
    page_range = get_pagination_range(paginator, page_obj.number, 5)
    
    context = {
        'guardians': page_obj,
        'total_guardians': total_guardians,
        'guardians_with_children': guardians_with_children,
        'mother_count': mother_count,
        'father_count': father_count,
        'relationship_choices': relationship_choices,
        'grades': grades,
        'sections': sections,
        'students': students,
        'search_query': search_query,
        'relationship_filter': relationship_filter,
        'grade_filter': grade_filter,
        'section_filter': section_filter,
        'page_range': page_range,
    }
    
    return render(request, 'principal/guardians.html', context)

@login_required
@user_passes_test(is_principal)
def principal_attendance(request):
    """View for attendance records"""
    from django.db.models import Q
    from datetime import date
    
    # Get query parameters
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')
    page_number = request.GET.get('page', 1)

    # Base queryset
    attendance_records = Attendance.objects.select_related('student', 'student__grade', 'student__section').all().order_by('-date', '-time_in')

    # Apply search
    if search_query:
        attendance_records = attendance_records.filter(
            Q(student__first_name__icontains=search_query) |
            Q(student__last_name__icontains=search_query) |
            Q(student__lrn__icontains=search_query) |
            Q(student__section__name__icontains=search_query)
        )
    
    # Apply filters
    if grade_filter:
        attendance_records = attendance_records.filter(student__grade_id=grade_filter)
    if section_filter:
        attendance_records = attendance_records.filter(student__section_id=section_filter)
    if status_filter:
        attendance_records = attendance_records.filter(status=status_filter)
    if date_filter:
        attendance_records = attendance_records.filter(date=date_filter)

    # Dashboard stats
    total_records = Attendance.objects.count()
    present_count = Attendance.objects.filter(status='ON TIME').count() + Attendance.objects.filter(status='LATE').count()
    absent_count = Attendance.objects.filter(status='ABSENT').count()
    excused_count = Attendance.objects.filter(status='EXCUSED').count()

    # For filter dropdowns
    grades = Grade.objects.all().order_by('name')
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    # Pagination
    from django.core.paginator import Paginator
    paginator = Paginator(attendance_records, 10)
    page_obj = paginator.get_page(page_number)
    
    from .admin_views import get_pagination_range
    page_range = get_pagination_range(paginator, page_obj.number, 5)

    context = {
        'attendance_records': page_obj,
        'grades': grades,
        'sections': sections,
        'search_query': search_query,
        'grade_filter': grade_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'date_filter': date_filter,
        'total_records': total_records,
        'present_count': present_count,
        'absent_count': absent_count,
        'excused_count': excused_count,
        'page_range': page_range,
    }
    return render(request, 'principal/attendance.html', context)

@login_required
@user_passes_test(is_principal)
def principal_excused(request):
    """View for excused absences"""
    from django.db.models import Q
    from datetime import date
    
    # Get query parameters
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')
    page_number = request.GET.get('page', 1)

    # Base queryset
    excused_absences = ExcusedAbsence.objects.select_related('student', 'student__grade', 'student__section').all().order_by('-date_absent')

    # Apply search
    if search_query:
        excused_absences = excused_absences.filter(
            Q(student__first_name__icontains=search_query) |
            Q(student__last_name__icontains=search_query) |
            Q(student__lrn__icontains=search_query) |
            Q(student__section__name__icontains=search_query)
        )
    
    # Apply filters
    if grade_filter:
        excused_absences = excused_absences.filter(student__grade_id=grade_filter)
    if section_filter:
        excused_absences = excused_absences.filter(student__section_id=section_filter)
    today = date.today()
    if status_filter:
        if status_filter == 'ACTIVE':
            excused_absences = excused_absences.filter(effective_date__lte=today, end_date__gte=today)
        elif status_filter == 'EXPIRED':
            excused_absences = excused_absences.filter(end_date__lt=today)
        elif status_filter == 'UPCOMING':
            excused_absences = excused_absences.filter(effective_date__gt=today)
    if date_filter:
        excused_absences = excused_absences.filter(date_absent=date_filter)

    # Dashboard stats
    total_excused = ExcusedAbsence.objects.count()
    active_excused = ExcusedAbsence.objects.filter(effective_date__lte=today, end_date__gte=today).count()
    expired_excused = ExcusedAbsence.objects.filter(end_date__lt=today).count()
    upcoming_excused = ExcusedAbsence.objects.filter(effective_date__gt=today).count()

    # For filter dropdowns
    grades = Grade.objects.all().order_by('name')
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    # Pagination
    from django.core.paginator import Paginator
    paginator = Paginator(excused_absences, 10)
    page_obj = paginator.get_page(page_number)
    
    from .admin_views import get_pagination_range
    page_range = get_pagination_range(paginator, page_obj.number, 5)

    context = {
        'excused_absences': page_obj,
        'grades': grades,
        'sections': sections,
        'search_query': search_query,
        'grade_filter': grade_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'date_filter': date_filter,
        'total_excused': total_excused,
        'active_excused': active_excused,
        'expired_excused': expired_excused,
        'upcoming_excused': upcoming_excused,
        'page_range': page_range,
    }
    return render(request, 'principal/excused.html', context)

@login_required
@user_passes_test(is_principal)
def principal_calendar(request):
    """View for calendar"""
    return render(request, 'principal/calendar.html')


@login_required
@user_passes_test(is_principal)
def get_principal_calendar_events(request):
    """API endpoint to get calendar events for principal"""
    from ..models import CalendarEvent
    from django.db.models import Q
    
    # Get events and announcements targeted to principal or all staff
    events = CalendarEvent.objects.filter(
        Q(event_type='EVENT') |  # All events
        Q(event_type='ANNOUNCEMENT', target_role__in=['PRINCIPAL', 'ADMIN'])  # Announcements for principal
    )
    
    events_data = []
    for event in events:
        events_data.append({
            'id': event.id,
            'title': event.title,
            'start': event.start_date.isoformat(),
            'end': event.end_date.isoformat(),
            'description': event.description,
            'location': event.location,
            'event_type': event.event_type,
            'target_role': event.target_role,
            'target_role_display': event.get_target_role_display() if event.target_role else None,
            'created_by': f"{event.created_by.first_name} {event.created_by.last_name}",
        })
    
    return JsonResponse(events_data, safe=False)


@login_required
@user_passes_test(is_principal)
def principal_announcements(request):
    """View for announcements"""
    return render(request, 'principal/announcements.html')

@login_required
@user_passes_test(is_principal)
def principal_messages(request):
    """View for messages - now using PostgreSQL with polling"""
    return render(request, 'principal/messages.html')

@login_required
@user_passes_test(is_principal)
def principal_settings(request):
    """View for settings"""
    return render(request, 'principal/settings.html')

# ============================================
# EXPORT FUNCTIONS FOR PRINCIPAL
# ============================================

from django.http import HttpResponse, JsonResponse
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_principal_data_to_excel(data, headers, filename, title):
    """Helper function to export principal data to Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = title
    
    # Add title
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    title_cell = ws.cell(row=1, column=1, value=title)
    title_cell.font = Font(size=16, bold=True, color="FFFFFF")
    title_cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 30
    
    # Add headers
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col_num, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Add data
    for row_num, row_data in enumerate(data, 3):
        for col_num, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num, value=value)
            cell.alignment = Alignment(horizontal='left', vertical='center')
    
    # Adjust column widths
    for col_num in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col_num)].width = 20
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    wb.save(response)
    return response

def export_principal_data_to_pdf(data, headers, filename, title):
    """Helper function to export principal data to PDF"""
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    doc = SimpleDocTemplate(response, pagesize=landscape(letter))
    elements = []
    
    # Title
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1F4E78'),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    elements.append(Paragraph(title, title_style))
    elements.append(Spacer(1, 0.2*inch))
    
    # Table
    table_data = [headers] + data
    table = Table(table_data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E78')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
    ]))
    elements.append(table)
    
    doc.build(elements)
    return response

def export_principal_data_to_word(data, headers, filename, title):
    """Helper function to export principal data to Word"""
    doc = Document()
    
    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(31, 78, 120)
    
    # Table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        # Set cell background color
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '1F4E78')
        hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
    
    # Data
    for row_data in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    doc.save(response)
    return response

@login_required
@user_passes_test(is_principal)
def export_principal_students(request):
    """Export students data with filter support"""
    try:
        # Get query parameters for filtering
        search_query = request.GET.get('search', '').strip()
        grade_filter = request.GET.get('grade', '')
        section_filter = request.GET.get('section', '')
        status_filter = request.GET.get('status', '')
        format_type = request.GET.get('format', 'excel')
        
        # Base queryset
        students = Student.objects.select_related('section', 'section__grade').order_by('id')
        
        # Apply search if provided
        if search_query:
            students = students.filter(
                Q(first_name__icontains=search_query) |
                Q(last_name__icontains=search_query) |
                Q(lrn__icontains=search_query)
            )
        
        # Apply grade filter if provided
        if grade_filter:
            students = students.filter(grade_id=grade_filter)
        
        # Apply section filter if provided
        if section_filter:
            students = students.filter(section_id=section_filter)
        
        # Apply status filter if provided
        if status_filter:
            students = students.filter(status=status_filter)
        
        headers = ['LRN', 'Last Name', 'First Name', 'Middle Name', 'Grade', 'Section', 'Status']
        data = []
        
        for student in students:
            data.append([
                student.lrn or '',
                student.last_name or '',
                student.first_name or '',
                student.middle_name or '',
                student.section.grade.name if student.section else '',
                student.section.name if student.section else '',
                student.get_status_display()
            ])
        
        title = 'Students List'
        filename = f'students_{timezone.now().strftime("%Y%m%d_%H%M%S")}'
        
        if format_type == 'pdf':
            return export_principal_data_to_pdf(data, headers, f'{filename}.pdf', title)
        elif format_type == 'word':
            return export_principal_data_to_word(data, headers, f'{filename}.docx', title)
        else:
            return export_principal_data_to_excel(data, headers, f'{filename}.xlsx', title)
    except Exception as e:
        import traceback
        print(f"Export error: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_principal)
def export_principal_teachers(request):
    """Export teachers data"""
    format_type = request.GET.get('format', 'excel')
    
    teachers = CustomUser.objects.filter(role=UserRole.TEACHER).select_related('section', 'section__grade').order_by('id')
    advisory_assignments = {a.teacher_id: a.section for a in AdvisoryAssignment.objects.select_related('section', 'section__grade')}
    
    headers = ['ID', 'Username', 'Last Name', 'First Name', 'Email', 'Advisory Section', 'Status']
    data = []
    
    for teacher in teachers:
        advisory_section = advisory_assignments.get(teacher.id)
        advisory_text = f"{advisory_section.grade.name} - {advisory_section.name}" if advisory_section else 'None'
        
        data.append([
            teacher.id,
            teacher.username,
            teacher.last_name or '',
            teacher.first_name or '',
            teacher.email or '',
            advisory_text,
            'Active' if teacher.is_active else 'Inactive'
        ])
    
    title = 'Teachers List'
    filename = f'teachers_{timezone.now().strftime("%Y%m%d_%H%M%S")}'
    
    if format_type == 'pdf':
        return export_principal_data_to_pdf(data, headers, f'{filename}.pdf', title)
    elif format_type == 'word':
        return export_principal_data_to_word(data, headers, f'{filename}.docx', title)
    else:
        return export_principal_data_to_excel(data, headers, f'{filename}.xlsx', title)

@login_required
@user_passes_test(is_principal)
def export_principal_guardians(request):
    """Export guardians data"""
    format_type = request.GET.get('format', 'excel')
    
    guardians = Guardian.objects.prefetch_related('students').order_by('id')
    
    headers = ['ID', 'Last Name', 'First Name', 'Middle Name', 'Email', 'Phone', 'Relationship', 'Children Count']
    data = []
    
    for guardian in guardians:
        data.append([
            guardian.id,
            guardian.last_name or '',
            guardian.first_name or '',
            guardian.middle_name or '',
            guardian.email or '',
            guardian.phone_number or '',
            guardian.get_relationship_display(),
            guardian.students.count()
        ])
    
    title = 'Guardians List'
    filename = f'guardians_{timezone.now().strftime("%Y%m%d_%H%M%S")}'
    
    if format_type == 'pdf':
        return export_principal_data_to_pdf(data, headers, f'{filename}.pdf', title)
    elif format_type == 'word':
        return export_principal_data_to_word(data, headers, f'{filename}.docx', title)
    else:
        return export_principal_data_to_excel(data, headers, f'{filename}.xlsx', title)

@login_required
@user_passes_test(is_principal)
def export_principal_grades(request):
    """Export grades data"""
    try:
        format_type = request.GET.get('format', 'excel')
        search_query = request.GET.get('search', '').strip()
        
        grades = Grade.objects.all().order_by('id')
        
        if search_query:
            grades = grades.filter(name__icontains=search_query)
        
        headers = ['Grade', 'Sections', 'Students']
        data = []
        
        for grade in grades:
            section_count = grade.sections.count()
            student_count = Student.objects.filter(grade=grade).count()
            
            data.append([
                grade.name,
                str(section_count),
                str(student_count)
            ])
        
        title = 'Grades List'
        filename = f'grades_{timezone.now().strftime("%Y%m%d_%H%M%S")}'
        
        if format_type == 'pdf':
            return export_principal_data_to_pdf(data, headers, f'{filename}.pdf', title)
        elif format_type == 'word':
            return export_principal_data_to_word(data, headers, f'{filename}.docx', title)
        else:
            return export_principal_data_to_excel(data, headers, f'{filename}.xlsx', title)
    except Exception as e:
        import traceback
        print(f"Export error: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_principal)
def export_principal_sections(request):
    """Export sections data"""
    try:
        format_type = request.GET.get('format', 'excel')
        
        sections = Section.objects.select_related('grade').annotate(student_count=Count('students')).order_by('id')
        advisory_assignments = {a.section_id: a.teacher for a in AdvisoryAssignment.objects.select_related('teacher')}
        
        headers = ["ID", "Section Name", "Grade", "Room Number", "Capacity"]
        data = []

        for section in sections:
            room_number = getattr(section, 'room_number', '') or ''
            capacity = getattr(section, 'capacity', '') or ''

            data.append([
                str(section.id),
                section.name,
                section.grade.name if section.grade else "",
                room_number,
                str(capacity)
            ])
        
        title = 'Sections List'
        filename = f'sections_{timezone.now().strftime("%Y%m%d_%H%M%S")}'
        
        if format_type == 'pdf':
            return export_principal_data_to_pdf(data, headers, f'{filename}.pdf', title)
        elif format_type == 'word':
            return export_principal_data_to_word(data, headers, f'{filename}.docx', title)
        else:
            return export_principal_data_to_excel(data, headers, f'{filename}.xlsx', title)
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        return JsonResponse({'status': 'error', 'message': str(e), 'trace': error_trace}, status=500)

@login_required
@user_passes_test(is_principal)
def export_principal_attendance(request):
    """Export attendance records"""
    format_type = request.GET.get('format', 'excel')
    date_filter = request.GET.get('date', '')
    
    attendance_records = Attendance.objects.select_related('student', 'student__section', 'student__section__grade').order_by('id')
    
    if date_filter:
        try:
            from datetime import datetime
            filter_date = datetime.strptime(date_filter, '%Y-%m-%d').date()
            attendance_records = attendance_records.filter(date=filter_date)
        except:
            pass
    
    headers = ['ID', 'Student Name', 'LRN', 'Grade', 'Section', 'Date', 'Time In', 'Status']
    data = []
    
    for record in attendance_records:
        data.append([
            record.id,
            f"{record.student.first_name} {record.student.last_name}",
            record.student.lrn or '',
            record.student.section.grade.name if record.student.section else '',
            record.student.section.name if record.student.section else '',
            record.date.strftime('%Y-%m-%d'),
            record.time_in.strftime('%H:%M:%S') if record.time_in else '',
            record.get_status_display()
        ])
    
    title = 'Attendance Records'
    filename = f'attendance_{timezone.now().strftime("%Y%m%d_%H%M%S")}'
    
    if format_type == 'pdf':
        return export_principal_data_to_pdf(data, headers, f'{filename}.pdf', title)
    elif format_type == 'word':
        return export_principal_data_to_word(data, headers, f'{filename}.docx', title)
    else:
        return export_principal_data_to_excel(data, headers, f'{filename}.xlsx', title)

@login_required
@user_passes_test(is_principal)
def export_principal_excused(request):
    """Export excused absences"""
    format_type = request.GET.get('format', 'excel')
    
    excused_absences = ExcusedAbsence.objects.select_related('student', 'student__section', 'student__section__grade').order_by('id')
    
    headers = ['ID', 'Student Name', 'LRN', 'Grade', 'Section', 'Date Absent', 'Effective Date', 'End Date', 'Has Letter']
    data = []

    for excuse in excused_absences:
        student_name = f"{excuse.student.first_name} {excuse.student.last_name}"
        has_letter = 'Yes' if excuse.excuse_letter else 'No'
        data.append([
            str(excuse.id),
            student_name,
            excuse.student.lrn or '',
            excuse.student.section.grade.name if excuse.student.section else '',
            excuse.student.section.name if excuse.student.section else '',
            excuse.date_absent.strftime('%Y-%m-%d') if excuse.date_absent else '',
            excuse.effective_date.strftime('%Y-%m-%d') if excuse.effective_date else '',
            excuse.end_date.strftime('%Y-%m-%d') if excuse.end_date else '',
            has_letter
        ])
    
    title = 'Excused Absences'
    filename = f'excused_{timezone.now().strftime("%Y%m%d_%H%M%S")}'
    
    if format_type == 'pdf':
        return export_principal_data_to_pdf(data, headers, f'{filename}.pdf', title)
    elif format_type == 'word':
        return export_principal_data_to_word(data, headers, f'{filename}.docx', title)
    else:
        return export_principal_data_to_excel(data, headers, f'{filename}.xlsx', title)
