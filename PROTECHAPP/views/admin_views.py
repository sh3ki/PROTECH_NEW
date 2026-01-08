from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.utils import timezone
from django.contrib import messages
from django.http import FileResponse, Http404, JsonResponse, HttpResponse, HttpResponseNotFound
from django.core.paginator import Paginator
from django.db.models import Q, Count
from PROTECHAPP.models import CustomUser, Student, Section, Grade, Guardian, Attendance, UserRole, UserStatus, ExcusedAbsence, AdvisoryAssignment, PasswordResetOTP
from django.views.decorators.http import require_http_methods, require_GET
from django.contrib.auth.hashers import make_password
from django.core.mail import send_mail
import json
import re
import os
import traceback
import secrets
import string
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from django.views.decorators.csrf import csrf_exempt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from datetime import datetime, timedelta
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT as WD_ORIENTATION
import io

def is_admin(user):
    """Check if the logged-in user is an admin"""
    return user.is_authenticated and user.role == UserRole.ADMIN

def is_admin_or_principal(user):
    """Check if the logged-in user is an admin or principal"""
    return user.is_authenticated and user.role in [UserRole.ADMIN, UserRole.PRINCIPAL]

def is_admin_or_teacher(user):
    """Allow admins, principals, registrars and teachers to perform certain actions like imports"""
    return user.is_authenticated and user.role in [UserRole.ADMIN, UserRole.PRINCIPAL, UserRole.REGISTRAR, UserRole.TEACHER]

def generate_random_password(length=12):
    """Generate a secure random password"""
    characters = string.ascii_letters + string.digits + "!@#$%^&*"
    password = ''.join(secrets.choice(characters) for _ in range(length))
    return password

def send_password_email(user_email, username, password, user_type="User"):
    """Send email with generated password to new user"""
    try:
        subject = f'Your PROTECH Account Has Been Created'
        message = f'''Hello,

Your PROTECH account has been created successfully!

Login Credentials:
Username: {username}
Password: {password}
Account Type: {user_type}

Please login at your earliest convenience and change your password in the settings.

Best regards,
PROTECH Admin Team'''
        
        send_mail(
            subject,
            message,
            settings.DEFAULT_FROM_EMAIL,
            [user_email],
            fail_silently=True,
        )
        return True
    except Exception as e:
        print(f"Error sending email to {user_email}: {str(e)}")
        return False


# ==========================
#  USER MANAGEMENT VIEWS
# ==========================

@login_required
@user_passes_test(is_admin)
def admin_users(request):
    """View for user management with filtering, searching and pagination"""
    # Get query parameters
    search_query = request.GET.get('search', '')
    role_filter = request.GET.get('role', '')
    status_filter = request.GET.get('status', '')
    page_number = request.GET.get('page', 1)
    items_per_page = request.GET.get('items_per_page', 10)
    sort_by = request.GET.get('sort_by', '-created_at')
    
    # Valid sorting fields - remove status and add is_active
    valid_sort_fields = ['username', '-username', 'email', '-email', 
                         'first_name', '-first_name', 'last_name', '-last_name',
                         'role', '-role', 'is_active', '-is_active', 
                         'created_at', '-created_at']
    
    if sort_by not in valid_sort_fields:
        sort_by = '-created_at'
    
    # Base queryset
    users = CustomUser.objects.all().order_by(sort_by)
    
    # Apply search if provided
    if search_query:
        users = users.filter(
            Q(username__icontains=search_query) | 
            Q(email__icontains=search_query) |
            Q(first_name__icontains=search_query) | 
            Q(last_name__icontains=search_query)
        )
    
    # Apply role filter if provided
    if role_filter:
        users = users.filter(role=role_filter)
    
    # Apply status filter (now using is_active instead of status)
    if status_filter:
        if status_filter == 'active':
            users = users.filter(is_active=True)
        elif status_filter == 'inactive':
            users = users.filter(is_active=False)
    
    # Get total users before pagination for stats
    total_filtered = users.count()
    total_users = CustomUser.objects.count()
    
    # Get role distribution for dashboard widgets
    role_stats = CustomUser.objects.values('role').annotate(count=Count('id'))
    active_stats = CustomUser.objects.values('is_active').annotate(count=Count('id'))
    
    # Create dictionaries for easy template access
    role_counts = {role[0]: 0 for role in UserRole.choices}
    
    for stat in role_stats:
        role_counts[stat['role']] = stat['count']
    
    # Create active/inactive counts
    active_count = 0
    inactive_count = 0
    
    for stat in active_stats:
        if stat['is_active']:
            active_count = stat['count']
        else:
            inactive_count = stat['count']
    
    # Get recently added users for the sidebar
    recent_users = CustomUser.objects.order_by('-created_at')[:5]
    
    # Parse items_per_page as integer with validation
    try:
        items_per_page = int(items_per_page)
        if items_per_page <= 0:
            items_per_page = 10
    except (ValueError, TypeError):
        items_per_page = 10
    
    # Pagination
    paginator = Paginator(users, items_per_page)
    page_obj = paginator.get_page(page_number)
    
    # Calculate page range for pagination UI
    page_range = get_pagination_range(paginator, page_obj.number, 5)
    
    # Create simple status choices for the template
    status_choices = [
        ('active', 'Active'),
        ('inactive', 'Inactive')
    ]
    
    context = {
        'users': page_obj,
        'role_choices': UserRole.choices,
        'status_choices': status_choices,  # Simple active/inactive choices
        'search_query': search_query,
        'role_filter': role_filter,
        'status_filter': status_filter,
        'sort_by': sort_by,
        'user_count': active_count,
        'total_filtered': total_filtered,
        'role_counts': role_counts,
        'active_count': active_count,
        'inactive_count': inactive_count,
        'recent_users': recent_users,
        'page_range': page_range,
        'items_per_page': items_per_page,
        'items_per_page_options': [10, 25, 50, 100],
        'sort_options': [
            {'field': 'username', 'label': 'Username (A-Z)'},
            {'field': '-username', 'label': 'Username (Z-A)'},
            {'field': 'email', 'label': 'Email (A-Z)'},
            {'field': '-email', 'label': 'Email (Z-A)'},
            {'field': 'first_name', 'label': 'First Name (A-Z)'},
            {'field': '-first_name', 'label': 'First Name (Z-A)'},
            {'field': 'created_at', 'label': 'Created (Oldest First)'},
            {'field': '-created_at', 'label': 'Created (Newest First)'},
        ],
        # Add sections for modal dropdown
        'sections': Section.objects.select_related('grade').order_by('grade__name', 'name'),
        # Order grades by numeric value rather than alphabetically
        'grades': Grade.objects.all().order_by('name'),
    }
    return render(request, 'admin/users.html', context)

def get_pagination_range(paginator, current_page, display_range=5):
    """
    Calculate pagination range to display in the template
    This creates a better UX by showing ellipsis for large page sets
    """
    total_pages = paginator.num_pages
    
    # If total pages is less than display range, show all pages
    if total_pages <= display_range:
        return range(1, total_pages + 1)
    
    # Calculate the range to display
    start_page = max(current_page - display_range // 2, 1)
    end_page = min(start_page + display_range - 1, total_pages)
    
    # Adjust if we're near the end
    if end_page - start_page + 1 < display_range:
        start_page = max(end_page - display_range + 1, 1)
    
    return range(start_page, end_page + 1)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def create_user(request):
    """API endpoint to create a new user"""
    try:
        import secrets
        import string
        from django.core.mail import EmailMultiAlternatives
        
        data = json.loads(request.body)
        
        # Validate required fields - removed 'password' from required fields
        required_fields = ['username', 'email', 'first_name', 'last_name', 'role']
        for field in required_fields:
            if not data.get(field):
                return JsonResponse({'status': 'error', 'message': f'{field} is required'}, status=400)
        
        # Check if username or email already exists
        if CustomUser.objects.filter(username=data['username']).exists():
            return JsonResponse({'status': 'error', 'message': 'Username already exists'}, status=400)
        
        if CustomUser.objects.filter(email=data['email']).exists():
            return JsonResponse({'status': 'error', 'message': 'Email already exists'}, status=400)
        
        # Generate a random secure password (12 characters)
        password_chars = string.ascii_letters + string.digits + '!@#$%^&*'
        generated_password = ''.join(secrets.choice(password_chars) for _ in range(12))
        
        # Create user with basic fields
        user = CustomUser(
            username=data['username'],
            email=data['email'],
            first_name=data['first_name'],
            last_name=data['last_name'],
            middle_name=data.get('middle_name', ''),
            role=data['role'],
            is_active=data.get('is_active', True)  # Default to active if not specified
        )
        
        # Set generated password
        user.password = make_password(generated_password)
        
        # Handle section assignment for teachers
        if data['role'] == UserRole.TEACHER and 'section' in data and data['section']:
            try:
                section = Section.objects.get(id=data['section'])
                user.section = section
            except Section.DoesNotExist:
                pass  # Silently ignore invalid section IDs
        
        user.save()
        
        # Send password to user's email
        email_sent = False
        email_error_msg = None
        
        try:
            subject = 'Your PROTECH Account Credentials'
            
            # Plain text version
            text_message = f'''Hello {user.get_full_name() or user.username},

Your account has been created successfully for the PROTECH Attendance Monitoring System.

Your login credentials:
Username: {user.username}
Password: {generated_password}

Please log in and change your password after your first login.

Login URL: {request.build_absolute_uri('/login/')}

Attendance Monitoring System PROTECH
'''
            
            # HTML version with modern dark design
            html_message = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Your PROTECH Account Credentials</title>
</head>
<body style="margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif; background-color: #1a1a1a;">
    <table width="100%" cellpadding="0" cellspacing="0" style="background-color: #1a1a1a; padding: 40px 20px;">
        <tr>
            <td align="center">
                <table width="600" cellpadding="0" cellspacing="0" style="background-color: #2a2a2a; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.3);">
                    <tr>
                        <td style="background: linear-gradient(135deg, #1e40af 0%, #3b82f6 100%); padding: 30px; text-align: center;">
                            <h1 style="color: white; font-size: 24px; font-weight: bold; margin: 0;">Welcome to PROTECH</h1>
                        </td>
                    </tr>
                    <tr>
                        <td style="padding: 20px 40px 0;">
                            <div style="background-color: #10b981; color: white; display: inline-block; padding: 4px 12px; border-radius: 4px; font-size: 11px; font-weight: 600; text-transform: uppercase;">Account Created</div>
                        </td>
                    </tr>
                    <tr>
                        <td style="padding: 20px 40px;">
                            <h2 style="color: #ffffff; font-size: 18px; font-weight: 600; margin: 0 0 15px;">Your Account Credentials</h2>
                            <p style="color: #d1d5db; font-size: 14px; line-height: 1.6; margin: 0 0 10px;">Hello {user.get_full_name() or user.username},</p>
                            <p style="color: #d1d5db; font-size: 14px; line-height: 1.6; margin: 0 0 25px;">Your account has been created successfully. Please use the following credentials to log in:</p>
                            
                            <div style="background-color: #1a1a1a; border: 2px solid #3f3f46; border-radius: 8px; padding: 20px; margin: 0 0 25px;">
                                <p style="color: #9ca3af; font-size: 13px; margin: 0 0 10px;">Username:</p>
                                <p style="color: #ffffff; font-size: 16px; font-weight: bold; margin: 0 0 20px; font-family: 'Courier New', monospace;">{user.username}</p>
                                <p style="color: #9ca3af; font-size: 13px; margin: 0 0 10px;">Password:</p>
                                <p style="color: #10b981; font-size: 18px; font-weight: bold; margin: 0; font-family: 'Courier New', monospace; letter-spacing: 1px;">{generated_password}</p>
                            </div>
                            
                            <p style="color: #ef4444; font-size: 13px; line-height: 1.5; margin: 0 0 10px;"><strong>Important:</strong> Please change your password after your first login.</p>
                            <p style="color: #9ca3af; font-size: 13px; line-height: 1.5; margin: 0;">Login URL: <a href="{request.build_absolute_uri('/login/')}" style="color: #3b82f6;">{request.build_absolute_uri('/login/')}</a></p>
                        </td>
                    </tr>
                    <tr>
                        <td style="padding: 20px 40px 30px; border-top: 1px solid #3f3f46;">
                            <p style="color: #6b7280; font-size: 12px; margin: 0;">Attendance Monitoring System PROTECH</p>
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
    </table>
</body>
</html>'''
            
            from django.conf import settings
            
            # Check email configuration
            print("\n" + "="*60)
            print("ATTEMPTING TO SEND EMAIL")
            print("="*60)
            print(f"EMAIL_BACKEND: {settings.EMAIL_BACKEND}")
            print(f"EMAIL_HOST: {settings.EMAIL_HOST}")
            print(f"EMAIL_PORT: {settings.EMAIL_PORT}")
            print(f"EMAIL_HOST_USER: {settings.EMAIL_HOST_USER}")
            print(f"DEFAULT_FROM_EMAIL: {settings.DEFAULT_FROM_EMAIL}")
            print(f"Sending email to: {user.email}")
            print(f"Username: {user.username}")
            print(f"Generated Password: {generated_password}")
            print("="*60)
            
            msg = EmailMultiAlternatives(
                subject,
                text_message,
                settings.DEFAULT_FROM_EMAIL,
                [user.email]
            )
            msg.attach_alternative(html_message, "text/html")
            
            print("Calling send()...")
            send_result = msg.send(fail_silently=False)
            print(f"Send result: {send_result}")
            
            email_sent = True
            print(f"✓ Email sent successfully to {user.email}")
            
        except Exception as email_error:
            email_error_msg = str(email_error)
            print(f"✗ Failed to send email to {user.email}: {email_error}")
            print(f"Error type: {type(email_error).__name__}")
            import traceback
            traceback.print_exc()
            # Don't fail the user creation if email fails
        
        # Prepare response message
        if email_sent:
            message = f'User created successfully. Login credentials have been sent to {user.email}.'
        else:
            message = f'User created successfully, but email could not be sent. Error: {email_error_msg}. Please manually provide credentials: Username: {user.username}, Password: {generated_password}'

        # Note: OTP verification will be triggered when the user initiates verification
        # (e.g. via the "Send Verification Code" action during first-time login).

        return JsonResponse({
            'status': 'success',
            'message': message,
            'user_id': user.id,
            'email_sent': email_sent,
            'generated_password': generated_password if not email_sent else None  # Include password in response if email failed
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
def get_user(request, user_id):
    """API endpoint to get a single user's details"""
    try:
        user = get_object_or_404(CustomUser, id=user_id)
        
        # Build the user data with additional fields
        user_data = {
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'middle_name': user.middle_name or '',
            'role': user.role,
            'is_active': user.is_active,
            'created_at': user.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'profile_pic': user.profile_pic,  # Add profile pic path
        }
        
        # Add section and grade data if the user is a teacher with assigned section
        if user.role == UserRole.TEACHER and user.section:
            user_data['section'] = user.section.id
            user_data['grade'] = user.section.grade.id
        
        return JsonResponse({'status': 'success', 'user': user_data})
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["PUT", "PATCH"])
def update_user(request, user_id):
    """API endpoint to update a user"""
    try:
        user = get_object_or_404(CustomUser, id=user_id)
        data = json.loads(request.body)
        
        # Update fields if provided
        if 'username' in data and data['username'] != user.username:
            if CustomUser.objects.filter(username=data['username']).exists():
                return JsonResponse({'status': 'error', 'message': 'Username already exists'}, status=400)
            user.username = data['username']
            
        if 'email' in data and data['email'] != user.email:
            if CustomUser.objects.filter(email=data['email']).exists():
                return JsonResponse({'status': 'error', 'message': 'Email already exists'}, status=400)
            user.email = data['email']
        
        # Update other fields
        user.first_name = data.get('first_name', user.first_name)
        user.last_name = data.get('last_name', user.last_name)
        user.middle_name = data.get('middle_name', user.middle_name)
        user.role = data.get('role', user.role)
        
        # Update is_active instead of status
        if 'is_active' in data:
            user.is_active = data['is_active']
        
        # Handle section assignment for teachers
        if user.role == UserRole.TEACHER:
            if 'section' in data and data['section']:
                try:
                    section = Section.objects.get(id=data['section'])
                    user.section = section
                except Section.DoesNotExist:
                    pass  # Silently ignore invalid section IDs
            else:
                # Clear section if not provided or empty
                user.section = None
        else:
            # Non-teachers should not have sections
            user.section = None
        
        user.save()
        
        return JsonResponse({
            'status': 'success',
            'message': 'User updated successfully'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["DELETE"])
def delete_user(request, user_id):
    """API endpoint to delete a user"""
    try:
        user = get_object_or_404(CustomUser, id=user_id)
        
        # Don't allow deleting yourself
        if request.user.id == user.id:
            return JsonResponse({
                'status': 'error', 
                'message': 'You cannot delete your own account'
            }, status=400)
            
        # Store the username for the success message
        username = user.username
        
        # Delete the user
        user.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': f'User {username} deleted successfully'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def reset_user_password(request, user_id):
    """API endpoint to reset a user's password"""
    try:
        user = get_object_or_404(CustomUser, id=user_id)
        data = json.loads(request.body)
        
        # Validate new password
        new_password = data.get('new_password')
        if not new_password or len(new_password) < 8:
            return JsonResponse({
                'status': 'error', 
                'message': 'Password must be at least 8 characters long'
            }, status=400)
            
        # Update password
        user.password = make_password(new_password)
        user.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Password for {user.username} has been reset successfully'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
def search_users(request):
    """API endpoint to search and filter users for AJAX requests"""
    # Get query parameters
    search_query = request.GET.get('search', '')
    role_filter = request.GET.get('role', '')
    status_filter = request.GET.get('status', '')
    page_number = request.GET.get('page', 1)
    items_per_page = request.GET.get('items_per_page', 10)
    sort_by = request.GET.get('sort_by', '-created_at')
    get_all = request.GET.get('all', 'false').lower() == 'true'  # For printing all users
    
    # Base queryset
    users = CustomUser.objects.all().order_by(sort_by)
    
    # Apply search if provided
    if search_query:
        users = users.filter(
            Q(username__icontains=search_query) | 
            Q(email__icontains=search_query) |
            Q(first_name__icontains=search_query) | 
            Q(last_name__icontains=search_query)
        )
    
    # Apply role filter if provided
    if role_filter:
        users = users.filter(role=role_filter)
    
    # Apply status filter
    if status_filter:
        if status_filter == 'active':
            users = users.filter(is_active=True)
        elif status_filter == 'inactive':
            users = users.filter(is_active=False)
    
    # Get total count for pagination
    total_count = users.count()
    
    # If 'all' parameter is true, return all users without pagination (for printing)
    if get_all:
        user_data = []
        for user in users:
            user_data.append({
                'id': user.id,
                'username': user.username,
                'email': user.email,
                'first_name': user.first_name,
                'middle_name': user.middle_name or '',
                'last_name': user.last_name,
                'role': user.role,
                'role_display': user.get_role_display(),
                'is_active': user.is_active,
                'created_at': user.created_at.isoformat() if user.created_at else '',
                'profile_pic': user.profile_pic,
            })
        
        return JsonResponse({
            'status': 'success',
            'users': user_data,
            'total_count': total_count,
        })
    
    # Parse page number and items_per_page
    try:
        page_number = int(page_number)
        items_per_page = int(items_per_page)
    except (ValueError, TypeError):
        page_number = 1
        items_per_page = 10
    
    # Create paginator
    paginator = Paginator(users, items_per_page)
    page_obj = paginator.get_page(page_number)
    
    # Prepare user data for JSON response
    user_data = []
    for user in page_obj:
        user_data.append({
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'role': user.role,
            'role_display': user.get_role_display(),
            'is_active': user.is_active,
            'created_at': user.created_at.strftime('%Y-%m-%d'),
            'created_at_display': user.created_at.strftime('%b %d, %Y'),
            'profile_pic': user.profile_pic,
        })
    
    # Prepare pagination data
    pagination = {
        'current_page': page_obj.number,
        'num_pages': paginator.num_pages,
        'has_next': page_obj.has_next(),
        'has_previous': page_obj.has_previous(),
        'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
        'start_index': page_obj.start_index(),
        'end_index': page_obj.end_index(),
        'total_count': total_count,
    }
    
    return JsonResponse({
        'status': 'success',
        'users': user_data,
        'pagination': pagination,
        'total_count': total_count,
    })


@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def import_users(request):
    """API endpoint to import users from CSV/Excel file - matches export format"""
    try:
        # Accept both legacy 'file' and newer 'import_file' form keys
        if 'file' in request.FILES:
            uploaded_file = request.FILES['file']
        elif 'import_file' in request.FILES:
            uploaded_file = request.FILES['import_file']
        else:
            return JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)

        # Validate file type
        file_extension = uploaded_file.name.split('.')[-1].lower() if uploaded_file.name and '.' in uploaded_file.name else ''
        if file_extension not in ['csv', 'xlsx', 'xls']:
            return JsonResponse({'status': 'error', 'message': 'Invalid file type. Only CSV and Excel files are allowed.'}, status=400)

        # Validate file size (10MB max to match modal description)
        if uploaded_file.size > 10 * 1024 * 1024:
            return JsonResponse({'status': 'error', 'message': 'File too large. Maximum file size is 10MB.'}, status=400)

        # Read file content robustly
        import pandas as pd
        import re
        import re
        import re
        import re
        import io

        file_bytes = uploaded_file.read()
        bio = io.BytesIO(file_bytes)
        df = None
        read_error = None
        try:
            if file_extension == 'csv':
                # Try utf-8 then fallback to latin1
                try:
                    df = pd.read_csv(io.BytesIO(file_bytes))
                except Exception:
                    df = pd.read_csv(io.BytesIO(file_bytes), encoding='latin1')
            else:
                # Prefer openpyxl for xlsx
                try:
                    df = pd.read_excel(bio, engine='openpyxl')
                except Exception:
                    # Retry with default engine/bytes buffer
                    bio.seek(0)
                    df = pd.read_excel(bio)
        except Exception as e:
            read_error = str(e)

        if df is None:
            return JsonResponse({'status': 'error', 'message': f'Error reading file: {read_error or "Could not parse file"}'}, status=400)

        # Normalize column names (strip whitespace)
        df.columns = df.columns.str.strip()
        # Map common header variants case-insensitively
        column_mapping = {
            'username': 'username',
            'user name': 'username',
            'user': 'username',
            'email': 'email',
            'first name': 'first_name',
            'firstname': 'first_name',
            'last name': 'last_name',
            'lastname': 'last_name',
            'middle name': 'middle_name',
            'middlename': 'middle_name',
            'role': 'role',
            'status': 'status',
            'password': 'password',
            'lrn': 'lrn',
            'id': 'id',
            'created date': 'created_date'
        }

        # Build a mapping from existing df columns to canonical names
        new_columns = {}
        for col in df.columns:
            key = str(col).strip().lower()
            if key in column_mapping:
                new_columns[col] = column_mapping[key]
            else:
                # Keep original column name for unknowns
                new_columns[col] = col

        df.rename(columns=new_columns, inplace=True)
        
        # Validate required columns (password is optional for updating existing users)
        required_columns = ['username', 'email', 'first_name', 'last_name', 'role']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            return JsonResponse({
                'status': 'error', 
                'message': f'Missing required columns: {", ".join(missing_columns)}. The file must include at least: Username, Email, First Name, Last Name, Role'
            }, status=400)
        
        # Check if password column exists
        has_password_column = 'password' in df.columns
        
        # Process data
        created_count = 0
        updated_count = 0
        skipped_count = 0
        skipped_rows = []
        errors = []
        skipped_rows = []
        skipped_rows = []
        emails_sent = 0
        
        for index, row in df.iterrows():
            try:
                # Clean and validate data
                username = str(row['username']).strip() if pd.notna(row['username']) else ''
                email = str(row['email']).strip() if pd.notna(row['email']) else ''
                first_name = str(row['first_name']).strip() if pd.notna(row['first_name']) else ''
                last_name = str(row['last_name']).strip() if pd.notna(row['last_name']) else ''
                role = str(row['role']).strip().upper() if pd.notna(row['role']) else ''
                middle_name = str(row['middle_name']).strip() if 'middle_name' in row and pd.notna(row['middle_name']) else ''
                lrn = str(row['lrn']).strip() if 'lrn' in row and pd.notna(row['lrn']) else ''
                
                # Handle Status column (export format: "Active" or "Inactive")
                is_active = True
                if 'status' in row and pd.notna(row['status']):
                    status_value = str(row['status']).strip().upper()
                    is_active = status_value in ['ACTIVE', 'TRUE', '1', 'YES']
                elif 'is_active' in row and pd.notna(row['is_active']):  # Support old format
                    is_active_value = str(row['is_active']).strip().upper()
                    is_active = is_active_value in ['TRUE', '1', 'YES', 'ACTIVE']
                
                # Validate required fields
                if not all([username, email, first_name, last_name, role]):
                    errors.append(f'Row {index + 2}: Missing required field(s)')
                    skipped_count += 1
                    continue
                
                # Validate and normalize role (handle both display names and codes)
                role_mapping = {
                    'ADMIN': UserRole.ADMIN,
                    'TEACHER': UserRole.TEACHER,
                    'PRINCIPAL': UserRole.PRINCIPAL,
                    'REGISTRAR': UserRole.REGISTRAR,
                    'ADMINISTRATOR': UserRole.ADMIN,
                }
                
                role_upper = role.upper()
                if role_upper in role_mapping:
                    role = role_mapping[role_upper]
                else:
                    valid_roles = [choice[0] for choice in UserRole.choices]
                    if role not in valid_roles:
                        errors.append(f'Row {index + 2}: Invalid role "{role}". Valid roles: ADMIN, TEACHER, PRINCIPAL, REGISTRAR')
                        skipped_count += 1
                        continue
                
                # Check if ID is provided in Excel
                user_id = row.get('id') if 'id' in row and pd.notna(row['id']) else None
                
                if user_id:
                    # Skip rows with existing ID - don't import or update
                    try:
                        if CustomUser.objects.filter(id=int(user_id)).exists():
                            skipped_count += 1
                            skipped_rows.append({
                                'row': index + 2,
                                'id': user_id,
                                'username': username,
                                'email': email,
                                'lrn': lrn,
                                'reason': 'Existing ID'
                            })
                            continue
                    except ValueError:
                        pass

                # If LRN column is present and matches an existing student, skip the row
                if lrn:
                    if Student.objects.filter(lrn=lrn).exists():
                        skipped_count += 1
                        skipped_rows.append({
                            'row': index + 2,
                            'username': username,
                            'email': email,
                            'lrn': lrn,
                            'reason': 'LRN already exists'
                        })
                        continue
                
                # Only import new users (blank ID)
                # Check if username or email already exists
                if CustomUser.objects.filter(username=username).exists():
                    skipped_count += 1
                    skipped_rows.append({
                        'row': index + 2,
                        'username': username,
                        'email': email,
                        'lrn': lrn,
                        'reason': 'Username already exists'
                    })
                    continue
                if CustomUser.objects.filter(email=email).exists():
                    skipped_count += 1
                    skipped_rows.append({
                        'row': index + 2,
                        'username': username,
                        'email': email,
                        'lrn': lrn,
                        'reason': 'Email already exists'
                    })
                    continue
                
                # Generate random password for new user
                generated_password = generate_random_password()
                
                # Create new user
                user = CustomUser(
                    username=username,
                    email=email,
                    first_name=first_name,
                    last_name=last_name,
                    middle_name=middle_name,
                    role=role,
                    is_active=is_active
                )
                user.password = make_password(generated_password)
                user.save()
                created_count += 1
                
                # Send email with generated password
                role_display = dict(UserRole.choices).get(role, role)
                if send_password_email(email, username, generated_password, role_display):
                    emails_sent += 1
                
            except Exception as e:
                errors.append(f'Row {index + 2}: {str(e)}')
                skipped_count += 1
                continue
        
        # Prepare response message
        message_parts = []
        if created_count > 0:
            message_parts.append(f'{created_count} user(s) created')
            if emails_sent > 0:
                message_parts.append(f'{emails_sent} password email(s) sent')
        if skipped_count > 0:
            message_parts.append(f'{skipped_count} user(s) skipped (existing IDs or errors)')
        
        message = 'Import completed! ' + ', '.join(message_parts) + '.'
        
        response_data = {
            'status': 'success',
            'message': message,
            'created_count': created_count,
            'skipped_count': skipped_count,
            'emails_sent': emails_sent,
            'total_processed': created_count + skipped_count
        }
        
        # Include errors if any (limit to first 10 errors)
        if errors:
            response_data['errors'] = errors[:10]
            if len(errors) > 10:
                response_data['additional_errors'] = len(errors) - 10

        # Include skipped row details for the UI to display which rows were ignored
        if skipped_rows:
            response_data['skipped_rows'] = skipped_rows
        
        return JsonResponse(response_data)
        
    except Exception as e:
        import traceback
        print("Error in import_users:", str(e))
        print(traceback.format_exc())
        return JsonResponse({
            'status': 'error',
            'message': f'An unexpected error occurred: {str(e)}'
        }, status=500)


@login_required
@user_passes_test(is_admin)
@require_http_methods(["GET"])
def download_import_template(request):
    """Generate and download a properly formatted Excel template for user import"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import io
    
    # Create workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Users Import Template"
    
    # Define styles (matching export format)
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    border_style = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    center_alignment = Alignment(horizontal="center", vertical="center")
    left_alignment = Alignment(horizontal="left", vertical="center")
    
    # Set column widths
    column_widths = [8, 15, 30, 15, 15, 15, 15, 12]
    for idx, width in enumerate(column_widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    
    # Headers (with ID column)
    headers = ["ID", "Username", "Email", "First Name", "Last Name", "Middle Name", "Role", "Status"]
    
    # Add headers with styling
    ws.append(headers)
    ws.row_dimensions[1].height = 25
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = border_style
    
    # Add instruction row
    instruction_row = ["Instructions:", "Leave ID blank for new users (auto-generates password & sends email). Rows with existing IDs will be SKIPPED (not imported).", "", "", "", "", "", ""]
    ws.append(instruction_row)
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
    instruction_cell = ws.cell(row=2, column=1)
    instruction_cell.font = Font(italic=True, color="666666", size=10)
    instruction_cell.alignment = left_alignment
    ws.row_dimensions[2].height = 30
    
    # Sample data rows (ID, Username, Email, First Name, Last Name, Middle Name, Role, Status)
    sample_data = [
        ["", "john.teacher", "john.teacher@protech.com", "John", "Teacher", "M", "Teacher", "Active"],
        ["", "jane.admin", "jane.admin@protech.com", "Jane", "Admin", "S", "Administrator", "Active"],
        ["", "robert.principal", "robert.principal@protech.com", "Robert", "Principal", "K", "Principal", "Active"],
        ["", "mary.registrar", "mary.registrar@protech.com", "Mary", "Registrar", "L", "Registrar", "Active"],
    ]
    
    # Add data rows with styling
    for row_data in sample_data:
        ws.append(row_data)
    
    # Apply formatting to data rows
    for row_idx in range(3, ws.max_row + 1):
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.border = border_style
            
            # Apply alignment (ID, Role, Status - center, rest left)
            if col_idx in [1, 7, 8]:
                cell.alignment = center_alignment
            else:
                cell.alignment = left_alignment
    
    # Freeze header row
    ws.freeze_panes = ws['A2']
    
    # Save to buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    # Create HTTP response
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="users_import_template_{datetime.now().strftime("%Y%m%d")}.xlsx"'
    
    return response


# ==========================
#  TEACHERS IMPORT & TEMPLATE
# ==========================

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def import_teachers(request):
    """Import teachers from CSV or Excel file with [UNCHANGED] password support"""
    try:
        import pandas as pd
        from django.contrib.auth.hashers import make_password
        
        # Accept both 'import_file' and legacy 'file' keys for robustness
        if 'import_file' not in request.FILES and 'file' not in request.FILES:
            # Helpful debug: report what keys were received (when running locally)
            received_keys = list(request.FILES.keys())
            print(f"import_sections: no file found in request.FILES, keys present: {received_keys}")
            return JsonResponse({'status': 'error', 'message': 'No file uploaded', 'received_files': received_keys}, status=400)

        file = request.FILES.get('import_file') or request.FILES.get('file')
        
        # Validate file type
        allowed_extensions = ['.csv', '.xlsx', '.xls']
        file_ext = os.path.splitext(file.name)[1].lower()
        if file_ext not in allowed_extensions:
            return JsonResponse({
                'status': 'error',
                'message': f'Invalid file type. Allowed: {", ".join(allowed_extensions)}'
            }, status=400)
        
        # Read file
        try:
            if file_ext == '.csv':
                df = pd.read_csv(file)
            else:
                df = pd.read_excel(file)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': f'Error reading file: {str(e)}'}, status=400)
        
        # Normalize column names
        df.columns = df.columns.str.strip()
        column_mapping = {
            'Username': 'username',
            'Email': 'email',
            'First Name': 'first_name',
            'Last Name': 'last_name',
            'Middle Name': 'middle_name',
            'Password': 'password',
            'Active Status': 'is_active',
            'Status': 'is_active',
            'ID': 'id',
            'Created Date': 'created_date'
        }
        df.rename(columns=column_mapping, inplace=True)
        
        # Validate required columns
        required_columns = ['username', 'email', 'first_name', 'last_name']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            return JsonResponse({
                'status': 'error',
                'message': f'Missing required columns: {", ".join(missing_columns)}'
            }, status=400)
        
        created_count = 0
        updated_count = 0
        skipped_count = 0
        errors = []
        emails_sent = 0
        
        for index, row in df.iterrows():
            try:
                username = str(row['username']).strip() if pd.notna(row['username']) else ''
                email = str(row['email']).strip() if pd.notna(row['email']) else ''
                first_name = str(row['first_name']).strip() if pd.notna(row['first_name']) else ''
                last_name = str(row['last_name']).strip() if pd.notna(row['last_name']) else ''
                middle_name = str(row['middle_name']).strip() if 'middle_name' in row and pd.notna(row['middle_name']) else ''
                
                is_active = True
                if 'is_active' in row and pd.notna(row['is_active']):
                    status_value = str(row['is_active']).strip().upper()
                    is_active = status_value in ['ACTIVE', 'TRUE', '1', 'YES']
                
                if not all([username, email, first_name, last_name]):
                    errors.append(f'Row {index + 2}: Missing required field(s)')
                    skipped_count += 1
                    continue
                
                # Check if ID is provided in Excel
                teacher_id = row.get('id') if 'id' in row and pd.notna(row['id']) else None
                
                if teacher_id:
                    # Skip rows with existing ID - don't import or update
                    try:
                        if CustomUser.objects.filter(id=int(teacher_id), role=UserRole.TEACHER).exists():
                            skipped_count += 1
                            continue
                    except ValueError:
                        pass
                
                # Only import new teachers (blank ID)
                # Check if username or email already exists
                if CustomUser.objects.filter(username=username).exists():
                    errors.append(f'Row {index + 2}: Username "{username}" already exists')
                    skipped_count += 1
                    continue
                if CustomUser.objects.filter(email=email).exists():
                    errors.append(f'Row {index + 2}: Email "{email}" already exists')
                    skipped_count += 1
                    continue
                
                # Generate random password for new teacher
                generated_password = generate_random_password()
                
                # Create new teacher
                teacher = CustomUser(
                    username=username,
                    email=email,
                    first_name=first_name,
                    last_name=last_name,
                    middle_name=middle_name,
                    role=UserRole.TEACHER,
                    is_active=is_active
                )
                teacher.password = make_password(generated_password)
                teacher.save()
                created_count += 1
                
                # Send email with generated password
                if send_password_email(email, username, generated_password, "Teacher"):
                    emails_sent += 1
                
            except Exception as e:
                errors.append(f'Row {index + 2}: {str(e)}')
                skipped_count += 1
                continue
        
        message_parts = []
        if created_count > 0:
            message_parts.append(f'{created_count} teacher(s) created')
            if emails_sent > 0:
                message_parts.append(f'{emails_sent} password email(s) sent')
        if skipped_count > 0:
            message_parts.append(f'{skipped_count} row(s) skipped (existing IDs or errors)')
        
        return JsonResponse({
            'status': 'success',
            'message': ', '.join(message_parts) if message_parts else 'No changes made',
            'created': created_count,
            'skipped': skipped_count,
            'emails_sent': emails_sent,
            'errors': errors[:10]
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'An unexpected error occurred: {str(e)}'}, status=500)


@login_required
@user_passes_test(is_admin)
@require_http_methods(["GET"])
def download_teachers_template(request):
    """Generate and download Excel template for teacher import"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import io
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Teachers Import Template"
    
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    border_style = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    center_alignment = Alignment(horizontal="center", vertical="center")
    left_alignment = Alignment(horizontal="left", vertical="center")
    
    column_widths = [8, 15, 30, 15, 15, 15, 12]
    for idx, width in enumerate(column_widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    
    headers = ["ID", "Username", "Email", "First Name", "Last Name", "Middle Name", "Status"]
    
    ws.append(headers)
    ws.row_dimensions[1].height = 25
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = border_style
    
    # Add instruction row
    instruction_row = ["Instructions:", "Leave ID blank for new teachers (auto-generates password & sends email). Rows with existing IDs will be SKIPPED (not imported).", "", "", "", "", ""]
    ws.append(instruction_row)
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
    instruction_cell = ws.cell(row=2, column=1)
    instruction_cell.font = Font(italic=True, color="666666", size=10)
    instruction_cell.alignment = left_alignment
    ws.row_dimensions[2].height = 30
    
    sample_data = [
        ["", "john.teacher", "john.teacher@protech.com", "John", "Doe", "M", "Active"],
        ["", "jane.teacher", "jane.teacher@protech.com", "Jane", "Smith", "S", "Active"],
    ]
    
    for row_data in sample_data:
        ws.append(row_data)
    
    for row_idx in range(3, ws.max_row + 1):
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.border = border_style
            if col_idx in [1, 7]:
                cell.alignment = center_alignment
            else:
                cell.alignment = left_alignment
    
    ws.freeze_panes = ws['A2']
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="teachers_import_template_{datetime.now().strftime("%Y%m%d")}.xlsx"'
    
    return response


# ==========================
#  TEACHER MANAGEMENT VIEWS
# ==========================

@login_required
@user_passes_test(is_admin)
def admin_teachers(request):
    """View for teacher management with filtering and section assignment"""
    # Get query parameters
    search_query = request.GET.get('search', '')
    advisory_filter = request.GET.get('advisory', '')
    status_filter = request.GET.get('status', '')
    section_filter = request.GET.get('section', '')
    page_number = request.GET.get('page', 1)

    # Base queryset - only teachers, order by created_at DESC (newest first)
    teachers = CustomUser.objects.filter(role=UserRole.TEACHER).order_by('-created_at')

    # Apply search if provided
    if search_query:
        teachers = teachers.filter(
            Q(username__icontains=search_query) | 
            Q(email__icontains=search_query) |
            Q(first_name__icontains=search_query) | 
            Q(last_name__icontains=search_query)
        )

    # Apply status filter using is_active
    if status_filter:
        if status_filter == 'active':
            teachers = teachers.filter(is_active=True)
        elif status_filter == 'inactive':
            teachers = teachers.filter(is_active=False)

    # Get all advisory assignments from database
    advisory_assignments = AdvisoryAssignment.objects.select_related('section', 'section__grade')

    # Create a dictionary of teacher_id -> section for quick lookup
    teacher_sections = {}
    for assignment in advisory_assignments:
        teacher_sections[assignment.teacher_id] = {
            'section': assignment.section,
            'section_id': assignment.section_id,
            'section_name': assignment.section.name,
            'grade_id': assignment.section.grade.id,
            'grade_name': assignment.section.grade.name
        }

    # Get all available sections for assignment dropdown
    all_sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    # Filter by advisory status - check if teacher has advisory assignment
    if advisory_filter == 'advisory':
        teachers = teachers.filter(id__in=list(teacher_sections.keys()))
    elif advisory_filter == 'non-advisory':
        teachers = teachers.exclude(id__in=list(teacher_sections.keys()))

    # Filter by section
    if section_filter:
        try:
            section_id = int(section_filter)
            teacher_ids = [teacher_id for teacher_id, data in teacher_sections.items() 
                           if data['section_id'] == section_id]
            teachers = teachers.filter(id__in=teacher_ids)
        except Exception:
            pass

    # Get counts for dashboard
    total_teachers = CustomUser.objects.filter(role=UserRole.TEACHER).count()
    advisory_teacher_ids = AdvisoryAssignment.objects.values_list('teacher_id', flat=True).distinct()
    advisory_teachers_count = len(advisory_teacher_ids)
    non_advisory_teachers_count = total_teachers - advisory_teachers_count
    active_teachers_count = CustomUser.objects.filter(role=UserRole.TEACHER, is_active=True).count()

    # Pagination
    paginator = Paginator(teachers, 10)
    page_obj = paginator.get_page(page_number)

    # Calculate page range for pagination UI
    page_range = get_pagination_range(paginator, page_obj.number, 5)

    # Annotate each teacher in the page with section/grade info from AdvisoryAssignment
    for teacher in page_obj:
        section_info = teacher_sections.get(teacher.id)
        # is_advisory is now a property that auto-calculates from teacher.section
        if section_info:
            teacher.section_name = section_info['section_name']
            teacher.grade_name = section_info['grade_name']
            teacher.section_id = section_info['section_id']
            teacher.grade_id = section_info['grade_id']
        else:
            teacher.section_name = None
            teacher.grade_name = None
            teacher.section_id = None
            teacher.grade_id = None
        # Add status for frontend
        teacher.status = 'Active' if teacher.is_active else 'Inactive'

    context = {
        'teachers': page_obj,
        'teacher_sections': teacher_sections,
        'sections': all_sections,
        'grades': Grade.objects.all().order_by('name'), 
        'status_choices': [('active', 'Active'), ('inactive', 'Inactive')],
        'search_query': search_query,
        'advisory_filter': advisory_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'total_teachers': total_teachers,
        'advisory_teachers_count': advisory_teachers_count,
        'non_advisory_teachers_count': non_advisory_teachers_count,
        'active_teachers_count': active_teachers_count,
        'page_range': page_range
    }
    return render(request, 'admin/teachers.html', context)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def assign_teacher_section(request):
    """API endpoint to assign a teacher to a section as advisor"""
    try:
        data = json.loads(request.body)
        
        teacher_id = data.get('teacher_id')
        section_id = data.get('section_id')
        
        if not teacher_id or not section_id:
            return JsonResponse({'status': 'error', 'message': 'Teacher ID and Section ID are required'}, status=400)
        
        # Validate teacher is a teacher
        teacher = get_object_or_404(CustomUser, id=teacher_id, role=UserRole.TEACHER)
        section = get_object_or_404(Section, id=section_id)
        
        # Check if section already has an advisor
        existing_assignment = AdvisoryAssignment.objects.filter(section=section).first()
        if existing_assignment:
            return JsonResponse({
                'status': 'error', 
                'message': f'Section already has an advisor: {existing_assignment.teacher.first_name} {existing_assignment.teacher.last_name}'
            }, status=400)
        
        # Check if teacher is already assigned to another section
        existing_teacher_assignment = AdvisoryAssignment.objects.filter(teacher=teacher).first()
        if existing_teacher_assignment:
            return JsonResponse({
                'status': 'error', 
                'message': f'Teacher is already assigned to {existing_teacher_assignment.section.name}'
            }, status=400)
            
        # Create assignment
        assignment = AdvisoryAssignment.objects.create(
            teacher=teacher,
            section=section
        )
        
        # Also update teacher's section field
        teacher.section = section
        teacher.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Teacher {teacher.first_name} {teacher.last_name} assigned as advisor to {section.name}',
            'assignment_id': assignment.id
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def remove_teacher_section(request):
    """API endpoint to remove a teacher from advisory role"""
    try:
        data = json.loads(request.body)
        
        teacher_id = data.get('teacher_id')
        
        if not teacher_id:
            return JsonResponse({'status': 'error', 'message': 'Teacher ID is required'}, status=400)
        
        # Validate teacher is a teacher
        teacher = get_object_or_404(CustomUser, id=teacher_id, role=UserRole.TEACHER)
        
        # Find and delete assignment
        assignment = AdvisoryAssignment.objects.filter(teacher=teacher).first()
        if assignment:
            section_name = assignment.section.name
            assignment.delete()
            
            # Also update teacher's section field
            teacher.section = None
            teacher.save()
            
            return JsonResponse({
                'status': 'success',
                'message': f'Teacher {teacher.first_name} {teacher.last_name} removed as advisor from {section_name}'
            })
        else:
            return JsonResponse({
                'status': 'error',
                'message': 'Teacher is not an advisor to any section'
            }, status=400)
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@login_required
@require_GET
def search_teachers(request):
    """API endpoint to search/filter/paginate teachers for AJAX requests"""
    try:
        # Check if user is admin or principal
        if not (request.user.role == UserRole.ADMIN or request.user.role == UserRole.PRINCIPAL):
            return JsonResponse({'status': 'error', 'message': 'Unauthorized access'}, status=403)
        
        search_query = request.GET.get('search', '')
        advisory_filter = request.GET.get('advisory', '')
        status_filter = request.GET.get('status', '')
        section_filter = request.GET.get('section', '')
        page_number = request.GET.get('page', 1)
        items_per_page = request.GET.get('items_per_page', 10)

        teachers = CustomUser.objects.filter(role=UserRole.TEACHER).order_by('-created_at')

        if search_query:
            teachers = teachers.filter(
                Q(username__icontains=search_query) |
                Q(email__icontains=search_query) |
                Q(first_name__icontains=search_query) |
                Q(last_name__icontains=search_query)
            )

        # Status filter
        if status_filter:
            if status_filter == 'active':
                teachers = teachers.filter(is_active=True)
            elif status_filter == 'inactive':
                teachers = teachers.filter(is_active=False)

        advisory_assignments = AdvisoryAssignment.objects.select_related('section', 'section__grade')
        teacher_sections = {}
        for assignment in advisory_assignments:
            teacher_sections[assignment.teacher_id] = {
                'section': assignment.section,
                'section_id': assignment.section_id,
                'section_name': assignment.section.name,
                'grade_id': assignment.section.grade.id,
                'grade_name': assignment.section.grade.name
            }

        # Advisory filter - check if teacher has advisory assignment
        if advisory_filter == 'advisory':
            teachers = teachers.filter(id__in=list(teacher_sections.keys()))
        elif advisory_filter == 'non-advisory':
            teachers = teachers.exclude(id__in=list(teacher_sections.keys()))

        if section_filter:
            try:
                section_id = int(section_filter)
                teacher_ids = [tid for tid, data in teacher_sections.items() if data['section_id'] == section_id]
                teachers = teachers.filter(id__in=teacher_ids)
            except Exception:
                pass

        paginator = Paginator(teachers, items_per_page)
        page_obj = paginator.get_page(page_number)

        teacher_data = []
        for teacher in page_obj:
            section_info = teacher_sections.get(teacher.id)
            grade = None
            section = None
            grade_id = None
            section_id = None
            is_advisory = bool(section_info)
            if section_info:
                grade = section_info['grade_name']
                section = section_info['section_name']
                grade_id = section_info['grade_id']
                section_id = section_info['section_id']
            teacher_data.append({
                'id': teacher.id,
                'username': teacher.username,
                'email': teacher.email,
                'first_name': teacher.first_name,
                'last_name': teacher.last_name,
                'profile_pic': teacher.profile_pic,
                'created_at': teacher.created_at.strftime('%Y-%m-%d'),
                'created_at_display': teacher.created_at.strftime('%b %d, %Y'),
                'advisory': is_advisory,
                'section': section,
                'section_id': section_id,
                'grade': grade,
                'grade_id': grade_id,
                'is_active': teacher.is_active,  # Include status for frontend
                'status': 'Active' if teacher.is_active else 'Inactive'
            })

        pagination = {
            'current_page': page_obj.number,
            'num_pages': paginator.num_pages,
            'has_next': page_obj.has_next(),
            'has_previous': page_obj.has_previous(),
            'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
            'start_index': page_obj.start_index(),
            'end_index': page_obj.end_index(),
            'total_count': paginator.count,
        }

        return JsonResponse({
            'status': 'success',
            'teachers': teacher_data,
            'pagination': pagination,
            'total_count': paginator.count,
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e),
            'teachers': [],
            'pagination': {
                'current_page': 1,
                'num_pages': 1,
                'has_next': False,
                'has_previous': False,
                'page_range': [1],
                'start_index': 0,
                'end_index': 0,
                'total_count': 0,
            },
            'total_count': 0
        }, status=500)


# ==========================
#  GRADE MANAGEMENT VIEWS
# ==========================

@login_required
@user_passes_test(is_admin)
def admin_grades(request):
    """View for grade management"""
    # Get query parameters
    search_query = request.GET.get('search', '')
    page_number = request.GET.get('page', 1)
    
    # Base queryset - order by created_at DESC (newest first)
    grades = Grade.objects.all().order_by('-created_at')
    
    # Apply search if provided
    if search_query:
        grades = grades.filter(name__icontains=search_query)
    
    # Annotate with section and student counts
    grades = grades.annotate(
        sections_count=Count('sections'),
        students_count=Count('sections__students', distinct=True)
    )
    
    # Get stats for dashboard
    total_grades = Grade.objects.count()
    grades_with_sections = Grade.objects.filter(sections__isnull=False).distinct().count()
    total_sections = Section.objects.count()
    total_students = Student.objects.count()
    
    # Pagination
    paginator = Paginator(grades, 10)
    page_obj = paginator.get_page(page_number)
    
    # Calculate page range for pagination UI
    page_range = get_pagination_range(paginator, page_obj.number, 5)
    
    context = {
        'grades': page_obj,
        'search_query': search_query,
        'total_grades': total_grades,
        'grades_with_sections': grades_with_sections,
        'total_sections': total_sections,
        'total_students': total_students,
        'page_range': page_range,
    }
    return render(request, 'admin/grades.html', context)

@login_required
@user_passes_test(is_admin_or_principal)
@require_GET
def search_grades(request):
    """AJAX search/filter for grades"""
    search_query = request.GET.get('search', '')
    page_number = request.GET.get('page', 1)
    items_per_page = request.GET.get('items_per_page', 10)
    
    # Base queryset - order by created_at DESC (newest first)
    grades = Grade.objects.all().order_by('-created_at')
    
    # Apply search if provided
    if search_query:
        grades = grades.filter(name__icontains=search_query)
    
    # Annotate with counts (fix: use distinct for sections_count)
    grades = grades.annotate(
        sections_count=Count('sections', distinct=True),
        students_count=Count('sections__students', distinct=True)
    )
    
    # Get total count for pagination
    total_count = grades.count()
    
    # Parse page number and items_per_page
    try:
        page_number = int(page_number)
        items_per_page = int(items_per_page)
    except (ValueError, TypeError):
        page_number = 1
        items_per_page = 10
    
    # Create paginator
    paginator = Paginator(grades, items_per_page)
    page_obj = paginator.get_page(page_number)
    
    # Prepare grade data for JSON response
    grade_data = []
    for grade in page_obj:
        grade_data.append({
            'id': grade.id,
            'name': grade.name,
            'sections_count': grade.sections_count,
            'students_count': grade.students_count,
            'created_at': grade.created_at.strftime('%Y-%m-%d'),
            'created_at_display': grade.created_at.strftime('%b %d, %Y'),
        })
    
    # Prepare pagination data
    pagination = {
        'current_page': page_obj.number,
        'num_pages': paginator.num_pages,
        'has_next': page_obj.has_next(),
        'has_previous': page_obj.has_previous(),
        'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
        'start_index': page_obj.start_index(),
        'end_index': page_obj.end_index(),
        'total_count': total_count,
    }
    
    return JsonResponse({
        'status': 'success',
        'grades': grade_data,
        'pagination': pagination,
        'total_count': total_count,
    })

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def create_grade(request):
    """Create a new grade"""
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        
        if not name:
            return JsonResponse({'status': 'error', 'message': 'Grade name is required.'}, status=400)
        
        if Grade.objects.filter(name__iexact=name).exists():
            return JsonResponse({'status': 'error', 'message': 'Grade with this name already exists.'}, status=400)
        
        grade = Grade.objects.create(name=name)
        
        return JsonResponse({
            'status': 'success',
            'message': f'Grade "{name}" created successfully.',
            'grade': {
                'id': grade.id,
                'name': grade.name,
                'sections_count': 0,
                'students_count': 0,
                'created_at': grade.created_at.strftime('%Y-%m-%d'),
                'created_at_display': grade.created_at.strftime('%b %d, %Y'),
            }
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["PATCH"])
def update_grade(request, grade_id):
    """Update an existing grade"""
    try:
        grade = get_object_or_404(Grade, id=grade_id)
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        
        if not name:
            return JsonResponse({'status': 'error', 'message': 'Grade name is required.'}, status=400)
        
        if Grade.objects.filter(name__iexact=name).exclude(id=grade_id).exists():
            return JsonResponse({'status': 'error', 'message': 'Another grade with this name already exists.'}, status=400)
        
        old_name = grade.name
        grade.name = name
        grade.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Grade "{old_name}" updated to "{name}" successfully.'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["DELETE"])
def delete_grade(request, grade_id):
    """Delete a grade"""
    try:
        grade = get_object_or_404(Grade, id=grade_id)
        
        # Check if grade has sections
        sections_count = grade.sections.count()
        students_count = Student.objects.filter(grade=grade).count()
        
        if sections_count > 0 or students_count > 0:
            return JsonResponse({
                'status': 'error', 
                'message': f'Cannot delete grade "{grade.name}" because it has {sections_count} section(s) and {students_count} student(s). Please remove all sections and students first.'
            }, status=400)
        
        grade_name = grade.name
        grade.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Grade "{grade_name}" deleted successfully.'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_GET
def get_grade_sections(request, grade_id):
    """Get sections for a specific grade"""
    try:
        grade = get_object_or_404(Grade, id=grade_id)
        sections = Section.objects.filter(grade=grade).annotate(
            students_count=Count('students')
        ).order_by('name')
        
        sections_data = []
        for section in sections:
            # Get the advisor for this section
            advisor = CustomUser.objects.filter(section=section, role=UserRole.TEACHER).first()
            advisor_name = f"{advisor.first_name} {advisor.last_name}" if advisor else "No advisor"
            
            sections_data.append({
                'id': section.id,
                'name': section.name,
                'students_count': section.students_count,
                'advisor': advisor_name,
                'created_at': section.created_at.strftime('%b %d, %Y'),
            })
        
        return JsonResponse({
            'status': 'success',
            'grade': {
                'id': grade.id,
                'name': grade.name
            },
            'sections': sections_data,
            'total_sections': len(sections_data)
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


# ==========================
#  SECTION MANAGEMENT VIEWS
# ==========================

@login_required
@user_passes_test(is_admin)
def admin_sections(request):
    """View for section management"""
    # Get query parameters
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    advisor_filter = request.GET.get('advisor', '')
    page_number = request.GET.get('page', 1)
    
    # Base queryset - order by created_at DESC (newest first)
    sections = Section.objects.select_related('grade').order_by('-created_at')
    
    # Apply search if provided
    if search_query:
        sections = sections.filter(
            Q(name__icontains=search_query) | 
            Q(grade__name__icontains=search_query)
        )
    
    # Apply grade filter if provided
    if grade_filter:
        try:
            grade_id = int(grade_filter)
            sections = sections.filter(grade_id=grade_id)
        except (ValueError, TypeError):
            pass
    
    # Apply advisor filter if provided
    if advisor_filter:
        if advisor_filter == 'with_advisor':
            # Filter sections that have advisory assignments
            sections = sections.filter(advisory_assignments__isnull=False).distinct()
        elif advisor_filter == 'without_advisor':
            # Filter sections that have no advisory assignments
            sections = sections.filter(advisory_assignments__isnull=True)
    
    # Annotate with student counts and prefetch advisory assignments
    sections = sections.annotate(
        students_count=Count('students', distinct=True)
    ).prefetch_related('advisory_assignments__teacher')
    
    # Get stats for dashboard
    total_sections = Section.objects.count()
    sections_with_students = Section.objects.filter(students__isnull=False).distinct().count()
    # Count sections that have advisory assignments
    sections_with_advisors = Section.objects.filter(advisory_assignments__isnull=False).distinct().count()
    total_students = Student.objects.count()
    
    # Get all grades for filter dropdown
    all_grades = Grade.objects.all().order_by('name')
    
    # Pagination
    paginator = Paginator(sections, 10)
    page_obj = paginator.get_page(page_number)
    
    # Calculate page range for pagination UI
    page_range = get_pagination_range(paginator, page_obj.number, 5)
    
    context = {
        'sections': page_obj,
        'grades': all_grades,
        'search_query': search_query,
        'grade_filter': grade_filter,
        'advisor_filter': advisor_filter,
        'total_sections': total_sections,
        'sections_with_students': sections_with_students,
        'sections_with_advisors': sections_with_advisors,
        'total_students': total_students,
        'page_range': page_range,
    }
    return render(request, 'admin/sections.html', context)

@login_required
@user_passes_test(is_admin_or_principal)
@require_GET
def search_sections(request):
    """AJAX search/filter for sections"""
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    advisor_filter = request.GET.get('advisor', '')
    page_number = request.GET.get('page', 1)
    items_per_page = request.GET.get('items_per_page', 10)
    
    # Base queryset - order by created_at DESC (newest first)
    sections = Section.objects.select_related('grade').order_by('-created_at')
    
    # Apply search if provided
    if search_query:
        sections = sections.filter(
            Q(name__icontains=search_query) | 
            Q(grade__name__icontains=search_query)
        )
    
    # Apply grade filter if provided
    if grade_filter:
        try:
            grade_id = int(grade_filter)
            sections = sections.filter(grade_id=grade_id)
        except (ValueError, TypeError):
            pass
    
    # Apply advisor filter if provided
    if advisor_filter:
        if advisor_filter == 'with_advisor':
            # Filter sections that have advisory assignments
            sections = sections.filter(advisory_assignments__isnull=False).distinct()
        elif advisor_filter == 'without_advisor':
            # Filter sections that have no advisory assignments
            sections = sections.filter(advisory_assignments__isnull=True)
    
    # Annotate with counts and prefetch related data
    sections = sections.annotate(
        students_count=Count('students', distinct=True)
    ).prefetch_related('advisory_assignments__teacher')
    
    # Get total count for pagination
    total_count = sections.count()
    
    # Parse page number and items_per_page
    try:
        page_number = int(page_number)
        items_per_page = int(items_per_page)
    except (ValueError, TypeError):
        page_number = 1
        items_per_page = 10
    
    # Create paginator
    paginator = Paginator(sections, items_per_page)
    page_obj = paginator.get_page(page_number)

    # Prepare section data for JSON response
    section_data = []
    for section in page_obj:
        # Get advisor info from AdvisoryAssignment table
        advisor = None
        advisory_assignment = AdvisoryAssignment.objects.filter(
            section_id=section.id
        ).select_related('teacher').first()
        
        if advisory_assignment:
            advisor = {
                'id': advisory_assignment.teacher.id,
                'name': f"{advisory_assignment.teacher.first_name} {advisory_assignment.teacher.last_name}",
                'username': advisory_assignment.teacher.username
            }
        
        section_data.append({
            'id': section.id,
            'name': section.name,
            'grade_id': section.grade.id,
            'grade_name': section.grade.name,
            'students_count': section.students_count,
            'advisor': advisor,
            'created_at': section.created_at.strftime('%Y-%m-%d'),
            'created_at_display': section.created_at.strftime('%b %d, %Y'),
        })
    
    # Prepare pagination data
    pagination = {
        'current_page': page_obj.number,
        'num_pages': paginator.num_pages,
        'has_next': page_obj.has_next(),
        'has_previous': page_obj.has_previous(),
        'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
        'start_index': page_obj.start_index(),
        'end_index': page_obj.end_index(),
        'total_count': total_count,
    }
    
    return JsonResponse({
        'status': 'success',
        'sections': section_data,
        'pagination': pagination,
        'total_count': total_count,
    })

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def create_section(request):
    """Create a new section"""
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        grade_id = data.get('grade_id')
        
        if not name:
            return JsonResponse({'status': 'error', 'message': 'Section name is required.'}, status=400)
        
        if not grade_id:
            return JsonResponse({'status': 'error', 'message': 'Grade is required.'}, status=400)
        
        # Check if grade exists
        try:
            grade = Grade.objects.get(id=grade_id)
        except Grade.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Selected grade does not exist.'}, status=400)
        
        # Check if section with same name exists in the same grade
        if Section.objects.filter(name__iexact=name, grade=grade).exists():
            return JsonResponse({'status': 'error', 'message': f'Section "{name}" already exists in {grade.name}.'}, status=400)
        
        section = Section.objects.create(name=name, grade=grade)
        
        return JsonResponse({
            'status': 'success',
            'message': f'Section "{name}" created successfully in {grade.name}.',
            'section': {
                'id': section.id,
                'name': section.name,
                'grade_id': section.grade.id,
                'grade_name': section.grade.name,
                'students_count': 0,
                'advisor': None,
                'created_at': section.created_at.strftime('%Y-%m-%d'),
                'created_at_display': section.created_at.strftime('%b %d, %Y'),
            }
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["PATCH"])
def update_section(request, section_id):
    """Update an existing section"""
    try:
        section = get_object_or_404(Section, id=section_id)
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        grade_id = data.get('grade_id')
        
        if not name:
            return JsonResponse({'status': 'error', 'message': 'Section name is required.'}, status=400)
        
        if not grade_id:
            return JsonResponse({'status': 'error', 'message': 'Grade is required.'}, status=400)
        
        # Check if grade exists
        try:
            grade = Grade.objects.get(id=grade_id)
        except Grade.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Selected grade does not exist.'}, status=400)
        
        # Check if section with same name exists in the same grade (excluding current section)
        if Section.objects.filter(name__iexact=name, grade=grade).exclude(id=section_id).exists():
            return JsonResponse({'status': 'error', 'message': f'Section "{name}" already exists in {grade.name}.'}, status=400)
        
        old_name = section.name
        old_grade = section.grade.name
        section.name = name
        section.grade = grade
        section.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Section updated from "{old_name}" ({old_grade}) to "{name}" ({grade.name}) successfully.'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["DELETE"])
def delete_section(request, section_id):
    """Delete a section"""
    try:
        section = get_object_or_404(Section, id=section_id)
        
        # Check if section has students
        students_count = section.students.count()
        
        if students_count > 0:
            return JsonResponse({
                'status': 'error', 
                'message': f'Cannot delete section "{section.name}" because it has {students_count} student(s). Please move or remove all students first.'
            }, status=400)
        
        section_name = section.name
        grade_name = section.grade.name
        section.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Section "{section_name}" from {grade_name} deleted successfully.'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin_or_principal)
@require_GET
def get_section_students(request, section_id):
    """Get students for a specific section"""
    try:
        section = get_object_or_404(Section, id=section_id)
        students = Student.objects.filter(section=section).order_by('last_name', 'first_name')
        
        students_data = []
        for student in students:
            students_data.append({
                'id': student.id,
                'lrn': student.lrn,
                'first_name': student.first_name,
                'last_name': student.last_name,
                'full_name': f"{student.first_name} {student.last_name}",
                'status': student.status,
                'created_at': student.created_at.strftime('%b %d, %Y'),
            })
        
        return JsonResponse({
            'status': 'success',
            'section': {
                'id': section.id,
                'name': section.name,
                'grade_name': section.grade.name
            },
            'students': students_data,
            'total_students': len(students_data)
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


# ==========================
#  STUDENTS MANAGEMENT VIEWS
# ==========================

@login_required
@user_passes_test(is_admin)
def admin_students(request):
    """View for student management"""
    # Get query parameters
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    page_number = request.GET.get('page', 1)

    # Base queryset
    students = Student.objects.select_related('grade', 'section').all().order_by('-created_at')

    # Apply search
    if search_query:
        students = students.filter(
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(lrn__icontains=search_query) |
            Q(email__icontains=search_query)
        )

    # Apply grade filter
    if grade_filter:
        students = students.filter(grade_id=grade_filter)

    # Apply section filter
    if section_filter:
        students = students.filter(section_id=section_filter)

    # Apply status filter
    if status_filter:
        students = students.filter(status=status_filter)

    # Dashboard stats
    total_students = Student.objects.count()
    active_students = Student.objects.filter(status='ACTIVE').count()
    inactive_students = Student.objects.filter(status='INACTIVE').count()
    face_enrolled_count = Student.objects.exclude(face_path__isnull=True).exclude(face_path__exact='').count()

    # For filter dropdowns
    grades = Grade.objects.all().order_by('name')
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    # Pagination
    paginator = Paginator(students, 10)
    page_obj = paginator.get_page(page_number)
    page_range = get_pagination_range(paginator, page_obj.number, 5)

    context = {
        'students': page_obj,
        'grades': grades,
        'sections': sections,
        'search_query': search_query,
        'grade_filter': grade_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'total_students': total_students,
        'active_students': active_students,
        'inactive_students': inactive_students,
        'face_enrolled_count': face_enrolled_count,
        'page_range': page_range,
    }
    return render(request, 'admin/students.html', context)

@login_required
@user_passes_test(is_admin_or_principal)
@require_http_methods(["GET"])
def search_students(request):
    """AJAX search/filter for students"""
    try:
        search_query = request.GET.get('search', '')
        grade_filter = request.GET.get('grade', '')
        section_filter = request.GET.get('section', '')
        status_filter = request.GET.get('status', '')
        page_number = request.GET.get('page', 1)
        items_per_page = request.GET.get('items_per_page', 10)

        students = Student.objects.select_related('grade', 'section').all().order_by('-created_at')

        if search_query:
            students = students.filter(
                Q(first_name__icontains=search_query) |
                Q(last_name__icontains=search_query) |
                Q(lrn__icontains=search_query)
            )
        if grade_filter:
            students = students.filter(grade_id=grade_filter)
        if section_filter:
            students = students.filter(section_id=section_filter)
        if status_filter:
            students = students.filter(status=status_filter)

        total_count = students.count()
        try:
            page_number = int(page_number)
            items_per_page = int(items_per_page)
        except (ValueError, TypeError):
            page_number = 1
            items_per_page = 10

        paginator = Paginator(students, items_per_page)
        page_obj = paginator.get_page(page_number)

        student_data = []
        for student in page_obj:
            # Complete the missing profile_pic handling code
            profile_pic = None
            if student.profile_pic:
                profile_pic = student.profile_pic

            student_data.append({
                'id': student.id,
                'lrn': student.lrn,
                'first_name': student.first_name,
                'last_name': student.last_name,
                'full_name': f"{student.first_name} {student.last_name}",
                'grade': student.grade.name if student.grade else '',
                'grade_id': student.grade.id if student.grade else None,
                'section': student.section.name if student.section else '',
                'section_id': student.section.id if student.section else None,
                'status': student.status,
                'profile_pic': profile_pic,
                'created_at': student.created_at.strftime('%Y-%m-%d'),
                'created_at_display': student.created_at.strftime('%b %d, %Y'),
            })

        pagination = {
            'current_page': page_obj.number,
            'num_pages': paginator.num_pages,
            'has_next': page_obj.has_next(),
            'has_previous': page_obj.has_previous(),
            'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
            'start_index': page_obj.start_index(),
            'end_index': page_obj.end_index(),
            'total_count': total_count,
        }

        return JsonResponse({
            'status': 'success',
            'students': student_data,
            'pagination': pagination,
            'total_count': total_count,
        })
    except Exception as e:
        import traceback
        return JsonResponse({
            'status': 'error',
            'message': str(e),
            'traceback': traceback.format_exc()
        }, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def create_student(request):
    """Create a new student with optional guardian and face enrollment"""
    try:
        import numpy as np
        from datetime import datetime
        
        # Handle form data with file upload
        lrn = request.POST.get('lrn', '')
        first_name = request.POST.get('first_name', '')
        middle_name = request.POST.get('middle_name', '')
        last_name = request.POST.get('last_name', '')
        grade_id = request.POST.get('grade_id', '')
        section_id = request.POST.get('section_id', '')
        status = request.POST.get('status', 'ACTIVE')
        
        # Validate required fields
        if not lrn or not first_name or not last_name or not grade_id or not section_id:
            return JsonResponse({'status': 'error', 'message': 'Missing required fields'}, status=400)
        
        # Check if LRN is unique
        if Student.objects.filter(lrn=lrn).exists():
            return JsonResponse({'status': 'error', 'message': 'LRN already exists'}, status=400)
        
        # Get grade and section objects
        try:
            grade = Grade.objects.get(id=grade_id)
            section = Section.objects.get(id=section_id)
        except (Grade.DoesNotExist, Section.DoesNotExist):
            return JsonResponse({'status': 'error', 'message': 'Invalid grade or section'}, status=400)
        
        # Create the student
        student = Student(
            lrn=lrn,
            first_name=first_name,
            middle_name=middle_name,
            last_name=last_name,
            grade=grade,
            section=section,
            status=status
        )
        
        # Handle profile picture if provided
        if 'profile_pic' in request.FILES:
            profile_pic = request.FILES['profile_pic']
            
            # Validate file size and type
            if profile_pic.size > 2 * 1024 * 1024:  # 2MB limit
                return JsonResponse({'status': 'error', 'message': 'File size must be less than 2MB'}, status=400)
            
            if not profile_pic.content_type.startswith('image/'):
                return JsonResponse({'status': 'error', 'message': 'Only image files are allowed'}, status=400)
            
            # Save the profile picture
            filename = f"student_{int(timezone.now().timestamp())}_{profile_pic.name}"
            fs = FileSystemStorage(location=PROFILE_PICS_DIR)
            filename = fs.save(filename, profile_pic)
            student.profile_pic = filename
        
        student.save()
        
        # Handle guardian creation if guardian data is provided
        has_guardian = request.POST.get('has_guardian') == 'true'
        if has_guardian:
            guardian_first_name = request.POST.get('guardian_first_name', '')
            guardian_last_name = request.POST.get('guardian_last_name', '')
            guardian_middle_name = request.POST.get('guardian_middle_name', '')
            guardian_email = request.POST.get('guardian_email', '')
            guardian_phone = request.POST.get('guardian_phone_number', '')
            guardian_relationship = request.POST.get('guardian_relationship', '')
            
            # Validate guardian required fields
            if guardian_first_name and guardian_last_name and guardian_email and guardian_phone and guardian_relationship:
                guardian = Guardian(
                    first_name=guardian_first_name,
                    middle_name=guardian_middle_name,
                    last_name=guardian_last_name,
                    email=guardian_email,
                    phone=guardian_phone,
                    relationship=guardian_relationship,
                    student=student
                )
                guardian.save()
        
        # Handle face enrollment if face data is provided
        has_face_enrollment = request.POST.get('has_face_enrollment') == 'true'
        if has_face_enrollment:
            # Get face images from request
            face_images = {}
            for angle in ['front', 'left', 'right']:
                image_key = f'face_image_{angle}'
                if image_key in request.FILES:
                    face_images[angle] = request.FILES[image_key]
            
            # Process face embeddings if all 3 images are provided
            if len(face_images) == 3:
                try:
                    # Note: We're storing the images temporarily and will process them on the frontend
                    # The frontend will send the face embeddings separately via the existing face enrollment endpoint
                    # For now, we just mark that face enrollment is pending
                    # The actual embedding extraction happens client-side using face-api.js
                    
                    # Store a placeholder to indicate face enrollment is in progress
                    student.face_path = f"pending_{student.id}"
                    student.save()
                    
                except Exception as face_error:
                    print(f"Error processing face enrollment: {str(face_error)}")
                    import traceback
                    traceback.print_exc()
                    # Continue even if face enrollment fails
        
        return JsonResponse({
            'status': 'success',
            'message': f'Student {first_name} {last_name} created successfully',
            'student_id': student.id
        })
        
    except Exception as e:
        import traceback
        print(f"Error creating student: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
def get_student(request, student_id):
    """
    Endpoint to retrieve a specific student's data for editing
    """
    try:
        student = get_object_or_404(Student, id=student_id)
        
        # Build the student data for response
        student_data = {
            'id': student.id,
            'lrn': student.lrn,
            'first_name': student.first_name,
            'middle_name': student.middle_name or '',
            'last_name': student.last_name,
            'grade_id': str(student.grade.id) if student.grade else '',
            'grade': student.grade.name if student.grade else '',
            'section_id': str(student.section.id) if student.section else '',
            'section': student.section.name if student.section else '',
            'status': student.status,
            'full_name': f"{student.first_name} {student.last_name}",
            'created_at': student.created_at.strftime('%Y-%m-%d'),
            'profile_pic': student.profile_pic or None
        };
        
        return JsonResponse({
            'status': 'success',
            'student': student_data
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST", "PATCH"])
def update_student(request, student_id):
    """Update an existing student (supports file upload via POST)"""
    try:
        student = get_object_or_404(Student, id=student_id)
        
        # Check if the request is a PATCH request
        if request.method == 'PATCH':
            # Check content type to determine how to process data
            if 'multipart/form-data' in request.content_type:
                # Handle form data the same way as POST request
                # For PATCH with form data, use request.POST and request.FILES
                lrn = request.POST.get('lrn', student.lrn)
                first_name = request.POST.get('first_name', student.first_name)
                middle_name = request.POST.get('middle_name', '')
                last_name = request.POST.get('last_name', student.last_name)
                grade_id = request.POST.get('grade_id', student.grade.id if student.grade else None)
                section_id = request.POST.get('section_id', student.section.id if student.section else None)
                status = request.POST.get('status', student.status)
                
                # Validate required fields
                if not lrn or not first_name or not last_name or not grade_id or not section_id:
                    return JsonResponse({
                        'status': 'error',
                        'message': 'Missing required fields'
                    }, status=400)
                
                # Check if LRN already exists for another student
                if Student.objects.filter(lrn=lrn).exclude(id=student_id).exists():
                    return JsonResponse({
                        'status': 'error',
                        'message': f'Student with LRN {lrn} already exists'
                    }, status=400)
                
                # Get grade and section objects
                try:
                    grade = Grade.objects.get(id=grade_id)
                    section = Section.objects.get(id=section_id)
                except (Grade.DoesNotExist, Section.DoesNotExist):
                    return JsonResponse({
                        'status': 'error',
                        'message': 'Invalid grade or section'
                    }, status=400)
                
                # Update student fields
                student.lrn = lrn
                student.first_name = first_name
                student.middle_name = middle_name
                student.last_name = last_name
                student.grade = grade
                student.section = section
                student.status = status
                
                # Handle profile picture if provided
                if 'profile_pic' in request.FILES:
                    profile_pic = request.FILES['profile_pic']
                    
                    # Validate file size and type
                    if profile_pic.size > 2 * 1024 * 1024:  # 2MB limit
                        return JsonResponse({'status': 'error', 'message': 'File size must be less than 2MB'}, status=400)
                    
                    if not profile_pic.content_type.startswith('image/'):
                        return JsonResponse({'status': 'error', 'message': 'Only image files are allowed'}, status=400)
                    
                    # Delete old profile picture if exists
                    if student.profile_pic:
                        old_profile_pic_path = os.path.join(PROFILE_PICS_DIR, student.profile_pic)
                        if os.path.exists(old_profile_pic_path):
                            os.remove(old_profile_pic_path)
                    
                    # Save the new profile picture
                    filename = f"student_{int(timezone.now().timestamp())}_{profile_pic.name}"
                    fs = FileSystemStorage(location=PROFILE_PICS_DIR)
                    filename = fs.save(filename, profile_pic)
                    student.profile_pic = filename
                
                student.save()
                
                return JsonResponse({
                    'status': 'success',
                    'message': f'Student {first_name} {last_name} updated successfully',
                    'student_id': student.id
                })
                
            elif request.content_type == 'application/json':
                # Handle JSON data
                data = json.loads(request.body)
                
                # Update student fields from JSON data
                if 'lrn' in data:
                    student.lrn = data.get('lrn')
                if 'first_name' in data:
                    student.first_name = data.get('first_name')
                if 'middle_name' in data:
                    student.middle_name = data.get('middle_name', '')
                if 'last_name' in data:
                    student.last_name = data.get('last_name')
                if 'grade_id' in data and data['grade_id']:
                    student.grade = get_object_or_404(Grade, id=data['grade_id'])
                if 'section_id' in data and data['section_id']:
                    student.section = get_object_or_404(Section, id=data['section_id'])
                if 'status' in data:
                    student.status = data.get('status')
                
                student.save()
                
                return JsonResponse({
                    'status': 'success',
                    'message': f'Student {student.first_name} {student.last_name} updated successfully',
                    'student_id': student.id
                })
            else:
                # Unsupported content-type
                return JsonResponse({
                    'status': 'error',
                    'message': 'Unsupported content type. Use multipart/form-data for file uploads or application/json for JSON data.'
                }, status=400)
        
        # Handle POST requests (usually form data with files)
        # Get values from POST data
        lrn = request.POST.get('lrn', student.lrn)
        first_name = request.POST.get('first_name', student.first_name)
        middle_name = request.POST.get('middle_name', '')
        last_name = request.POST.get('last_name', student.last_name)
        grade_id = request.POST.get('grade_id', student.grade.id if student.grade else None)
        section_id = request.POST.get('section_id', student.section.id if student.section else None)
        status = request.POST.get('status', student.status)
        
        # Validate required fields
        if not lrn or not first_name or not last_name or not grade_id or not section_id:
            return JsonResponse({
                'status': 'error',
                'message': 'Missing required fields'
            }, status=400)
        
        # Check if LRN already exists for another student
        if Student.objects.filter(lrn=lrn).exclude(id=student_id).exists():
            return JsonResponse({
                'status': 'error',
                'message': f'Student with LRN {lrn} already exists'
            }, status=400)
        
        # Get grade and section objects
        try:
            grade = Grade.objects.get(id=grade_id)
            section = Section.objects.get(id=section_id)
        except (Grade.DoesNotExist, Section.DoesNotExist):
            return JsonResponse({
                'status': 'error',
                'message': 'Invalid grade or section'
            }, status=400)
        
        # Update student fields
        student.lrn = lrn
        student.first_name = first_name
        student.middle_name = middle_name
        student.last_name = last_name
        student.grade = grade
        student.section = section
        student.status = status
        
        # Handle profile picture if provided
        if 'profile_pic' in request.FILES:
            profile_pic = request.FILES['profile_pic']
            
            # Validate file size and type
            if profile_pic.size > 2 * 1024 * 1024:  # 2MB limit
                return JsonResponse({'status': 'error', 'message': 'File size must be less than 2MB'}, status=400)
            
            if not profile_pic.content_type.startswith('image/'):
                return JsonResponse({'status': 'error', 'message': 'Only image files are allowed'}, status=400)
            
            # Delete old profile picture if exists
            if student.profile_pic:
                old_profile_pic_path = os.path.join(PROFILE_PICS_DIR, student.profile_pic)
                if os.path.exists(old_profile_pic_path):
                    os.remove(old_profile_pic_path)
        
            # Save the new profile picture
            filename = f"student_{int(timezone.now().timestamp())}_{profile_pic.name}"
            fs = FileSystemStorage(location=PROFILE_PICS_DIR)
            filename = fs.save(filename, profile_pic)
            student.profile_pic = filename
        
        student.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Student {first_name} {last_name} updated successfully',
            'student_id': student.id
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["DELETE"])
def delete_student(request, student_id):
    """Delete a student"""
    try:
        student = get_object_or_404(Student, id=student_id)
        
        # Store student name for success message
        student_name = f"{student.first_name} {student.last_name}"
        
        # Delete the profile picture file if it exists
        if student.profile_pic:
            profile_pic_path = os.path.join(PROFILE_PICS_DIR, student.profile_pic)
            if os.path.exists(profile_pic_path):
                os.remove(profile_pic_path)
        
        # Delete the student
        student.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Student {student_name} deleted successfully'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def reset_student_password(request, student_id):
    """Reset a student's password (if applicable)"""
    # If students have user accounts, implement password reset logic here.
    return JsonResponse({'status': 'error', 'message': 'Not implemented.'}, status=400)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def save_face_embedding(request):
    """Save face embedding for a student"""
    try:
        import json
        import numpy as np
        from datetime import datetime
        
        data = json.loads(request.body)
        student_id = data.get('student_id')
        face_embeddings = data.get('face_embeddings')  # Array of embeddings from multiple angles
        
        if not student_id or not face_embeddings:
            return JsonResponse({'status': 'error', 'message': 'Student ID and face embeddings are required'}, status=400)
        
        # Get student
        try:
            student = Student.objects.get(id=student_id)
        except Student.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Student not found'}, status=404)
        
        # Create face embeddings directory if it doesn't exist
        face_embeddings_dir = os.path.join(settings.BASE_DIR, 'face_embeddings')
        os.makedirs(face_embeddings_dir, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"student_{student.id}_{timestamp}.npy"
        file_path = os.path.join(face_embeddings_dir, filename)
        
        # Delete old face embedding file if exists
        if student.face_path and student.face_path.strip() and not student.face_path.startswith('pending'):
            old_file_path = os.path.join(settings.BASE_DIR, student.face_path)
            if os.path.exists(old_file_path):
                try:
                    os.remove(old_file_path)
                except Exception as e:
                    print(f"Error deleting old face embedding: {str(e)}")
        
        # Convert embeddings to numpy array and save
        embeddings_array = np.array(face_embeddings)
        np.save(file_path, embeddings_array)
        
        # Update student record with relative path
        relative_path = os.path.join('face_embeddings', filename)
        student.face_path = relative_path
        student.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Face enrolled successfully for {student.first_name} {student.last_name}',
            'student_id': student.id,
            'face_path': relative_path
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON data'}, status=400)
    except Exception as e:
        import traceback
        print("Error in save_face_embedding:", str(e))
        print(traceback.format_exc())
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin_or_principal)
def export_students(request):
    """Export students to Excel, PDF, or Word file"""
    try:
        # Get query parameters for filtering
        search_query = request.GET.get('search', '')
        grade_filter = request.GET.get('grade', '')
        section_filter = request.GET.get('section', '')
        status_filter = request.GET.get('status', '')
        export_format = request.GET.get('format', 'excel').lower()
        
        print(f"DEBUG: Export students called with format: {export_format}")
        
        # Base queryset (ordered by ID ascending for backup/restore)
        students = Student.objects.select_related('grade', 'section').all().order_by('id')
        
        # Apply search if provided
        if search_query:
            students = students.filter(
                Q(first_name__icontains=search_query) |
                Q(last_name__icontains=search_query) |
                Q(lrn__icontains=search_query)
            )
        
        # Apply grade filter if provided
        if grade_filter:
            students = students.filter(grade_id=grade_filter)
        
        # Apply section filter if provided
        if section_filter:
            students = students.filter(section_id=section_filter)
        
        # Apply status filter
        if status_filter:
            students = students.filter(status=status_filter)
        
        # Define headers
        headers = ["ID", "LRN", "Last Name", "First Name", "Middle Name", "Grade", "Section", "Status", "Face Enrolled", "Created Date"]
        
        # Prepare data rows
        data_rows = []
        for student in students:
            face_enrolled = "Yes" if student.face_path and student.face_path.strip() and not student.face_path.startswith('pending') else "No"
            data_rows.append([
                str(student.id),
                student.lrn,
                student.last_name,
                student.first_name,
                student.middle_name or "",
                student.grade.name if student.grade else "",
                student.section.name if student.section else "",
                student.get_status_display(),
                face_enrolled,
                student.created_at.strftime("%Y-%m-%d %H:%M:%S") if student.created_at else ""
            ])
        
        # Export based on format
        if export_format == 'pdf':
            return export_students_to_pdf(headers, data_rows)
        elif export_format == 'word':
            return export_students_to_word(headers, data_rows)
        else:  # Default to Excel
            return export_students_to_excel_format(headers, data_rows)
        
    except Exception as e:
        import traceback
        print(f"Export students error: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

def export_students_to_excel_format(headers, data_rows):
    """Export students to Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Students"
    
    # Define styles
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    border_style = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Add headers
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = border_style
    
    # Add data rows
    for row_num, row_data in enumerate(data_rows, 2):
        for col_num, cell_value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num, value=cell_value)
            cell.border = border_style
            cell.alignment = Alignment(vertical="center")
    
    # Auto-adjust column widths
    for col in ws.columns:
        max_length = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[col_letter].width = adjusted_width
    
    # Create response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    filename = f'students_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    print(f"DEBUG: Excel filename: {filename}")
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    wb.save(response)
    return response

def export_students_to_pdf(headers, data_rows):
    """Export students to PDF format"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=18)
    elements = []
    
    # Add title
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#1F4E78'),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    title = Paragraph(f"Students Report - {datetime.now().strftime('%B %d, %Y')}", title_style)
    elements.append(title)
    elements.append(Spacer(1, 12))
    
    # Prepare table data
    table_data = [headers] + data_rows
    
    # Create table
    table = Table(table_data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E78')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F0F0')]),
    ]))
    
    elements.append(table)
    doc.build(elements)
    
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    filename = f'students_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
    print(f"DEBUG: PDF filename: {filename}")
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

def export_students_to_word(headers, data_rows):
    """Export students to Word format"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Students Report - {datetime.now().strftime("%B %d, %Y")}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing
    doc.add_paragraph()
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set header background color
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '1F4E78')
        hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
        
        # Set header text color to white
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Add data rows
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = str(cell_value)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f'students_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    print(f"DEBUG: Word filename: {filename}")
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

@login_required
@user_passes_test(is_admin)
def get_student(request, student_id):
    """
    Endpoint to retrieve a specific student's data for editing
    """
    try:
        student = get_object_or_404(Student, id=student_id)
        
        # Build the student data for response
        student_data = {
            'id': student.id,
            'lrn': student.lrn,
            'first_name': student.first_name,
            'middle_name': student.middle_name or '',
            'last_name': student.last_name,
            'grade_id': str(student.grade.id) if student.grade else '',
            'grade': student.grade.name if student.grade else '',
            'section_id': str(student.section.id) if student.section else '',
            'section': student.section.name if student.section else '',
            'status': student.status,
            'full_name': f"{student.first_name} {student.last_name}",
            'created_at': student.created_at.strftime('%Y-%m-%d'),
            'profile_pic': student.profile_pic or None
        }
        
        return JsonResponse({
            'status': 'success',
            'student': student_data
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

# ==========================
#  GUARDIAN MANAGEMENT VIEWS
# ==========================

@login_required
@user_passes_test(is_admin)
def admin_guardians(request):
    """Guardian management view with dashboard cards"""
    # Calculate dashboard statistics
    total_guardians = Guardian.objects.count()
    
    # Count guardians with children (since each guardian has exactly one student)
    guardians_with_children = Guardian.objects.exclude(student=None).count()
    
    # Count mother guardians
    mother_count = Guardian.objects.filter(relationship='MOTHER').count()
    
    # Count father guardians
    father_count = Guardian.objects.filter(relationship='FATHER').count()
    
    # Get relationship choices for filter dropdown
    relationship_choices = [
        ('FATHER', 'Father'),
        ('MOTHER', 'Mother'),
        ('GUARDIAN', 'Guardian'),
        ('GRANDMOTHER', 'Grandmother'),
        ('GRANDFATHER', 'Grandfather'),
        ('AUNT', 'Aunt'),
        ('UNCLE', 'Uncle'),
        ('SIBLING', 'Sibling'),
        ('OTHER', 'Other'),
    ]
    
    # Get all grades for filter dropdown
    grades = Grade.objects.all().order_by('name')
    
    # Get all sections for filter dropdown
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')
    
    # Get students for children dropdown
    students = Student.objects.select_related('grade', 'section').filter(status='ACTIVE')
    
    context = {
        'total_guardians': total_guardians,
        'guardians_with_children': guardians_with_children,
        'mother_count': mother_count,
        'father_count': father_count,
        'relationship_choices': relationship_choices,
        'grades': grades,
        'sections': sections,
        'students': students,
    }
    
    return render(request, 'admin/guardians.html', context)

@login_required
@user_passes_test(is_admin_or_principal)
@require_http_methods(["GET"])
def search_guardians(request):
    """API endpoint to search and filter guardians for AJAX requests"""
    try:
        # Get query parameters
        search_query = request.GET.get('search', '')
        relationship_filter = request.GET.get('relationship', '')
        grade_filter = request.GET.get('grade', '')
        section_filter = request.GET.get('section', '')
        page_number = request.GET.get('page', 1)
        items_per_page = request.GET.get('items_per_page', 10)
        
        # Base queryset
        guardians = Guardian.objects.all().order_by('-created_at')
        
        # Apply search if provided
        if search_query:
            guardians = guardians.filter(
                Q(first_name__icontains=search_query) |
                Q(middle_name__icontains=search_query) |
                Q(last_name__icontains=search_query) |
                Q(email__icontains=search_query) |
                Q(phone__icontains=search_query)  # Use phone instead of phone_number
            )
        
        # Apply relationship filter if provided
        if relationship_filter:
            guardians = guardians.filter(relationship=relationship_filter)
        
        # Apply grade filter if provided
        if grade_filter:
            # Find students in this grade, then filter guardians related to them
            student_ids = Student.objects.filter(grade_id=grade_filter).values_list('id', flat=True)
            guardians = guardians.filter(student_id__in=student_ids)
        
        # Apply section filter if provided
        if section_filter:
            # Find students in this section, then filter guardians related to them
            student_ids = Student.objects.filter(section_id=section_filter).values_list('id', flat=True)
            guardians = guardians.filter(student_id__in=student_ids)
        
        # Get total count for pagination
        total_count = guardians.count()
        
        # Parse page number and items_per_page
        try:
            page_number = int(page_number)
            items_per_page = int(items_per_page)
        except (ValueError, TypeError):
            page_number = 1
            items_per_page = 10
        
        # Create paginator
        paginator = Paginator(guardians, items_per_page)
        page_obj = paginator.get_page(page_number)
        
        # Prepare guardian data for JSON response
        guardian_data = []
        for guardian in page_obj:
            try:
                # Create guardian data dict with safe fallbacks
                data = {
                    'id': guardian.id,
                    'full_name': f"{guardian.first_name} {guardian.last_name}",
                    'first_name': guardian.first_name,
                    'middle_name': guardian.middle_name or '',
                    'last_name': guardian.last_name,
                    'email': guardian.email or 'No email',
                    'phone_number': guardian.phone or 'No phone',  # Use phone instead of phone_number
                    'relationship': guardian.relationship,
                    'created_at_display': guardian.created_at.strftime('%B %d, %Y') if guardian.created_at else 'N/A'
                }
                
                # Safely handle relationship display and children data
                try:
                    data['relationship_display'] = guardian.get_relationship_display()
                except (AttributeError, TypeError):
                    data['relationship_display'] = guardian.relationship
                
                # Child info with safe fallbacks
                try:
                    student = Student.objects.filter(id=guardian.student_id).first() if hasattr(guardian, 'student_id') else None
                    if student:
                        data['children_count'] = 1
                        data['children_preview'] = f"{student.first_name} {student.last_name}"
                    else:
                        data['children_count'] = 0
                        data['children_preview'] = 'No children'

                except Exception as e:
                    data['children_count'] = 0
                    data['children_preview'] = 'Error loading children'



                    print(f"Error loading children for guardian {guardian.id}: {str(e)}")
                
                guardian_data.append(data)
            except Exception as e:
                print(f"Error processing guardian {guardian.id if hasattr(guardian, 'id') else 'unknown'}: {str(e)}")
                continue
        
        # Prepare pagination data
        pagination = {
            'current_page': page_obj.number,
            'num_pages': paginator.num_pages,
            'has_next': page_obj.has_next(),
            'has_previous': page_obj.has_previous(),
            'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
            'start_index': page_obj.start_index(),
            'end_index': page_obj.end_index(),
            'total_count': total_count,
        }
        
        return JsonResponse({
            'status': 'success',
            'guardians': guardian_data,
            'pagination': pagination,
            'total_count': total_count,
        })
    
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        print(f"Error in search_guardians: {str(e)}")
        print(error_traceback)
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}',
            'traceback': error_traceback
        }, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def create_guardian(request):
    """Create a new guardian"""
    try:
        data = json.loads(request.body)
        
        # Validate required fields
        required_fields = ['first_name', 'last_name', 'phone_number', 'relationship', 'student_id']
        for field in required_fields:
            if not field in data or not data[field]:
                return JsonResponse({
                    'status': 'error',
                    'message': f'Field {field} is required'
                }, status=400)
        
        # Check if student exists and doesn't already have a guardian
       

        try:
            student = Student.objects.get(id=data['student_id'])
        except Student.DoesNotExist:
            return JsonResponse({
                'status': 'error',
                'message': 'Student not found'
            }, status=404)
        
        # Validate email if provided
        email = data.get('email', '').strip()
        if email:
            # Add email validation if needed
            pass
        
        # Create guardian - Note: using phone_number from form but storing in phone field
        guardian = Guardian.objects.create(
            first_name=data['first_name'].strip(),
            middle_name=data.get('middle_name', '').strip(),
            last_name=data['last_name'].strip(),
            email=email,
            phone=data['phone_number'].strip(),  # Changed to map phone_number to phone
            relationship=data['relationship'],
            student=student
        )
        
        return JsonResponse({
            'status': 'success',
            'message': 'Guardian created successfully',
            'guardian': {
                'id': guardian.id,
                'full_name': f"{guardian.first_name} {guardian.last_name}",
                'email': guardian.email or '',
                'phone_number': guardian.phone,  # Return as phone_number for consistency with form
                'relationship': guardian.relationship,
                'student_name': f"{student.first_name} {student.last_name}",
            }
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'status': 'error',
            'message': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}'
        }, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["PATCH"])
def update_guardian(request, guardian_id):
    """Update an existing guardian"""
    try:
        data = json.loads(request.body)
        
        guardian = get_object_or_404(Guardian, id=guardian_id)
        
        # Validate required fields
        required_fields = ['first_name', 'last_name', 'phone_number', 'relationship', 'student_id']
        for field in required_fields:
            if not field in data or not data[field]:
                return JsonResponse({
                    'status': 'error',
                    'message': f'Field {field} is required'
                }, status=400)
        
        # Check if student exists and validate assignment
        try:
            student = Student.objects.get(id=data['student_id'])
        except Student.DoesNotExist:
            return JsonResponse({
                'status': 'error',
                'message': 'Student not found'
            }, status=404)
        
        # Validate email if provided
        email = data.get('email', '').strip()
        # Email validation could go here if needed
        
        # Update guardian fields
        guardian.first_name = data['first_name'].strip()
        guardian.middle_name = data.get('middle_name', '').strip()
        guardian.last_name = data['last_name'].strip()
        guardian.email = email
        guardian.phone = data['phone_number'].strip()  # Map phone_number from form to phone field in model
        guardian.relationship = data['relationship']
        guardian.student = student
        guardian.save()
        
        return JsonResponse({
            'status': 'success',
            'message': 'Guardian updated successfully',
            'guardian': {
                'id': guardian.id,
                'full_name': f"{guardian.first_name} {guardian.last_name}",
                'email': guardian.email or '',
                'phone_number': guardian.phone,  # Return as phone_number for consistency with form
                'relationship': guardian.relationship,
                'student_name': f"{student.first_name} {student.last_name}"
            }
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'status': 'error',
            'message': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        error_traceback = traceback.format_exc()
        print(f"Error in update_guardian: {str(e)}")
        print(error_traceback)
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}',
            'traceback': error_traceback
        }, status=500)  # Changed to 500 for server errors

@login_required
@user_passes_test(is_admin)
@require_http_methods(["DELETE"])
def delete_guardian(request, guardian_id):
    """Delete a guardian"""
    try:
        guardian = get_object_or_404(Guardian, id=guardian_id)
        guardian_name = f"{guardian.first_name} {guardian.last_name}"  # Use direct concatenation instead of get_full_name
        guardian.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Guardian {guardian_name} deleted successfully'
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}'
        }, status=500)

@login_required
@user_passes_test(is_admin)
@require_GET
def get_guardian_children(request, guardian_id):
    """Get children for a specific guardian"""
    try:
        guardian = get_object_or_404(Guardian, id=guardian_id)
        
        children = []
        if guardian.student:
            children.append({
                'id': guardian.student.id,
                'full_name': f"{guardian.student.first_name} {guardian.student.last_name}",  # Direct concatenation
                'lrn': guardian.student.lrn,
                'grade': guardian.student.grade.name if guardian.student.grade else 'N/A',
                'section': guardian.student.section.name if guardian.student.section else 'N/A',
                'status': guardian.student.status
            })
        
        return JsonResponse({
            'status': 'success',
            'guardian': {
                'name': f"{guardian.first_name} {guardian.last_name}",  # Direct concatenation
                'relationship': guardian.get_relationship_display() if hasattr(guardian, 'get_relationship_display') else guardian.relationship
            },
            'children': children
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}'
        }, status=500)

@login_required
@user_passes_test(is_admin)
@require_GET
def admin_get_sections_by_grade(request):
    """Get sections for a specific grade"""
    try:
        grade_id = request.GET.get('grade_id')
        if not grade_id:
            return JsonResponse({
                'status': 'error',
                'message': 'Grade ID is required'
            }, status=400)
        
        sections = Section.objects.filter(grade_id=grade_id).order_by('name')
        sections_data = [
            {
                'id': section.id,
                'name': section.name,
                'grade_id': section.grade.id,
                'grade_name': section.grade.name
            }
            for section in sections
        ]
        
        return JsonResponse({
            'status': 'success',
            'sections': sections_data
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}'
        }, status=500)

@login_required
@user_passes_test(is_admin)
@require_GET
def admin_get_students_by_section(request):
    """Get students for a specific section"""
    try:
        section_id = request.GET.get('section_id')
        if not section_id:
            return JsonResponse({
                'status': 'error',
                'message': 'Section ID is required'
            }, status=400)
        

        
        # Get students who don't already have a guardian assigned (for create) or are assigned to the current guardian (for edit)
        students = Student.objects.filter(
            section_id=section_id,
            status='ACTIVE'
        ).select_related('grade', 'section').order_by('first_name', 'last_name')
        
        students_data = [
            {
                'id': student.id,
                'first_name': student.first_name,
                'last_name': student.last_name,
                'lrn': student.lrn,
                'grade': student.grade.name if student.grade else 'N/A',
                'section': student.section.name if student.section else 'N/A',
                'grade_id': student.grade.id if student.grade else None,
                'section_id': student.section.id if student.section else None,
                'has_guardian': hasattr(student, 'guardian') and student.guardian is not None
            }
            for student in students
        ]
        
        return JsonResponse({
            'status': 'success',
            'students': students_data
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}'
        }, status=500)

@login_required
@user_passes_test(is_admin)
@require_GET
def get_guardian_details(request, guardian_id):
    """
    Endpoint to retrieve a specific guardian's data for editing
    """
    try:
        guardian = get_object_or_404(Guardian, id=guardian_id)
        
        # Prepare children data
        children = []
        if guardian.student:
            child = guardian.student
            children.append({
                'id': child.id,
                'grade_id': child.grade.id if child.grade else '',
                'grade': child.grade.name if child.grade else '',
                'section_id': child.section.id if child.section else '',
                'section': child.section.name if child.section else '',
                'first_name': child.first_name,
                'last_name': child.last_name,
                'full_name': f"{child.first_name} {child.last_name}"  # Direct concatenation
            })
        
        guardian_data = {
            'id': guardian.id,
            'first_name': guardian.first_name,
            'middle_name': guardian.middle_name or '',
            'last_name': guardian.last_name,
            'email': guardian.email or '',
            'phone_number': guardian.phone,  # Return phone field as phone_number for form
            'relationship': guardian.relationship,
            'children': children,
            'created_at': guardian.created_at.isoformat() if guardian.created_at else None
        }
        
        return JsonResponse({
            'status': 'success',
            'guardian': guardian_data
        })
        
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        print(f"Error in get_guardian_details: {str(e)}")
        print(error_traceback)
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}',
            'traceback': error_traceback
        }, status=500)


# ==========================
#  ATTENDANCE MANAGEMENT VIEWS
# ==========================

@login_required
@user_passes_test(is_admin)
def admin_attendance_records(request):
    """Attendance management view with dashboard cards and filters"""
    from django.db.models import Q
    from datetime import date

    # Dashboard stats
    total_records = Attendance.objects.count()
    present_count = Attendance.objects.filter(status='ON TIME').count() + Attendance.objects.filter(status='LATE').count()
    absent_count = Attendance.objects.filter(status='ABSENT').count()
    excused_count = Attendance.objects.filter(status='EXCUSED').count()

    grades = Grade.objects.all().order_by('name')
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    # Filters (for initial page load)
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')

    context = {
        'grades': grades,
        'sections': sections,
        'grade_filter': grade_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'date_filter': date_filter,
        'total_records': total_records,
        'present_count': present_count,
        'absent_count': absent_count,
        'excused_count': excused_count,
    }
    return render(request, 'admin/attendance.html', context)

@login_required
@user_passes_test(is_admin_or_principal)
@require_http_methods(["GET"])
def search_attendance_records(request):
    """AJAX search/filter for attendance records"""
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')
    page_number = request.GET.get('page', 1)
    items_per_page = request.GET.get('items_per_page', 10)

    records = Attendance.objects.select_related('student', 'student__grade', 'student__section').all().order_by('-date', '-student__last_name')

    if search_query:
        records = records.filter(
            Q(student__first_name__icontains=search_query) |
            Q(student__last_name__icontains=search_query) |
            Q(student__lrn__icontains=search_query) |
            Q(student__section__name__icontains=search_query)
        )
    if grade_filter:
        records = records.filter(student__grade_id=grade_filter)
    if section_filter:
        records = records.filter(student__section_id=section_filter)
    if status_filter:
        records = records.filter(status=status_filter)
    if date_filter:
        records = records.filter(date=date_filter)

    total_count = records.count()
    try:
        page_number = int(page_number)
        items_per_page = int(items_per_page)
    except (ValueError, TypeError):
        page_number = 1
        items_per_page = 10

    paginator = Paginator(records, items_per_page)
    page_obj = paginator.get_page(page_number)

    attendance_data = []
    for record in page_obj:
        attendance_data.append({
            'id': record.id,
            'lrn': record.student.lrn,
            'full_name': f"{record.student.first_name} {record.student.last_name}",
            'grade': record.student.grade.name if record.student.grade else '',
            'section': record.student.section.name if record.student.section else '',
            'date': record.date.strftime('%Y-%m-%d'),
            'time_in': record.time_in.strftime('%H:%M:%S') if record.time_in else '',
            'time_out': record.time_out.strftime('%H:%M:%S') if record.time_out else '',
            'status': record.status,
            'profile_pic': record.student.profile_pic or None,
        })

    pagination = {
        'current_page': page_obj.number,
        'num_pages': paginator.num_pages,
        'has_next': page_obj.has_next(),
        'has_previous': page_obj.has_previous(),
        'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
        'start_index': page_obj.start_index(),
        'end_index': page_obj.end_index(),
        'total_count': total_count,
    }

    return JsonResponse({
        'status': 'success',
        'records': attendance_data,
        'pagination': pagination,
        'total_count': total_count,
    })

@login_required
@user_passes_test(is_admin)
@require_GET
def get_attendance_record(request, attendance_id):
    """API endpoint to get details of a specific attendance record"""
    try:
        record = get_object_or_404(Attendance, id=attendance_id)
        
        data = {
            'id': record.id,
            'student_id': record.student.id,
            'student_name': f"{record.student.first_name} {record.student.last_name}",
            'lrn': record.student.lrn,
            'date': record.date.strftime('%Y-%m-%d'),
            'time_in': record.time_in.strftime('%H:%M:%S') if record.time_in else None,
            'time_out': record.time_out.strftime('%H:%M:%S') if record.time_out else None,
            'status': record.status,
            'notes': record.notes or '',
            'section': record.student.section.name if record.student.section else '',
            'section_id': record.student.section.id if record.student.section else None,
            'grade': record.student.grade.name if record.student.grade else '',
            'grade_id': record.student.grade.id if record.student.grade else None,
        }
        
        return JsonResponse({
            'status': 'success',
            'record': data
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def create_attendance_record(request):
    """API endpoint to create a new attendance record"""
    try:
        data = json.loads(request.body)
        
        # Validate required fields
        required_fields = ['student_id', 'date', 'status']
        for field in required_fields:
            if field not in data or not data[field]:
                return JsonResponse({
                    'status': 'error',
                    'message': f'{field} is required'
                }, status=400)
        
        # Check if student exists
        try:
            student = Student.objects.get(id=data['student_id'])
        except Student.DoesNotExist:
            return JsonResponse({
                'status': 'error',
                'message': 'Student not found'
            }, status=404)
        
        # Check if an attendance record already exists for this student on this date
        if Attendance.objects.filter(student=student, date=data['date']).exists():
            return JsonResponse({
                'status': 'error',
                'message': f'Attendance record already exists for {student.first_name} {student.last_name} on {data["date"]}'
            }, status=400)
        
        # Parse time fields if provided
        time_in = None
        if data.get('time_in'):
            try:
                from datetime import datetime
                time_in = datetime.strptime(data['time_in'], '%H:%M:%S').time()
            except ValueError:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Invalid time_in format. Use HH:MM:SS'
                }, status=400)
                
        time_out = None
        if data.get('time_out'):
            try:
                from datetime import datetime
                time_out = datetime.strptime(data['time_out'], '%H:%M:%S').time()
            except ValueError:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Invalid time_out format. Use HH:MM:SS'
                }, status=400)
        
        # Create attendance record
        record = Attendance(
            student=student,
            date=data['date'],
            time_in=time_in,
            time_out=time_out,
            status=data['status']
        )
        record.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Attendance record for {student.first_name} {student.last_name} created successfully',
            'record_id': record.id
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'status': 'error',
            'message': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        import traceback
        print("Error in create_attendance_record:", str(e))
        print(traceback.format_exc())
        return JsonResponse({
            'status': 'error',
            'message': str(e),
            'traceback': traceback.format_exc()
        }, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["PUT", "PATCH"])
def update_attendance_record(request, attendance_id):
    """API endpoint to update an attendance record"""
    try:
        record = get_object_or_404(Attendance, id=attendance_id)
        data = json.loads(request.body)
        
        # Update status if provided
        if 'status' in data:
            record.status = data['status']
        
        # Update date if provided
        if 'date' in data:
            # Check if another record exists for this student on the new date (excluding current record)
            if Attendance.objects.filter(student=record.student, date=data['date']).exclude(id=attendance_id).exists():
                return JsonResponse({
                    'status': 'error',
                    'message': f'Another attendance record exists for {record.student.first_name} {record.student.last_name} on {data["date"]}'
                }, status=400)
            record.date = data['date']
        
        # Update time_in if provided
        if 'time_in' in data:
            if data['time_in']:
                try:
                    from datetime import datetime
                    record.time_in = datetime.strptime(data['time_in'], '%H:%M:%S').time()
                except ValueError:
                    return JsonResponse({
                        'status': 'error',
                        'message': 'Invalid time_in format. Use HH:MM:SS'
                    }, status=400)
            else:
                record.time_in = None
        
        # Update time_out if provided
        if 'time_out' in data:
            if data['time_out']:
                try:
                    from datetime import datetime
                    record.time_out = datetime.strptime(data['time_out'], '%H:%M:%S').time()
                except ValueError:
                    return JsonResponse({
                        'status': 'error',
                        'message': 'Invalid time_out format. Use HH:MM:SS'
                    }, status=400)
            else:
                record.time_out = None
        
        # Update notes if provided
        if 'notes' in data:
            record.notes = data.get('notes', '')
        
        # Save the updated record
        record.save()
        
        return JsonResponse({
            'status': 'success',
            'message': 'Attendance record updated successfully'
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'status': 'error',
            'message': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["DELETE"])
def delete_attendance_record(request, attendance_id):
    """API endpoint to delete an attendance record"""
    try:
        record = get_object_or_404(Attendance, id=attendance_id)
        student_name = f"{record.student.first_name} {record.student.last_name}"
        record_date = record.date.strftime('%Y-%m-%d')
        
        record.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Attendance record for {student_name} on {record_date} deleted successfully'
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

@login_required
@user_passes_test(is_admin)
def export_attendance_to_excel(request):
    """Export attendance records to Excel, PDF, or Word file"""
    import sys
    import traceback
    sys.stdout.flush()  # Force flush output
    print("\n" + "="*80, flush=True)
    print("=== EXPORT ATTENDANCE CALLED ===", flush=True)
    print("="*80, flush=True)
    print(f"User: {request.user}, Is authenticated: {request.user.is_authenticated}", flush=True)
    print(f"User role: {request.user.role if hasattr(request.user, 'role') else 'No role'}", flush=True)
    
    # Simple test - return a basic response
    if request.GET.get('test') == 'true':
        print("TEST MODE - Returning test response", flush=True)
        return JsonResponse({'status': 'success', 'message': 'Export function is reachable'})
    
    try:
        # Get query parameters for filtering
        search_query = request.GET.get('search', '')
        grade_filter = request.GET.get('grade', '')
        section_filter = request.GET.get('section', '')
        status_filter = request.GET.get('status', '')
        date_filter = request.GET.get('date', '')
        export_format = request.GET.get('format', 'excel').lower()
        print(f"Export format: {export_format}, Date filter: {date_filter}")
        print(f"Filters - Search: {search_query}, Grade: {grade_filter}, Section: {section_filter}, Status: {status_filter}")
        
        # Base queryset
        print("Querying attendance records...")
        records = Attendance.objects.select_related('student', 'student__grade', 'student__section').all().order_by('-date', '-student__last_name')
        print(f"Initial query returned {records.count()} records")
        
        # Apply search if provided
        if search_query:
            print(f"Applying search filter: {search_query}")
            records = records.filter(
                Q(student__first_name__icontains=search_query) | 
                Q(student__last_name__icontains=search_query) |
                Q(student__lrn__icontains=search_query) |
                Q(student__section__name__icontains=search_query)
            )
        
        # Apply grade filter
        if grade_filter:
            print(f"Applying grade filter: {grade_filter}")
            records = records.filter(student__grade__id=grade_filter)
        
        # Apply section filter
        if section_filter:
            print(f"Applying section filter: {section_filter}")
            records = records.filter(student__section__id=section_filter)
        
        # Apply status filter
        if status_filter:
            print(f"Applying status filter: {status_filter}")
            records = records.filter(status=status_filter)
        
        # Apply date filter
        if date_filter:
            print(f"Applying date filter: {date_filter}")
            records = records.filter(date=date_filter)
        
        print(f"After filters: {records.count()} records")
        
        # Define headers
        headers = ["ID", "Student Name", "LRN", "Grade", "Section", "Date", "Time In", "Time Out", "Status", "Notes"]
        
        # Prepare data rows
        print("Preparing data rows...")
        data_rows = []
        
        try:
            for idx, record in enumerate(records):
                if idx % 100 == 0 and idx > 0:
                    print(f"Processed {idx} records...")
                
                # Safely get student info
                try:
                    student_name = f"{record.student.first_name} {record.student.last_name}" if record.student else "N/A"
                    student_lrn = record.student.lrn if record.student and record.student.lrn else ""
                    grade_name = record.student.grade.name if record.student and record.student.grade else ""
                    section_name = record.student.section.name if record.student and record.student.section else ""
                except Exception as e:
                    print(f"Error getting student info for record {record.id}: {e}")
                    student_name = "Error"
                    student_lrn = ""
                    grade_name = ""
                    section_name = ""
                    
                data_rows.append([
                    str(record.id),
                    student_name,
                    student_lrn,
                    grade_name,
                    section_name,
                    record.date.strftime("%Y-%m-%d") if record.date else "",
                    record.time_in.strftime("%H:%M:%S") if record.time_in else "",
                    record.time_out.strftime("%H:%M:%S") if record.time_out else "",
                    record.status or "",
                    record.notes or ""
                ])
        except Exception as row_error:
            print(f"Error processing records: {row_error}")
            import traceback
            print(traceback.format_exc())
            raise
        
        print(f"Total records to export: {len(data_rows)}")
        
        # Export based on format
        if export_format == 'pdf':
            print("Calling PDF export function...")
            return export_attendance_to_pdf(headers, data_rows)
        elif export_format == 'word':
            return export_attendance_to_word(headers, data_rows)
        else:  # Default to Excel
            return export_attendance_to_excel_format(headers, data_rows)
        
    except Exception as e:
        print("\n" + "!"*80)
        print("!!! EXPORT ATTENDANCE ERROR !!!")
        print("!"*80)
        error_trace = traceback.format_exc()
        print(error_trace)
        print("!"*80 + "\n")
        return JsonResponse({'status': 'error', 'message': str(e), 'traceback': error_trace}, status=500)


def export_attendance_to_excel_format(headers, data_rows):
    """Export attendance to Excel format"""
    # Create a workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Attendance Records"
    
    # Define styles
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    border_style = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    center_alignment = Alignment(horizontal="center", vertical="center")
    left_alignment = Alignment(horizontal="left", vertical="center")
    
    # Set column widths
    column_widths = [8, 20, 15, 10, 15, 12, 12, 12, 12, 25]
    for col_num, width in enumerate(column_widths, 1):
        ws.column_dimensions[chr(64 + col_num)].width = width
    
    # Add headers
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = header
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = border_style
    
    # Set header row height
    ws.row_dimensions[1].height = 25
    
    # Add data rows
    for row_num, row_data in enumerate(data_rows, 2):
        for col_num, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num)
            cell.value = value
            cell.border = border_style
            
            # Apply alignment based on column
            if col_num in [1, 4, 5, 6, 7, 8, 9]:  # ID, Grade, Section, Date, Times, Status - center
                cell.alignment = center_alignment
            else:  # Name, LRN, Notes - left
                cell.alignment = left_alignment
    
    # Create response
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    from datetime import datetime as dt
    response = HttpResponse(
        output.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename=attendance_{dt.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    
    return response


def export_attendance_to_pdf(headers, data_rows):
    """Export attendance to PDF format"""
    print("=== PDF Export Function Called ===")
    print(f"Received {len(headers)} headers and {len(data_rows)} data rows")
    
    try:
        print("Importing reportlab modules...")
        try:
            from reportlab.lib.pagesizes import letter, landscape
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
            print("Reportlab modules imported successfully")
        except ImportError as import_err:
            print(f"!!! REPORTLAB IMPORT ERROR: {import_err}")
            raise Exception(f"reportlab library not installed or not found: {import_err}")
        
        # Ensure all data is string and handle None values
        print("Converting data to strings...")
        headers = [str(h) if h is not None else "" for h in headers]
        data_rows = [[str(cell) if cell is not None else "" for cell in row] for row in data_rows]
        
        # If no data, add a message row
        if not data_rows:
            print("No data rows found, adding placeholder")
            data_rows = [["No attendance records found"] + [""] * (len(headers) - 1)]
        
        print(f"Processing {len(data_rows)} data rows")
        
        # Create PDF in memory
        print("Creating PDF buffer...")
        buffer = io.BytesIO()
        
        # Create simple PDF with landscape orientation for more width
        print("Creating PDF document...")
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=landscape(letter),
            rightMargin=30,
            leftMargin=30,
            topMargin=30,
            bottomMargin=18
        )
        
        elements = []
        
        # Prepare table data
        print("Preparing table data...")
        table_data = [headers] + data_rows
        
        # Create table without fixed column widths (let it auto-size)
        print("Creating table...")
        table = Table(table_data, repeatRows=1)
        
        # Simple table style
        print("Applying table style...")
        table_style = TableStyle([
            # Header styling
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            
            # Data rows styling
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ])
        
        table.setStyle(table_style)
        elements.append(table)
        
        # Build PDF
        print("Building PDF document...")
        doc.build(elements)
        print("PDF built successfully")
        
        buffer.seek(0)
        
        print("Creating HTTP response...")
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        from datetime import datetime as dt
        response['Content-Disposition'] = f'attachment; filename="attendance_{dt.now().strftime("%Y%m%d_%H%M%S")}.pdf"'
        
        print("PDF export completed successfully!")
        return response
        
    except Exception as e:
        error_trace = traceback.format_exc()
        print(f"!!! PDF export error !!!")
        print(error_trace)
        # Re-raise the exception to be caught by the parent function
        raise


def export_attendance_to_word(headers, data_rows):
    """Export attendance to Word format"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    def set_cell_background(cell, fill_color):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Attendance Records', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(31, 78, 120)
    
    # Add space
    doc.add_paragraph()
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Set column widths (in inches)
    col_widths = [0.5, 1.5, 1.2, 0.7, 1.2, 1.0, 0.9, 0.9, 0.9, 1.8]
    for idx, width in enumerate(col_widths):
        for cell in table.columns[idx].cells:
            cell.width = Inches(width)
    
    # Add headers
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        cell.text = header
        
        # Style header
        set_cell_background(cell, '1F4E78')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(11)
    
    # Add data rows
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.add_row().cells
        
        # Set alternating row colors
        if row_idx % 2 == 1:
            for cell in row_cells:
                set_cell_background(cell, 'F3F4F6')
        
        for col_idx, value in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = str(value)
            
            # Style cell
            for paragraph in cell.paragraphs:
                # Center alignment for certain columns
                if col_idx in [0, 3, 4, 5, 6, 7, 8]:  # ID, Grade, Section, Date, Times, Status
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:  # Name, LRN, Notes
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    from datetime import datetime as dt
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=attendance_{dt.now().strftime("%Y%m%d_%H%M%S")}.docx'
    
    return response


@login_required
@user_passes_test(is_admin)
@require_GET
def get_students_for_attendance(request):
    """API endpoint to get students for attendance form"""
    try:
        # Get filters
        section_id = request.GET.get('section', None)
        grade_id = request.GET.get('grade', None)
        
        # Base queryset - only active students
        students = Student.objects.filter(status='ACTIVE').select_related('grade', 'section')
        
        # Apply filters if provided
        if section_id:
            students = students.filter(section_id=section_id)
        elif grade_id:
            students = students.filter(grade_id=grade_id)
        
        # Order by name
        students = students.order_by('last_name', 'first_name')
        
        # Prepare data for response
        students_data = []
        for student in students:
            students_data.append({
                'id': student.id,
                'lrn': student.lrn,
                'name': f"{student.first_name} {student.last_name}",
                'grade': student.grade.name if student.grade else '',
                'section': student.section.name if student.section else '',
                'grade_id': student.grade.id if student.grade else None,
                'section_id': student.section.id if student.section else None
            })
        
        return JsonResponse({
            'status': 'success',
            'students': students_data,
            'total': len(students_data)
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'An error occurred: {str(e)}'
        }, status=500)
    

# ==========================
#  EXCUSED MANAGEMENT VIEWS
# ==========================

@login_required
@user_passes_test(is_admin)
def admin_excused(request):
    """View for excused absences management"""
    from datetime import date
    total_excused = ExcusedAbsence.objects.count()
    today = date.today()
    active_excused = ExcusedAbsence.objects.filter(effective_date__lte=today, end_date__gte=today).count()
    expired_excused = ExcusedAbsence.objects.filter(end_date__lt=today).count()
    upcoming_excused = ExcusedAbsence.objects.filter(effective_date__gt=today).count()  # <-- Add this line

    grades = Grade.objects.all().order_by('name')
    sections = Section.objects.select_related('grade').order_by('grade__name', 'name')

    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')

    context = {
        'grades': grades,
        'sections': sections,
        'grade_filter': grade_filter,
        'section_filter': section_filter,
        'status_filter': status_filter,
        'date_filter': date_filter,
        'total_excused': total_excused,
        'active_excused': active_excused,
        'expired_excused': expired_excused,
        'upcoming_excused': upcoming_excused,  # <-- Add this line
    }
    return render(request, 'admin/excused.html', context)

@login_required
@user_passes_test(is_admin_or_principal)
@require_http_methods(["GET"])
def search_excused_absences(request):
    """AJAX search/filter for excused absences"""
    search_query = request.GET.get('search', '')
    grade_filter = request.GET.get('grade', '')
    section_filter = request.GET.get('section', '')
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')
    page_number = request.GET.get('page', 1)
    items_per_page = request.GET.get('items_per_page', 10)

    excused = ExcusedAbsence.objects.select_related('student', 'student__grade', 'student__section').all().order_by('-date_absent', '-student__last_name')

    if search_query:
        excused = excused.filter(
            Q(student__first_name__icontains=search_query) |
            Q(student__last_name__icontains=search_query) |
            Q(student__lrn__icontains=search_query) |
            Q(student__section__name__icontains=search_query)
        )
    if grade_filter:
        excused = excused.filter(student__grade_id=grade_filter)
    if section_filter:
        excused = excused.filter(student__section_id=section_filter)
    if status_filter:
        from datetime import date
        today = date.today()
        if status_filter == 'ACTIVE':
            excused = excused.filter(effective_date__lte=today, end_date__gte=today)
        elif status_filter == 'EXPIRED':
            excused = excused.filter(end_date__lt=today)
        elif status_filter == 'UPCOMING':
            excused = excused.filter(effective_date__gt=today)
    if date_filter:
        excused = excused.filter(date_absent=date_filter)

    total_count = excused.count()
    try:
        page_number = int(page_number)
        items_per_page = int(items_per_page)
    except (ValueError, TypeError):
        page_number = 1
        items_per_page = 10

    paginator = Paginator(excused, items_per_page)
    page_obj = paginator.get_page(page_number)

    excused_data = []
    for record in page_obj:
        # Determine status
        from datetime import date
        today = date.today()
        if record.effective_date > today:
            status = 'UPCOMING'
        elif record.end_date < today:
            status = 'EXPIRED'
        else:
            status = 'ACTIVE'
        excused_data.append({
            'id': record.id,
            'lrn': record.student.lrn,
            'full_name': f"{record.student.first_name} {record.student.last_name}",
            'grade': record.student.grade.name if record.student.grade else '',
            'section': record.student.section.name if record.student.section else '',
            'date_absent': record.date_absent.strftime('%Y-%m-%d'),
            'effective_date': record.effective_date.strftime('%Y-%m-%d'),
            'end_date': record.end_date.strftime('%Y-%m-%d'),
            'excuse_letter': record.excuse_letter,
            'status': status,
            'profile_pic': record.student.profile_pic or None,
        })

    pagination = {
        'current_page': page_obj.number,
        'num_pages': paginator.num_pages,
        'has_next': page_obj.has_next(),
        'has_previous': page_obj.has_previous(),
        'page_range': list(get_pagination_range(paginator, page_obj.number, 5)),
        'start_index': page_obj.start_index(),
        'end_index': page_obj.end_index(),
        'total_count': total_count,
    }

    return JsonResponse({
        'status': 'success',
        'records': excused_data,
        'pagination': pagination,
        'total_count': total_count,
    })

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def create_excused_absence(request):
    """API endpoint to create a new excused absence"""
    try:
        # Accept multipart/form-data for file upload
        student_id = request.POST.get('student_id')
        date_absent = request.POST.get('date_absent')
        effective_date = request.POST.get('effective_date')
        end_date = request.POST.get('end_date')
        excuse_letter_file = request.FILES.get('excuse_letter')

        # Validate required fields
        if not student_id or not date_absent or not effective_date or not end_date or not excuse_letter_file:
            return JsonResponse({'status': 'error', 'message': 'All fields are required'}, status=400)

        # Validate student
        try:
            student = Student.objects.get(id=student_id)
        except Student.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Student not found'}, status=404)

        # Validate file type and size
        if excuse_letter_file.size > 5 * 1024 * 1024:
            return JsonResponse({'status': 'error', 'message': 'File too large. Max 5MB.'}, status=400)
        if not excuse_letter_file.content_type.startswith('image/'):
            return JsonResponse({'status': 'error', 'message': 'Only image files allowed.'}, status=400)

        # Save file in private_excuse_letters
        import uuid
        filename = f"{uuid.uuid4()}_{excuse_letter_file.name}"
        fs = FileSystemStorage(location=PRIVATE_EXCUSE_LETTERS_DIR)
        filename = fs.save(filename, excuse_letter_file)

        # Create record
        record = ExcusedAbsence.objects.create(
            student=student,
            date_absent=date_absent,
            excuse_letter=filename,
            effective_date=effective_date,
            end_date=end_date
        )

        return JsonResponse({
            'status': 'success',
            'message': f'Excused absence for {student.first_name} {student.last_name} created successfully',
            'record_id': record.id
        })
    except Exception as e:
        import traceback
        print("Error in create_excused_absence:", str(e))
        print(traceback.format_exc())
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["GET"])
def get_excused_absence(request, excused_id):
    """API endpoint to get details of a specific excused absence"""
    try:
        record = get_object_or_404(ExcusedAbsence, id=excused_id)
        
        data = {
            'id': record.id,
            'student_id': record.student.id,
            'student_name': f"{record.student.first_name} {record.student.last_name}",
            'lrn': record.student.lrn,
            'date_absent': record.date_absent.strftime('%Y-%m-%d'),
            'effective_date': record.effective_date.strftime('%Y-%m-%d'),
            'end_date': record.end_date.strftime('%Y-%m-%d'),
            'excuse_letter': record.excuse_letter,
            'section': record.student.section.name if record.student.section else '',
            'section_id': record.student.section.id if record.student.section else None,
            'grade': record.student.grade.name if record.student.grade else '',
            'grade_id': record.student.grade.id if record.student.grade else None,
        }
        return JsonResponse({'status': 'success', 'record': data})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST", "PATCH"])
def update_excused_absence(request, excused_id):
    """API endpoint to update an excused absence"""
    try:
        record = get_object_or_404(ExcusedAbsence, id=excused_id)
        # Accept multipart/form-data for file upload
        student_id = request.POST.get('student_id', record.student.id)
        date_absent = request.POST.get('date_absent', record.date_absent)
        effective_date = request.POST.get('effective_date', record.effective_date)
        end_date = request.POST.get('end_date', record.end_date)
        excuse_letter_file = request.FILES.get('excuse_letter', None)

        # Validate student
        try:
            student = Student.objects.get(id=student_id)
        except Student.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Student not found'}, status=404)
        
        record.student = student
        record.date_absent = date_absent
        record.effective_date = effective_date
        record.end_date = end_date

        # Handle new file upload
        if excuse_letter_file:
            if excuse_letter_file.size > 5 * 1024 * 1024:
                return JsonResponse({'status': 'error', 'message': 'File too large. Max 5MB.'}, status=400)
            if not excuse_letter_file.content_type.startswith('image/'):
                return JsonResponse({'status': 'error', 'message': 'Only image files allowed.'}, status=400)
            # Delete old file
            if record.excuse_letter:
                old_path = os.path.join(PRIVATE_EXCUSE_LETTERS_DIR, record.excuse_letter)
                if os.path.exists(old_path):
                    os.remove(old_path)
            # Save new file
            import uuid
            filename = f"{uuid.uuid4()}_{excuse_letter_file.name}"
            fs = FileSystemStorage(location=PRIVATE_EXCUSE_LETTERS_DIR)
            filename = fs.save(filename, excuse_letter_file)
            record.excuse_letter = filename
        record.save()
        return JsonResponse({'status': 'success', 'message': 'Excused absence updated successfully'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["DELETE"])
def delete_excused_absence(request, excused_id):
    """API endpoint to delete an excused absence"""
    try:
        record = get_object_or_404(ExcusedAbsence, id=excused_id)
        # Delete file
        if record.excuse_letter:
            file_path = os.path.join(PRIVATE_EXCUSE_LETTERS_DIR, record.excuse_letter)
            if os.path.exists(file_path):
                os.remove(file_path)
        record.delete()
        return JsonResponse({'status': 'success', 'message': 'Excused absence deleted successfully'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


# ==========================
#  OTHER MANAGEMENT VIEWS
# ==========================

@login_required
@user_passes_test(is_admin)
def admin_dashboard(request):
    """View for admin dashboard"""
    context = {
        'current_date': timezone.now(),
        'user_count': CustomUser.objects.count(),
        'student_count': Student.objects.count(),
        'teacher_count': CustomUser.objects.filter(role=UserRole.TEACHER).count(),
    }
    return render(request, 'admin/dashboard.html', context)

@login_required
@user_passes_test(is_admin)
@login_required
@user_passes_test(is_admin)
def admin_settings(request):
    """View for system settings"""
    from PROTECHAPP.models import SystemSettings
    from django.utils import timezone
    import pytz
    
    # Get or create system settings
    settings_obj, created = SystemSettings.objects.get_or_create(pk=1)
    
    # Convert UTC late_time_cutoff to Asia/Manila for display
    manila_tz = pytz.timezone('Asia/Manila')
    utc_tz = pytz.UTC
    
    # Create a datetime object with today's date and the stored UTC time
    if settings_obj.late_time_cutoff:
        # Convert to time object if it's a string
        late_time = settings_obj.late_time_cutoff
        if isinstance(late_time, str):
            from datetime import datetime
            late_time = datetime.strptime(late_time, '%H:%M:%S').time()
        
        # Combine today's date with the time and make it timezone-aware in UTC
        utc_datetime = timezone.datetime.combine(
            timezone.now().date(),
            late_time
        )
        utc_datetime = utc_tz.localize(utc_datetime)
        
        # Convert to Manila timezone
        manila_datetime = utc_datetime.astimezone(manila_tz)
        late_cutoff_display = manila_datetime.strftime('%H:%M')
    else:
        late_cutoff_display = '08:00'
    
    context = {
        'current_mode': settings_obj.attendance_mode,
        'late_cutoff_time': late_cutoff_display
    }
    return render(request, 'admin/settings.html', context)


@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def save_attendance_mode(request):
    """Save the attendance mode setting"""
    from PROTECHAPP.models import SystemSettings
    
    attendance_mode = request.POST.get('attendance_mode')
    
    if attendance_mode in ['SEPARATE', 'HYBRID']:
        settings_obj, created = SystemSettings.objects.get_or_create(pk=1)
        settings_obj.attendance_mode = attendance_mode
        settings_obj.save()
        
        messages.success(request, f'Attendance mode successfully changed to {attendance_mode.title()} mode!')
    else:
        messages.error(request, 'Invalid attendance mode selected.')
    
    return redirect('admin_settings')


@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def save_late_time_cutoff(request):
    """Save the late time cutoff setting (converts Manila time to UTC)"""
    from PROTECHAPP.models import SystemSettings
    from django.utils import timezone
    import pytz
    from datetime import datetime
    
    late_time = request.POST.get('late_time_cutoff')
    
    if late_time:
        try:
            # Parse the time input (format: HH:MM)
            time_obj = datetime.strptime(late_time, '%H:%M').time()
            
            # Create a datetime object with today's date in Manila timezone
            manila_tz = pytz.timezone('Asia/Manila')
            manila_datetime = manila_tz.localize(
                datetime.combine(timezone.now().date(), time_obj)
            )
            
            # Convert to UTC
            utc_datetime = manila_datetime.astimezone(pytz.UTC)
            utc_time = utc_datetime.time()
            
            # Save to database
            settings_obj, created = SystemSettings.objects.get_or_create(pk=1)
            settings_obj.late_time_cutoff = utc_time
            settings_obj.save()
            
            messages.success(request, f'Late time cutoff successfully set to {late_time} (Manila Time)!')
        except ValueError:
            messages.error(request, 'Invalid time format. Please use HH:MM format.')
    else:
        messages.error(request, 'Please provide a valid time.')
    
    return redirect('admin_settings')


@login_required
@user_passes_test(is_admin)
@require_http_methods(["GET"])
def download_database_backup(request):
    """Generate and download database backup"""
    from PROTECHAPP.backup_utils import create_database_backup
    from django.http import FileResponse
    import mimetypes
    
    try:
        # Create backup with user info
        success, filepath, error_msg = create_database_backup(backup_type='MANUAL', user=request.user)
        
        if not success:
            # Format error message for better display
            if 'pg_dump command not found' in error_msg:
                messages.error(
                    request, 
                    'Backup failed: pg_dump command not found. Please ensure PostgreSQL client tools are installed.'
                )
            else:
                messages.error(request, f'Backup failed: {error_msg}')
            return redirect('admin_settings')
        
        # Serve file for download
        file_handle = open(filepath, 'rb')
        response = FileResponse(file_handle, content_type='application/sql')
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(filepath)}"'
        response['Content-Length'] = os.path.getsize(filepath)
        
        return response
        
    except Exception as e:
        messages.error(request, f'Error creating backup: {str(e)}')
        return redirect('admin_settings')


@login_required
@user_passes_test(is_admin)
@require_http_methods(["GET"])
def get_backup_status(request):
    """Get backup status and list of recent backups from database"""
    from PROTECHAPP.models import BackupLog
    
    try:
        # Get last 5 successful backups from database
        backups = BackupLog.objects.filter(status='SUCCESS').order_by('-created_at')[:5]
        
        # Format for JSON response
        backup_data = []
        for backup in backups:
            backup_data.append({
                'filename': backup.filename,
                'size_mb': float(backup.file_size_mb),
                'created_at': backup.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'backup_type': backup.get_backup_type_display(),
                'initiated_by': backup.initiated_by.get_full_name() if backup.initiated_by else 'System'
            })
        
        # Get total count of all backups
        total_count = BackupLog.objects.filter(status='SUCCESS').count()
        
        return JsonResponse({
            'status': 'success',
            'backups': backup_data,
            'total_count': total_count
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)


@login_required
@user_passes_test(is_admin)
def admin_messages(request):
    """View for messages"""
    from django.conf import settings
    context = {
        'firebase_config': settings.FIREBASE_WEB_CONFIG
    }
    return render(request, 'admin/messages.html', context)


@login_required
@user_passes_test(is_admin)
def admin_announcements(request):
    """View for announcements"""
    return render(request, 'admin/announcements.html')


# Create profile pics directory if it doesn't exist
PROFILE_PICS_DIR = os.path.join(settings.BASE_DIR, 'private_profile_pics')
os.makedirs(PROFILE_PICS_DIR, exist_ok=True)

# Create private_excuse_letters directory if it doesn't exist
PRIVATE_EXCUSE_LETTERS_DIR = os.path.join(settings.BASE_DIR, 'private_excuse_letters')
os.makedirs(PRIVATE_EXCUSE_LETTERS_DIR, exist_ok=True)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def upload_profile_pic(request):
    """API endpoint to upload a profile picture"""
    try:
        import uuid
        import traceback
        
        if 'profile_pic' not in request.FILES:
            return JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)
            
        # Check for user_id in POST data (required)
        user_id = request.POST.get('user_id')
        if not user_id:
            return JsonResponse({'status': 'error', 'message': 'User ID is required'}, status=400)
            
        # Verify user exists
        try:
            user = CustomUser.objects.get(id=user_id)
        except CustomUser.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'User not found'}, status=404)
            
        uploaded_file = request.FILES['profile_pic']
        
        # Validate file type
        allowed_types = ['image/jpeg', 'image/png', 'image/gif']
        if uploaded_file.content_type not in allowed_types:
            return JsonResponse({'status': 'error', 'message': 'Invalid file type. Only JPEG, PNG, and GIF images are allowed.'}, status=400)
            
        # Validate file size (max 5MB)
        if uploaded_file.size > 5 * 1024 * 1024:
            return JsonResponse({'status': 'error', 'message': 'File too large. Maximum file size is 5MB.'}, status=400)
        
        # Create a safe filename
        file_name = f"{uuid.uuid4()}_{uploaded_file.name}"
        
        # Save file to private directory
        fs = FileSystemStorage(location=PROFILE_PICS_DIR)
        filename = fs.save(file_name, uploaded_file)
        
        # Store just the filename in the user model, not the full path
        user.profile_pic = filename
        user.save()
        
        return JsonResponse({
            'status': 'success',
            'message': 'Profile picture uploaded successfully',
            'file_path': filename
        })
    
    except Exception as e:
        import traceback
        print(f"Error uploading profile picture: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
def serve_profile_pic(request, path):
    """View to securely serve profile pictures from private directory"""
    try:
        # Validate the path to prevent directory traversal
        if '..' in path or path.startswith('/'):
            return HttpResponseNotFound("Not found")
        
        full_path = os.path.join(PROFILE_PICS_DIR, path)
        if os.path.exists(full_path) and os.path.isfile(full_path):
            with open(full_path, 'rb') as f:
                return HttpResponse(f.read(), content_type='image/jpeg')  # Adjust content type if needed
        else:
            return HttpResponseNotFound("Image not found")
    except Exception as e:
        return HttpResponseNotFound(f"Error: {str(e)}")

@login_required
def serve_profile_pic_default(request):
    """Serve a default profile picture when none is specified"""
    # You can either redirect to a static default image or serve a placeholder
    default_image_path = os.path.join(settings.STATIC_ROOT, 'images', 'default_profile.png')
    
    if os.path.exists(default_image_path):
        with open(default_image_path, 'rb') as f:
            return HttpResponse(f.read(), content_type='image/png')
    else:
        # If default image doesn't exist, return a 404 or an empty response
        return HttpResponse(status=204)  # 204 No Content

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def upload_excuse_letter(request):
    """API endpoint to upload/replace an excuse letter for an excused absence"""
    try:
        excused_id = request.POST.get('id')
        file = request.FILES.get('excuse_letter')
        if not excused_id or not file:
            return JsonResponse({'status': 'error', 'message': 'Missing required fields'}, status=400)
        record = get_object_or_404(ExcusedAbsence, id=excused_id)
        # Validate file
        if file.size > 5 * 1024 * 1024:
            return JsonResponse({'status': 'error', 'message': 'File too large. Max 5MB.'}, status=400)
        if not file.content_type.startswith('image/'):
            return JsonResponse({'status': 'error', 'message': 'Only image files allowed.'}, status=400)
        # Delete old file if exists
        if record.excuse_letter:
            old_path = os.path.join(PRIVATE_EXCUSE_LETTERS_DIR, record.excuse_letter)
            if os.path.exists(old_path):
                os.remove(old_path)
        # Save new file
        import uuid
        filename = f"{uuid.uuid4()}_{file.name}"
        fs = FileSystemStorage(location=PRIVATE_EXCUSE_LETTERS_DIR)
        filename = fs.save(filename, file)
        record.excuse_letter = filename
        record.save()
        return JsonResponse({'status': 'success', 'filename': filename})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def delete_excuse_letter(request):
    """API endpoint to delete an excuse letter file for an excused absence"""
    try:
        data = json.loads(request.body)
        excused_id = data.get('id')
        if not excused_id:
            return JsonResponse({'status': 'error', 'message': 'Missing required fields'}, status=400)
        record = get_object_or_404(ExcusedAbsence, id=excused_id)
        if record.excuse_letter:
            file_path = os.path.join(PRIVATE_EXCUSE_LETTERS_DIR, record.excuse_letter)
            if os.path.exists(file_path):
                os.remove(file_path)
            record.excuse_letter = ''
            record.save()
        return JsonResponse({'status': 'success'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST", "PATCH"])
def update_excused_absence(request, excused_id):
    """API endpoint to update an excused absence"""
    try:
        record = get_object_or_404(ExcusedAbsence, id=excused_id)
        # Accept multipart/form-data for file upload
        student_id = request.POST.get('student_id', record.student.id)
        date_absent = request.POST.get('date_absent', record.date_absent)
        effective_date = request.POST.get('effective_date', record.effective_date)
        end_date = request.POST.get('end_date', record.end_date)
        excuse_letter_file = request.FILES.get('excuse_letter', None)

        # Validate student
        try:
            student = Student.objects.get(id=student_id)
        except Student.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Student not found'}, status=404)
        
        record.student = student
        record.date_absent = date_absent
        record.effective_date = effective_date
        record.end_date = end_date

        # Handle new file upload
        if excuse_letter_file:
            if excuse_letter_file.size > 5 * 1024 * 1024:
                return JsonResponse({'status': 'error', 'message': 'File too large. Max 5MB.'}, status=400)
            if not excuse_letter_file.content_type.startswith('image/'):
                return JsonResponse({'status': 'error', 'message': 'Only image files allowed.'}, status=400)
            # Delete old file
            if record.excuse_letter:
                old_path = os.path.join(PRIVATE_EXCUSE_LETTERS_DIR, record.excuse_letter)
                if os.path.exists(old_path):
                    os.remove(old_path)
            # Save new file
            import uuid
            filename = f"{uuid.uuid4()}_{excuse_letter_file.name}"
            fs = FileSystemStorage(location=PRIVATE_EXCUSE_LETTERS_DIR)
            filename = fs.save(filename, excuse_letter_file)
            record.excuse_letter = filename
        record.save()
        return JsonResponse({'status': 'success', 'message': 'Excused absence updated successfully'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["DELETE"])
def delete_excused_absence(request, excused_id):
    """API endpoint to delete an excused absence"""
    try:
        record = get_object_or_404(ExcusedAbsence, id=excused_id)
        # Delete file
        if record.excuse_letter:
            file_path = os.path.join(PRIVATE_EXCUSE_LETTERS_DIR, record.excuse_letter)
            if os.path.exists(file_path):
                os.remove(file_path)
        record.delete()
        return JsonResponse({'status': 'success', 'message': 'Excused absence deleted successfully'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def upload_excuse_letter(request):
    """API endpoint to upload an excuse letter"""
    try:
        excused_id = request.POST.get('id')
        file = request.FILES.get('excuse_letter')
        if not excused_id or not file:
            return JsonResponse({'status': 'error', 'message': 'Missing required fields'}, status=400)
        record = get_object_or_404(ExcusedAbsence, id=excused_id)
        # Validate file
        if file.size > 5 * 1024 * 1024:
            return JsonResponse({'status': 'error', 'message': 'File too large. Max 5MB.'}, status=400)
        if not file.content_type.startswith('image/'):
            return JsonResponse({'status': 'error', 'message': 'Only image files allowed.'}, status=400)
        # Delete old file if exists
        if record.excuse_letter:
            old_path = os.path.join(PRIVATE_EXCUSE_LETTERS_DIR, record.excuse_letter)
            if os.path.exists(old_path):
                os.remove(old_path)
        # Save new file
        import uuid
        filename = f"{uuid.uuid4()}_{file.name}"
        fs = FileSystemStorage(location=PRIVATE_EXCUSE_LETTERS_DIR)
        filename = fs.save(filename, file)
        record.excuse_letter = filename
        record.save()
        return JsonResponse({'status': 'success', 'filename': filename})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
# Serve files from private_excuse_letters directory
def serve_private_excuse_letter(request, filename):
    # Optionally add authentication/authorization checks here
    path = os.path.join(settings.MEDIA_ROOT, 'private_excuse_letters', filename)
    if not os.path.isfile(path):
        raise Http404("Excuse letter not found")
    return FileResponse(open(path, 'rb'), content_type='application/octet-stream')


# ==========================
#  EXPORT USERS TO MULTIPLE FORMATS
# ==========================

@login_required
@user_passes_test(is_admin)
@require_http_methods(["GET"])
def export_users_to_excel(request):
    """Export users/principals/registrars/teachers to Excel, PDF, or Word file"""
    try:
        # Get query parameters for filtering
        search_query = request.GET.get('search', '')
        role_filter = request.GET.get('role', '')
        status_filter = request.GET.get('status', '')
        export_format = request.GET.get('format', 'excel').lower()
        
        # Base queryset - order by created_at descending (latest to oldest) to match table display
        users = CustomUser.objects.all().order_by('-created_at')
        
        # Apply search if provided
        if search_query:
            users = users.filter(
                Q(username__icontains=search_query) | 
                Q(email__icontains=search_query) |
                Q(first_name__icontains=search_query) | 
                Q(last_name__icontains=search_query)
            )
        
        # Apply role filter if provided
        if role_filter:
            users = users.filter(role=role_filter)
        
        # Apply status filter
        if status_filter:
            if status_filter == 'active':
                users = users.filter(is_active=True)
            elif status_filter == 'inactive':
                users = users.filter(is_active=False)
        
        # Define headers (professional report format - no password for security)
        headers = ["First Name", "Middle Name", "Last Name", "Username", "Email", "Role", "Status", "Created At"]
        
        # Prepare data rows
        data_rows = []
        for user in users:
            data_rows.append([
                user.first_name,
                user.middle_name or "",
                user.last_name,
                user.username,
                user.email,
                user.get_role_display(),
                "Active" if user.is_active else "Inactive",
                user.created_at.strftime("%Y-%m-%d %H:%M:%S") if user.created_at else ""
            ])
        
        # Export based on format
        if export_format == 'pdf':
            return export_users_to_pdf(headers, data_rows)
        elif export_format == 'word':
            return export_users_to_word(headers, data_rows)
        else:  # Default to Excel
            return export_users_to_excel_format(headers, data_rows)
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


def export_users_to_excel_format(headers, data_rows):
    """Export users to Excel format"""
    # Create a workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Users"
    
    # Define styles
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    border_style = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    center_alignment = Alignment(horizontal="center", vertical="center")
    left_alignment = Alignment(horizontal="left", vertical="center")
    
    # Set column widths for professional layout
    column_widths = [15, 15, 15, 15, 30, 15, 12, 20]
    for col_num, width in enumerate(column_widths, 1):
        ws.column_dimensions[chr(64 + col_num)].width = width
    
    # Add institution header (row 1)
    institution_cell = ws.cell(row=1, column=1)
    institution_cell.value = "PROTECH - DETECT TO PROTECT"
    institution_cell.font = Font(bold=True, color="1F4E78", size=14)
    institution_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    ws.row_dimensions[1].height = 25
    
    # Add report title (row 2)
    title_cell = ws.cell(row=2, column=1)
    title_cell.value = "USER MANAGEMENT REPORT"
    title_cell.font = Font(bold=True, color="1F4E78", size=16)
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
    ws.row_dimensions[2].height = 30
    
    # Add timestamp (row 3)
    timestamp_cell = ws.cell(row=3, column=1)
    timestamp_cell.value = f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
    timestamp_cell.font = Font(color="666666", size=10)
    timestamp_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=len(headers))
    ws.row_dimensions[3].height = 20
    
    # Add summary statistics (row 4)
    total_users = len(data_rows)
    active_count = sum(1 for row in data_rows if row[6] == 'Active')
    inactive_count = total_users - active_count
    summary_cell = ws.cell(row=4, column=1)
    summary_cell.value = f"Total Users: {total_users} | Active: {active_count} | Inactive: {inactive_count}"
    summary_cell.font = Font(color="333333", size=10)
    summary_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.merge_cells(start_row=4, start_column=1, end_row=4, end_column=len(headers))
    ws.row_dimensions[4].height = 20
    
    # Add empty row for spacing (row 5)
    ws.row_dimensions[5].height = 10
    
    # Add table headers (row 6)
    header_row = 6
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col_num)
        cell.value = header
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = border_style
    
    # Set header row height
    ws.row_dimensions[header_row].height = 25
    
    # Add data rows (starting from row 7)
    for row_num, row_data in enumerate(data_rows, 7):
        for col_num, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num)
            cell.value = value
            cell.border = border_style
            
            # Apply alignment based on column type
            if col_num in [6, 7]:  # Role, Status
                cell.alignment = center_alignment
            else:
                cell.alignment = left_alignment
    
    # Freeze the header row (row 6)
    ws.freeze_panes = "A7"
    
    # Create HTTP response
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename=users_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    
    # Save workbook to response
    wb.save(response)
    
    return response


def export_users_to_pdf(headers, data_rows):
    """Export users to PDF format"""
    buffer = io.BytesIO()
    
    # Create PDF document with landscape orientation
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=30, leftMargin=30, topMargin=50, bottomMargin=30)
    
    # Container for PDF elements
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Institution header style
    institution_style = ParagraphStyle(
        'Institution',
        parent=styles['Normal'],
        fontSize=16,
        textColor=colors.HexColor('#1F4E78'),
        spaceAfter=4,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1F4E78'),
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#666666'),
        spaceAfter=4,
        alignment=TA_CENTER,
        fontName='Helvetica'
    )
    
    info_style = ParagraphStyle(
        'Info',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#333333'),
        spaceAfter=15,
        alignment=TA_CENTER,
        fontName='Helvetica'
    )
    
    # Add institution header
    institution = Paragraph("PROTECH - DETECT TO PROTECT", institution_style)
    elements.append(institution)
    
    # Add title
    title = Paragraph("USER MANAGEMENT REPORT", title_style)
    elements.append(title)
    
    # Add generation timestamp
    timestamp = Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", subtitle_style)
    elements.append(timestamp)
    
    # Add summary statistics
    total_users = len(data_rows)
    active_count = sum(1 for row in data_rows if row[6] == 'Active')
    inactive_count = total_users - active_count
    summary = Paragraph(f"Total Users: {total_users} | Active: {active_count} | Inactive: {inactive_count}", info_style)
    elements.append(summary)
    
    elements.append(Spacer(1, 0.2*inch))
    
    # Prepare table data
    table_data = [headers] + data_rows
    
    # Create table with adjusted column widths for A4 (landscape orientation)
    col_widths = [1.0*inch, 1.0*inch, 1.0*inch, 1.0*inch, 1.8*inch, 0.9*inch, 0.8*inch, 1.3*inch]
    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    
    # Define table style
    table_style = TableStyle([
        # Header style
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E78')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 14),
        ('TOPPADDING', (0, 0), (-1, 0), 14),
        
        # Data cells style - Left align text columns
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 1), (4, -1), 'LEFT'),  # First 5 columns left-aligned
        ('ALIGN', (5, 1), (-1, -1), 'CENTER'),  # Role, Status, Created At centered
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        
        # Grid
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('WORDWRAP', (0, 0), (-1, -1), True),
        
        # Alternating row colors
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F9F9F9')]),
    ])
    
    table.setStyle(table_style)
    elements.append(table)
    
    # Add footer with summary
    elements.append(Spacer(1, 0.4*inch))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#666666'),
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    footer = Paragraph(f"Total Users: {len(data_rows)}", footer_style)
    elements.append(footer)
    
    # Add separator line
    elements.append(Spacer(1, 0.1*inch))
    line_style = ParagraphStyle(
        'Line',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.HexColor('#999999'),
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique'
    )
    line = Paragraph("PROTECH - User Management System", line_style)
    elements.append(line)
    
    # Build PDF
    doc.build(elements)
    
    # Get PDF from buffer
    pdf = buffer.getvalue()
    buffer.close()
    
    # Create HTTP response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename=users_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
    response.write(pdf)
    
    return response


def export_users_to_word(headers, data_rows):
    """Export users to Word format"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    def set_cell_background(cell, fill_color):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def set_cell_border(cell, **kwargs):
        """Set cell borders with more visible styling to match PDF"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_elm = OxmlElement(f'w:{edge}')
            edge_elm.set(qn('w:val'), 'single')
            edge_elm.set(qn('w:sz'), '12')  # Increased from 4 to 12 for more visible borders
            edge_elm.set(qn('w:space'), '0')
            edge_elm.set(qn('w:color'), 'CCCCCC')  # Light gray border to match PDF
            tcBorders.append(edge_elm)
        tcPr.append(tcBorders)
    
    # Create Word document
    doc = Document()
    
    # Set narrow margins and landscape orientation to match PDF
    sections = doc.sections
    for section in sections:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = Inches(11.69)  # A4 landscape width
        section.page_height = Inches(8.27)  # A4 landscape height
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add institution header
    institution = doc.add_paragraph('PROTECH - DETECT TO PROTECT')
    institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
    institution_run = institution.runs[0]
    institution_run.font.color.rgb = RGBColor(31, 78, 120)  # #1F4E78
    institution_run.font.size = Pt(16)
    institution_run.font.bold = True
    institution_run.font.name = 'Arial'
    
    # Add title
    title = doc.add_heading('USER MANAGEMENT REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(31, 78, 120)  # #1F4E78
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = 'Arial'
    
    # Add timestamp
    timestamp = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp_run = timestamp.runs[0]
    timestamp_run.font.size = Pt(10)
    timestamp_run.font.color.rgb = RGBColor(102, 102, 102)  # #666666
    timestamp_run.font.name = 'Arial'
    
    # Add summary statistics
    total_users = len(data_rows)
    active_count = sum(1 for row in data_rows if row[6] == 'Active')
    inactive_count = total_users - active_count
    summary = doc.add_paragraph(f"Total Users: {total_users} | Active: {active_count} | Inactive: {inactive_count}")
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_run = summary.runs[0]
    summary_run.font.size = Pt(10)
    summary_run.font.color.rgb = RGBColor(51, 51, 51)  # #333333
    summary_run.font.name = 'Arial'
    summary_run.font.bold = False
    
    # Add spacing before table
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Create table with professional styling to match PDF
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Add headers with EXACT styling as PDF
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        cell.text = header
        
        # Set cell background color to match PDF exactly (#1F4E78)
        set_cell_background(cell, '1F4E78')
        
        # Style header text to match PDF
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                run.font.name = 'Arial'
        
        # Add cell border with visible styling
        set_cell_border(cell)
        
        # Add padding and alignment
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set cell margins for proper padding
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            node = OxmlElement(f'w:{margin_name}')
            node.set(qn('w:w'), '100')
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)
    
    # Add data rows
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.add_row().cells
        
        # Set alternating row colors to match PDF exactly (#F9F9F9)
        row_color = 'FFFFFF' if row_idx % 2 == 0 else 'F9F9F9'
        
        for idx, value in enumerate(row_data):
            cell = row_cells[idx]
            cell.text = str(value)
            
            # Set row background color
            set_cell_background(cell, row_color)
            
            # Add cell border to match PDF
            set_cell_border(cell)
            
            # Set cell margins for proper padding
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin_name in ['top', 'left', 'bottom', 'right']:
                node = OxmlElement(f'w:{margin_name}')
                node.set(qn('w:w'), '100')
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)
            
            # Style data cells
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.name = 'Arial'
                
                # Center align Role and Status columns
                if idx in [5, 6]:  # Role, Status
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Set column widths to match PDF layout
    column_widths = [Inches(1.1), Inches(1.1), Inches(1.1), Inches(1.1), Inches(2.0), Inches(1.0), Inches(0.8), Inches(1.4)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(column_widths):
                cell.width = column_widths[idx]
    
    # Add footer to match PDF exactly
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Total Users: {len(data_rows)}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(102, 102, 102)  # #666666
    footer_run.font.bold = True
    footer_run.font.name = 'Arial'
    
    # Add separator line
    doc.add_paragraph()
    system_name = doc.add_paragraph("PROTECH - User Management System")
    system_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    system_run = system_name.runs[0]
    system_run.font.size = Pt(8)
    system_run.font.color.rgb = RGBColor(153, 153, 153)  # #999999
    system_run.font.italic = True
    system_run.font.name = 'Arial'
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create HTTP response
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=users_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    
    return response


@login_required
@user_passes_test(is_admin_or_principal)
def export_teachers_to_excel(request):
    """Export teachers to Excel, PDF, or Word file"""
    try:
        # Get query parameters for filtering
        search_query = request.GET.get('search', '')
        advisory_filter = request.GET.get('advisory', '')
        status_filter = request.GET.get('status', '')
        export_format = request.GET.get('format', 'excel').lower()
        
        # Base queryset - only teachers (ordered by ID ascending)
        teachers = CustomUser.objects.filter(role=UserRole.TEACHER).order_by('id')
        
        # Apply search if provided
        if search_query:
            teachers = teachers.filter(
                Q(username__icontains=search_query) | 
                Q(email__icontains=search_query) |
                Q(first_name__icontains=search_query) | 
                Q(last_name__icontains=search_query)
            )
        
        # Apply advisory filter
        if advisory_filter:
            advisory_assignments = AdvisoryAssignment.objects.values_list('teacher_id', flat=True)
            if advisory_filter == 'advisory':
                teachers = teachers.filter(id__in=advisory_assignments)
            elif advisory_filter == 'non-advisory':
                teachers = teachers.exclude(id__in=advisory_assignments)
        
        # Apply status filter
        if status_filter:
            if status_filter == 'active':
                teachers = teachers.filter(is_active=True)
            elif status_filter == 'inactive':
                teachers = teachers.filter(is_active=False)
        
        # Define headers (include password column with [UNCHANGED] for backup/restore)
        headers = ["ID", "Username", "Email", "Password", "First Name", "Last Name", "Middle Name", "Advisory Status", "Grade", "Section", "Active Status", "Created Date"]
        
        # Get all advisory assignments for quick lookup
        advisory_assignments = AdvisoryAssignment.objects.values_list('teacher_id', flat=True).distinct()
        advisory_set = set(advisory_assignments)
        
        # Prepare data rows
        data_rows = []
        for teacher in teachers:
            # Get advisory/section info
            advisory_status = "Advisory" if teacher.id in advisory_set else "Non-Advisory"
            grade_name = teacher.section.grade.name if teacher.section else "-"
            section_name = teacher.section.name if teacher.section else "-"
            
            data_rows.append([
                str(teacher.id),
                teacher.username,
                teacher.email,
                "[UNCHANGED]",  # Password placeholder
                teacher.first_name,
                teacher.last_name,
                teacher.middle_name or "",
                advisory_status,
                grade_name,
                section_name,
                "Active" if teacher.is_active else "Inactive",
                teacher.created_at.strftime("%Y-%m-%d %H:%M:%S") if teacher.created_at else ""
            ])
        
        # Export based on format
        if export_format == 'pdf':
            return export_teachers_to_pdf(headers, data_rows)
        elif export_format == 'word':
            return export_teachers_to_word(headers, data_rows)
        else:  # Default to Excel
            return export_teachers_to_excel_format(headers, data_rows)
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


def export_teachers_to_excel_format(headers, data_rows):
    """Export teachers to Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Teachers"
    
    # Define styles
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    border_style = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    center_alignment = Alignment(horizontal="center", vertical="center")
    left_alignment = Alignment(horizontal="left", vertical="center")
    
    # Set column widths (updated for password column)
    column_widths = [8, 15, 25, 15, 15, 15, 15, 15, 12, 15, 12, 18]
    for col_num, width in enumerate(column_widths, 1):
        ws.column_dimensions[chr(64 + col_num)].width = width
    
    # Add headers
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = header
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = border_style
    
    ws.row_dimensions[1].height = 25
    
    # Add data rows
    for row_num, row_data in enumerate(data_rows, 2):
        for col_num, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num)
            cell.value = value
            cell.border = border_style
            
            # Apply alignment based on column type
            if col_num in [1, 4, 8, 10, 11]:  # ID, Password, Advisory Status, Section, Active Status
                cell.alignment = center_alignment
            else:
                cell.alignment = left_alignment
    
    # Freeze the header row
    ws.freeze_panes = "A2"
    
    # Create HTTP response
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename=teachers_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    
    # Save workbook to response
    wb.save(response)
    
    return response


def export_teachers_to_pdf(headers, data_rows):
    """Export teachers to PDF format"""
    buffer = io.BytesIO()
    
    # Create PDF document
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=20, leftMargin=20, topMargin=50, bottomMargin=30)
    
    # Container for PDF elements
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1F4E78'),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    # Add title
    title = Paragraph(f"Teachers Report - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", title_style)
    elements.append(title)
    elements.append(Spacer(1, 0.2*inch))
    
    # Prepare table data
    table_data = [headers] + data_rows
    
    # Create table with adjusted column widths for A4 - smaller for more columns
    col_widths = [0.3*inch, 0.7*inch, 1.2*inch, 0.7*inch, 0.7*inch, 0.7*inch, 0.8*inch, 0.5*inch, 0.6*inch, 0.6*inch, 0.9*inch]
    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    
    # Define table style
    table_style = TableStyle([
        # Header style
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E78')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 8),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('TOPPADDING', (0, 0), (-1, 0), 10),
        
        # Data cells style
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 1), (-1, -1), 'CENTER'),  # All columns centered
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 7),
        ('TOPPADDING', (0, 1), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 5),
        
        # Grid
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('WORDWRAP', (0, 0), (-1, -1), True),
        
        # Alternating row colors
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F5F5')]),
    ])
    
    table.setStyle(table_style)
    elements.append(table)
    
    # Add footer
    elements.append(Spacer(1, 0.3*inch))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.grey,
        alignment=TA_CENTER
    )
    footer = Paragraph(f"Total Teachers: {len(data_rows)}", footer_style)
    elements.append(footer)
    
    # Build PDF
    doc.build(elements)
    
    # Get PDF from buffer
    pdf = buffer.getvalue()
    buffer.close()
    
    # Create HTTP response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename=teachers_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
    response.write(pdf)
    
    return response


def export_teachers_to_word(headers, data_rows):
    """Export teachers to Word format"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    def set_cell_background(cell, fill_color):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Create Word document
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add title
    title = doc.add_heading('Teachers Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(31, 78, 120)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    
    # Add timestamp
    timestamp = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp_run = timestamp.runs[0]
    timestamp_run.font.size = Pt(10)
    timestamp_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Add spacing
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        cell.text = header
        
        # Set cell background color
        set_cell_background(cell, '1F4E78')
        
        # Style header text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Add data rows
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.add_row().cells
        
        # Set alternating row colors
        row_color = 'FFFFFF' if row_idx % 2 == 0 else 'F5F5F5'
        
        for idx, value in enumerate(row_data):
            cell = row_cells[idx]
            cell.text = str(value)
            
            # Set row background color
            set_cell_background(cell, row_color)
            
            # Style data cells
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Center align ID, Advisory Status, and Active Status columns
                if idx in [0, 6, 9]:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Set column widths
    column_widths = [Inches(0.4), Inches(1.0), Inches(1.5), Inches(0.9), Inches(0.9), Inches(0.9), Inches(1.0), Inches(0.6), Inches(0.9), Inches(0.8), Inches(1.2)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(column_widths):
                cell.width = column_widths[idx]
    
    # Add footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Total Teachers: {len(data_rows)}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.font.italic = True
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create HTTP response
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=teachers_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    
    return response


# ==========================
#  EXPORT GRADES TO MULTIPLE FORMATS
# ==========================

@login_required
@user_passes_test(is_admin)
@require_http_methods(["GET"])
def export_grades(request):
    """Export grades to Excel, PDF, or Word file"""
    try:
        export_format = request.GET.get('format', 'excel')
        search_query = request.GET.get('search', '').strip()
        
        grades = Grade.objects.all().order_by('id')
        
        if search_query:
            grades = grades.filter(name__icontains=search_query)
        
        headers = ["ID", "Grade Name", "Total Sections", "Total Students", "Created Date"]
        
        data_rows = []
        for grade in grades:
            section_count = grade.sections.count()
            student_count = Student.objects.filter(grade=grade).count()
            
            data_rows.append([
                str(grade.id),
                grade.name,
                str(section_count),
                str(student_count),
                grade.created_at.strftime("%Y-%m-%d %H:%M:%S") if hasattr(grade, 'created_at') and grade.created_at else ""
            ])
        
        if export_format == 'pdf':
            return export_grades_to_pdf(headers, data_rows)
        elif export_format == 'word':
            return export_grades_to_word(headers, data_rows)
        else:
            return export_grades_to_excel_format(headers, data_rows)
        
    except Exception as e:
        import traceback
        print(f"Export error: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def import_grades(request):
    """Import grades from CSV or Excel file"""
    try:
        import pandas as pd
        
        if 'import_file' not in request.FILES:
            return JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)
        
        file = request.FILES['import_file']
        file_ext = os.path.splitext(file.name)[1].lower()
        
        if file_ext not in ['.csv', '.xlsx', '.xls']:
            return JsonResponse({'status': 'error', 'message': 'Invalid file type'}, status=400)
        
        try:
            df = pd.read_csv(file) if file_ext == '.csv' else pd.read_excel(file)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': f'Error reading file: {str(e)}'}, status=400)
        
        df.columns = df.columns.str.strip()
        column_mapping = {
            'Grade Name': 'name',
            'ID': 'id',
            'Total Sections': 'total_sections',
            'Total Students': 'total_students',
            'Created Date': 'created_date'
        }
        df.rename(columns=column_mapping, inplace=True)
        
        if 'name' not in df.columns:
            return JsonResponse({'status': 'error', 'message': 'Missing required column: Grade Name'}, status=400)
        
        created_count = 0
        updated_count = 0
        skipped_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                name = str(row['name']).strip() if pd.notna(row['name']) else ''
                
                if not name:
                    errors.append(f'Row {index + 2}: Missing grade name')
                    skipped_count += 1
                    continue
                
                existing_grade = Grade.objects.filter(name__iexact=name).first()
                
                if existing_grade:
                    # Grade exists, just update (name doesn't change typically)
                    updated_count += 1
                else:
                    # Create new grade
                    Grade.objects.create(name=name)
                    created_count += 1
                    
            except Exception as e:
                errors.append(f'Row {index + 2}: {str(e)}')
                skipped_count += 1
                continue
        
        message_parts = []
        if created_count > 0:
            message_parts.append(f'{created_count} grade(s) created')
        if updated_count > 0:
            message_parts.append(f'{updated_count} grade(s) verified')
        if skipped_count > 0:
            message_parts.append(f'{skipped_count} row(s) skipped')
        
        return JsonResponse({
            'status': 'success',
            'message': ', '.join(message_parts) if message_parts else 'No changes made',
            'created': created_count,
            'updated': updated_count,
            'skipped': skipped_count,
            'skipped_rows': skipped_rows,
            'errors': errors[:10]
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'An unexpected error occurred: {str(e)}'}, status=500)


@login_required
@user_passes_test(is_admin)
@require_http_methods(["GET"])
def download_grades_template(request):
    """Generate and download Excel template for grade import"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import io
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Grades Import Template"
    
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    border_style = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    center_alignment = Alignment(horizontal="center", vertical="center")
    left_alignment = Alignment(horizontal="left", vertical="center")
    
    column_widths = [8, 20, 18, 18, 20]
    for idx, width in enumerate(column_widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    
    headers = ["ID", "Grade Name", "Total Sections", "Total Students", "Created Date"]
    ws.append(headers)
    ws.row_dimensions[1].height = 25
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = border_style
    
    sample_data = [
        ["", "Grade 7", "3", "90", ""],
        ["10", "Grade 8", "3", "85", "2025-11-20 10:00:00"],
        ["", "Grade 9", "2", "60", ""],
    ]
    
    for row_data in sample_data:
        ws.append(row_data)
    
    for row_idx in range(2, ws.max_row + 1):
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.border = border_style
            cell.alignment = center_alignment if col_idx in [1, 3, 4] else left_alignment
    
    ws.freeze_panes = ws['A2']
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="grades_import_template_{datetime.now().strftime("%Y%m%d")}.xlsx"'
    return response


def export_grades_to_excel_format(headers, data_rows):
    """Export grades to Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Grades"
    
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    border_style = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    center_alignment = Alignment(horizontal="center", vertical="center")
    left_alignment = Alignment(horizontal="left", vertical="center")
    
    column_widths = [8, 20, 18, 18, 20]
    for idx, width in enumerate(column_widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    
    ws.append(headers)
    ws.row_dimensions[1].height = 25
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = border_style
    
    for row_data in data_rows:
        ws.append(row_data)
    
    for row_idx in range(2, ws.max_row + 1):
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.border = border_style
            
            if col_idx in [1, 3, 4]:
                cell.alignment = center_alignment
            else:
                cell.alignment = left_alignment
    
    ws.freeze_panes = ws['A2']
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename=grades_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    
    return response


def export_grades_to_pdf(headers, data_rows):
    """Export grades to PDF format"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    elements = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#1F4E78'),
        spaceAfter=30,
        alignment=1
    )
    
    elements.append(Paragraph("Grades Report", title_style))
    elements.append(Spacer(1, 0.2*inch))
    
    table_data = [headers] + data_rows
    table = Table(table_data, colWidths=[0.5*inch, 1.5*inch, 1.2*inch, 1.2*inch, 1.5*inch])
    
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E78')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('WORDWRAP', (0, 0), (-1, -1), True),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F5F5')]),
    ]))
    
    elements.append(table)
    elements.append(Spacer(1, 0.3*inch))
    elements.append(Paragraph(f"Total Grades: {len(data_rows)}", styles['Normal']))
    
    doc.build(elements)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename=grades_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
    
    return response


def export_grades_to_word(headers, data_rows):
    """Export grades to Word format"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    def set_cell_background(cell, fill_color):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    doc = Document()
    
    title = doc.add_heading('Grades Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    for idx, header_text in enumerate(headers):
        cell = header_cells[idx]
        cell.text = header_text
        set_cell_background(cell, '1F4E78')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.add_row().cells
        row_color = 'FFFFFF' if row_idx % 2 == 0 else 'F5F5F5'
        
        for idx, value in enumerate(row_data):
            cell = row_cells[idx]
            cell.text = str(value)
            set_cell_background(cell, row_color)
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                if idx in [0, 2, 3]:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Total Grades: {len(data_rows)}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=grades_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    
    return response


# ==========================
#  EXPORT SECTIONS TO MULTIPLE FORMATS
# ==========================

@login_required
@user_passes_test(is_admin_or_principal)
@require_http_methods(["GET"])
def export_sections(request):
    """Export sections to Excel, PDF, or Word file"""
    try:
        export_format = request.GET.get('format', 'excel')
        search_query = request.GET.get('search', '').strip()
        grade_filter = request.GET.get('grade', '').strip()
        advisor_filter = request.GET.get('advisor', '').strip()
        
        # Use select_related for grade and prefetch advisory assignments (ordered by ID for backup/restore)
        sections = Section.objects.select_related('grade').prefetch_related(
            'advisory_assignments__teacher'
        ).all().order_by('id')
        
        if search_query:
            sections = sections.filter(
                Q(name__icontains=search_query) |
                Q(grade__name__icontains=search_query)
            )
        
        if grade_filter:
            sections = sections.filter(grade__id=grade_filter)
        
        if advisor_filter:
            if advisor_filter == 'with_advisor':
                sections = sections.filter(advisory_assignments__isnull=False).distinct()
            elif advisor_filter == 'without_advisor':
                sections = sections.filter(advisory_assignments__isnull=True)
        
        headers = ["ID", "Section Name", "Grade", "Room Number", "Capacity"]
        
        data_rows = []
        for section in sections:
            # Room number and capacity are optional fields on the import template.
            # The `Section` model does not currently store these by default, so export blank values.
            room_number = getattr(section, 'room_number', '') or ''
            capacity = getattr(section, 'capacity', '') or ''

            data_rows.append([
                str(section.id),
                section.name,
                section.grade.name if section.grade else "",
                room_number,
                str(capacity)
            ])
        
        if export_format == 'pdf':
            return export_sections_to_pdf(headers, data_rows)
        elif export_format == 'word':
            return export_sections_to_word(headers, data_rows)
        else:
            return export_sections_to_excel_format(headers, data_rows)
        
    except Exception as e:
        import traceback
        print(f"Export sections error: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


def export_sections_to_excel_format(headers, data_rows):
    """Export sections to Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Sections"
    
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    border_style = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    center_alignment = Alignment(horizontal="center", vertical="center")
    left_alignment = Alignment(horizontal="left", vertical="center")
    
    column_widths = [8, 20, 15, 12, 10]
    for idx, width in enumerate(column_widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    
    ws.append(headers)
    ws.row_dimensions[1].height = 25
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = border_style
    
    for row_data in data_rows:
        ws.append(row_data)
    
    for row_idx in range(2, ws.max_row + 1):
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.border = border_style
            
            if col_idx in [1, 5]:
                cell.alignment = center_alignment
            else:
                cell.alignment = left_alignment
    
    ws.freeze_panes = ws['A2']
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename=sections_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    
    return response


def export_sections_to_pdf(headers, data_rows):
    """Export sections to PDF format"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    elements = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#1F4E78'),
        spaceAfter=30,
        alignment=1
    )
    
    elements.append(Paragraph("Sections Report", title_style))
    elements.append(Spacer(1, 0.2*inch))
    
    table_data = [headers] + data_rows
    table = Table(table_data, colWidths=[0.4*inch, 1.2*inch, 0.8*inch, 1.5*inch, 1*inch, 1.2*inch])
    
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E78')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('WORDWRAP', (0, 0), (-1, -1), True),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F5F5')]),
    ]))
    
    elements.append(table)
    elements.append(Spacer(1, 0.3*inch))
    elements.append(Paragraph(f"Total Sections: {len(data_rows)}", styles['Normal']))
    
    doc.build(elements)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename=sections_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
    
    return response


def export_sections_to_word(headers, data_rows):
    """Export sections to Word format"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    def set_cell_background(cell, fill_color):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    doc = Document()
    
    title = doc.add_heading('Sections Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    for idx, header_text in enumerate(headers):
        cell = header_cells[idx]
        cell.text = header_text
        set_cell_background(cell, '1F4E78')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.add_row().cells
        row_color = 'FFFFFF' if row_idx % 2 == 0 else 'F5F5F5'
        
        for idx, value in enumerate(row_data):
            cell = row_cells[idx]
            cell.text = str(value)
            set_cell_background(cell, row_color)
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                if idx in [0, 4]:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Total Sections: {len(data_rows)}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=sections_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    
    return response


@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def import_sections(request):
    """Import sections from CSV or Excel file"""
    try:
        import pandas as pd
        
        if 'import_file' not in request.FILES:
            return JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)
        
        file = request.FILES['import_file']
        file_ext = os.path.splitext(file.name)[1].lower()
        
        if file_ext not in ['.csv', '.xlsx', '.xls']:
            return JsonResponse({'status': 'error', 'message': 'Invalid file type'}, status=400)
        
        try:
            df = pd.read_csv(file) if file_ext == '.csv' else pd.read_excel(file)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': f'Error reading file: {str(e)}'}, status=400)
        
        df.columns = df.columns.str.strip()
        column_mapping = {
            'Section Name': 'name',
            'Grade': 'grade',
            'Room Number': 'room_number',
            'Capacity': 'capacity',
            'Advisor': 'advisor',
            'Total Students': 'total_students',
            'ID': 'id',
            'Created Date': 'created_date'
        }
        df.rename(columns=column_mapping, inplace=True)
        
        required_columns = ['name', 'grade']
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            return JsonResponse({'status': 'error', 'message': f'Missing required columns: {", ".join(missing_columns)}'}, status=400)
        
        created_count = 0
        updated_count = 0
        skipped_count = 0
        skipped_rows = []
        errors = []
        
        for index, row in df.iterrows():
            try:
                name = str(row['name']).strip() if pd.notna(row['name']) else ''
                grade_name = str(row['grade']).strip() if pd.notna(row['grade']) else ''
                
                if not all([name, grade_name]):
                    errors.append(f'Row {index + 2}: Missing required field(s)')
                    skipped_count += 1
                    continue
                
                grade = Grade.objects.filter(name__iexact=grade_name).first()
                if not grade:
                    errors.append(f'Row {index + 2}: Grade "{grade_name}" not found')
                    skipped_count += 1
                    continue
                
                # Optional fields (not yet stored on model) - capture if present in file
                room_number = str(row['room_number']).strip() if 'room_number' in df.columns and pd.notna(row.get('room_number')) else ''
                capacity = None
                if 'capacity' in df.columns and pd.notna(row.get('capacity')):
                    try:
                        capacity = int(float(row.get('capacity')))
                    except Exception:
                        capacity = None

                # If an explicit ID column is provided and matches an existing Section, treat as existing
                row_id = None
                if 'id' in df.columns and pd.notna(row.get('id')):
                    try:
                        row_id = int(float(row.get('id')))
                    except Exception:
                        row_id = None

                existing_section = Section.objects.filter(name__iexact=name, grade=grade).first()

                if row_id and Section.objects.filter(id=row_id).exists():
                    skipped_count += 1
                    skipped_rows.append({
                        'row': index + 2,
                        'id': row_id,
                        'name': name,
                        'grade': grade_name,
                        'reason': 'ID already exists'
                    })
                elif existing_section:
                    # Respect uniqueness: do not create duplicate Section for same name+grade
                    skipped_count += 1
                    skipped_rows.append({
                        'row': index + 2,
                        'existing_id': existing_section.id,
                        'name': name,
                        'grade': grade_name,
                        'reason': 'Already exists'
                    })
                else:
                    # Create a new Section record
                    Section.objects.create(name=name, grade=grade)
                    created_count += 1
                    
            except Exception as e:
                errors.append(f'Row {index + 2}: {str(e)}')
                skipped_count += 1
                continue
        
        message_parts = []
        if created_count > 0:
            message_parts.append(f'{created_count} section(s) created')
        if updated_count > 0:
            message_parts.append(f'{updated_count} section(s) verified')
        if skipped_count > 0:
            message_parts.append(f'{skipped_count} section(s) skipped')
        
        return JsonResponse({
            'status': 'success',
            'message': ', '.join(message_parts) if message_parts else 'No changes made',
            'created': created_count,
            'updated': updated_count,
            'skipped': skipped_count,
            'skipped_rows': skipped_rows,
            'errors': errors[:10]
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'An unexpected error occurred: {str(e)}'}, status=500)


@login_required
@user_passes_test(is_admin)
@require_http_methods(["GET"])
def download_sections_template(request):
    """Generate and download Excel template for section import"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import io
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Sections Import Template"
    
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    border_style = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    center_alignment = Alignment(horizontal="center", vertical="center")
    left_alignment = Alignment(horizontal="left", vertical="center")
    
    column_widths = [8, 18, 15, 12, 10]
    for idx, width in enumerate(column_widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    
    headers = ["ID", "Section Name", "Grade", "Room Number", "Capacity"]
    ws.append(headers)
    ws.row_dimensions[1].height = 25
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = border_style
    
    sample_data = [
        ["", "Section A", "Grade 7", "101", "40"],
        ["25", "Section B", "Grade 7", "102", "40"],
        ["", "Section A", "Grade 8", "201", "35"],
    ]
    
    for row_data in sample_data:
        ws.append(row_data)
    
    for row_idx in range(2, ws.max_row + 1):
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.border = border_style
            # center ID and Capacity (cols A and E)
            cell.alignment = center_alignment if col_idx in [1, 5] else left_alignment
    
    ws.freeze_panes = ws['A2']
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="sections_import_template_{datetime.now().strftime("%Y%m%d")}.xlsx"'
    return response


# ==========================
#  EXPORT GUARDIANS TO MULTIPLE FORMATS
# ==========================

@login_required
@user_passes_test(is_admin_or_principal)
@require_http_methods(["GET"])
def export_guardians(request):
    """Export guardians to Excel, PDF, or Word file"""
    try:
        export_format = request.GET.get('format', 'excel').lower()
        search_query = request.GET.get('search', '').strip()
        relationship_filter = request.GET.get('relationship', '').strip()
        grade_filter = request.GET.get('grade', '').strip()
        section_filter = request.GET.get('section', '').strip()
        
        guardians = Guardian.objects.select_related('student', 'student__grade', 'student__section').all().order_by('id')
        
        if search_query:
            guardians = guardians.filter(
                Q(first_name__icontains=search_query) |
                Q(last_name__icontains=search_query) |
                Q(email__icontains=search_query) |
                Q(phone__icontains=search_query)
            )
        
        if relationship_filter:
            guardians = guardians.filter(relationship=relationship_filter)
        
        if grade_filter:
            guardians = guardians.filter(student__grade__id=grade_filter)
        
        if section_filter:
            guardians = guardians.filter(student__section__id=section_filter)
        
        headers = ["ID", "Guardian Name", "Email", "Phone", "Relationship", "Student", "Created Date"]
        
        data_rows = []
        for guardian in guardians:
            full_name = f"{guardian.first_name} {guardian.last_name}"
            student_name = f"{guardian.student.first_name} {guardian.student.last_name}" if guardian.student else "N/A"
            
            data_rows.append([
                str(guardian.id),
                full_name,
                guardian.email or "",
                guardian.phone or "",
                guardian.relationship,
                student_name,
                guardian.created_at.strftime("%Y-%m-%d %H:%M:%S") if guardian.created_at else ""
            ])
        
        if export_format == 'pdf':
            return export_guardians_to_pdf(headers, data_rows)
        elif export_format == 'word':
            return export_guardians_to_word(headers, data_rows)
        else:  # Default to Excel
            return export_guardians_to_excel_format(headers, data_rows)
        
    except Exception as e:
        import traceback
        print(f"Export guardians error: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


def export_guardians_to_excel_format(headers, data_rows):
    """Export guardians to Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Guardians"
    
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    border_style = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    center_alignment = Alignment(horizontal="center", vertical="center")
    left_alignment = Alignment(horizontal="left", vertical="center")
    
    column_widths = [8, 25, 30, 18, 18, 18, 20]
    for idx, width in enumerate(column_widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    
    ws.append(headers)
    ws.row_dimensions[1].height = 25
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = border_style
    
    for row_data in data_rows:
        ws.append(row_data)
    
    for row_idx in range(2, ws.max_row + 1):
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.border = border_style
            
            if col_idx in [1, 6]:
                cell.alignment = center_alignment
            else:
                cell.alignment = left_alignment
    
    ws.freeze_panes = ws['A2']
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="guardians_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx"'
    
    return response


def export_guardians_to_pdf(headers, data_rows):
    """Export guardians to PDF format"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    elements = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#1F4E78'),
        spaceAfter=30,
        alignment=1
    )
    
    elements.append(Paragraph("Guardians Report", title_style))
    elements.append(Spacer(1, 0.2*inch))
    
    table_data = [headers] + data_rows
    table = Table(table_data, colWidths=[0.3*inch, 1.5*inch, 1.5*inch, 1*inch, 1*inch, 0.7*inch, 1.1*inch])
    
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E78')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('WORDWRAP', (0, 0), (-1, -1), True),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F5F5')]),
    ]))
    
    elements.append(table)
    elements.append(Spacer(1, 0.3*inch))
    elements.append(Paragraph(f"Total Guardians: {len(data_rows)}", styles['Normal']))
    
    doc.build(elements)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="guardians_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf"'
    
    return response


def export_guardians_to_word(headers, data_rows):
    """Export guardians to Word format"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    def set_cell_background(cell, fill_color):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    doc = Document()
    
    title = doc.add_heading('Guardians Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    for idx, header_text in enumerate(headers):
        cell = header_cells[idx]
        cell.text = header_text
        set_cell_background(cell, '1F4E78')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.add_row().cells
        row_color = 'FFFFFF' if row_idx % 2 == 0 else 'F5F5F5'
        
        for idx, value in enumerate(row_data):
            cell = row_cells[idx]
            cell.text = str(value)
            set_cell_background(cell, row_color)
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                if idx in [0, 5]:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Total Guardians: {len(data_rows)}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="guardians_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
    
    return response


@login_required
@user_passes_test(is_admin_or_principal)
@require_http_methods(["POST"])
def import_guardians(request):
    """Import guardians from CSV or Excel file"""
    try:
        import pandas as pd
        
        if 'import_file' not in request.FILES and 'file' not in request.FILES:
            return JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)

        file = request.FILES.get('import_file') or request.FILES.get('file')
        file_ext = os.path.splitext(file.name)[1].lower()

        if file_ext not in ['.csv', '.xlsx', '.xls']:
            return JsonResponse({'status': 'error', 'message': 'Invalid file type'}, status=400)

        try:
            df = pd.read_csv(file) if file_ext == '.csv' else pd.read_excel(file)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': f'Error reading file: {str(e)}'}, status=400)

        # Normalize column names for flexible matching
        df.columns = [str(c).strip() for c in df.columns]
        normalized_cols = {c: re.sub(r'[^a-z0-9]', '', c.strip().lower()) for c in df.columns}

        def find_col(*candidates):
            for cand in candidates:
                for orig, nk in normalized_cols.items():
                    if nk == re.sub(r'[^a-z0-9]', '', cand.strip().lower()):
                        return orig
            # also try direct normalized candidate names
            for orig, nk in normalized_cols.items():
                for cand in candidates:
                    if nk == cand:
                        return orig
            return None

        # Candidate header names
        name_col = find_col('guardian name', 'name', 'guardian')
        email_col = find_col('email', 'guardian email')
        phone_col = find_col('phone', 'guardian phone', 'mobile')
        rel_col = find_col('relationship', 'relation', 'guardian relationship')
        student_col = find_col('student lrn', 'studentlrn', 'lrn', 'student')
        id_col = find_col('id')

        if not name_col or not rel_col or not student_col:
            return JsonResponse({'status': 'error', 'message': 'Missing required columns. Expected: Guardian Name, Relationship, and Student (LRN or full name).'}, status=400)

        created_count = 0
        updated_count = 0
        skipped_rows = []
        errors = []

        # Cache students by LRN for fast lookup
        students_by_lrn = {s.lrn: s for s in Student.objects.all()}

        for index, row in df.iterrows():
            row_num = int(index) + 2
            try:
                def _val(col):
                    if not col or col not in df.columns:
                        return ''
                    v = row[col]
                    return str(v).strip() if pd.notna(v) else ''

                full_name = _val(name_col)
                email = _val(email_col)
                phone = _val(phone_col)
                relationship = _val(rel_col)
                student_val = _val(student_col)

                if not full_name or not relationship:
                    skipped_rows.append({'row': row_num, 'reason': 'Missing guardian name or relationship'})
                    continue

                # Prefer LRN lookup if the student column looks like LRN (12 digits) or matches a known LRN
                student = None
                candidate_lrn = re.sub(r'[^0-9]', '', student_val or '')
                if candidate_lrn and len(candidate_lrn) == 12 and candidate_lrn in students_by_lrn:
                    student = students_by_lrn[candidate_lrn]
                else:
                    # try matching by student full name if provided in the cell
                    sval = student_val
                    if sval and sval.upper() != 'N/A':
                        parts = sval.split()
                        if len(parts) >= 2:
                            first = parts[0]
                            last = parts[-1]
                            student = Student.objects.filter(first_name__iexact=first, last_name__iexact=last).first()

                if not student:
                    skipped_rows.append({'row': row_num, 'student': student_val, 'reason': 'Student not found'})
                    continue

                # split guardian name
                name_parts = full_name.split()
                first_name = name_parts[0] if len(name_parts) > 0 else ''
                last_name = name_parts[-1] if len(name_parts) > 1 else ''

                if not first_name or not last_name:
                    skipped_rows.append({'row': row_num, 'reason': 'Invalid guardian name'})
                    continue

                # Check if guardian already exists for this student (prefer ID lookup if provided)
                existing_guardian = None
                try:
                    # If ID column exists and has a value, prefer that lookup
                    if id_col and id_col in df.columns and row.get(id_col) and pd.notna(row.get(id_col)):
                        try:
                            gid = int(str(row.get(id_col)).strip())
                            existing_guardian = Guardian.objects.filter(id=gid).first()
                        except Exception:
                            existing_guardian = None

                    if not existing_guardian:
                        # Try matching by email (strong match)
                        if email:
                            existing_guardian = Guardian.objects.filter(email__iexact=email, student=student).first()

                    if not existing_guardian:
                        # Try matching by phone (normalized digits)
                        if phone:
                            candidate_digits = re.sub(r'[^0-9]', '', phone)
                            if candidate_digits:
                                for g in Guardian.objects.filter(student=student):
                                    gphone = getattr(g, 'phone', None) or getattr(g, 'phone_number', None) or ''
                                    gphone_digits = re.sub(r'[^0-9]', '', str(gphone) or '')
                                    if gphone_digits and gphone_digits == candidate_digits:
                                        existing_guardian = g
                                        break

                    if not existing_guardian:
                        # Try normalized full-name match against existing guardians for this student
                        try:
                            norm_full = re.sub(r'[^a-z0-9]', '', full_name.strip().lower())
                            for g in Guardian.objects.filter(student=student):
                                gname = f"{(g.first_name or '').strip()} {(g.last_name or '').strip()}".strip()
                                if gname:
                                    if re.sub(r'[^a-z0-9]', '', gname.lower()) == norm_full:
                                        existing_guardian = g
                                        break
                        except Exception:
                            pass

                    if not existing_guardian:
                        existing_guardian = Guardian.objects.filter(first_name__iexact=first_name, last_name__iexact=last_name, student=student).first()
                except Exception:
                    existing_guardian = None

                # Prepare create/update data
                model_field_names = {f.name for f in Guardian._meta.get_fields() if getattr(f, 'concrete', True) and not (hasattr(f, 'many_to_many') and f.many_to_many)}
                row_data = {'first_name': first_name, 'last_name': last_name, 'email': email, 'phone': phone, 'relationship': relationship, 'student': student}

                if existing_guardian:
                    # Update only non-empty fields that differ from existing values
                    update_fields = {}
                    # email
                    if email and (not existing_guardian.email or existing_guardian.email.strip().lower() != email.strip().lower()):
                        update_fields['email'] = email
                    # phone
                    if phone and (not getattr(existing_guardian, 'phone', None) or str(getattr(existing_guardian, 'phone')).strip() != phone.strip()):
                        # attempt to set common phone field names
                        if 'phone' in model_field_names:
                            update_fields['phone'] = phone
                        elif 'phone_number' in model_field_names:
                            update_fields['phone_number'] = phone
                    # relationship
                    if relationship and (not getattr(existing_guardian, 'relationship', None) or str(getattr(existing_guardian, 'relationship')).strip().lower() != relationship.strip().lower()):
                        update_fields['relationship'] = relationship

                    if update_fields:
                        for k, v in update_fields.items():
                            setattr(existing_guardian, k, v)
                        existing_guardian.save()
                        updated_count += 1
                    else:
                        skipped_rows.append({'row': row_num, 'guardian': full_name, 'reason': 'No new data'})
                    continue

                # Create new guardian when none exists
                create_kwargs = {k: v for k, v in row_data.items() if k in model_field_names}
                Guardian.objects.create(**create_kwargs)
                created_count += 1

            except Exception as e:
                errors.append(f'Row {row_num}: {str(e)}')
                continue

        # Build response
        message_parts = []
        if created_count:
            message_parts.append(f'{created_count} guardian(s) created')
        if updated_count:
            message_parts.append(f'{updated_count} guardian(s) updated')
        if skipped_rows:
            message_parts.append(f'{len(skipped_rows)} row(s) skipped')

        return JsonResponse({
            'status': 'success',
            'message': ', '.join(message_parts) if message_parts else 'No changes made',
            'created': created_count,
            'updated': updated_count,
            'skipped_rows': skipped_rows,
            'errors': errors[:10]
        })

    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'An unexpected error occurred: {str(e)}'}, status=500)


@login_required
@user_passes_test(is_admin_or_principal)
@require_http_methods(["GET"])
def download_guardians_template(request):
    """Generate and download Excel template for guardian import"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import io
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Guardians Import Template"
    
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    border_style = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    center_alignment = Alignment(horizontal="center", vertical="center")
    left_alignment = Alignment(horizontal="left", vertical="center")
    
    column_widths = [8, 25, 30, 18, 15, 25, 20]
    for idx, width in enumerate(column_widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    
    headers = ["ID", "Guardian Name", "Email", "Phone", "Relationship", "Student", "Created Date"]
    ws.append(headers)
    ws.row_dimensions[1].height = 25
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = border_style
    
    sample_data = [
        ["", "John Doe", "john.doe@email.com", "+1234567890", "Father", "Jane Doe", ""],
        ["50", "Maria Smith", "maria@email.com", "+0987654321", "Mother", "Bob Smith", "2025-11-20 10:00:00"],
        ["", "Robert Wilson", "robert@email.com", "+1122334455", "Guardian", "Alice Wilson", ""],
    ]
    
    for row_data in sample_data:
        ws.append(row_data)
    
    for row_idx in range(2, ws.max_row + 1):
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.border = border_style
            cell.alignment = center_alignment if col_idx in [1, 5] else left_alignment
    
    ws.freeze_panes = ws['A2']
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="guardians_import_template_{datetime.now().strftime("%Y%m%d")}.xlsx"'
    return response


# ==========================
#  EXPORT EXCUSED ABSENCES TO MULTIPLE FORMATS
# ==========================

@login_required
@user_passes_test(is_admin)
@require_http_methods(["GET"])
def export_excused(request):
    """Export excused absences to Excel, PDF, or Word file"""
    try:
        export_format = request.GET.get('format', 'excel')
        search_query = request.GET.get('search', '').strip()
        grade_filter = request.GET.get('grade', '').strip()
        section_filter = request.GET.get('section', '').strip()
        status_filter = request.GET.get('status', '').strip()
        date_filter = request.GET.get('date', '').strip()
        
        from django.utils import timezone
        
        excused = ExcusedAbsence.objects.select_related('student', 'student__grade', 'student__section').all().order_by('-date_absent')
        
        if search_query:
            excused = excused.filter(
                Q(student__first_name__icontains=search_query) |
                Q(student__last_name__icontains=search_query) |
                Q(student__lrn__icontains=search_query) |
                Q(student__section__name__icontains=search_query)
            )
        
        if grade_filter:
            excused = excused.filter(student__grade__id=grade_filter)
        
        if section_filter:
            excused = excused.filter(student__section__id=section_filter)
        
        if status_filter:
            today = timezone.now().date()
            if status_filter == 'ACTIVE':
                excused = excused.filter(date_absent__lte=today, return_date__gte=today)
            elif status_filter == 'UPCOMING':
                excused = excused.filter(date_absent__gt=today)
            elif status_filter == 'EXPIRED':
                excused = excused.filter(return_date__lt=today)
        
        if date_filter:
            excused = excused.filter(date_absent=date_filter)
        
        headers = ["ID", "Student Name", "LRN", "Grade", "Section", "Date Absent", "Effective Date", "End Date", "Has Letter"]
        
        data_rows = []
        for record in excused:
            student_name = f"{record.student.first_name} {record.student.last_name}"
            has_letter = "Yes" if record.excuse_letter else "No"
            
            data_rows.append([
                str(record.id),
                student_name,
                record.student.lrn or "",
                record.student.grade.name if record.student.grade else "",
                record.student.section.name if record.student.section else "",
                record.date_absent.strftime("%Y-%m-%d") if record.date_absent else "",
                record.effective_date.strftime("%Y-%m-%d") if record.effective_date else "",
                record.end_date.strftime("%Y-%m-%d") if record.end_date else "",
                has_letter
            ])
        
        if export_format == 'pdf':
            return export_excused_to_pdf(headers, data_rows)
        elif export_format == 'word':
            return export_excused_to_word(headers, data_rows)
        else:
            return export_excused_to_excel_format(headers, data_rows)
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


def export_excused_to_excel_format(headers, data_rows):
    """Export excused absences to Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Excused Absences"
    
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    border_style = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    center_alignment = Alignment(horizontal="center", vertical="center")
    left_alignment = Alignment(horizontal="left", vertical="center")
    
    column_widths = [8, 25, 15, 12, 15, 15, 15, 35, 12]
    for idx, width in enumerate(column_widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    
    ws.append(headers)
    ws.row_dimensions[1].height = 25
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = border_style
    
    for row_data in data_rows:
        ws.append(row_data)
    
    for row_idx in range(2, ws.max_row + 1):
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.border = border_style
            
            if col_idx in [1, 9]:
                cell.alignment = center_alignment
            else:
                cell.alignment = left_alignment
    
    ws.freeze_panes = ws['A2']
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename=excused_absences_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    
    return response


def export_excused_to_pdf(headers, data_rows):
    """Export excused absences to PDF format"""
    from reportlab.lib.pagesizes import landscape
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), topMargin=0.5*inch, bottomMargin=0.5*inch)
    elements = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#1F4E78'),
        spaceAfter=30,
        alignment=1
    )
    
    elements.append(Paragraph("Excused Absences Report", title_style))
    elements.append(Spacer(1, 0.2*inch))
    
    table_data = [headers] + data_rows
    table = Table(table_data, colWidths=[0.4*inch, 1.5*inch, 0.9*inch, 0.7*inch, 0.9*inch, 1*inch, 1*inch, 2*inch, 0.7*inch])
    
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E78')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('WORDWRAP', (0, 0), (-1, -1), True),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F5F5')]),
    ]))
    
    elements.append(table)
    elements.append(Spacer(1, 0.3*inch))
    elements.append(Paragraph(f"Total Excused Absences: {len(data_rows)}", styles['Normal']))
    
    doc.build(elements)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename=excused_absences_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
    
    return response


def export_excused_to_word(headers, data_rows):
    """Export excused absences to Word format"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    def set_cell_background(cell, fill_color):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    doc = Document()
    
    title = doc.add_heading('Excused Absences Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    for idx, header_text in enumerate(headers):
        cell = header_cells[idx]
        cell.text = header_text
        set_cell_background(cell, '1F4E78')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.add_row().cells
        row_color = 'FFFFFF' if row_idx % 2 == 0 else 'F5F5F5'
        
        for idx, value in enumerate(row_data):
            cell = row_cells[idx]
            cell.text = str(value)
            set_cell_background(cell, row_color)
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                if idx in [0, 8]:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Total Excused Absences: {len(data_rows)}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=excused_absences_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    
    return response

@login_required
@user_passes_test(is_admin)
def download_students_import_template(request):
    """Generate Excel template for importing students"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Students"
        
        # Headers (match export format: Last Name before First Name)
        headers = ['ID', 'LRN', 'Last Name', 'First Name', 'Middle Name', 'Grade', 'Section', 'Status',
               'Face Enrolled', 'Created Date']
        
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.value = header
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Add instruction row (pro tips / column hints) matching export columns
        instruction_values = [
            'Leave blank for new',
            '12 digits',
            'Required',
            'Required',
            'Optional',
            'Required (e.g., Grade 7)',
            'Required (existing section)',
            'Active or Inactive (case-insensitive)',
            'Yes or No',
            'YYYY-MM-DD HH:MM:SS'
        ]

        for col_num, instruction in enumerate(instruction_values, 1):
            cell = ws.cell(row=2, column=col_num)
            cell.value = instruction
            cell.font = Font(italic=True, color="808080")
            cell.fill = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")

        # Add sample data rows (match exported view: ID present, Title Case Status, Face Enrolled, Created Date)
        sample_rows = [
            ['2139', '097787337401', 'Robertson', 'Kristi', '', 'Grade 7', 'Mabini', 'Active', 'No', '2025-10-01 04:03:27'],
            ['2140', '034505903688', 'Holloway', 'Megan', 'Lisa', 'Grade 7', 'Mabini', 'Active', 'No', '2025-10-01 04:03:27']
        ]

        start_sample_row = 3
        for r_idx, row in enumerate(sample_rows, start_sample_row):
            for c_idx, value in enumerate(row, 1):
                cell = ws.cell(row=r_idx, column=c_idx)
                cell.value = value
                cell.alignment = Alignment(horizontal="left", vertical="center")

        # Add a Pro Tips merged block below the sample rows
        tips_row = start_sample_row + len(sample_rows) + 1
        tips_text = (
            "Pro Tips:\n"
            "- ID is exported for existing records; leave blank for new students.\n"
            "- LRN must be exactly 12 digits.\n"
            "- Status may be 'Active' or 'Inactive' (case-insensitive).\n"
            "- Face Enrolled will be 'Yes' or 'No' and is informational only.\n"
            "- Created Date is informational (do not modify when re-importing)."
        )
        ws.merge_cells(start_row=tips_row, start_column=1, end_row=tips_row, end_column=len(headers))
        tips_cell = ws.cell(row=tips_row, column=1)
        tips_cell.value = tips_text
        tips_cell.font = Font(italic=True, color="808080")
        tips_cell.alignment = Alignment(wrap_text=True, horizontal="left", vertical="top")
        
        # Adjust column widths for export-style layout
        ws.column_dimensions['A'].width = 10  # ID
        ws.column_dimensions['B'].width = 15  # LRN
        ws.column_dimensions['C'].width = 20  # Last Name
        ws.column_dimensions['D'].width = 18  # First Name
        ws.column_dimensions['E'].width = 14  # Middle Name
        ws.column_dimensions['F'].width = 12  # Grade
        ws.column_dimensions['G'].width = 12  # Section
        ws.column_dimensions['H'].width = 12  # Status
        ws.column_dimensions['I'].width = 12  # Face Enrolled
        ws.column_dimensions['J'].width = 22  # Created Date
        
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        response = HttpResponse(
            output.read(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename="student_import_template.xlsx"'
        return response
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Error generating template: {str(e)}'}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def import_students(request):
    """Import students from Excel file"""
    try:
        import openpyxl
        import re
        from django.db import transaction
        
        # Accept both 'import_file' and legacy 'file' keys from different templates
        if 'import_file' in request.FILES:
            excel_file = request.FILES['import_file']
        elif 'file' in request.FILES:
            excel_file = request.FILES['file']
        else:
            return JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)
        
        # Validate file type
        if not excel_file.name.endswith(('.xlsx', '.xls', '.csv')):
            return JsonResponse({
                'status': 'error',
                'message': 'Invalid file type. Please upload Excel (.xlsx, .xls) or CSV files only.'
            }, status=400)
        
        # Validate file size (max 10MB)
        if excel_file.size > 10 * 1024 * 1024:
            return JsonResponse({
                'status': 'error',
                'message': 'File too large. Maximum file size is 10MB.'
            }, status=400)
        
        try:
            # Read Excel file
            wb = openpyxl.load_workbook(excel_file)
            ws = wb.active
            
            # Read header row and build flexible, order-independent mapping
            headers = []
            for col in range(1, ws.max_column + 1):
                cell_value = ws.cell(row=1, column=col).value
                if cell_value:
                    headers.append(str(cell_value).strip())

            # Helper to normalize header names
            def _normalize(h):
                return re.sub(r'[^a-z0-9]', '', h.strip().lower()) if h else ''

            # Synonyms mapping for student import headers -> canonical keys
            synonyms = {
                'id': 'id',
                'studentlrn': 'lrn', 'lrn': 'lrn', 'learnerreference': 'lrn', 'learnerref': 'lrn',
                'firstname': 'first_name', 'first name': 'first_name', 'givenname': 'first_name',
                'middlename': 'middle_name', 'middle name': 'middle_name',
                'lastname': 'last_name', 'last name': 'last_name', 'surname': 'last_name',
                'grade': 'grade',
                'section': 'section',
                'status': 'status',
                'guardianfirstname': 'guardian_first_name', 'guardian first name': 'guardian_first_name',
                'guardianmiddlename': 'guardian_middle_name', 'guardian middle name': 'guardian_middle_name',
                'guardianlastname': 'guardian_last_name', 'guardian last name': 'guardian_last_name',
                'guardianemail': 'guardian_email', 'guardian phone': 'guardian_phone', 'guardianphone': 'guardian_phone',
                'guardianrelationship': 'guardian_relationship', 'alignment': 'alignment'
            }

            # Build header_to_col mapping using normalized keys
            header_to_col = {}
            for col in range(1, ws.max_column + 1):
                raw = ws.cell(row=1, column=col).value
                if not raw:
                    continue
                nk = _normalize(str(raw))
                header_to_col[nk] = col

            # Determine which canonical keys are present
            canonical_present = set()
            for nk in header_to_col.keys():
                if nk in synonyms and synonyms[nk]:
                    canonical_present.add(synonyms[nk])
                else:
                    canonical_present.add(nk)

            # Required canonical keys for importing students
            required = set(['lrn', 'first_name', 'last_name', 'grade', 'section', 'status'])
            missing = [r for r in required if r not in canonical_present]
            if missing:
                expected_examples = [
                    'ID', 'LRN', 'First Name', 'Middle Name', 'Last Name', 'Grade', 'Section', 'Status',
                    'Guardian First Name', 'Guardian Last Name', 'Guardian Phone', 'Guardian Relationship', 'Alignment'
                ]
                return JsonResponse({
                    'status': 'error',
                    'message': f'Missing required columns: {", ".join(missing)}. Expected (examples): {", ".join(expected_examples)}'
                }, status=400)

            # Helper to get value by canonical key for a given row
            def get_val(row_num, canonical_key):
                # try direct matches (normalized)
                for nk, col in header_to_col.items():
                    mapped = synonyms.get(nk, None)
                    if mapped == canonical_key or nk == canonical_key:
                        val = ws.cell(row=row_num, column=col).value
                        return str(val).strip() if val is not None else ''
                # also try matching canonical_key against synonym keys (e.g., 'first_name' -> 'firstname')
                for syn, mapped in synonyms.items():
                    if mapped == canonical_key and syn in header_to_col:
                        val = ws.cell(row=row_num, column=header_to_col[syn]).value
                        return str(val).strip() if val is not None else ''
                return ''

            # Process rows using mapping
            students_data = []
            errors = []
            skipped_rows = []

            # Get valid grades and sections for validation
            valid_grades = {grade.name: grade for grade in Grade.objects.all()}
            valid_sections = {section.name: section for section in Section.objects.select_related('grade').all()}
            valid_relationships = ['FATHER', 'MOTHER', 'GUARDIAN', 'GRANDMOTHER', 'GRANDFATHER', 'AUNT', 'UNCLE', 'SIBLING', 'OTHER']
            existing_lrns = set(Student.objects.values_list('lrn', flat=True))

            for row_num in range(2, ws.max_row + 1):
                # Read values using canonical keys
                student_id = get_val(row_num, 'id')
                lrn = get_val(row_num, 'lrn')
                first_name = get_val(row_num, 'first_name')
                middle_name = get_val(row_num, 'middle_name')
                last_name = get_val(row_num, 'last_name')
                grade_name = get_val(row_num, 'grade')
                section_name = get_val(row_num, 'section')
                status = get_val(row_num, 'status')
                guardian_first = get_val(row_num, 'guardian_first_name')
                guardian_middle = get_val(row_num, 'guardian_middle_name')
                guardian_last = get_val(row_num, 'guardian_last_name')
                guardian_email = get_val(row_num, 'guardian_email')
                guardian_phone = get_val(row_num, 'guardian_phone')
                guardian_relationship = get_val(row_num, 'guardian_relationship')
                alignment = get_val(row_num, 'alignment')

                # Skip empty rows (based on core student columns)
                if not any([lrn, first_name, last_name, grade_name, section_name, status]):
                    continue

                row_errors = []

                # Skip rows with existing ID if provided
                if student_id and student_id not in ['nan', '']:
                    try:
                        if Student.objects.filter(id=int(float(student_id))).exists():
                            continue
                    except Exception:
                        pass

                # Validate required student fields
                if not lrn:
                    row_errors.append('LRN is required')
                elif not re.match(r'^\d{12}$', lrn):
                    row_errors.append('LRN must be exactly 12 digits')
                elif lrn in existing_lrns:
                    # Instead of failing validation, skip rows whose LRN already exists
                    skipped_rows.append({
                        'row': row_num,
                        'lrn': lrn,
                        'reason': 'LRN already exists'
                    })
                    # Do not process this row further
                    continue

                if not first_name:
                    row_errors.append('First Name is required')
                if not last_name:
                    row_errors.append('Last Name is required')
                if not grade_name:
                    row_errors.append('Grade is required')
                elif grade_name not in valid_grades:
                    row_errors.append(f"Invalid Grade '{grade_name}'")

                if not section_name:
                    row_errors.append('Section is required')
                elif section_name not in valid_sections:
                    row_errors.append(f"Invalid Section '{section_name}'")
                elif grade_name and grade_name in valid_grades and section_name in valid_sections:
                    section_obj = valid_sections[section_name]
                    grade_obj = valid_grades[grade_name]
                    if section_obj.grade.id != grade_obj.id:
                        row_errors.append(f"Section '{section_name}' does not belong to grade '{grade_name}'")

                # Normalize status to accept Title Case and common variants
                def _normalize_status(s):
                    if not s:
                        return ''
                    sk = str(s).strip().upper()
                    if sk in ['ACTIVE', 'TRUE', '1', 'YES', 'A']:
                        return 'ACTIVE'
                    if sk in ['INACTIVE', 'FALSE', '0', 'NO', 'I']:
                        return 'INACTIVE'
                    return sk

                status = _normalize_status(status)
                if status not in ['ACTIVE', 'INACTIVE']:
                    row_errors.append("Status must be 'ACTIVE' or 'INACTIVE'")

                # Validate guardian data if provided
                guardian_data = None
                has_guardian_data = any([guardian_first, guardian_last, guardian_phone, guardian_relationship])
                if has_guardian_data:
                    if not guardian_first:
                        row_errors.append('Guardian First Name is required if guardian data provided')
                    if not guardian_last:
                        row_errors.append('Guardian Last Name is required if guardian data provided')
                    if not guardian_phone:
                        row_errors.append('Guardian Phone is required if guardian data provided')
                    elif not re.match(r'^09\d{9}$', guardian_phone):
                        row_errors.append('Guardian Phone must be valid format (09XXXXXXXXX)')
                    if not guardian_relationship:
                        row_errors.append('Guardian Relationship is required if guardian data provided')
                    elif guardian_relationship not in valid_relationships:
                        row_errors.append(f"Invalid Guardian Relationship '{guardian_relationship}'")
                    if guardian_email and not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', guardian_email):
                        row_errors.append('Invalid Guardian Email format')
                    if not row_errors:
                        guardian_data = {
                            'first_name': guardian_first,
                            'middle_name': guardian_middle,
                            'last_name': guardian_last,
                            'email': guardian_email if guardian_email else '',
                            'phone': guardian_phone,
                            'relationship': guardian_relationship
                        }

                if row_errors:
                    errors.append(f"Row {row_num}: {'; '.join(row_errors)}")
                else:
                    students_data.append({
                        'lrn': lrn,
                        'first_name': first_name,
                        'middle_name': middle_name,
                        'last_name': last_name,
                        'grade': valid_grades[grade_name],
                        'section': valid_sections[section_name],
                        'status': status,
                        'alignment': alignment if alignment else '',
                        'guardian': guardian_data
                    })
                    existing_lrns.add(lrn)
            
            # Return errors if any
            if errors:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Validation failed',
                    'errors': errors[:10]  # Limit to first 10 errors
                }, status=400)
            
            if not students_data:
                # If we have skipped rows, return success with zero created but include skipped info
                if skipped_rows:
                    return JsonResponse({
                        'status': 'success',
                        'message': 'No new student data to import; existing rows were skipped',
                        'students_created': 0,
                        'guardians_created': 0,
                        'skipped_rows': skipped_rows
                    })
                return JsonResponse({
                    'status': 'error',
                    'message': 'No valid student data found in file'
                }, status=400)
            
            # Import data in transaction
            with transaction.atomic():
                created_students = 0
                created_guardians = 0

                for student_info in students_data:
                    guardian_info = student_info.pop('guardian', None)

                    # Filter out any keys that are not actual Student model fields
                    model_field_names = {f.name for f in Student._meta.get_fields()
                                         if getattr(f, 'concrete', True) and not (hasattr(f, 'many_to_many') and f.many_to_many) and not getattr(f, 'one_to_many', False)}
                    create_kwargs = {k: v for k, v in student_info.items() if k in model_field_names}

                    # Create student using only valid model fields
                    student = Student.objects.create(**create_kwargs)
                    created_students += 1

                    # Create guardian if data provided
                    if guardian_info:
                        Guardian.objects.create(
                            student=student,
                            **guardian_info
                        )
                        created_guardians += 1
            
            response = {
                'status': 'success',
                'message': f'Successfully imported {created_students} student(s) and {created_guardians} guardian(s)',
                'students_created': created_students,
                'guardians_created': created_guardians
            }
            if skipped_rows:
                response['skipped_rows'] = skipped_rows
            return JsonResponse(response)
            
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': f'Error processing file: {str(e)}'
            }, status=500)
            
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Error importing students: {str(e)}'
        }, status=500)

@login_required
@user_passes_test(is_admin)
def download_excused_import_template(request):
    """Generate Excel template for importing excused absences"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Excused Absences"
        
        # Headers
        headers = ['ID', 'Student LRN', 'Start Date', 'End Date', 'Reason', 'Status', 'Notes']
        
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.value = header
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Add instruction row
        instruction_values = [
            'Leave blank for new',
            '12 digits (required)',
            'YYYY-MM-DD',
            'YYYY-MM-DD',
            'Text reason',
            'ACTIVE, UPCOMING, or EXPIRED',
            'Optional notes'
        ]
        
        for col_num, instruction in enumerate(instruction_values, 1):
            cell = ws.cell(row=2, column=col_num)
            cell.value = instruction
            cell.font = Font(italic=True, color="808080")
            cell.fill = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 10
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 15
        ws.column_dimensions['E'].width = 30
        ws.column_dimensions['F'].width = 15
        ws.column_dimensions['G'].width = 40
        
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        response = HttpResponse(
            output.read(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename="excused_absences_import_template.xlsx"'
        return response
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Error generating template: {str(e)}'}, status=500)

@login_required
@user_passes_test(is_admin)
@require_http_methods(["POST"])
def import_excused_absences(request):
    """Import excused absences from Excel file"""
    try:
        import openpyxl
        from datetime import datetime, date
        from django.db import transaction
        # Accept both 'import_file' and legacy 'file' keys from different templates
        if 'import_file' in request.FILES:
            excel_file = request.FILES['import_file']
        elif 'file' in request.FILES:
            excel_file = request.FILES['file']
        else:
            return JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)
        
        # Validate file type
        if not excel_file.name.endswith(('.xlsx', '.xls', '.csv')):
            return JsonResponse({
                'status': 'error',
                'message': 'Invalid file type. Please upload Excel (.xlsx, .xls) or CSV files only.'
            }, status=400)
        
        # Validate file size (max 10MB)
        if excel_file.size > 10 * 1024 * 1024:
            return JsonResponse({
                'status': 'error',
                'message': 'File too large. Maximum file size is 10MB.'
            }, status=400)
        
        try:
            # Read Excel file
            wb = openpyxl.load_workbook(excel_file)
            ws = wb.active

            # Read header row (all columns) and build flexible mapping
            headers = []
            for col in range(1, ws.max_column + 1):
                cell_value = ws.cell(row=1, column=col).value
                if cell_value:
                    headers.append(str(cell_value).strip())

            # Helper to normalize header names
            import re
            def _normalize(h):
                return re.sub(r'[^a-z0-9]', '', h.strip().lower()) if h else ''

            # Synonyms mapping for excused import
            # Map common header variants to canonical keys used by this import
            synonyms = {
                'id': 'id',
                'studentlrn': 'lrn', 'lrn': 'lrn', 'learnerreference': 'lrn', 'learnerref': 'lrn',
                'absencedate': 'date_absent', 'absencedate': 'date_absent', 'dateabsent': 'date_absent', 'date': 'date_absent',
                'startdate': 'effective_date', 'start date': 'effective_date', 'effectivedate': 'effective_date', 'effective date': 'effective_date',
                'enddate': 'end_date', 'end date': 'end_date',
                'reason': 'reason', 'status': 'status', 'notes': 'notes',
                'excuseletter': 'excuse_letter', 'excuse letter': 'excuse_letter', 'letter': 'excuse_letter'
            }

            # Build header->column mapping using normalized keys
            header_to_col = {}
            for col in range(1, ws.max_column + 1):
                raw = ws.cell(row=1, column=col).value
                if not raw:
                    continue
                nk = _normalize(str(raw))
                header_to_col[nk] = col

            # Required canonical keys: LRN, effective_date, end_date
            required_targets = set([_normalize('Student LRN'), _normalize('Effective Date'), _normalize('End Date')])

            # Determine present canonical keys (accept synonyms)
            present = set()
            for nk in header_to_col.keys():
                if nk in synonyms and synonyms[nk]:
                    mapped = synonyms[nk]
                    present.add(_normalize(mapped))
                else:
                    # if header already matches common canonical form
                    present.add(nk)

            missing = [k for k in required_targets if k not in present]
            if missing:
                # Build friendly expected headers message
                expected_examples = ['ID', 'Student LRN', 'Absence Date', 'Effective Date', 'End Date', 'Reason', 'Status', 'Notes']
                return JsonResponse({
                    'status': 'error',
                    'message': f"Missing required columns. Expected one of: Student LRN, Effective Date, End Date. Example headers: {', '.join(expected_examples)}"
                }, status=400)

            # Helper to get cell value by canonical mapping
            def get_val(row_num, canonical_targets):
                # canonical_targets: list of possible normalized keys in order
                for target in canonical_targets:
                    # direct match
                    if target in header_to_col:
                        val = ws.cell(row=row_num, column=header_to_col[target]).value
                        return str(val).strip() if val is not None else ''
                    # try synonyms mapping
                    for nk, col in header_to_col.items():
                        if nk in synonyms and synonyms[nk] == target:
                            val = ws.cell(row=row_num, column=col).value
                            return str(val).strip() if val is not None else ''
                return ''

            # Prepare validation helpers
            valid_students = {student.lrn: student for student in Student.objects.all()}
            excused_data = []
            errors = []

            # Iterate rows
            for row_num in range(2, ws.max_row + 1):
                # Read common fields using mapping
                excused_id = get_val(row_num, [_normalize('ID')])
                student_lrn = get_val(row_num, [_normalize('Student LRN'), 'lrn'])
                # Try multiple header variants for dates
                date_absent_str = get_val(row_num, [_normalize('Absence Date'), _normalize('Date'), _normalize('Start Date')])
                effective_date_str = get_val(row_num, [_normalize('Effective Date'), _normalize('Start Date'), _normalize('StartDate')])
                end_date_str = get_val(row_num, [_normalize('End Date'), _normalize('EndDate')])
                reason = get_val(row_num, [_normalize('Reason')])
                status = get_val(row_num, [_normalize('Status')])
                notes = get_val(row_num, [_normalize('Notes')])
                excuse_letter = get_val(row_num, [_normalize('Excuse Letter'), _normalize('Letter')])

                # Skip completely empty rows
                if not (student_lrn or date_absent_str or effective_date_str or end_date_str):
                    continue

                row_errors = []

                # Skip rows with existing ID if provided
                if excused_id and excused_id not in ['nan', '']:
                    try:
                        if ExcusedAbsence.objects.filter(id=int(float(excused_id))).exists():
                            continue
                    except Exception:
                        pass

                # Validate LRN
                if not student_lrn:
                    row_errors.append('Student LRN is required')
                elif student_lrn not in valid_students:
                    row_errors.append(f"Student with LRN '{student_lrn}' not found")

                # Parse dates
                parsed_effective = None
                parsed_end = None
                parsed_date_absent = None
                if effective_date_str:
                    try:
                        parsed_effective = datetime.strptime(effective_date_str, '%Y-%m-%d').date()
                    except Exception:
                        row_errors.append('Effective Date must be in YYYY-MM-DD format')
                else:
                    row_errors.append('Effective Date is required')

                if end_date_str:
                    try:
                        parsed_end = datetime.strptime(end_date_str, '%Y-%m-%d').date()
                    except Exception:
                        row_errors.append('End Date must be in YYYY-MM-DD format')
                else:
                    row_errors.append('End Date is required')

                if date_absent_str:
                    try:
                        parsed_date_absent = datetime.strptime(date_absent_str, '%Y-%m-%d').date()
                    except Exception:
                        # Do not fail if absence date invalid; we'll fallback to effective date later
                        parsed_date_absent = None

                # Validate date logic
                if parsed_effective and parsed_end and parsed_effective > parsed_end:
                    row_errors.append('Effective Date cannot be after End Date')

                if row_errors:
                    errors.append(f"Row {row_num}: {'; '.join(row_errors)}")
                else:
                    # If absence date not provided, default to effective_date
                    if not parsed_date_absent:
                        parsed_date_absent = parsed_effective

                    excused_data.append({
                        'student': valid_students[student_lrn],
                        'date_absent': parsed_date_absent,
                        'excuse_letter': excuse_letter if excuse_letter else '',
                        'effective_date': parsed_effective,
                        'end_date': parsed_end
                    })

            # Return errors if any
            if errors:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Validation failed',
                    'errors': errors[:10]
                }, status=400)

            if not excused_data:
                return JsonResponse({
                    'status': 'error',
                    'message': 'No valid excused absence data found in file'
                }, status=400)

            # Import data in transaction
            with transaction.atomic():
                created_count = 0

                for excused_info in excused_data:
                    ExcusedAbsence.objects.create(**excused_info)
                    created_count += 1

            return JsonResponse({
                'status': 'success',
                'message': f'Successfully imported {created_count} excused absence(s)',
                'created_count': created_count
            })

        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': f'Error processing file: {str(e)}'
            }, status=500)
            
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Error importing excused absences: {str(e)}'
        }, status=500)

@login_required
@user_passes_test(is_admin)
def download_attendance_import_template(request):
    """Generate Excel template for importing attendance records"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Attendance"
        
        # Headers
        headers = ['ID', 'Student LRN', 'Date', 'Time In', 'Time Out', 'Status']
        
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.value = header
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Add instruction row
        instruction_values = [
            'Leave blank for new',
            '12 digits (required)',
            'YYYY-MM-DD',
            'HH:MM (optional)',
            'HH:MM (optional)',
            'ON TIME, LATE, ABSENT, or EXCUSED'
        ]
        
        for col_num, instruction in enumerate(instruction_values, 1):
            cell = ws.cell(row=2, column=col_num)
            cell.value = instruction
            cell.font = Font(italic=True, color="808080")
            cell.fill = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 10
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 12
        ws.column_dimensions['E'].width = 12
        ws.column_dimensions['F'].width = 20
        
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        response = HttpResponse(
            output.read(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename="attendance_import_template.xlsx"'
        return response
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Error generating template: {str(e)}'}, status=500)

@login_required
@user_passes_test(is_admin_or_teacher)
@require_http_methods(["POST"])
def import_attendance(request):
    """Import attendance records from Excel file"""
    try:
        import openpyxl
        from datetime import datetime, time
        from django.db import transaction
        
        if 'import_file' not in request.FILES:
            return JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)
        
        excel_file = request.FILES['import_file']
        
        # Validate file type
        if not excel_file.name.endswith(('.xlsx', '.xls', '.csv')):
            return JsonResponse({
                'status': 'error',
                'message': 'Invalid file type. Please upload Excel (.xlsx, .xls) or CSV files only.'
            }, status=400)
        
        # Validate file size (max 10MB)
        if excel_file.size > 10 * 1024 * 1024:
            return JsonResponse({
                'status': 'error',
                'message': 'File too large. Maximum file size is 10MB.'
            }, status=400)
        
        try:
            # Read Excel file
            wb = openpyxl.load_workbook(excel_file)
            ws = wb.active
            
            # Read header row (all columns) and build flexible mapping
            headers = []
            for col in range(1, ws.max_column + 1):
                cell_value = ws.cell(row=1, column=col).value
                if cell_value:
                    headers.append(str(cell_value).strip())

            # Helper to normalize header names
            import re
            def _normalize(h):
                return re.sub(r'[^a-z0-9]', '', h.strip().lower()) if h else ''

            # Canonical expected headers for attendance (human-friendly)
            expected_headers = ['ID', 'Student LRN', 'Date', 'Time In', 'Time Out', 'Status']
            normalized_expected = {_normalize(h): h for h in expected_headers}

            # Common synonyms for flexible header matching
            synonyms = {
                'id': 'id',
                'studentlrn': 'lrn', 'lrn': 'lrn', 'learnerreference': 'lrn', 'learnerref': 'lrn',
                'date': 'date', 'attendance_date': 'date',
                'timein': 'timein', 'timein24': 'timein', 'timeinhhmm': 'timein',
                'timeout': 'timeout', 'timeouthhmm': 'timeout',
                'status': 'status', 'attendance': 'status', 'presentstatus': 'status'
            }

            # Build header->column mapping
            header_to_col = {}
            for col in range(1, ws.max_column + 1):
                raw = ws.cell(row=1, column=col).value
                if not raw:
                    continue
                nk = _normalize(str(raw))
                header_to_col[nk] = col

            # Determine required minimal headers (LRN, Date, Status)
            required_keys = set([_normalize('Student LRN'), _normalize('Date'), _normalize('Status')])

            # Figure out which required headers are present (accept synonyms)
            present_keys = set()
            for nk in header_to_col.keys():
                if nk in normalized_expected:
                    present_keys.add(nk)
                elif nk in synonyms:
                    # map synonym to normalized target
                    mapped = synonyms[nk]
                    mapped_nk = _normalize(mapped) if mapped else ''
                    if mapped_nk in normalized_expected:
                        present_keys.add(mapped_nk)

            missing_required = [normalized_expected[k] for k in required_keys if k not in present_keys]
            if missing_required:
                return JsonResponse({
                    'status': 'error',
                    'message': f'Missing required columns: {", ".join(missing_required)}. Expected headers (examples): {", ".join(expected_headers)}'
                }, status=400)

            # Build helper to get cell value by canonical key
            def get_val(row_num, canonical):
                # canonical: 'id','lrn','date','timein','timeout','status'
                # try direct matches
                for nk, col in header_to_col.items():
                    if nk == canonical:
                        val = ws.cell(row=row_num, column=col).value
                        return str(val).strip() if val is not None else ''
                # try synonyms
                for nk, col in header_to_col.items():
                    if nk in synonyms and synonyms[nk] == canonical:
                        val = ws.cell(row=row_num, column=col).value
                        return str(val).strip() if val is not None else ''
                return ''

            # Prepare validation helpers
            valid_students = {student.lrn: student for student in Student.objects.all()}
            attendance_data = []
            errors = []

            # Accept these canonical status values (map common variants)
            def normalize_status(s):
                if not s:
                    return ''
                sk = s.strip().upper()
                if sk in ['PRESENT', 'ON TIME', 'ONTIME']:
                    return 'ON TIME'
                if sk == 'LATE':
                    return 'LATE'
                if sk in ['ABSENT', 'A']:
                    return 'ABSENT'
                if sk in ['EXCUSED', 'E']:
                    return 'EXCUSED'
                return sk

            for row_num in range(2, ws.max_row + 1):
                # read values using header mapping
                attendance_id = get_val(row_num, _normalize('ID'))
                student_lrn = get_val(row_num, _normalize('Student LRN'))
                date_str = get_val(row_num, _normalize('Date'))
                time_in_str = get_val(row_num, _normalize('Time In'))
                time_out_str = get_val(row_num, _normalize('Time Out'))
                status_raw = get_val(row_num, _normalize('Status'))

                # Skip completely empty rows
                if not (student_lrn or date_str or status_raw or time_in_str or time_out_str):
                    continue

                row_errors = []

                # Skip rows with existing attendance ID if provided
                if attendance_id and attendance_id not in ['nan', '']:
                    try:
                        if Attendance.objects.filter(id=int(float(attendance_id))).exists():
                            continue
                    except Exception:
                        pass

                # Validate LRN
                if not student_lrn:
                    row_errors.append('Student LRN is required')
                elif student_lrn not in valid_students:
                    row_errors.append(f"Student with LRN '{student_lrn}' not found")

                # Validate date
                attendance_date = None
                if not date_str:
                    row_errors.append('Date is required')
                else:
                    try:
                        attendance_date = datetime.strptime(date_str, '%Y-%m-%d').date()
                    except Exception:
                        row_errors.append('Date must be in YYYY-MM-DD format')

                # Validate times (optional)
                time_in = None
                time_out = None
                if time_in_str:
                    try:
                        time_in = datetime.strptime(time_in_str, '%H:%M').time()
                    except Exception:
                        row_errors.append('Time In must be in HH:MM format')
                if time_out_str:
                    try:
                        time_out = datetime.strptime(time_out_str, '%H:%M').time()
                    except Exception:
                        row_errors.append('Time Out must be in HH:MM format')

                # Validate status
                status = normalize_status(status_raw)
                if not status:
                    row_errors.append('Status is required')
                elif status not in ['ON TIME', 'LATE', 'ABSENT', 'EXCUSED']:
                    row_errors.append('Status must be one of: ON TIME, LATE, ABSENT, EXCUSED')

                if row_errors:
                    errors.append("Row {}: {}".format(row_num, '; '.join(row_errors)))
                else:
                    attendance_data.append({
                        'student': valid_students[student_lrn],
                        'date': attendance_date,
                        'time_in': time_in,
                        'time_out': time_out,
                        'status': status
                    })
            
            # Return errors if any
            if errors:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Validation failed',
                    'errors': errors[:10]  # Limit to first 10 errors
                }, status=400)
            
            if not attendance_data:
                return JsonResponse({
                    'status': 'error',
                    'message': 'No valid attendance data found in file'
                }, status=400)
            
            # Import data in transaction
            with transaction.atomic():
                created_count = 0
                skipped_count = 0
                # track (student_id, date) pairs already created in this import to avoid duplicates within file
                created_pairs = set()

                for attendance_info in attendance_data:
                    student = attendance_info['student']
                    date = attendance_info['date']
                    pair = (student.id, date)

                    # If an attendance with same student and date already exists in DB, skip
                    if Attendance.objects.filter(student=student, date=date).exists():
                        skipped_count += 1
                        continue

                    # If we've already created a record for this pair in this import, skip duplicate row
                    if pair in created_pairs:
                        skipped_count += 1
                        continue

                    # Create new attendance
                    Attendance.objects.create(**attendance_info)
                    created_pairs.add(pair)
                    created_count += 1

            return JsonResponse({
                'status': 'success',
                'message': f'Successfully imported {created_count} attendance record(s)',
                'created_count': created_count,
                'skipped_count': skipped_count
            })
            
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': f'Error processing file: {str(e)}'
            }, status=500)
            
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Error importing attendance: {str(e)}'
        }, status=500)

